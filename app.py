# -*- coding: utf-8 -*-
"""
================================================================
  실험 데이터 자동 통계 분석 시스템  (스마트 통계 에이전트)
================================================================
실행: streamlit run app.py
"""
import io, copy
from collections import Counter
import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import seaborn as sns
from lxml import etree
from scipy import stats
from scipy.stats import studentized_range
import statsmodels.api as sm
from statsmodels.formula.api import ols
from statsmodels.stats.multicomp import pairwise_tukeyhsd
import scikit_posthocs as sp
from sklearn.linear_model import LogisticRegression, Ridge, Lasso, ElasticNet
from sklearn.ensemble import (RandomForestClassifier, RandomForestRegressor,
                              ExtraTreesClassifier, ExtraTreesRegressor,
                              GradientBoostingClassifier, GradientBoostingRegressor,
                              HistGradientBoostingClassifier, HistGradientBoostingRegressor,
                              AdaBoostClassifier, AdaBoostRegressor)
from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor
from sklearn.svm import SVC, SVR
from sklearn.neighbors import KNeighborsClassifier, KNeighborsRegressor
from sklearn.naive_bayes import GaussianNB
from sklearn.pipeline import Pipeline
from sklearn.inspection import permutation_importance
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler, MinMaxScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, r2_score
try:
    import anthropic
    _HAS_ANTHROPIC = True
except Exception:
    _HAS_ANTHROPIC = False
try:
    import requests as _requests
    _HAS_REQUESTS = True
except Exception:
    _HAS_REQUESTS = False
try:
    import docx as _docx_probe  # 설치 여부 확인용
    _ = _docx_probe
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False

_KOREAN_FONT = None


def set_korean_font():
    """그래프 한글 폰트를 잡는다. (윈도우·맥·리눅스 서버 배포 모두 대응)

    리눅스 서버(클라우드 배포)에서는 한글 폰트를 나중에 설치하는 경우가 많아
    matplotlib 캐시에 폰트가 없을 수 있다. 그때는 폰트 파일을 직접 등록하고
    캐시를 다시 만들어 본다. 이걸 안 하면 그래프 글자가 전부 □로 나온다.
    """
    global _KOREAN_FONT
    cands = ["Malgun Gothic", "AppleGothic", "NanumGothic", "NanumBarunGothic",
             "Noto Sans CJK KR", "Noto Sans KR", "UnDotum"]
    names = {f.name for f in fm.fontManager.ttflist}
    if not names & set(cands):
        import glob as _glob
        for _pat in ("/usr/share/fonts/**/Nanum*.tt[fc]",
                     "/usr/share/fonts/**/NotoSansCJK*.tt[fc]",
                     "/usr/share/fonts/**/NotoSansKR*.tt[fc]"):
            for _p in _glob.glob(_pat, recursive=True):
                try:
                    fm.fontManager.addfont(_p)
                except Exception:
                    pass
        names = {f.name for f in fm.fontManager.ttflist}
        if not names & set(cands):
            try:                      # 마지막 수단: 폰트 캐시를 통째로 다시 만든다
                fm._load_fontmanager(try_read_cache=False)
                names = {f.name for f in fm.fontManager.ttflist}
            except Exception:
                pass
    for c in cands:
        if c in names:
            plt.rcParams["font.family"] = c
            _KOREAN_FONT = c
            break
    plt.rcParams["axes.unicode_minus"] = False
set_korean_font()

# ================================================================ 공통 그래프 디자인
# 원클릭 보고서의 차분한 블루 톤을 앱 전체 그래프의 기본값으로 사용한다.
# 개별 그래프에서 별도 색을 지정하지 않아도 같은 분위기로 보이도록 rcParams에 반영한다.
from cycler import cycler as _cycler
_SMART_CHART_BLUE = ["#DCE9F5", "#C2D9EE", "#A3C4E2", "#82ACD3",
                     "#6291C2", "#4576AB", "#2D5A8E", "#1F4569"]
_SMART_CHART_RED = "#C96767"
_SMART_CHART_RED_LIGHT = "#E9B5B5"
_SMART_CHART_NEUTRAL = "#E3E9EF"
plt.rcParams.update({
    "figure.facecolor": "white",
    "axes.facecolor": "white",
    "savefig.facecolor": "white",
    "axes.edgecolor": "#AEBECD",
    "axes.spines.top": False,
    "axes.spines.right": False,
    "axes.labelcolor": "#43576A",
    "axes.titlecolor": "#23394D",
    "axes.titleweight": "bold",
    "axes.titlesize": 11.5,
    "axes.labelsize": 9.5,
    "xtick.color": "#5B6F82",
    "ytick.color": "#5B6F82",
    "xtick.labelsize": 9,
    "ytick.labelsize": 9,
    "grid.color": "#DCE7F0",
    "grid.linewidth": 0.8,
    "grid.alpha": 0.9,
    "legend.frameon": False,
    "legend.fontsize": 8,
    "axes.prop_cycle": _cycler(color=["#4576AB", "#6F9FC8", "#91B8D8", "#2D5A8E",
                                       "#B2CEE5", "#5A86B3", "#789FC3", "#355F8A"]),
})

st.set_page_config(page_title="스마트 통계 에이전트", page_icon="📊", layout="wide")


# ================================================================ 공통 화면/엑셀 표 디자인
# 모든 분석 화면이 같은 "스마트 블루" 표 디자인을 쓰도록 한 곳에서 관리한다.
# 나중에 색을 바꾸고 싶으면 아래 팔레트만 수정하면 앱 전체에 반영된다.
_ST_DATAFRAME = st.dataframe
_SMART_BLUE = {
    "navy": "#244A73", "header": "#3D6F9F", "mid": "#9EC5E5",
    "light": "#EAF3FA", "pale": "#F7FBFF", "line": "#C9DCEB",
}

# Streamlit 기본 표 주변도 카드처럼 보이게 한다. 셀 색은 pandas Styler가 담당한다.
st.markdown("""
<style>
[data-testid="stDataFrame"] {border:1px solid #d7e5f1; border-radius:10px; overflow:hidden;}
[data-testid="stDataEditor"] {border:1px solid #d7e5f1; border-radius:10px; overflow:hidden;}
</style>
""", unsafe_allow_html=True)


def _smart_excluded_gradient_col(name):
    """값의 크기가 '좋고 나쁨'을 뜻하지 않는 통계 열은 값 기반 그라데이션에서 제외."""
    s = str(name).lower().replace(" ", "")
    keys = ("p-value", "pvalue", "p값", "p(", "유의", "통계량", "t값", "t통계",
            "f값", "f통계", "df", "자유도", "ci", "신뢰구간", "표준오차", "se(",
            "표준편차", "sd(", "검정", "판정", "반복", "번호", "순번")
    return any(k in s for k in keys)


def smart_table(data, *args, **kwargs):
    """st.dataframe 호환 래퍼.

    - 원클릭 보고서의 푸른 계열 분위기를 모든 표에 통일한다.
    - 일반 결과표에는 값 크기에 따른 자동 색상(그라데이션)을 넣지 않는다.
      숫자가 크다는 이유만으로 더 중요하거나 더 좋은 값처럼 보이는 오해를 막기 위함이다.
    - 머리행과 아주 옅은 행 구분만 유지한다.
    - 이미 Styler가 넘어온 경우(결측치 강조 등) 기존 의미 기반 스타일은 보존한다.
    """
    try:
        is_styler = data.__class__.__name__ == "Styler"
        if is_styler:
            sty = data
            try:
                sty = sty.set_table_styles([
                    {"selector": "th", "props": [("background-color", _SMART_BLUE["header"]),
                                                    ("color", "white"), ("font-weight", "700"),
                                                    ("border", f"1px solid {_SMART_BLUE['line']}")]},
                    {"selector": "td", "props": [("border", f"1px solid {_SMART_BLUE['line']}")]},
                ], overwrite=False)
            except Exception:
                pass
            return _ST_DATAFRAME(sty, *args, **kwargs)

        if isinstance(data, pd.DataFrame):
            shown = sup_display(data)
            sty = shown.style
            try:
                # 연한 행 구분 + 파란 머리행
                def _band_rows(row):
                    bg = _SMART_BLUE["pale"] if (row.name % 2 == 0 if isinstance(row.name, (int, np.integer)) else False) else "white"
                    return [f"background-color:{bg}" for _ in row]
                sty = sty.apply(_band_rows, axis=1)
            except Exception:
                pass

            try:
                sty = sty.set_table_styles([
                    {"selector": "th", "props": [("background-color", _SMART_BLUE["header"]),
                                                    ("color", "white"), ("font-weight", "700"),
                                                    ("border", f"1px solid {_SMART_BLUE['line']}")]},
                    {"selector": "td", "props": [("border", f"1px solid {_SMART_BLUE['line']}")]},
                ], overwrite=False)
            except Exception:
                pass
            return _ST_DATAFRAME(sty, *args, **kwargs)
    except Exception:
        pass
    return _ST_DATAFRAME(data, *args, **kwargs)


def dataframe_to_styled_xlsx(df, title="스마트 통계 에이전트 분석 결과", sheet_name="분석결과"):
    """화면의 스마트 블루 디자인을 실제 .xlsx에도 반영한다."""
    from openpyxl import Workbook
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.formatting.rule import ColorScaleRule
    from openpyxl.utils import get_column_letter
    import datetime as _dt

    frame = sup_display(df.copy() if isinstance(df, pd.DataFrame) else pd.DataFrame(df))
    wb = Workbook()
    ws = wb.active
    ws.title = str(sheet_name)[:31] or "분석결과"
    ws.sheet_properties.tabColor = "3D6F9F"
    ws.sheet_view.showGridLines = False
    ncol = max(len(frame.columns), 1)

    # 제목/메타
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncol)
    c = ws.cell(1, 1, title)
    c.fill = PatternFill("solid", fgColor="244A73")
    c.font = Font(color="FFFFFF", bold=True, size=14)
    c.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 27
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncol)
    c2 = ws.cell(2, 1, f"생성일: {_dt.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    c2.font = Font(color="5B6F82", size=9, italic=True)
    c2.fill = PatternFill("solid", fgColor="F7FBFF")
    c2.alignment = Alignment(horizontal="left")

    header_row = 4
    thin = Side(style="thin", color="C9DCEB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for j, col in enumerate(frame.columns, 1):
        cell = ws.cell(header_row, j, str(col))
        cell.fill = PatternFill("solid", fgColor="3D6F9F")
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    ws.row_dimensions[header_row].height = 24

    for i, row in enumerate(frame.itertuples(index=False, name=None), header_row + 1):
        for j, val in enumerate(row, 1):
            cell = ws.cell(i, j)
            if pd.isna(val):
                cell.value = None
            elif isinstance(val, (np.integer,)):
                cell.value = int(val)
            elif isinstance(val, (np.floating,)):
                cell.value = float(val)
            else:
                cell.value = val
            cell.border = border
            cell.alignment = Alignment(vertical="center")
            if (i - header_row) % 2 == 0:
                cell.fill = PatternFill("solid", fgColor="F7FBFF")
            if isinstance(cell.value, (int, float)) and not isinstance(cell.value, bool):
                cell.number_format = '#,##0.###'

    if len(frame) > 0 and len(frame.columns) > 0:
        # Excel 내장 TableStyle은 Office 테마에 따라 색이 달라져 앱 화면과 어긋날 수 있다.
        # 필터 기능만 유지하고 셀 색은 전부 스마트 블루 팔레트로 직접 지정한다.
        ref = f"A{header_row}:{get_column_letter(len(frame.columns))}{header_row + len(frame)}"
        ws.auto_filter.ref = ref

    ws.freeze_panes = f"A{header_row + 1}"
    ws.auto_filter.ref = (f"A{header_row}:{get_column_letter(len(frame.columns))}{header_row + len(frame)}"
                          if len(frame.columns) and len(frame) else None)
    for j, col in enumerate(frame.columns, 1):
        vals = [str(col)] + ["" if pd.isna(v) else str(v) for v in frame[col].head(200)]
        width = min(max(max((len(v) for v in vals), default=8) + 3, 10), 34)
        ws.column_dimensions[get_column_letter(j)].width = width

    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()

MM = 7200 / 25.4          # 1mm = 283.46 HWPUNIT
BODY_W = int(150 * MM)    # 본문 폭 150mm

# ---------------------------------------------------------------- 통계 헬퍼
def compact_letter_display(means_sorted, not_sig_pairs):
    """유의성 문자(a,b,c) 생성. Piepho(2004) insert-and-absorb 방식.
    means_sorted: 평균 내림차순 그룹 리스트, not_sig_pairs: 유의차 없는 쌍(frozenset)들.
    같은 문자를 공유하면 두 처리 간 유의차가 없음을 뜻한다(ab, bc 등 중간 그룹 지원)."""
    def diff(a, b):  # 유의차 있음 = 다른 문자여야 함
        return a != b and frozenset({a, b}) not in not_sig_pairs
    items = list(means_sorted)
    if not items:
        return {}
    cols = [set(items)]  # 모든 그룹을 한 열에서 시작
    changed = True
    while changed:
        changed = False
        for col in list(cols):
            broke = False
            members = list(col)
            for i in range(len(members)):
                for j in range(i + 1, len(members)):
                    a, b = members[i], members[j]
                    if diff(a, b):  # 유의차 있는 쌍이 한 열에 → 열을 쪼갬
                        cols.remove(col)
                        c1 = set(col); c1.discard(a)
                        c2 = set(col); c2.discard(b)
                        cols.append(c1); cols.append(c2)
                        changed = True; broke = True
                        break
                if broke: break
            if broke: break
    cols = [c for c in cols if c]
    # 다른 열의 부분집합인 열 흡수(제거)
    keep = [c for i, c in enumerate(cols)
            if not any(i != j and c < o for j, o in enumerate(cols))]
    uniq = []
    for c in keep:
        if c not in uniq: uniq.append(c)
    uniq.sort(key=lambda col: min(items.index(g) for g in col))
    letters = {g: "" for g in items}
    for i, col in enumerate(uniq):
        for g in col: letters[g] += chr(97 + i)
    return {g: "".join(sorted(v)) for g, v in letters.items()}

def _model_emmeans(model, data, treatment_col):
    """적합된 statsmodels 모형에서 처리별 추정주변평균(EMM)과 설계벡터를 계산한다.

    블록이 포함된 RCBD/불균형 자료에서는 원자료 평균이 아니라 모형이 보정한 평균을
    사용해야 ANOVA와 사후검정이 같은 오차구조를 공유한다.
    """
    from patsy import build_design_matrices
    design_info = model.model.data.design_info
    levels = list(pd.unique(data[treatment_col].dropna()))
    rows, means = {}, {}
    params = np.asarray(model.params, dtype=float)
    for level in levels:
        tmp = data.copy()
        tmp[treatment_col] = level
        mat = build_design_matrices([design_info], tmp, return_type="dataframe")[0]
        xbar = np.asarray(mat, dtype=float).mean(axis=0)
        rows[level] = xbar
        means[level] = float(xbar @ params)
    return levels, means, rows


def _simulate_dunnett_adjustment(t_values, corr, df_resid, alpha=0.05,
                                 n_sim=60000, seed=20260726):
    """모형 기반 대조들의 상관을 반영한 Dunnett 단일단계 보정.

    다변량 t 분포를 몬테카를로로 재현해 최대 |t| 분포를 구한다. 난괴법·불균형
    자료에서도 같은 적합모형의 공분산을 사용하므로 블록을 무시한 원자료 Dunnett보다
    ANOVA와 일관된 결과를 낸다.
    """
    tv = np.asarray(t_values, dtype=float)
    if tv.size == 0:
        return np.array([], dtype=float), np.nan
    if tv.size == 1:
        p = 2 * stats.t.sf(abs(tv[0]), df_resid)
        crit = stats.t.ppf(1 - alpha / 2, df_resid)
        return np.array([float(p)]), float(crit)
    corr = np.asarray(corr, dtype=float)
    corr = (corr + corr.T) / 2
    np.fill_diagonal(corr, 1.0)
    # 수치 오차로 음의 고유값이 생기면 아주 작게 보정
    vals, vecs = np.linalg.eigh(corr)
    vals = np.clip(vals, 1e-10, None)
    corr = (vecs * vals) @ vecs.T
    d = np.sqrt(np.diag(corr))
    corr = corr / np.outer(d, d)
    rng = np.random.default_rng(seed)
    z = rng.multivariate_normal(np.zeros(tv.size), corr, size=int(n_sim))
    chi = rng.chisquare(float(df_resid), size=int(n_sim))
    sims = z / np.sqrt(chi[:, None] / float(df_resid))
    max_abs = np.max(np.abs(sims), axis=1)
    p_adj = np.array([(1 + np.count_nonzero(max_abs >= abs(t))) / (len(max_abs) + 1)
                      for t in tv], dtype=float)
    crit = float(np.quantile(max_abs, 1 - alpha))
    return p_adj, crit


def posthoc_from_model(model, data, treatment_col, method, control=None,
                       alpha=0.05, random_state=20260726):
    """적합모형 기반 사후검정.

    반환값: {not_sig, table, means}. Dunnett는 대조구 비교만 table에 담고 CLD는 만들지 않는다.
    Tukey/Bonferroni/Duncan은 모든 처리쌍을 모형의 잔차 공분산으로 비교한다.
    """
    levels, emmeans, xrows = _model_emmeans(model, data, treatment_col)
    covb = np.asarray(model.cov_params(), dtype=float)
    dfe = float(model.df_resid)
    k = len(levels)
    rows = []
    not_sig = set()

    def contrast(a, b):
        c = np.asarray(xrows[a]) - np.asarray(xrows[b])
        diff = float(c @ np.asarray(model.params, dtype=float))
        var = float(c @ covb @ c)
        se = float(np.sqrt(max(var, 0.0)))
        tval = diff / se if se > 0 else (np.inf if diff else 0.0)
        p_raw = float(2 * stats.t.sf(abs(tval), dfe)) if np.isfinite(tval) else 0.0
        return c, diff, se, tval, p_raw

    if method.startswith("던넷") or method.startswith("Dunnett"):
        if control not in levels:
            _match = next((g for g in levels if str(g) == str(control)), None)
            control = _match if _match is not None else (levels[0] if levels else None)
        others = [g for g in levels if g != control]
        if control is None or not others:
            return {"not_sig": set(), "table": pd.DataFrame(), "means": emmeans,
                    "control": control, "method": method}
        contrasts, vals = [], []
        for g in others:
            c, diff, se, tval, p_raw = contrast(g, control)
            contrasts.append(c)
            vals.append((g, diff, se, tval, p_raw))
        cmat = np.vstack(contrasts)
        ccov = cmat @ covb @ cmat.T
        ses = np.sqrt(np.clip(np.diag(ccov), 0, None))
        denom = np.outer(ses, ses)
        corr = np.divide(ccov, denom, out=np.eye(len(others)), where=denom > 0)
        p_adj, crit = _simulate_dunnett_adjustment(
            [v[3] for v in vals], corr, dfe, alpha=alpha, seed=random_state)
        for i, (g, diff, se, tval, p_raw) in enumerate(vals):
            lo = diff - crit * se if np.isfinite(crit) else np.nan
            hi = diff + crit * se if np.isfinite(crit) else np.nan
            padj = float(p_adj[i])
            if padj >= alpha:
                not_sig.add(frozenset({control, g}))
            rows.append({
                "대조구": control, "처리구": g,
                "대조구 평균(보정)": emmeans[control], "처리 평균(보정)": emmeans[g],
                "평균 차이": diff, "t 통계량": tval,
                "p(동시보정)": padj, "95% 동시CI 하한": lo, "95% 동시CI 상한": hi,
                "판정": "유의(*)" if padj < alpha else "n.s.",
            })
        return {"not_sig": not_sig, "table": pd.DataFrame(rows), "means": emmeans,
                "control": control, "method": method, "critical": crit}

    pair_data = []
    ordered = sorted(levels, key=lambda g: emmeans[g], reverse=True)
    m = max(k * (k - 1) // 2, 1)
    for i in range(k):
        for j in range(i + 1, k):
            a, b = levels[i], levels[j]
            c, diff, se, tval, p_raw = contrast(a, b)
            if method == "Tukey HSD":
                q = abs(tval) * np.sqrt(2)
                p_adj = float(studentized_range.sf(q, k, dfe))
                significant = p_adj < alpha
                crit = float(studentized_range.ppf(1 - alpha, k, dfe) / np.sqrt(2))
            elif method == "던컨(Duncan)":
                ia, ib = ordered.index(a), ordered.index(b)
                rng_size = abs(ia - ib) + 1
                alpha_range = 1 - (1 - alpha) ** max(rng_size - 1, 1)
                q = abs(tval) * np.sqrt(2)
                qcrit = float(studentized_range.ppf(1 - alpha_range, rng_size, dfe))
                significant = q > qcrit
                p_adj = float(studentized_range.sf(q, rng_size, dfe))
                crit = qcrit / np.sqrt(2)
            else:  # Bonferroni
                p_adj = min(float(p_raw) * m, 1.0)
                significant = p_adj < alpha
                crit = float(stats.t.ppf(1 - alpha / (2 * m), dfe))
            if not significant:
                not_sig.add(frozenset({a, b}))
            pair_data.append({
                "그룹1": a, "그룹2": b, "평균차": diff, "표준오차": se,
                "t 통계량": tval, "p(보정)": p_adj,
                "95% 하한": diff - crit * se, "95% 상한": diff + crit * se,
                "판정": "유의(*)" if significant else "n.s.",
            })
    return {"not_sig": not_sig, "table": pd.DataFrame(pair_data),
            "means": emmeans, "method": method}


def posthoc_not_sig(data, group_col, value_col, method, alpha=0.05,
                    model=None, control=None):
    """하위호환용 래퍼. model이 주어지면 RCBD/불균형을 반영한 모형 기반 비교를 사용한다."""
    if model is not None:
        return posthoc_from_model(model, data, group_col, method,
                                  control=control, alpha=alpha)["not_sig"]
    not_sig = set()
    if method == "Tukey HSD":
        res = pairwise_tukeyhsd(data[value_col], data[group_col])
        for row in res._results_table.data[1:]:
            if not row[-1]:
                not_sig.add(frozenset({row[0], row[1]}))
    elif method == "던컨(Duncan)":
        groups = data.groupby(group_col)[value_col]
        means = groups.mean().sort_values(ascending=False)
        counts = groups.count()
        k, n_all = len(means), len(data)
        ssw = sum(((groups.get_group(g) - groups.get_group(g).mean()) ** 2).sum()
                  for g in means.index)
        dfe = n_all - k
        mse = ssw / dfe
        nh = k / sum(1 / counts[g] for g in means.index)
        se = np.sqrt(mse / nh)
        order = means.index.tolist()
        for i in range(k):
            for j in range(i + 1, k):
                rng_size = j - i + 1
                alpha_range = 1 - (1 - alpha) ** (rng_size - 1)
                rp = studentized_range.ppf(1 - alpha_range, rng_size, dfe)
                if abs(means[order[i]] - means[order[j]]) <= rp * se:
                    not_sig.add(frozenset({order[i], order[j]}))
    elif method.startswith("던넷") or method.startswith("Dunnett"):
        ctrl = control if control is not None else st.session_state.get("dunnett_ctrl")
        groups = data.groupby(group_col)[value_col]
        names = list(groups.groups.keys())
        if ctrl not in names:
            ctrl = names[0]
        try:
            from scipy.stats import dunnett as _dunnett
            others = [g for g in names if g != ctrl]
            res = _dunnett(*[groups.get_group(g).values for g in others],
                           control=groups.get_group(ctrl).values)
            for g, p in zip(others, np.atleast_1d(res.pvalue)):
                if p >= alpha:
                    not_sig.add(frozenset({ctrl, g}))
        except Exception as ex:
            st.warning(f"던넷 검정을 수행하지 못했습니다: {str(ex)[:80]}")
    else:
        pmat = sp.posthoc_ttest(data, val_col=value_col, group_col=group_col,
                               p_adjust="bonferroni")
        for a in pmat.index:
            for b in pmat.columns:
                if a != b and pmat.loc[a, b] >= alpha:
                    not_sig.add(frozenset({a, b}))
    return not_sig

_SUP_MAP = {"a": "ᵃ", "b": "ᵇ", "c": "ᶜ", "d": "ᵈ", "e": "ᵉ", "f": "ᶠ",
            "g": "ᵍ", "h": "ʰ", "i": "ⁱ", "j": "ʲ", "k": "ᵏ", "l": "ˡ",
            "m": "ᵐ", "n": "ⁿ", "*": "*"}

def sup_text(s):
    """'607.6^a' → '607.6ᵃ' (유니코드 위첨자). 문서 생성이 실패했을 때의 안전망이자
    화면·CSV 표시용 공통 변환기. '^'가 화면이나 문서에 그대로 남지 않게 한다."""
    t = str(s)
    if "^" not in t:
        return s
    base, _, sup = t.partition("^")
    if not sup or any(ch not in _SUP_MAP for ch in sup):
        return s          # 'm^2' 처럼 유의성 문자가 아닌 경우는 건드리지 않는다
    return base + "".join(_SUP_MAP[ch] for ch in sup)

def sup_display(df):
    """화면 표시용: '607.6^a' → '607.6ᵃ' (문서 저장 시엔 ^ 그대로 유지)"""
    try:
        out = df.copy()
    except Exception:
        return df
    for c in out.columns:
        # pandas 3.x는 문자열 열의 dtype이 object가 아니라 str이라, dtype 비교 대신
        # '숫자·날짜가 아니면 훑는다'로 두어야 버전이 올라가도 계속 동작한다.
        try:
            if (pd.api.types.is_numeric_dtype(out[c])
                    or pd.api.types.is_datetime64_any_dtype(out[c])):
                continue
        except Exception:
            pass
        try:
            out[c] = out[c].map(lambda v: sup_text(v) if isinstance(v, str) and "^" in v else v)
        except Exception:
            pass
    out.columns = [sup_text(c) if isinstance(c, str) else c for c in out.columns]
    return out

# ---------------------------------------------------------------- 오류 도우미
def error_help(err, context="", key="err"):
    """오류가 났을 때 (1) 앱 안에서 AI에게 바로 물어보고 답을 화면에 띄우고,
    (2) 구글·ChatGPT 링크도 함께 제공한다.

    스트림릿이 기본으로 붙여주는 구글/ChatGPT 링크는 '검색창에 붙여넣기'까지만 해 준다.
    (ChatGPT 쪽이 자동 전송을 막아서 예전처럼 바로 답이 뜨지 않는다.)
    그래서 앱 안에서 바로 답을 받는 버튼을 따로 만든다.
    """
    import urllib.parse as _up, traceback as _tb
    msg = (f"{type(err).__name__}: {err}" if isinstance(err, BaseException) else str(err))
    trace = ""
    if isinstance(err, BaseException):
        try:
            trace = "".join(_tb.format_exception(type(err), err, err.__traceback__))
        except Exception:
            trace = ""
    prompt = ("파이썬 Streamlit 앱에서 아래 오류가 났습니다. "
              "원인을 한국어로 쉽게 설명하고, 사용자가 바로 할 수 있는 해결 방법을 "
              "1·2·3 단계로 알려주세요.\n\n[오류]\n" + msg
              + (f"\n\n[상황]\n{context}" if context else "")
              + (f"\n\n[상세]\n{trace[-1500:]}" if trace else ""))
    with st.container(border=True):
        st.markdown("###### 🆘 이 오류, 도움받기")
        c1, c2, c3 = st.columns([1.6, 1, 1])
        _ans_key = f"errans_{key}"
        if c1.button("🤖 앱 안에서 바로 물어보기", key=f"errai_{key}", width="stretch",
                     help="사이드바에 넣어 둔 AI 키를 사용합니다. 답이 이 화면에 바로 나옵니다."):
            if not st.session_state.get("api_key"):
                st.session_state[_ans_key] = ("⚠️ 사이드바 **🤖 AI 기능 켜기**에서 "
                                              "API 키를 먼저 넣어 주세요.")
            else:
                try:
                    with st.spinner("AI가 오류를 살펴보는 중..."):
                        st.session_state[_ans_key] = ai_call(
                            prompt, st.session_state.get("api_key"),
                            st.session_state.get("ai_model_g"), max_tokens=900)
                except Exception as _ex:
                    st.session_state[_ans_key] = f"⚠️ AI 호출 실패: {_ex}"
        c2.link_button("🔎 구글 검색",
                       "https://www.google.com/search?q=" + _up.quote_plus(msg),
                       width="stretch")
        c3.link_button("💬 ChatGPT",
                       "https://chatgpt.com/?hints=search&q=" + _up.quote_plus(prompt[:1800]),
                       width="stretch")
        if st.session_state.get(_ans_key):
            st.markdown(st.session_state[_ans_key])
            st.caption("※ AI 답변은 참고용입니다.")


def _install_error_helper():
    """스트림릿이 잡아 주는 '예기치 못한 오류' 아래에도 도움받기 상자를 붙인다."""
    try:
        from streamlit import error_util as _eu
    except Exception:
        return
    if getattr(_eu, "_smart_agent_patched", False):
        return
    _orig = _eu.handle_uncaught_app_exception

    def _patched(ex):
        # 세션 복원이 버튼 키를 건드리면 스트림릿이 막는다. 그 키를 기억해 두고
        # 세션에서 빼 두면 다음 실행부터는 같은 오류가 나지 않는다.
        try:
            if type(ex).__name__ == "StreamlitValueAssignmentNotAllowedError":
                import re as _re_k
                _m = _re_k.search(r"key[`'\s]{0,3}'?([^'`]+)'", str(ex))
                if _m:
                    _deny = set(st.session_state.get("_pin_deny", set()))
                    _deny.add(_m.group(1))
                    st.session_state["_pin_deny"] = _deny
                    st.session_state.pop(_m.group(1), None)
        except Exception:
            pass
        try:
            _orig(ex)
        except Exception:
            try: st.error(f"⚠️ {type(ex).__name__}: {ex}")
            except Exception: pass
        try:
            error_help(ex, context="앱 실행 중 예기치 못한 오류", key="uncaught")
        except Exception:
            pass

    _eu.handle_uncaught_app_exception = _patched
    _eu._smart_agent_patched = True


_install_error_helper()


def strip_md(text):
    """AI가 만든 마크다운 기호(**, ##, - 등)를 문서용 평문으로 정리"""
    import re as _re
    out = []
    for ln in str(text).split("\n"):
        t = ln.rstrip()
        # 제목(#, ##, ###) → 앞 기호 제거
        m = _re.match(r"^\s*#{1,6}\s*(.*)$", t)
        if m: t = m.group(1)
        # 굵게/기울임 제거
        t = _re.sub(r"\*\*(.+?)\*\*", r"\1", t)
        t = _re.sub(r"__(.+?)__", r"\1", t)
        t = _re.sub(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)", r"\1", t)
        t = _re.sub(r"`(.+?)`", r"\1", t)
        # 마크다운 불릿 → 보고서 기호. AI가 앞에 공백을 넣어 오는 경우가 많은데
        # 그대로 두면 '    - ' 처럼 자꾸 깊어지므로 한 단계('  - ')로 통일한다.
        t = _re.sub(r"^\s*[-*+]\s+", "  - ", t)
        # 표 구분선 제거
        if _re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", t) and "-" in t and t.count("-") > 2:
            continue
        # 빈 줄이 연달아 나오면 문서에서 문단 사이가 크게 벌어져 정렬이 어긋나 보인다
        if not t.strip() and (not out or not out[-1].strip()):
            continue
        out.append(t)
    while out and not out[-1].strip():
        out.pop()
    return "\n".join(out)

from economic_core import (
    is_excluded_cost, validate_economic_inputs,
    calculate_row_economics, validate_economic_results,
    calculate_partial_budget, validate_partial_budget_results,
    perform_dominance_analysis, calculate_mrr_table,
    validate_manual_budget_table,
    round_half_up, run_economic_self_test,
    calculate_investment_analysis, investment_sensitivity_table,
)


def validate_repeated_measure_balance(data, subject_col, time_col):
    """개체별 조사시기 집합과 개체×시기 중복을 함께 검사한다."""
    d = data[[subject_col, time_col]].dropna().copy()
    expected = set(d[time_col].unique())
    bad_subjects, duplicate_subjects = [], []
    for subject, group in d.groupby(subject_col):
        if set(group[time_col].unique()) != expected:
            bad_subjects.append(subject)
        if group.duplicated([time_col], keep=False).any():
            duplicate_subjects.append(subject)
    return {
        "ok": not bad_subjects and not duplicate_subjects,
        "expected_times": sorted(expected, key=str),
        "bad_subjects": bad_subjects,
        "duplicate_subjects": duplicate_subjects,
    }


def scale_observed_value_to_10a(value, source_area_a):
    """원자료가 source_area_a 기준일 때 10a 기준으로 환산."""
    area = float(source_area_a)
    if area <= 0:
        raise ValueError("기준 면적은 0보다 커야 합니다.")
    return value * (10.0 / area)


def dataframe_signature(df):
    """세션에 남은 분석 결과가 현재 데이터에서 나온 것인지 확인하는 서명."""
    if df is None:
        return None
    try:
        h = int(pd.util.hash_pandas_object(df, index=True).sum())
    except Exception:
        h = hash((tuple(df.shape), tuple(map(str, df.columns))))
    return (tuple(df.shape), tuple(map(str, df.columns)), h)


def q_ref(col):
    """Patsy 수식에서 열 이름을 안전하게 참조 (작은따옴표·특수문자 포함 대응)"""
    name = str(col).replace("\\", "\\\\").replace("'", "\\'")
    return f"Q('{name}')"

def safe_formula(dep, factors=(), covars=(), interactions=()):
    """ANOVA·회귀·ANCOVA 공통 수식 생성.
    factors: 범주형(C()), covars: 연속형, interactions: [(a, b), ...]"""
    rhs = [f"C({q_ref(f)})" for f in factors if f]
    rhs += [q_ref(c) for c in covars if c]
    rhs += [f"C({q_ref(a)}):C({q_ref(b)})" for a, b in interactions if a and b]
    if not rhs:
        rhs = ["1"]
    return f"{q_ref(dep)} ~ " + " + ".join(rhs)

def clean_columns(df):
    """열 이름 중복·공백 문제를 자동으로 정리 (중복이면 뒤에 _2, _3 붙임)"""
    df = df.copy()
    cols, seen = [], {}
    for c in df.columns:
        name = str(c).strip()
        if name == "" or name.lower().startswith("unnamed"):
            name = "열"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 1
        cols.append(name)
    df.columns = cols
    return df

def validate_anova_data(data, group_col, value_col, min_rep=2):
    """분산분석 전 자료가 분석 가능한지 확인. (실행가능여부, 안내메시지들)"""
    msgs = []
    if data.empty:
        return False, ["❌ 분석할 자료가 없습니다. 결측치를 확인하거나 다른 열을 선택하세요."]
    counts = data.groupby(group_col)[value_col].count()
    counts = counts[counts > 0]
    ng = len(counts)
    if ng < 2:
        return False, [f"❌ 처리구가 {ng}개뿐입니다. 분산분석은 **2개 이상**의 처리구가 필요합니다. "
                       "처리구 열을 다시 선택하거나, 결측치로 자료가 빠지지 않았는지 확인하세요."]
    small = counts[counts < min_rep]
    if len(small) == ng:
        return False, [f"❌ 모든 처리구의 반복이 {min_rep}개 미만입니다(각 1개). "
                       "분산분석은 처리구마다 반복이 2개 이상 있어야 오차를 계산할 수 있습니다."]
    if len(small) > 0:
        msgs.append(f"⚠️ 반복이 {min_rep}개 미만인 처리구가 있습니다: "
                    f"{', '.join(f'{k}({v}개)' for k, v in small.items())}. 결과 해석에 주의하세요.")
    if counts.min() < 3:
        msgs.append("ℹ️ 반복이 3개 미만인 처리구는 정규성 검정을 생략합니다.")
    if data[value_col].nunique() == 1:
        msgs.append("⚠️ 측정값이 모두 동일합니다. 처리 간 차이를 검정할 수 없습니다.")
    return True, msgs

def calc_cv_lsd(model, data, group_col, value_col, alpha=0.05):
    """분산분석 모형에서 CV(%)와 LSD를 계산.
    CV(%) = √(오차평균제곱) ÷ 전체평균 × 100
    LSD   = t(α/2, 오차자유도) × √(2×MSE/r)"""
    try:
        aov = sm.stats.anova_lm(model, typ=2)
        mse = aov.loc["Residual", "sum_sq"] / aov.loc["Residual", "df"]
        dfe = aov.loc["Residual", "df"]
        grand = data[value_col].mean()
        cv = np.sqrt(mse) / grand * 100 if grand else np.nan
        counts = data.groupby(group_col)[value_col].count()
        # 반복수가 다르면 조화평균 사용
        r = len(counts) / np.sum(1.0 / counts) if len(counts) else np.nan
        lsd = stats.t.ppf(1 - alpha/2, dfe) * np.sqrt(2 * mse / r) if r and r > 0 else np.nan
        return {"CV": cv, "LSD": lsd, "MSE": mse, "dfe": dfe, "r": r}
    except Exception:
        return {"CV": np.nan, "LSD": np.nan, "MSE": np.nan, "dfe": np.nan, "r": np.nan}

def cv_grade(cv):
    """포장시험 CV% 판정"""
    if np.isnan(cv): return "-"
    if cv < 10: return "매우 우수"
    if cv < 20: return "양호"
    if cv < 30: return "다소 높음"
    return "재검토 필요"

def find_numeric_like(df, min_ratio=0.6):
    """문자로 읽혔지만 사실상 숫자인 열을 찾음 (콤마·단위·공백 포함)"""
    cands = {}
    for c in df.columns:
        if pd.api.types.is_numeric_dtype(df[c]):
            continue
        ser = df[c].dropna().astype(str).str.strip()
        if ser.empty:
            continue
        cleaned = (ser.str.replace(",", "", regex=False)
                      .str.replace(r"[^\d.\-]", "", regex=True))
        conv = pd.to_numeric(cleaned, errors="coerce")
        ratio = conv.notna().mean()
        if ratio >= min_ratio:
            cands[c] = round(ratio * 100, 1)
    return cands

def to_numeric_clean(ser):
    """'1,200 kg' → 1200.0 처럼 숫자만 추출"""
    cleaned = (ser.astype(str).str.strip()
                  .str.replace(",", "", regex=False)
                  .str.replace(r"[^\d.\-]", "", regex=True))
    return pd.to_numeric(cleaned, errors="coerce")

def _polish_figure(fig):
    """공통 그래프 마감.

    검은 윤곽선은 '그래프 전체'가 아니라 실제 데이터 요소에만 적용한다.
    - 막대그래프: 막대 자체에만 얇은 검은 선
    - 원형/도넛: 조각 자체에만 얇은 검은 선
    - 선그래프/히트맵/산점도: 검은 외곽 프레임 없음
    단일 계열 막대는 원클릭 보고서와 같은 블루 그라데이션으로 자동 통일한다.
    """
    try:
        from matplotlib.container import BarContainer
        from matplotlib.patches import Wedge
        fig.patch.set_facecolor("white")
        elem_border = bool(st.session_state.get("fig_border", True))
        for ax in fig.axes:
            ax.set_facecolor("white")
            ax.tick_params(colors="#4B5F73", labelsize=9, length=3, direction="out")
            ax.xaxis.label.set_color("#31485E")
            ax.yaxis.label.set_color("#31485E")
            if ax.title:
                ax.title.set_color("#23394D")
                ax.title.set_fontweight("bold")

            # 전체 사각 프레임 금지. 좌·하단 축선만 연하게 유지한다.
            if getattr(ax, "name", "rectilinear") == "rectilinear":
                for side in ("top", "right"):
                    if side in ax.spines:
                        ax.spines[side].set_visible(False)
                for side in ("left", "bottom"):
                    if side in ax.spines:
                        ax.spines[side].set_visible(True)
                        ax.spines[side].set_color("#AEBECD")
                        ax.spines[side].set_linewidth(0.75)

            bars = [c for c in getattr(ax, "containers", []) if isinstance(c, BarContainer)]
            # 단일 계열 막대는 옅은→진한 블루 그라데이션.
            if len(bars) == 1 and len(bars[0].patches) > 1:
                n = len(bars[0].patches)
                grad = ["#BFD7EA", "#93B9D8", "#6F9FC8", "#4F7FAF",
                        "#35658F", "#244A73"]
                for i, p in enumerate(bars[0].patches):
                    idx = round((len(grad)-1) * i / max(n-1, 1))
                    p.set_facecolor(grad[idx])

            # 검은 윤곽선은 '단일 계열의 일반 막대'에만 적용한다.
            # 리커트·누적경영비처럼 여러 계열을 쌓는 그래프까지 각 조각을 검게 두르면
            # 표가 잘게 끊겨 보이므로 해당 그래프는 원래의 흰 구분선을 유지한다.
            if elem_border:
                if len(bars) == 1:
                    for p in bars[0].patches:
                        try:
                            p.set_edgecolor("#111111")
                            p.set_linewidth(0.55)
                        except Exception:
                            pass
                # 원형/도넛 조각 윤곽선.
                for p in ax.patches:
                    if isinstance(p, Wedge):
                        try:
                            p.set_edgecolor("#111111")
                            p.set_linewidth(0.60)
                        except Exception:
                            pass

            lg = ax.get_legend()
            if lg is not None:
                try:
                    lg.get_frame().set_linewidth(0)
                    lg.get_frame().set_facecolor("white")
                except Exception:
                    pass
    except Exception:
        pass
    return fig

def _figure_png(fig, dpi=150):
    """Matplotlib Figure를 화면/다운로드 공용 PNG bytes로 변환한다."""
    _polish_figure(fig)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", facecolor="white")
    buf.seek(0)
    return buf.getvalue()


def show_plot(fig, max_width=660):
    """그래프를 브라우저 전체 폭으로 억지 확대하지 않고 보고서 크기로 표시한다.

    Streamlit 버전에 따라 st.pyplot의 기본 폭 정책이 달라졌기 때문에, 화면 표시만큼은
    PNG로 고정해 CSS stretch의 영향을 받지 않게 한다. 그래프 가로 설정은 반영하되
    일반 단일 그래프는 660px, 2패널 이상은 최대 920px까지만 표시한다.
    """
    data = _figure_png(fig, dpi=145)
    try:
        fig_w = float(fig.get_figwidth())
    except Exception:
        fig_w = 6.0
    px = int(max(430, min(int(max_width), round(fig_w * 92))))
    st.image(data, width=px)
    return data


def fig_to_png(fig, show=True):
    """그래프를 스마트 블루 스타일로 마감해 PNG로 반환한다."""
    data = _figure_png(fig, dpi=160)
    if show:
        try:
            fig_w = float(fig.get_figwidth())
        except Exception:
            fig_w = 6.0
        # 1패널은 660px, 매우 넓은 다중패널 그림도 920px을 넘기지 않는다.
        cap = 920 if fig_w >= 10 else 660
        px = int(max(430, min(cap, round(fig_w * 92))))
        st.image(data, width=px)
    plt.close(fig)
    return data

def cronbach_alpha(df_items):
    k = df_items.shape[1]
    if k < 2: return np.nan
    item_var = df_items.var(axis=0, ddof=1).sum()
    total_var = df_items.sum(axis=1).var(ddof=1)
    if total_var == 0: return np.nan
    return (k / (k - 1)) * (1 - item_var / total_var)


def likert_cutoffs(scale_max):
    """척도 범위에 맞는 부정/긍정 경계. 5점이면 <=2 / >=4, 7점이면 <=3 / >=5."""
    m = max(int(scale_max), 2)
    center = (m + 1) / 2.0
    neg = int(np.ceil(center) - 1)
    pos = int(np.floor(center) + 1)
    return pos, neg

# ---------------------------------------------------------------- hwpx 스타일
_HH = "http://www.hancom.co.kr/hwpml/2011/head"
_HP = "http://www.hancom.co.kr/hwpml/2011/paragraph"
def _q(t): return f"{{{_HH}}}{t}"
_LANG = {"HANGUL":"hangul","LATIN":"latin","HANJA":"hanja","JAPANESE":"japanese",
         "OTHER":"other","SYMBOL":"symbol","USER":"user"}

def _reg_font(hdr, name):
    ids = {}
    for group in hdr.iter(_q("fontfaces")):
        for ff in group:
            cnt = int(ff.get("fontCnt")); f = etree.SubElement(ff, _q("font"))
            f.set("id", str(cnt)); f.set("face", name); f.set("type", "TTF"); f.set("isEmbedded", "0")
            for ch in ff.find(_q("font")): f.append(copy.deepcopy(ch))
            ff.set("fontCnt", str(cnt+1)); ids[ff.get("lang")] = str(cnt)
        break
    return ids

def _mk_charpr(hdr, cp0, fids, size, bold=False, sup=False, color=None, spacing=None):
    """color: '#0000FF' 같은 글자색 / spacing: 자간(%). 음수면 글자를 좁혀 한 줄에 더 넣는다."""
    cps = next(hdr.iter(_q("charProperties"))); new = copy.deepcopy(cp0); nid = str(len(cps))
    new.set("id", nid); new.set("height", str(int(size*100)))
    if bold: new.set("bold", "1")
    if color: new.set("textColor", color)
    if spacing is not None:
        el = next((c for c in new if c.tag == _q("spacing")), None)
        if el is not None:
            for k in list(el.attrib): el.set(k, str(int(spacing)))
    if sup:   # 위첨자: 작게 + 위로 올림
        # 한글(HWP)에서는 offset 이 '음수일 때 위로' 올라간다. (양수로 주면 아래첨자가 됨 —
        # 실제 한글에서 확인. 일부 문서에는 반대로 적혀 있으니 값을 바꾸지 말 것)
        rel = str(int(st.session_state.get("sup_size", 65)))
        off = str(-abs(int(st.session_state.get("sup_off", 35))))
        for tag, val in [("relSz", rel), ("offset", off)]:
            el = next((c for c in new if c.tag == _q(tag)), None)
            if el is not None:
                for k in list(el.attrib): el.set(k, val)
    fr = new.find(_q("fontRef"))
    for lang, fid in fids.items(): fr.set(_LANG.get(lang, lang.lower()), fid)
    cps.append(new); cps.set("itemCnt", str(len(cps))); return nid

def _mk_parapr(hdr, align="CENTER", left=0, intent=0):
    """left = 왼쪽 여백, intent = 첫 줄 들여쓰기(음수면 내어쓰기).
    left=L, intent=-L 로 주면 둘째 줄부터 첫 줄 글자 시작 위치에 맞춰 정렬된다."""
    pp = next(hdr.iter(_q("paraProperties"))); new = copy.deepcopy(pp[0]); nid = str(len(pp)); new.set("id", nid)
    al = next((c for c in new.iter() if c.tag == _q("align")), None)
    if al is None: al = etree.SubElement(new, _q("align"))
    al.set("horizontal", align); al.set("vertical", "CENTER")
    if left or intent:
        for mg in new.iter():
            if not isinstance(mg.tag, str) or not mg.tag.endswith("}margin"):
                continue
            for ch in mg:
                if not isinstance(ch.tag, str):
                    continue
                if ch.tag.endswith("}intent"): ch.set("value", str(int(intent)))
                elif ch.tag.endswith("}left"): ch.set("value", str(int(left)))
    pp.append(new); pp.set("itemCnt", str(len(pp))); return nid


def _bullet_layout(line):
    """'  - 내용' → (앞여백 1.0글자, 글머리 1.0글자, '- 내용')

    앞의 공백을 글자로 찍으면 '- ' 가 오른쪽으로 밀려 보기 싫으므로,
    공백은 지우고 **문단 왼쪽 여백**으로 옮긴다. 글머리 폭만큼 내어쓰기를 주면
    줄이 넘어갔을 때 둘째 줄이 본문 첫 글자에 맞춰 정렬된다.
    """
    import re as _re
    s = str(line)
    m = _re.match(r"^([ \t\u00a0]*)(([○◦●□■▪▶–—-])[ \t]*)?(.*)$", s, _re.S)
    if not m:
        return 0.0, 0.0, s.strip()
    ws, mark, rest = m.group(1), m.group(2) or "", m.group(4)
    def cells(t):
        return sum(1.0 if ord(ch) > 0x1100 else 0.5 for ch in t)
    if not mark:
        return 0.0, 0.0, s.strip()
    mark = mark.rstrip() + " "          # 글머리 뒤 공백은 한 칸으로 통일
    return cells(ws), cells(mark), mark + rest.strip()


def _prefix_cells(line):
    """'○ ', '  - ' 같은 글머리 부분의 폭을 '글자 수'로 잰다 (전각=1, 반각=0.5)."""
    ws, mark, _ = _bullet_layout(line)
    return ws + mark

def _fills(doc, shade, line_color, lw="0.1 mm", side_lines=False):
    def bf(borders, fill=None):
        return doc.ensure_border_fill(border_color=line_color, border_width=lw,
                                      fill_color=fill, active_borders=borders)
    full = ["top", "bottom", "left", "right"]
    edge = full if side_lines else ["top", "bottom"]
    return (bf(full), bf(full, shade), bf(edge), bf(edge, shade))


def _fills_plain(doc, shade, line_color, lw="0.1 mm"):
    """줄글이 많은 표(부분예산표 등)용 — 안쪽 가로선만 없애고
    표의 맨 위·머리행 아래·**맨 아래 선은 남긴다.**"""
    def bf(borders, fill=None):
        return doc.ensure_border_fill(border_color=line_color, border_width=lw,
                                      fill_color=fill, active_borders=borders)
    none_, top_bottom, bottom = [], ["top", "bottom"], ["bottom"]
    # (본문 안쪽, 머리행, 본문 가장자리, 머리행 가장자리, 마지막 행)
    return (bf(none_), bf(top_bottom, shade), bf(none_), bf(top_bottom, shade),
            bf(bottom))

def _selected_hwp_font():
    """한글 표/보고서에 실제 적용할 글꼴명. 목록 밖 글꼴은 직접 입력할 수 있다."""
    ss = st.session_state
    choice = str(ss.get("hwp_font", "휴먼명조"))
    if choice == "직접 입력…":
        custom = str(ss.get("hwp_font_custom", "")).strip()
        return custom or "휴먼명조"
    return choice


def _doc_opts():
    ss = st.session_state
    return dict(font=_selected_hwp_font(), size=ss.get("hwp_size", 10),
                shade=ss.get("hwp_shade", "#D9D9D9"), line=ss.get("hwp_line", "#000000"),
                lw=ss.get("hwp_lw", "0.1 mm"), sides=ss.get("hwp_sides", False),
                row_h=float(ss.get("hwp_rowh", 6.5)),
                tight=int(ss.get("hwp_tight", -14)),
                ai_color=ss.get("hwp_aicolor", "#0000FF"))

def _cells_of(s):
    """문자열의 표시 폭을 '반각 칸 수'로 센다 (한글·전각=2, 영문·숫자=1)."""
    return sum(2 if ord(ch) > 0x1100 else 1 for ch in str(s))


def _col_widths(tdf, total, min_cells=5, max_cells=46):
    """열마다 들어가는 글자 길이에 비례해 폭을 나눈다.
    (모든 열을 똑같이 나누면 긴 글이 든 열만 여러 줄로 접혀 표가 지저분해진다)"""
    n = tdf.shape[1]
    if n <= 0:
        return []
    need = []
    for c in range(n):
        vals = [tdf.columns[c]] + list(tdf.iloc[:, c])
        longest = max((_cells_of(v) for v in vals), default=min_cells)
        need.append(min(max(longest, min_cells), max_cells))
    s = float(sum(need)) or 1.0
    out = [max(int(total * w / s), int(total * 0.06)) for w in need]
    out[-1] = total - sum(out[:-1])          # 반올림 오차는 마지막 열이 흡수
    return out


def _fit_size(table, n_cols, row_h_mm=6.5, widths=None, row_lines=None):
    """표를 본문 폭에 맞추고 행 높이를 촘촘하게.
    widths: 열별 폭(HWPUNIT) / row_lines: 행별 줄 수(줄이 접히는 행은 높게)"""
    ws = list(widths) if widths else [BODY_W // n_cols] * n_cols
    h1 = int(row_h_mm * MM)
    try:
        total_h = 0
        for r in range(table.row_count):
            lines = 1
            if row_lines and r < len(row_lines):
                lines = max(1, int(row_lines[r]))
            h = h1 * lines
            total_h += h
            for c in range(n_cols):
                table.cell(r, c).set_size(width=ws[c], height=h)
        sz = next((e for e in table.element.iter(f"{{{_HP}}}sz")), None)
        if sz is not None:
            sz.set("width", str(sum(ws))); sz.set("height", str(total_h))
    except Exception:
        pass

# 한 표가 한 쪽을 넘으면 한글에서 표가 잘려 보이는 일이 잦다. 이 행 수를 넘으면
# 여러 개의 표로 나눠서 넣는다(제목에 (1/3) 표시).
_MAX_TABLE_ROWS = 24


def _split_long_table(tb, max_rows=_MAX_TABLE_ROWS):
    """행이 너무 많은 표를 쪽 단위로 나눈다. 나눌 필요가 없으면 원본 하나만 돌려준다."""
    if tb is None or len(tb) <= max_rows:
        return [tb]
    return [tb.iloc[i:i + max_rows] for i in range(0, len(tb), max_rows)]


def _write_table(doc, tdf, center, cp_body, cp_head, fills, cp_sup=None, row_h_mm=6.5,
                 size_pt=10, tight=0):
    """셀 값에 '^'가 있으면 뒤쪽을 위첨자로 (예: '13.5^a')
    tight: 셀 글자에 적용한 자간(%). 음수면 그만큼 글자가 좁아지므로 줄 수 계산에 반영한다."""
    inner, header, b_edge, h_edge = fills[:4]
    last_row_fill = fills[4] if len(fills) > 4 else None
    nr, nc = tdf.shape[0]+1, tdf.shape[1]
    table = doc.add_table(nr, nc)  # nr: 머리행 포함 행 수
    try:
        # 쪽이 바뀌어도 표가 셀 단위로 이어지고, 다음 쪽에 머리행이 다시 나오게 한다.
        table.element.set("pageBreak", "CELL")
        table.element.set("repeatHeader", "1")
    except Exception:
        pass
    widths = _col_widths(tdf, BODY_W)
    # 열 폭 대비 글자 길이로 '몇 줄이 될지'를 미리 재서 행 높이를 잡는다
    per_cell = max(size_pt * 100 / 2.0 * (1 + min(0, tight) / 100.0), 1.0)  # 반각 한 칸 폭
    cap = [max(int(w / per_cell) - 1, 4) for w in widths]
    row_lines = []
    for r in range(nr):
        vals = list(tdf.columns) if r == 0 else list(tdf.iloc[r-1])
        row_lines.append(max(
            [max(1, -(-_cells_of(v) // cap[c])) for c, v in enumerate(vals)] or [1]))

    def fill(r, c, txt, is_h):
        cell = table.cell(r, c); edge = (c == 0 or c == nc-1)
        bfid = (h_edge if edge else header) if is_h else (b_edge if edge else inner)
        if (not is_h) and last_row_fill is not None and r == nr - 1:
            bfid = last_row_fill
        try: cell.element.set("borderFillIDRef", bfid)
        except Exception: pass
        s = str(txt)
        try:
            paras = cell.paragraphs; pp = paras[0] if paras else cell.add_paragraph()
            pp.element.set("paraPrIDRef", center)
            if cp_sup and "^" in s:
                base, _, sup = s.partition("^")
                pp.add_run(base, char_pr_id_ref=(cp_head if is_h else cp_body))
                pp.add_run(sup, char_pr_id_ref=cp_sup)
            else:
                # cp_sup 를 못 만든 경우에도 '^'가 그대로 남지 않게 유니코드 위첨자로
                pp.add_run(sup_text(s), char_pr_id_ref=(cp_head if is_h else cp_body))
        except Exception:
            cell.text = str(sup_text(s))
    for c, name in enumerate(tdf.columns): fill(0, c, name, True)
    for r in range(tdf.shape[0]):
        for c in range(tdf.shape[1]): fill(r+1, c, tdf.iloc[r, c], False)
    _fit_size(table, nc, row_h_mm, widths=widths, row_lines=row_lines)

def _setup(doc, font):
    hdr = doc.headers[0].element
    cp0 = next(e for e in hdr.iter(_q("charPr")) if e.get("id") == "0")
    fids = _reg_font(hdr, font)
    return hdr, cp0, fids

def dataframe_to_hwpx(df, title="분석 결과표", **kw):
    from hwpx import HwpxDocument
    import tempfile, os
    o = _doc_opts(); o.update(kw)
    doc = HwpxDocument.new()
    hdr, cp0, fids = _setup(doc, o["font"])
    cp_body = _mk_charpr(hdr, cp0, fids, o["size"])
    cp_head = _mk_charpr(hdr, cp0, fids, o["size"], bold=True)
    cp_title = _mk_charpr(hdr, cp0, fids, o["size"] + 3, bold=True)
    cp_sup = _mk_charpr(hdr, cp0, fids, o["size"], sup=True)
    center = _mk_parapr(hdr, "CENTER"); left = _mk_parapr(hdr, "LEFT")
    p = doc.add_paragraph(); p.element.set("paraPrIDRef", left)
    p.add_run(title, char_pr_id_ref=cp_title)
    _write_table(doc, df, center, cp_body, cp_head,
                 _fills(doc, o["shade"], o["line"], o["lw"], o["sides"]),
                 cp_sup=cp_sup, row_h_mm=o["row_h"], size_pt=o["size"])
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".hwpx"); tmp.close()
    doc.save_to_path(tmp.name)
    with open(tmp.name, "rb") as f: data = f.read()
    os.unlink(tmp.name); return data

def collect_captions(items):
    """보고서에 들어갈 표·그림 목차를 미리 계산"""
    tno = fno = 0
    tabs, figs = [], []
    for it in items:
        blocks = it.get("blocks")
        if blocks:
            for b in blocks:
                if b.get("table") is not None:
                    tno += 1; tabs.append(f"<표 {tno}> {b.get('caption','')}".rstrip())
                if b.get("image"):
                    fno += 1; figs.append(f"<그림 {fno}> {b.get('caption','')}".rstrip())
        else:
            if it.get("table") is not None:
                tno += 1; tabs.append(f"<표 {tno}> {it.get('heading','')}".rstrip())
            if it.get("image"):
                fno += 1; figs.append(f"<그림 {fno}> {it.get('heading','')}".rstrip())
    return tabs, figs

def build_report_hwpx(items, doc_title="실험 통계 분석 보고서", **kw):
    from hwpx import HwpxDocument
    import tempfile, os
    o = _doc_opts(); o.update(kw)
    sz = o["size"]
    doc = HwpxDocument.new()
    hdr, cp0, fids = _setup(doc, o["font"])
    cp_title = _mk_charpr(hdr, cp0, fids, sz + 6, bold=True)
    cp_head = _mk_charpr(hdr, cp0, fids, sz + 2, bold=True)
    cp_cap = _mk_charpr(hdr, cp0, fids, sz, bold=True)
    cp_body = _mk_charpr(hdr, cp0, fids, sz)
    cp_th = _mk_charpr(hdr, cp0, fids, sz, bold=True)
    cp_sup = _mk_charpr(hdr, cp0, fids, sz, sup=True)
    # AI 해석은 사람이 쓴 문장과 구분되도록 파란색
    cp_ai = _mk_charpr(hdr, cp0, fids, sz, color=o.get("ai_color", "#0000FF"))
    # 줄글이 많은 표는 자간을 좁혀 한 줄에 담는다(두 줄로 접히면 보기 나쁨)
    _tight = int(o.get("tight", -14))
    cp_body_t = _mk_charpr(hdr, cp0, fids, sz, spacing=_tight)
    cp_th_t = _mk_charpr(hdr, cp0, fids, sz, bold=True, spacing=_tight)
    center = _mk_parapr(hdr, "CENTER"); left = _mk_parapr(hdr, "LEFT")
    _hang_cache = {(0.0, 0.0): left}

    def _hang(ws, mark):
        """앞여백 ws글자 + 글머리 mark글자 → 문단모양 id (같은 값은 재사용)."""
        key = (round(ws, 1), round(mark, 1))
        if key not in _hang_cache:
            L = int((key[0] + key[1]) * sz * 100)
            I = -int(key[1] * sz * 100)
            _hang_cache[key] = (_mk_parapr(hdr, "LEFT", left=L, intent=I) if L else left)
        return _hang_cache[key]

    fills = _fills(doc, o["shade"], o["line"], o["lw"], o["sides"])
    fills_plain = _fills_plain(doc, o["shade"], o["line"], o["lw"])

    p = doc.add_paragraph(); p.element.set("paraPrIDRef", center)
    p.add_run(doc_title, char_pr_id_ref=cp_title)
    doc.add_paragraph()

    tno, fno = 0, 0
    for idx, it in enumerate(items, 1):
        p = doc.add_paragraph(); p.element.set("paraPrIDRef", left)
        p.add_run(f"□ {it.get('heading','분석 결과')}", char_pr_id_ref=cp_head)

        def render_text(t, ai=False):
            if not t:
                return
            for line in strip_md(t).split("\n"):
                if not line.strip():
                    continue          # 빈 줄은 문서에서 문단 사이만 벌어져 보기 나쁘다
                ws, mark, body = _bullet_layout(line)
                pp = doc.add_paragraph()
                # 앞 공백은 글자로 찍지 않고 문단 왼쪽 여백으로 옮기고,
                # 글머리 폭만큼 내어쓰기를 줘서 둘째 줄이 본문에 맞춰지게 한다.
                pp.element.set("paraPrIDRef", _hang(ws, mark))
                pp.add_run(body, char_pr_id_ref=(cp_ai if ai else cp_body))

        def render_table(tb, cap="", plain=False):
            nonlocal tno
            if tb is None:
                return
            tno += 1
            parts = _split_long_table(tb)
            for _i, _part in enumerate(parts):
                doc.add_paragraph()          # 표 위 한 줄 띄우기
                pp = doc.add_paragraph(); pp.element.set("paraPrIDRef", left)
                _cap = f"<표 {tno}> {cap}".rstrip()
                if len(parts) > 1:
                    _cap += f" ({_i + 1}/{len(parts)})"
                pp.add_run(_cap, char_pr_id_ref=cp_cap)
                _write_table(doc, _part, left if plain else center,
                             cp_body_t if plain else cp_body,
                             cp_th_t if plain else cp_th,
                             fills_plain if plain else fills,
                             cp_sup=cp_sup, size_pt=sz,
                             # 줄글 표는 행 높이를 낮춰 위아래 여백을 줄인다
                             row_h_mm=(o["row_h"] * 0.7 if plain else o["row_h"]),
                             tight=_tight if plain else 0)

        def render_image(im, cap=""):
            nonlocal fno
            if im:
                fno += 1
                doc.add_paragraph()          # 그림 위 한 줄 띄우기
                iid = doc.add_image(im, "png")
                pp = doc.add_paragraph(); pp.element.set("paraPrIDRef", center)
                pp.add_picture(iid, width=int(120*MM), height=int(80*MM), align="CENTER")
                # 그림 제목은 가운데 정렬(표 제목은 왼쪽 정렬 유지)
                pp = doc.add_paragraph(); pp.element.set("paraPrIDRef", center)
                pp.add_run(f"<그림 {fno}> {cap}".rstrip(), char_pr_id_ref=cp_cap)

        blocks = it.get("blocks")
        if blocks:
            # 여러 블록을 순서대로 렌더링 (설문처럼 표·그림이 여러 개인 경우)
            for blk in blocks:
                render_text(blk.get("text"), ai=bool(blk.get("ai")))
                render_table(blk.get("table"), blk.get("caption", ""),
                             plain=bool(blk.get("plain")))
                render_image(blk.get("image"), blk.get("caption", ""))
        else:
            render_text(it.get("text"))
            render_table(it.get("table"), it.get("heading", ""))
            render_image(it.get("image"), it.get("heading", ""))
        doc.add_paragraph()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".hwpx"); tmp.close()
    doc.save_to_path(tmp.name)
    with open(tmp.name, "rb") as f: data = f.read()
    os.unlink(tmp.name); return data

# ---------------------------------------------------------------- 워드(docx) 생성
def _docx_table(doc, df, sup=True):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    nc = df.shape[1]
    t = doc.add_table(rows=1, cols=nc)
    try: t.style = "Light Grid Accent 1"
    except Exception:
        try: t.style = "Table Grid"
        except Exception: pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    def setcell(cell, val, bold=False):
        cell.text = ""
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s = str(val)
        if sup and "^" in s:
            base, _, sp = s.partition("^")
            r = p.add_run(base); r.font.size = Pt(9); r.bold = bold
            r2 = p.add_run(sp); r2.font.size = Pt(9); r2.bold = bold
            r2.font.superscript = True
        else:
            # sup=False 로 불렸더라도 '^'가 그대로 찍히지 않게 유니코드 위첨자로 대체
            r = p.add_run(str(sup_text(s))); r.font.size = Pt(9); r.bold = bold
    for i, c in enumerate(df.columns): setcell(t.rows[0].cells[i], c, bold=True)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row): setcell(cells[i], v)
    try:
        # 표가 여러 쪽에 걸칠 때 다음 쪽에도 머리행이 나오게 한다.
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _Ox
        _el = _Ox("w:tblHeader"); _el.set(_qn("w:val"), "true")
        t.rows[0]._tr.get_or_add_trPr().append(_el)
    except Exception:
        pass
    return t

def dataframe_to_docx(df, title="분석 결과표"):
    if not _HAS_DOCX:
        raise RuntimeError("python-docx 미설치")
    import docx, tempfile, os
    from docx.shared import Pt
    doc = docx.Document()
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)
    doc.add_heading(title, level=1)
    _docx_table(doc, df)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx"); tmp.close()
    doc.save(tmp.name)
    with open(tmp.name, "rb") as f: data = f.read()
    os.unlink(tmp.name); return data

def build_report_docx(items, doc_title="실험 통계 분석 보고서"):
    if not _HAS_DOCX:
        raise RuntimeError("python-docx 미설치")
    import docx, io, tempfile, os
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = docx.Document()
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)
    h = doc.add_heading(doc_title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tno, fno = [0], [0]
    def add_text(t, ai=False):
        if not t:
            return
        from docx.shared import Pt as _Pt, RGBColor as _RGB
        for line in strip_md(t).split("\n"):
            if not line.strip():
                continue                     # 빈 줄은 문단 사이만 벌어져 보기 나쁘다
            ws, mark, body = _bullet_layout(line)
            p = doc.add_paragraph()
            r = p.add_run(body)
            if ai:
                r.font.color.rgb = _RGB(0x00, 0x00, 0xFF)
            if ws or mark:
                # 앞 공백은 글자로 찍지 않고 여백으로, 글머리 폭만큼 내어쓰기
                p.paragraph_format.left_indent = _Pt((ws + mark) * 10)
                p.paragraph_format.first_line_indent = _Pt(-mark * 10)
    def add_table(tb, cap=""):
        if tb is None:
            return
        tno[0] += 1
        parts = _split_long_table(tb)
        for _i, _part in enumerate(parts):
            _cap = f"<표 {tno[0]}> {cap}".rstrip()
            if len(parts) > 1:
                _cap += f" ({_i + 1}/{len(parts)})"
            c = doc.add_paragraph(_cap)
            c.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in c.runs: r.bold = True
            _docx_table(doc, _part)
    def add_image(im, cap=""):
        if im:
            fno[0] += 1
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(im), width=Mm(120))
            c = doc.add_paragraph(f"<그림 {fno[0]}> {cap}".rstrip())
            c.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in c.runs: r.bold = True
    for idx, it in enumerate(items, 1):
        doc.add_heading(f"□ {it.get('heading','분석 결과')}", level=1)
        blocks = it.get("blocks")
        if blocks:
            for blk in blocks:
                add_text(blk.get("text"), ai=bool(blk.get("ai")))
                add_table(blk.get("table"), blk.get("caption", ""))
                add_image(blk.get("image"), blk.get("caption", ""))
        else:
            add_text(it.get("text")); add_table(it.get("table"), it.get("heading", ""))
            add_image(it.get("image"), it.get("heading", ""))
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx"); tmp.close()
    doc.save(tmp.name)
    with open(tmp.name, "rb") as f: data = f.read()
    os.unlink(tmp.name); return data

# ---------------------------------------------------------------- 보고서 담기
@st.cache_data(show_spinner=False, max_entries=80)
def _make_docs(csv_text, title, fmt, opts_sig):
    _df = pd.read_csv(io.StringIO(csv_text), dtype=str).fillna("")
    return dataframe_to_hwpx(_df, title) if fmt == "hwpx" else dataframe_to_docx(_df, title)

_XL_NOT_VALUE = ("표준편차", "표준오차", "표준 편차", "표준 오차", "sd", "se", "std",
                 "p-value", "p값", "pvalue", "유의", "cv", "lsd", "df", "자유도",
                 "f값", "t값", "n수", "반복수", "순위", "rank", "개수")


def _xl_is_value_col(name):
    """차트에 그릴 '값' 열인지 판단. 표준편차·p값·n 같은 보조 열은 제외한다."""
    t = str(name).strip().lower()
    if t in ("n", "N".lower()):
        return False
    return not any(k in t for k in _XL_NOT_VALUE)


def _xl_split(v):
    """'345.500^a' 또는 '345.500ᵃ' → (345.5, 'a'). 숫자로 볼 수 없으면 (None, None).

    엑셀 차트는 **숫자 셀**만 그릴 수 있다. 이 앱의 표는 유의성 문자를 값에 붙여
    문자열로 만들어 두기 때문에, 그대로 내보내면 그래프가 그려지지 않는다.
    """
    import re as _re
    if isinstance(v, bool):
        return (None, None)
    if isinstance(v, (int, float)):
        return (float(v), "") if pd.notna(v) else (None, None)
    t = str(v).strip()
    if not t or t.lower() in ("nan", "none", "-", "―"):
        return (None, None)
    _SC = "ᵃᵇᶜᵈᵉᶠᵍʰⁱʲᵏˡᵐⁿ"
    letters = ""
    while t and t[-1] in _SC:                     # 유니코드 위첨자를 보통 글자로
        letters = "abcdefghijklmn"[_SC.index(t[-1])] + letters
        t = t[:-1]
    m = _re.match(r"^\s*(-?[\d,]+(?:\.\d+)?)\s*\^?\s*([a-zA-Z]{1,3}\**|\*+)?\s*$", t)
    if not m:
        return (None, None)
    return float(m.group(1).replace(",", "")), (m.group(2) or "") + letters


def _xlsx_blue_fill(value, vmin, vmax):
    """숫자 셀용 옅은 파랑 그라데이션 색상(HEX)을 반환."""
    try:
        if value is None or pd.isna(value) or vmax <= vmin:
            return None
        ratio = max(0.0, min(1.0, (float(value) - float(vmin)) / (float(vmax) - float(vmin))))
        lo = (247, 251, 255)   # F7FBFF
        hi = (158, 197, 229)   # 9EC5E5
        rgb = tuple(round(lo[i] + (hi[i] - lo[i]) * ratio) for i in range(3))
        return ''.join(f'{x:02X}' for x in rgb)
    except Exception:
        return None


def _xlsx_sig_specs_for_df(df, sheet_name='데이터', hdr=4):
    """make_xlsx의 출력 열 구조를 그대로 재현해 차트용 유의성 문자 위치를 계산한다.

    반환 key: (시트명, 값 열 문자) -> {sig_col, sig_values, start_row}
    """
    from openpyxl.utils import get_column_letter
    out_cols = []
    sig_map = {}
    for c in df.columns:
        raw_vals = list(df[c])
        pairs = [_xl_split(v) for v in raw_vals]
        nonempty = [i for i, v in enumerate(raw_vals)
                    if not pd.isna(v) and str(v).strip().lower() not in ('', 'nan', 'none', '-', '―')]
        all_parseable = bool(nonempty) and all(pairs[i][0] is not None for i in nonempty)
        is_native_num = pd.api.types.is_numeric_dtype(df[c]) and not pd.api.types.is_bool_dtype(df[c])
        is_num_col = is_native_num or all_parseable
        value_pos = len(out_cols) + 1
        out_cols.append((c, is_num_col))
        letters = [p[1] or '' for p in pairs]
        if is_num_col and any(letters):
            sig_pos = len(out_cols) + 1
            out_cols.append((f'{c} 유의성', False))
            sig_map[(str(sheet_name), get_column_letter(value_pos))] = {
                'sig_col': get_column_letter(sig_pos),
                'sig_values': letters,
                'value_values': [p[0] for p in pairs],
                'start_row': hdr + 1,
            }
    return sig_map


def _inject_xlsx_significance_labels(xlsx_bytes, specs):
    """Excel 막대 위에 ``165.25ᵃ`` 형태의 유의성 레이블을 넣는다.

    Excel의 일반 데이터 레이블과 셀 참조 텍스트를 동시에 켜면 일부 버전에서
    작은 '범례 키 사각형 + a'처럼 렌더링되는 문제가 있다. 따라서 각 막대에
    **값과 유의성 문자를 합친 하나의 사용자 지정 텍스트 레이블**만 넣는다.
    범례 키·계열명·범주명은 모두 명시적으로 끈다.
    """
    if not specs:
        return xlsx_bytes
    import zipfile, re as _re
    from lxml import etree as _ET

    src = io.BytesIO(xlsx_bytes)
    out = io.BytesIO()
    CURI = 'http://schemas.openxmlformats.org/drawingml/2006/chart'
    AURI = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    ns = {'c': CURI, 'a': AURI}
    C = '{%s}' % CURI
    A = '{%s}' % AURI

    def _parse_formula(f):
        m = _re.match(r"(?:'((?:[^']|'')+)'|([^!]+))!\$?([A-Z]+)\$?\d+", str(f or ''))
        if not m:
            return None, None
        sheet = (m.group(1) or m.group(2) or '').replace("''", "'")
        return sheet, m.group(3)

    def _fmt_num(v):
        try:
            x = float(v)
            if abs(x) >= 1000:
                return f"{x:,.2f}".rstrip('0').rstrip('.')
            return f"{x:.2f}".rstrip('0').rstrip('.')
        except Exception:
            return ''

    def _rich_value_sig(parent, value_txt, sig_txt):
        """Excel 데이터 레이블에 값 + 실제 위첨자 a/b/c를 rich text로 넣는다.

        작은 범례키 사각형이나 Unicode 위첨자 글꼴 깨짐 없이
        165.25ᵃ처럼 보이되, a는 DrawingML baseline 속성으로 진짜 위첨자 처리한다.
        """
        tx = _ET.SubElement(parent, C + 'tx')
        rich = _ET.SubElement(tx, C + 'rich')
        _ET.SubElement(rich, A + 'bodyPr')
        _ET.SubElement(rich, A + 'lstStyle')
        p = _ET.SubElement(rich, A + 'p')
        r1 = _ET.SubElement(p, A + 'r')
        _ET.SubElement(r1, A + 'rPr', lang='ko-KR', sz='900')
        _ET.SubElement(r1, A + 't').text = value_txt
        if sig_txt:
            r2 = _ET.SubElement(p, A + 'r')
            # baseline=30000은 본문 기준 약 30% 위로 올리는 DrawingML superscript 효과.
            _ET.SubElement(r2, A + 'rPr', lang='en-US', sz='720', baseline='30000')
            _ET.SubElement(r2, A + 't').text = str(sig_txt)
        _ET.SubElement(p, A + 'endParaRPr', lang='ko-KR', sz='900')

    with zipfile.ZipFile(src, 'r') as zin, zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            raw = zin.read(info.filename)
            if not (info.filename.startswith('xl/charts/chart') and info.filename.endswith('.xml')):
                zout.writestr(info, raw)
                continue
            try:
                root = _ET.fromstring(raw)
                changed = False
                for ser in root.xpath('.//c:barChart/c:ser', namespaces=ns):
                    fnode = ser.find('.//c:val/c:numRef/c:f', namespaces=ns)
                    if fnode is None:
                        continue
                    sh, vcol = _parse_formula(fnode.text)
                    spec = specs.get((sh, vcol))
                    if not spec:
                        continue
                    old = ser.find(C + 'dLbls')
                    if old is not None:
                        ser.remove(old)
                    dlbls = _ET.Element(C + 'dLbls')
                    sig_values = list(spec.get('sig_values') or [])
                    num_values = list(spec.get('value_values') or [])
                    for i, sigv in enumerate(sig_values):
                        if not sigv:
                            continue
                        dl = _ET.SubElement(dlbls, C + 'dLbl')
                        _ET.SubElement(dl, C + 'idx', val=str(i))
                        _ET.SubElement(dl, C + 'layout')
                        _value_txt = _fmt_num(num_values[i] if i < len(num_values) else None)
                        _rich_value_sig(dl, _value_txt, sigv)
                        _ET.SubElement(dl, C + 'dLblPos', val='outEnd')
                        _ET.SubElement(dl, C + 'showLegendKey', val='0')
                        _ET.SubElement(dl, C + 'showVal', val='0')
                        _ET.SubElement(dl, C + 'showCatName', val='0')
                        _ET.SubElement(dl, C + 'showSerName', val='0')
                        _ET.SubElement(dl, C + 'showPercent', val='0')
                    _ET.SubElement(dlbls, C + 'showLegendKey', val='0')
                    _ET.SubElement(dlbls, C + 'showVal', val='0')
                    _ET.SubElement(dlbls, C + 'showCatName', val='0')
                    _ET.SubElement(dlbls, C + 'showSerName', val='0')
                    _ET.SubElement(dlbls, C + 'showPercent', val='0')
                    _ET.SubElement(dlbls, C + 'dLblPos', val='outEnd')
                    children = list(ser)
                    pos = next((i for i, el in enumerate(children)
                                if el.tag in (C + 'cat', C + 'val')), len(children))
                    ser.insert(pos, dlbls)
                    changed = True
                if changed:
                    raw = _ET.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
            except Exception:
                pass
            zout.writestr(info, raw)
    return out.getvalue()



def make_xlsx(df, title, chart=True, error_bars=False):
    """스마트 블루 표 + 엑셀에서 직접 편집 가능한 차트가 든 xlsx를 만든다.

    app(9)의 편집 가능한 Excel 차트 기능을 보존하면서, 화면과 같은 푸른 계열
    머리행·행 구분·숫자 그라데이션·테두리·필터·틀 고정을 적용한다.
    기본 차트는 처리구명 + 평균값을 명확히 보여주고 오차막대는 자동으로 넣지 않는다.
    """
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, Reference
    from openpyxl.chart.error_bar import ErrorBars
    from openpyxl.chart.label import DataLabelList
    from openpyxl.chart.data_source import NumDataSource, NumRef, AxDataSource, StrRef, StrData, StrVal
    from openpyxl.chart.marker import DataPoint
    from openpyxl.chart.shapes import GraphicalProperties
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    # ① 값과 유의성 문자를 분리해 숫자 셀로 만든다.
    # 문자열 열은 비어 있지 않은 값이 모두 숫자로 해석될 때만 숫자열로 취급한다.
    # 일부만 숫자인 열을 숫자열로 바꾸면 나머지 문자값이 빈칸으로 사라질 수 있다.
    cols, sig = [], {}
    for c in df.columns:
        raw_vals = list(df[c])
        pairs = [_xl_split(v) for v in raw_vals]
        nonempty = [i for i, v in enumerate(raw_vals)
                    if not pd.isna(v) and str(v).strip().lower() not in ('', 'nan', 'none', '-', '―')]
        all_parseable = bool(nonempty) and all(pairs[i][0] is not None for i in nonempty)
        is_native_num = pd.api.types.is_numeric_dtype(df[c]) and not pd.api.types.is_bool_dtype(df[c])
        is_num_col = is_native_num or all_parseable
        if is_num_col:
            vals = [p[0] if p[0] is not None else None for p in pairs]
            cols.append((c, vals, True))
            if any(p[1] for p in pairs):
                sig[c] = [p[1] or '' for p in pairs]
        else:
            cols.append((c, [('' if pd.isna(v) else str(v)) for v in raw_vals], False))

    wb = Workbook()
    ws = wb.active
    ws.title = '데이터'
    ws.sheet_properties.tabColor = '3D6F9F'
    ws.sheet_view.showGridLines = False
    FONT = '맑은 고딕'
    NAVY, HEADER, MID, LIGHT, PALE, LINE = '244A73', '3D6F9F', '9EC5E5', 'EAF3FA', 'F7FBFF', 'C9DCEB'
    thin = Side(style='thin', color=LINE)
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    HDR = 4
    out_cols = []
    for name, vals, is_num in cols:
        out_cols.append((name, vals, is_num))
        if name in sig:
            out_cols.append((f'{name} 유의성', sig[name], False))

    n_out = max(len(out_cols), 1)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=n_out)
    a1 = ws.cell(1, 1, title)
    a1.fill = PatternFill('solid', fgColor=NAVY)
    a1.font = Font(name=FONT, size=14, bold=True, color='FFFFFF')
    a1.alignment = Alignment(horizontal='left', vertical='center')
    ws.row_dimensions[1].height = 30

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=n_out)
    a2 = ws.cell(2, 1, '숫자를 고치면 그래프가 바로 따라 바뀝니다. 그래프를 눌러 색·글꼴·축을 자유롭게 바꾸세요.')
    a2.font = Font(name=FONT, size=9, italic=True, color='5B6F82')
    a2.fill = PatternFill('solid', fgColor='F7FBFF')
    a2.alignment = Alignment(horizontal='left')

    for j, (name, vals, is_num) in enumerate(out_cols, start=1):
        h = ws.cell(row=HDR, column=j, value=str(name))
        h.font = Font(name=FONT, bold=True, color='FFFFFF', size=10)
        h.fill = PatternFill('solid', fgColor=HEADER)
        h.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        h.border = border
        for i, v in enumerate(vals, start=HDR + 1):
            cell = ws.cell(row=i, column=j, value=v)
            cell.font = Font(name=FONT)
            cell.border = border
            cell.alignment = Alignment(horizontal='center' if not is_num else 'right', vertical='center')
            # 값 크기와 무관한 아주 옅은 행 구분만 사용한다.
            # 평균·CV·순위 같은 숫자가 크다는 이유로 더 중요해 보이지 않게 한다.
            fill_color = PALE if (i - HDR) % 2 == 0 else 'FFFFFF'
            cell.fill = PatternFill('solid', fgColor=fill_color)
            if is_num and v is not None:
                try:
                    fv = float(v)
                    cell.number_format = '#,##0.00' if abs(fv) < 1000 else '#,##0.###'
                except Exception:
                    pass
        w = max([len(str(name))] + [len(str(v)) for v in vals[:200]]) * 1.55 + 4
        ws.column_dimensions[get_column_letter(j)].width = min(max(w, 11), 32)

    nrow = len(df)
    if out_cols:
        ws.freeze_panes = f'A{HDR+1}'
        ws.auto_filter.ref = f'A{HDR}:{get_column_letter(len(out_cols))}{HDR+nrow}' if nrow else f'A{HDR}:{get_column_letter(len(out_cols))}{HDR}'

    if not chart or nrow == 0:
        buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

    # ② 범주축은 원본 표의 첫 번째 열을 우선 사용한다.
    # 처리구가 1·2·3 같은 숫자 코드여도 유의성 문자 열을 X축으로 잘못 잡지 않는다.
    lab_idx = 1 if out_cols else None
    val_idx = [j for j, (nm, _, isn) in enumerate(out_cols, 1)
               if j != lab_idx and isn and _xl_is_value_col(nm)]
    err_idx = next((j for j, (nm, _, isn) in enumerate(out_cols, 1)
                    if isn and any(k in str(nm) for k in ('표준편차', '표준오차', 'SD', 'SE'))), None)
    if lab_idx is None or not val_idx:
        buf = io.BytesIO(); wb.save(buf); return buf.getvalue()
    val_idx = val_idx[:3]

    ch = BarChart()
    ch.type = 'col'
    # Excel 기본 테마(style 번호)는 앱의 색을 덮어쓸 수 있어 사용하지 않는다.
    ch.title = title
    # X축은 처리구명이 바로 보이므로 '처리구' 같은 축 제목은 따로 넣지 않는다.
    # 축 제목이 범주명 자리를 먹어 처리구명이 안 보이는 문제를 방지한다.
    # 단위/측정항목은 차트 제목에 이미 포함된다. Excel의 세로축 제목은
    # 폭이 좁을 때 눈금과 겹치므로 표시하지 않는다.
    ch.y_axis.title = None
    ch.x_axis.title = None
    ch.gapWidth = 82
    ch.height, ch.width = 9.2, 17
    ch.roundedCorners = False
    ch.x_axis.delete = False
    ch.y_axis.delete = False
    ch.x_axis.axPos = "b"
    ch.y_axis.axPos = "l"
    ch.x_axis.tickLblPos = "nextTo"
    ch.y_axis.tickLblPos = "nextTo"
    ch.x_axis.majorTickMark = "none"
    ch.y_axis.majorTickMark = "out"
    ch.x_axis.noMultiLvlLbl = True

    cats = Reference(ws, min_col=lab_idx, min_row=HDR + 1, max_row=HDR + nrow)
    for j in val_idx:
        data = Reference(ws, min_col=j, min_row=HDR, max_row=HDR + nrow)
        ch.add_data(data, titles_from_data=True, from_rows=False)
    ch.set_categories(cats)
    # 처리구가 문자일 때 openpyxl의 기본 set_categories()가 numRef를 만들면
    # Excel에서 X축 처리구명이 통째로 안 보일 수 있다. 문자 범주는 strRef로 강제한다.
    if not out_cols[lab_idx - 1][2]:
        _cat_formula = "'데이터'!${}${}:${}${}".format(
            get_column_letter(lab_idx), HDR + 1, get_column_letter(lab_idx), HDR + nrow)
        _cat_values = [str(v) if v is not None else "" for v in out_cols[lab_idx - 1][1]]
        _cache = StrData(
            ptCount=len(_cat_values),
            pt=[StrVal(idx=i, v=v) for i, v in enumerate(_cat_values)]
        )
        for _ser in ch.series:
            _ser.cat = AxDataSource(strRef=StrRef(f=_cat_formula, strCache=_cache))

    # 차트도 화면과 같은 스마트 블루. 단일 계열 막대는 처리별로 밝기 그라데이션을 준다.
    _excel_series = ('3D6F9F', '6FA3CF', '9EC5E5')
    _excel_points = ('C2D9EE', 'A3C4E2', '82ACD3', '6291C2', '4576AB', '2D5A8E', '1F4569')
    for _si, (ser, col) in enumerate(zip(ch.series, _excel_series)):
        try:
            ser.graphicalProperties.solidFill = col
            ser.graphicalProperties.line.solidFill = '000000'
            ser.graphicalProperties.line.width = 6350
            if len(ch.series) == 1 and nrow > 1:
                # 화면의 원클릭 막대처럼 각 처리별로 옅은→진한 블루를 준다.
                pts = []
                for _i in range(nrow):
                    _c = _excel_points[round((len(_excel_points)-1) * _i / max(nrow-1, 1))]
                    _dp = DataPoint(idx=_i)
                    _dp.graphicalProperties = GraphicalProperties(solidFill=_c)
                    _dp.graphicalProperties.line.solidFill = '000000'
                    _dp.graphicalProperties.line.width = 4763
                    pts.append(_dp)
                ser.dPt = pts
        except Exception:
            pass
    try:
        # 차트/플롯 전체를 둘러싸는 검은 사각 프레임은 넣지 않는다.
        ch.graphical_properties = GraphicalProperties(solidFill='FFFFFF')
        ch.graphical_properties.line.noFill = True
        ch.plot_area.graphicalProperties = GraphicalProperties(solidFill='FFFFFF')
        ch.plot_area.graphicalProperties.line.noFill = True

        # 사용자가 요청한 대로 차트 전체/축/격자선의 불필요한 선은 모두 제거한다.
        # 검은 윤곽선은 막대 자체에만 남는다.
        ch.x_axis.spPr = GraphicalProperties()
        ch.x_axis.spPr.line.noFill = True
        ch.y_axis.spPr = GraphicalProperties()
        ch.y_axis.spPr.line.noFill = True
        from openpyxl.chart.axis import ChartLines
        _nogrid_x = ChartLines()
        _nogrid_x.spPr = GraphicalProperties()
        _nogrid_x.spPr.line.noFill = True
        _nogrid_y = ChartLines()
        _nogrid_y.spPr = GraphicalProperties()
        _nogrid_y.spPr.line.noFill = True
        ch.x_axis.majorGridlines = _nogrid_x
        ch.y_axis.majorGridlines = _nogrid_y
    except Exception:
        pass
    if len(ch.series) == 1:
        ch.legend = None
        ch.varyColors = False
        try:
            # 평균값은 막대 위에 바로 표시한다.
            ch.dLbls = DataLabelList()
            ch.dLbls.showVal = True
            ch.dLbls.showCatName = False
            ch.dLbls.showSerName = False
            ch.dLbls.showLegendKey = False
            ch.dLbls.position = "outEnd"
            ch.dLbls.numFmt = '#,##0.00'
        except Exception:
            pass
    else:
        try:
            ch.legend.position = "r"
            ch.legend.overlay = False
            # 응답자 특성처럼 빈도·비율이 함께 있는 그룹 막대도 값이 바로 보이게 한다.
            ch.dLbls = DataLabelList()
            ch.dLbls.showVal = True
            ch.dLbls.showCatName = False
            ch.dLbls.showSerName = False
            ch.dLbls.showLegendKey = False
            ch.dLbls.position = "outEnd"
            ch.dLbls.numFmt = '0.##'
        except Exception:
            pass

    if error_bars and err_idx and len(val_idx) == 1:
        try:
            ref = NumRef("'데이터'!${}${}:${}${}".format(
                get_column_letter(err_idx), HDR + 1, get_column_letter(err_idx), HDR + nrow))
            ch.series[0].errBars = ErrorBars(
                errDir='y', errBarType='both', errValType='cust',
                plus=NumDataSource(numRef=ref), minus=NumDataSource(numRef=ref))
            try:
                ch.series[0].errBars.spPr = GraphicalProperties()
                ch.series[0].errBars.spPr.line.solidFill = '000000'
            except Exception:
                pass
        except Exception:
            pass

    ws.add_chart(ch, f'{get_column_letter(len(out_cols) + 2)}{HDR}')

    note = HDR + nrow + 2
    ws.cell(row=note, column=1, value='※ 유의성 문자(a, b, c)가 있는 결과는 그래프의 막대 위에도 자동 표시됩니다.').font = Font(name=FONT, size=9, color='5B6F82')
    if error_bars and err_idx and len(val_idx) == 1:
        ws.cell(row=note + 1, column=1,
                value=f"※ 오차막대는 '{out_cols[err_idx - 1][0]}' 열을 사용했습니다.").font = Font(name=FONT, size=9, color='5B6F82')

    buf = io.BytesIO(); wb.save(buf)
    _raw = buf.getvalue()
    _raw = _inject_xlsx_significance_labels(_raw, _xlsx_sig_specs_for_df(df, '데이터', HDR))
    return _raw


def _rewrite_chart_sheet_refs(chart, new_sheet, old_sheet='데이터'):
    """openpyxl 차트의 값·범주·오차막대 참조 시트명을 바꾼다."""
    old_tokens = (f"'{old_sheet}'", old_sheet)
    new_token = f"'{new_sheet}'"
    for ser in getattr(chart, 'series', []):
        refs = [getattr(ser, 'val', None), getattr(ser, 'cat', None),
                getattr(ser, 'tx', None), getattr(ser, 'errBars', None)]
        for ref in refs:
            if ref is None:
                continue
            for sub in ('numRef', 'strRef'):
                r = getattr(ref, sub, None)
                if r is not None and getattr(r, 'f', None):
                    f = r.f
                    for old in old_tokens:
                        f = f.replace(old + '!', new_token + '!')
                    r.f = f
            for side in ('plus', 'minus'):
                nd = getattr(ref, side, None)
                r = getattr(nd, 'numRef', None) if nd is not None else None
                if r is not None and getattr(r, 'f', None):
                    f = r.f
                    for old in old_tokens:
                        f = f.replace(old + '!', new_token + '!')
                    r.f = f
    return chart


def make_xlsx_multi(blocks, doc_title='분석 결과'):
    """여러 분석 결과를 항목별 시트로 나누고, 각 시트에 스마트 블루 표/편집가능 차트를 담는다."""
    from openpyxl import load_workbook
    used, sheets = set(), []
    for b in blocks:
        tb = b.get('table')
        if tb is None or not len(tb):
            continue
        nm = str(b.get('caption') or '결과')
        for bad in ':\\/?*[]':
            nm = nm.replace(bad, ' ')
        nm = (nm.strip() or '결과')[:28]
        base, i = nm, 2
        while nm in used:
            nm = f'{base[:26]}_{i}'; i += 1
        used.add(nm)
        sheets.append((nm, tb, str(b.get('caption') or doc_title)))
    if not sheets:
        return None

    first = io.BytesIO(make_xlsx(sheets[0][1], sheets[0][2]))
    wb = load_workbook(first)
    first_ws = wb.active
    first_name = sheets[0][0]
    first_ws.title = first_name
    try:
        first_ws.sheet_properties.tabColor = '3D6F9F'
        first_ws.sheet_view.showGridLines = False
    except Exception:
        pass
    for ch in getattr(first_ws, '_charts', []):
        _rewrite_chart_sheet_refs(ch, first_name)

    for nm, tb, cap in sheets[1:]:
        src = load_workbook(io.BytesIO(make_xlsx(tb, cap)))
        ws_src = src.active
        ws = wb.create_sheet(nm)
        # 병합 셀/행 높이/열 너비/셀 스타일을 최대한 그대로 복사한다.
        for mr in ws_src.merged_cells.ranges:
            ws.merge_cells(str(mr))
        for r, dim in ws_src.row_dimensions.items():
            ws.row_dimensions[r].height = dim.height
        for k, v in ws_src.column_dimensions.items():
            ws.column_dimensions[k].width = v.width
        for row in ws_src.iter_rows():
            for c in row:
                if c.value is None and not c.has_style:
                    continue
                nc = ws.cell(row=c.row, column=c.column, value=c.value)
                nc.font = c.font.copy(); nc.fill = c.fill.copy(); nc.border = c.border.copy()
                nc.alignment = c.alignment.copy(); nc.number_format = c.number_format
                nc.protection = c.protection.copy()
        ws.freeze_panes = ws_src.freeze_panes
        ws.auto_filter.ref = ws_src.auto_filter.ref
        try:
            ws.sheet_properties.tabColor = '3D6F9F'
            ws.sheet_view.showGridLines = False
        except Exception:
            pass
        for ch in getattr(ws_src, '_charts', []):
            try:
                _rewrite_chart_sheet_refs(ch, nm)
                ws.add_chart(ch, ch.anchor)
            except Exception:
                pass
    buf = io.BytesIO(); wb.save(buf)
    _raw = buf.getvalue()
    _specs = {}
    for _nm, _tb, _cap in sheets:
        _specs.update(_xlsx_sig_specs_for_df(_tb, _nm, 4))
    _raw = _inject_xlsx_significance_labels(_raw, _specs)
    return _raw

def dl_table(df, title, key, fname="table", image=None):
    """표(+선택적으로 그래프) 내려받기 — 체크했을 때만 파일을 만들어 화면이 빨라집니다.
    image(PNG bytes)를 넘기면 한글/워드 파일에 표 아래 그래프도 함께 들어갑니다."""
    want = st.checkbox(f"📥 '{title}' 파일로 저장", key=f"dlchk_{key}")
    if not want:
        return
    c1, c2 = st.columns(2)
    csv_text = df.to_csv(index=False)
    opts_sig = tuple(sorted((k, str(v)) for k, v in st.session_state.items()
                            if str(k).startswith("hwp_")))
    try:
        if image:
            item = [{"heading": title, "table": df, "image": image}]
            hwpx_bytes = build_report_hwpx(item, doc_title=title)
        else:
            hwpx_bytes = _make_docs(csv_text, title, "hwpx", opts_sig)
        c1.download_button("📄 한글(hwpx)" + (" (그래프 포함)" if image else ""), hwpx_bytes,
                           f"{fname}.hwpx", key=f"hwx_{key}", width="stretch")
    except Exception as _e:
        c1.caption(f"한글 파일 생성 실패 ({type(_e).__name__})")
    if _HAS_DOCX:
        try:
            if image:
                item = [{"heading": title, "table": df, "image": image}]
                docx_bytes = build_report_docx(item, doc_title=title)
            else:
                docx_bytes = _make_docs(csv_text, title, "docx", opts_sig)
            c2.download_button("📝 워드(docx)" + (" (그래프 포함)" if image else ""), docx_bytes,
                               f"{fname}.docx", key=f"dcx_{key}", width="stretch")
        except Exception as _e:
            c2.caption(f"워드 파일 생성 실패 ({type(_e).__name__})")
    else:
        c2.caption("워드 저장: pip install python-docx 필요")
    # app(9)의 "편집 가능한 Excel 차트" 기능을 유지하면서 스마트 블루 디자인을 적용한다.
    try:
        c2.download_button("📈 엑셀(xlsx) — 스마트 블루 디자인 + 편집 가능한 그래프", make_xlsx(df, title),
                           f"{fname}.xlsx", key=f"xls_{key}", width="stretch",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           help="화면과 같은 스마트 블루 표·차트 디자인으로 저장됩니다. 막대 색·글꼴·축 범위·차트 종류도 엑셀에서 직접 바꿀 수 있습니다.")
    except Exception as _e:
        c2.caption(f"엑셀 파일 생성 실패 ({type(_e).__name__})")
    c1.download_button("📊 CSV (표만)", csv_text.encode("utf-8-sig"),
                       f"{fname}.csv", key=f"csv_{key}", width="stretch")

def report_capture(slot, heading, text=None, table=None, image=None, blocks=None):
    """단일 표/그림, 또는 여러 개(blocks)를 담을 수 있음.
    blocks = [{'text':.., 'table':df, 'image':png}, ...]"""
    st.session_state[slot] = {"heading": heading, "text": text,
                              "table": table, "image": image, "blocks": blocks}

def report_button(slot, label="➕ 이 결과를 보고서에 담기"):
    if st.session_state.get(slot):
        if st.button(label, key="btn_" + slot):
            st.session_state.report_items.append(st.session_state[slot])
            st.success(f"보고서에 담았습니다! (현재 {len(st.session_state.report_items)}개) — '📑 보고서'에서 생성하세요.")


def survey_download_panel(slot, key, fname):
    """설문 결과의 한글/엑셀 다운로드를 체크박스에 숨기지 않고 바로 보여준다."""
    item = st.session_state.get(slot)
    if not item:
        return
    st.markdown("#### 📥 설문 분석 결과 다운로드")
    st.caption("한글은 표·그래프를 묶은 보고서로, Excel은 결과표를 항목별 시트로 저장합니다.")

    # 이전 버전의 상태 보존 로직이 download_button key를 session_state에 써 둔 세션에서는
    # StreamlitValueAssignmentNotAllowedError가 계속 재현될 수 있다. 렌더 직전에 정리한다.
    for _wk in (f"dl_svyhwp_{key}", f"dl_svyxls_{key}"):
        try:
            if _wk in st.session_state:
                del st.session_state[_wk]
        except Exception:
            pass

    c1, c2 = st.columns(2)
    try:
        hwp = build_report_hwpx([item], doc_title=item.get("heading", "설문조사 분석 결과"))
        c1.download_button("📘 한글 보고서(hwpx)", hwp, f"{fname}.hwpx",
                           key=f"dl_svyhwp_{key}", width="stretch")
    except Exception as ex:
        c1.caption(f"한글 파일 생성 실패 ({type(ex).__name__})")
    blocks = item.get("blocks") or []
    if not blocks and item.get("table") is not None:
        blocks = [{"caption": item.get("heading", "설문 결과"), "table": item.get("table")}]
    xblocks = [b for b in blocks if b.get("table") is not None]
    try:
        if xblocks:
            xls = make_xlsx_multi(xblocks, doc_title=item.get("heading", "설문조사 분석 결과"))
            c2.download_button("📈 Excel(xlsx) — 항목별 시트", xls, f"{fname}.xlsx",
                               key=f"dl_svyxls_{key}", width="stretch",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            c2.caption("Excel로 저장할 표가 없습니다.")
    except Exception as ex:
        c2.caption(f"Excel 파일 생성 실패 ({type(ex).__name__})")


# ================================================================ 회원가입/로그인 (Supabase Auth)
# 별도 Python 패키지 없이 Supabase의 공식 Auth/REST endpoint를 requests로 호출한다.
def _secret(name, default=""):
    import os
    try:
        v = st.secrets.get(name, None)
        if v is not None:
            return str(v)
    except Exception:
        pass
    return str(os.environ.get(name, default))


def _truthy(v):
    return str(v).strip().lower() in ("1", "true", "yes", "y", "on")


def _auth_config():
    # 2026년 Supabase 권장 키: publishable / secret.
    # 기존 프로젝트의 anon / service_role 키도 호환되도록 fallback을 둔다.
    publishable = _secret("SUPABASE_PUBLISHABLE_KEY") or _secret("SUPABASE_ANON_KEY")
    secret = _secret("SUPABASE_SECRET_KEY") or _secret("SUPABASE_SERVICE_ROLE_KEY")
    return {
        "url": _secret("SUPABASE_URL").rstrip("/"),
        "anon": publishable,      # 기존 코드 호환용 이름
        "service": secret,        # 기존 코드 호환용 이름
        "required": _truthy(_secret("AUTH_REQUIRED", "false")),
        "admins": {x.strip().lower() for x in _secret("ADMIN_EMAILS", "").split(",") if x.strip()},
    }


def _supabase_headers(token=None, service=False, content_type=False):
    """Supabase 새 API key 모델과 legacy key를 모두 지원하는 헤더."""
    cfg = _auth_config()
    key = cfg["service"] if service and cfg["service"] else cfg["anon"]
    headers = {"apikey": key} if key else {}
    # 새 sb_publishable_/sb_secret_ 키는 API key 자체를 Bearer JWT로 보내면 안 된다.
    # 사용자 세션 JWT가 있을 때만 Authorization에 넣는다.
    if token:
        headers["Authorization"] = f"Bearer {token}"
    elif key and not str(key).startswith("sb_"):
        # legacy anon/service_role은 JWT 형태라 기존 방식과 호환된다.
        headers["Authorization"] = f"Bearer {key}"
    if content_type:
        headers["Content-Type"] = "application/json"
    return headers


def _supabase_post(path, payload, token=None, service=False, params=None, timeout=25):
    cfg = _auth_config()
    if not cfg["url"] or not cfg["anon"]:
        return None, "Supabase 설정이 없습니다."
    if service and not cfg["service"]:
        return None, "Supabase Secret key 설정이 없습니다."
    headers = _supabase_headers(token=token, service=service, content_type=True)
    try:
        r = _requests.post(cfg["url"] + path, headers=headers, json=payload,
                           params=params, timeout=timeout)
    except Exception as ex:
        return None, f"네트워크 오류: {type(ex).__name__}"
    if r.status_code not in (200, 201, 204):
        try:
            body = r.json() or {}
            msg = body.get("msg") or body.get("message") or body.get("error_description") or body.get("error") or r.text
        except Exception:
            msg = r.text
        return None, f"{r.status_code}: {str(msg)[:180]}"
    if r.status_code == 204 or not r.text.strip():
        return {}, None
    try:
        return r.json(), None
    except Exception:
        return {}, None


def _supabase_get(path, token=None, service=False, params=None, timeout=25):
    cfg = _auth_config()
    if not cfg["url"] or not cfg["anon"]:
        return None, "Supabase 설정이 없습니다."
    if service and not cfg["service"]:
        return None, "Supabase Secret key 설정이 없습니다."
    headers = _supabase_headers(token=token, service=service)
    try:
        r = _requests.get(cfg["url"] + path, headers=headers, params=params, timeout=timeout)
    except Exception as ex:
        return None, f"네트워크 오류: {type(ex).__name__}"
    if r.status_code != 200:
        return None, f"{r.status_code}: {r.text[:180]}"
    try:
        return r.json(), None
    except Exception:
        return None, "응답을 읽지 못했습니다."

def _supabase_signup(email, password, name, org_type, organization, department):
    payload = {
        "email": email.strip(), "password": password,
        "data": {"name": name.strip(), "organization_type": org_type,
                 "organization": organization.strip(), "department": department.strip()}
    }
    return _supabase_post("/auth/v1/signup", payload)


def _supabase_login(email, password):
    return _supabase_post("/auth/v1/token", {"email": email.strip(), "password": password},
                          params={"grant_type": "password"})


def _supabase_refresh(refresh_token):
    return _supabase_post("/auth/v1/token", {"refresh_token": refresh_token},
                          params={"grant_type": "refresh_token"})


def _supabase_recover(email):
    return _supabase_post("/auth/v1/recover", {"email": email.strip()})


def _save_auth_session(js):
    import time
    if not js:
        return False
    user = js.get("user") or {}
    token = js.get("access_token")
    if not token or not user:
        return False
    st.session_state["auth_user"] = user
    st.session_state["auth_access_token"] = token
    st.session_state["auth_refresh_token"] = js.get("refresh_token", "")
    st.session_state["auth_expires_at"] = time.time() + float(js.get("expires_in") or 3600)
    return True


def _current_auth_user():
    import time
    user = st.session_state.get("auth_user")
    if not user:
        return None
    if time.time() > float(st.session_state.get("auth_expires_at", 0)) - 60:
        rt = st.session_state.get("auth_refresh_token")
        if rt:
            js, err = _supabase_refresh(rt)
            if not err and _save_auth_session(js):
                user = st.session_state.get("auth_user")
    return user


def _record_login(user):
    """관리자 대시보드용 프로필/접속 기록. service role secret은 서버 안에서만 사용."""
    cfg = _auth_config()
    if not user or not cfg["service"]:
        return
    import datetime as _dt
    meta = user.get("user_metadata") or {}
    now = _dt.datetime.now(_dt.timezone.utc).isoformat()
    profile = {
        "id": user.get("id"), "email": user.get("email", ""),
        "name": meta.get("name", ""), "organization_type": meta.get("organization_type", ""),
        "organization": meta.get("organization", ""), "department": meta.get("department", ""),
        "last_login_at": now,
    }
    # profiles: id unique/PK 전제. merge-duplicates로 가입자 프로필을 갱신한다.
    cfg2 = _auth_config(); key = cfg2["service"]
    headers = _supabase_headers(service=True, content_type=True)
    headers["Prefer"] = "resolution=merge-duplicates,return=minimal"
    try:
        _requests.post(cfg2["url"] + "/rest/v1/profiles", headers=headers,
                       params={"on_conflict": "id"}, json=profile, timeout=15)
        event = {"user_id": profile.get("id"), "email": profile.get("email", ""),
                 "name": profile.get("name", ""), "organization_type": profile.get("organization_type", ""),
                 "organization": profile.get("organization", ""), "department": profile.get("department", ""),
                 "logged_in_at": now}
        _requests.post(cfg2["url"] + "/rest/v1/login_events", headers=headers,
                       json=event, timeout=15)
    except Exception:
        pass


def _auth_logout():
    for k in ("auth_user", "auth_access_token", "auth_refresh_token", "auth_expires_at"):
        st.session_state.pop(k, None)
    st.rerun()


def _record_usage(action):
    """로그인 사용자가 실제 분석/보고서 기능을 실행했을 때 기관별 이용량을 남긴다."""
    cfg = _auth_config(); user = _current_auth_user()
    if not user or not cfg["service"]:
        return
    import datetime as _dt
    meta = user.get("user_metadata") or {}
    key = cfg["service"]
    headers = _supabase_headers(service=True, content_type=True)
    payload = {
        "user_id": user.get("id"), "email": user.get("email", ""),
        "name": meta.get("name", ""), "organization": meta.get("organization", ""),
        "department": meta.get("department", ""), "action": str(action)[:300],
        "used_at": _dt.datetime.now(_dt.timezone.utc).isoformat(),
    }
    try:
        _requests.post(cfg["url"] + "/rest/v1/usage_events", headers=headers, json=payload, timeout=10)
    except Exception:
        pass


def render_auth_gate():
    """AUTH_REQUIRED=true이고 Supabase가 설정된 경우 회원만 앱에 진입하게 한다."""
    cfg = _auth_config()
    configured = bool(cfg["url"] and cfg["anon"])
    if not (configured and cfg["required"]):
        return True
    user = _current_auth_user()
    if user:
        return True

    st.markdown("<br><br>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1.35, 1])
    with c2:
        st.markdown("## 📊 스마트 통계 에이전트")
        st.caption("회원가입 후 연구 데이터를 쉽고 정확하게 분석하세요.")
        login_tab, signup_tab, reset_tab = st.tabs(["🔐 로그인", "✨ 회원가입", "🔑 비밀번호 찾기"])
        with login_tab:
            em = st.text_input("이메일", key="auth_login_email")
            pw = st.text_input("비밀번호", type="password", key="auth_login_pw")
            if st.button("로그인", type="primary", width="stretch", key="auth_login_btn"):
                if not em or not pw:
                    st.warning("이메일과 비밀번호를 입력해 주세요.")
                else:
                    with st.spinner("로그인 중..."):
                        js, err = _supabase_login(em, pw)
                    if err:
                        st.error("로그인에 실패했습니다. 이메일/비밀번호 또는 이메일 인증 여부를 확인해 주세요.")
                    elif _save_auth_session(js):
                        _record_login(st.session_state.get("auth_user"))
                        st.rerun()
                    else:
                        st.error("로그인 응답을 확인하지 못했습니다.")

        with signup_tab:
            nm = st.text_input("이름", key="auth_name")
            em2 = st.text_input("이메일", key="auth_signup_email")
            pw2 = st.text_input("비밀번호 (8자 이상 권장)", type="password", key="auth_signup_pw")
            org_type = st.selectbox("기관 유형", ["도·특광역시 농업기술원", "농촌진흥청/소속기관",
                                                  "시·군 농업기술센터", "대학교/연구기관",
                                                  "농업 관련 기업/단체", "기타"], key="auth_org_type")
            org = st.text_input("소속기관", placeholder="예: 경상북도농업기술원", key="auth_org")
            dept = st.text_input("부서/연구소 (선택)", placeholder="예: 영양고추연구소", key="auth_dept")
            consent = st.checkbox("이름·이메일·소속기관 및 서비스 접속기록을 운영 목적으로 저장하는 것에 동의합니다.",
                                  key="auth_consent")
            if st.button("회원가입", type="primary", width="stretch", key="auth_signup_btn"):
                if not all([nm, em2, pw2, org]) or not consent:
                    st.warning("이름·이메일·비밀번호·소속기관과 개인정보 안내 동의를 확인해 주세요.")
                elif len(pw2) < 8:
                    st.warning("비밀번호는 8자 이상을 권장합니다.")
                else:
                    with st.spinner("회원가입 중..."):
                        js, err = _supabase_signup(em2, pw2, nm, org_type, org, dept)
                    if err:
                        st.error(f"회원가입에 실패했습니다: {err}")
                    elif _save_auth_session(js):
                        _record_login(st.session_state.get("auth_user"))
                        st.success("가입과 로그인이 완료되었습니다.")
                        st.rerun()
                    else:
                        st.success("가입 신청이 완료되었습니다. 이메일 인증 메일을 확인한 뒤 로그인해 주세요.")

        with reset_tab:
            rem = st.text_input("가입한 이메일", key="auth_reset_email")
            if st.button("비밀번호 재설정 메일 보내기", width="stretch", key="auth_reset_btn"):
                if not rem:
                    st.warning("이메일을 입력해 주세요.")
                else:
                    _, err = _supabase_recover(rem)
                    if err:
                        st.error("재설정 메일 요청에 실패했습니다.")
                    else:
                        st.success("재설정 안내 메일을 보냈습니다.")
        st.caption("🔒 비밀번호는 스마트 통계 에이전트 코드가 직접 저장하지 않고 Supabase Auth가 처리합니다.")
    st.stop()


def _is_admin_user(user=None):
    user = user or _current_auth_user()
    email = str((user or {}).get("email", "")).lower()
    return bool(email and email in _auth_config()["admins"])


def render_admin_dashboard():
    st.title("👑 관리자 — 이용 현황")
    cfg = _auth_config()
    if not _is_admin_user():
        st.error("관리자 계정만 접근할 수 있습니다.")
        return
    if not cfg["service"]:
        st.warning("SUPABASE_SECRET_KEY(또는 기존 SERVICE_ROLE_KEY)가 설정되어야 관리자 통계를 볼 수 있습니다.")
        return
    profiles, e1 = _supabase_get("/rest/v1/profiles", service=True,
                                 params={"select": "*", "order": "last_login_at.desc"})
    events, e2 = _supabase_get("/rest/v1/login_events", service=True,
                               params={"select": "*", "order": "logged_in_at.desc", "limit": 5000})
    usage, e3 = _supabase_get("/rest/v1/usage_events", service=True,
                              params={"select": "*", "order": "used_at.desc", "limit": 10000})
    if e1:
        st.error(f"회원 목록을 불러오지 못했습니다: {e1}")
        return
    p = pd.DataFrame(profiles or [])
    ev = pd.DataFrame(events or [])
    uv = pd.DataFrame(usage or [])
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("가입/프로필 사용자", f"{len(p):,}명")
    c2.metric("확인된 소속기관", f"{p['organization'].replace('', np.nan).nunique() if 'organization' in p else 0:,}곳")
    c3.metric("로그인 기록", f"{len(ev):,}회")
    c4.metric("기능 이용 기록", f"{len(uv):,}회")
    if not p.empty and "organization" in p:
        st.markdown("### 🏢 기관별 사용자")
        g = (p.assign(소속기관=p["organization"].fillna("미입력").replace("", "미입력"))
               .groupby("소속기관", dropna=False).size().reset_index(name="사용자 수")
               .sort_values("사용자 수", ascending=False))
        smart_table(g, width="stretch", hide_index=True)
    if not p.empty:
        st.markdown("### 👥 사용자 목록")
        cols = [c for c in ["name", "email", "organization_type", "organization", "department", "last_login_at"] if c in p]
        show = p[cols].rename(columns={"name":"이름", "email":"이메일", "organization_type":"기관유형",
                                      "organization":"소속기관", "department":"부서", "last_login_at":"최근 로그인"})
        smart_table(show, width="stretch", hide_index=True)
        st.download_button("📊 사용자 목록 Excel", dataframe_to_styled_xlsx(show, "스마트 통계 에이전트 사용자 목록"),
                           "사용자목록.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    if not uv.empty:
        st.markdown("### 📈 기관별 기능 이용")
        _orguse = (uv.assign(소속기관=uv.get("organization", pd.Series(index=uv.index, dtype=object)).fillna("미입력").replace("", "미입력"))
                     .groupby("소속기관", dropna=False).size().reset_index(name="기능 이용 횟수")
                     .sort_values("기능 이용 횟수", ascending=False))
        smart_table(_orguse, width="stretch", hide_index=True)
        st.markdown("### 🧭 최근 기능 이용 기록")
        _ucols = [c for c in ["used_at", "name", "email", "organization", "department", "action"] if c in uv]
        _ushow = uv[_ucols].head(500).rename(columns={"used_at":"이용시각", "name":"이름", "email":"이메일",
                                                        "organization":"소속기관", "department":"부서", "action":"기능"})
        smart_table(_ushow, width="stretch", hide_index=True)

    if not ev.empty:
        st.markdown("### 🕘 최근 로그인")
        cols = [c for c in ["logged_in_at", "name", "email", "organization", "department"] if c in ev]
        show2 = ev[cols].head(300).rename(columns={"logged_in_at":"접속시각", "name":"이름", "email":"이메일",
                                                   "organization":"소속기관", "department":"부서"})
        smart_table(show2, width="stretch", hide_index=True)

# ---------------------------------------------------------------- AI 호출
_AI_SYS = (
    "당신은 농촌진흥청 및 도 농업기술원의 수석 응용통계 전문가입니다. "
    "농업 시험연구보고서와 농학 논문 작성을 20년간 지원해 왔습니다.\n\n"
    "【통계 해석 지침】\n"
    "- p-value: 단순히 '유의하다'로 끝내지 말고, 처리가 실제 생육·수량 반응으로 이어졌는지를 "
    "농학적으로 서술합니다. p<0.05는 '처리 간 차이가 우연으로 보기 어렵다'는 의미이며, "
    "p>=0.05는 '차이가 관찰되었더라도 오차 범위 내'임을 분명히 밝힙니다.\n"
    "- 자유도(df): 반복수·처리수와 연결해 시험 규모의 타당성을 언급합니다.\n"
    "- 변이계수(CV%): 포장시험 정밀도 지표로 해석합니다. 10% 미만 매우 우수, "
    "10~20% 양호, 20% 초과 시 포장 불균일·조사 오차 가능성을 지적합니다.\n"
    "- 사후검정 문자(a, b, c): 같은 문자를 공유하면 통계적으로 동등한 수준임을 뜻합니다. "
    "'ab'처럼 두 군에 걸친 처리는 중간 수준으로 해석합니다.\n"
    "- LSD: 두 평균의 차이가 이 값보다 클 때 유의하다고 서술합니다.\n\n"
    "【농학적 연결】\n"
    "- 수량 증가는 초장·엽수·생체중 등 생육 형질의 변화와 연결해 설명합니다.\n"
    "- 통계적 유의성과 농업적 실용성(증수량이 농가 소득에 미치는 영향)을 구분해 서술합니다.\n"
    "- 확인할 수 없는 원인(품종 특성, 기상, 토양)은 단정하지 말고 '~로 추정된다'로 씁니다.\n\n"
    "【서식 규칙】\n"
    "- 농촌진흥청 시험연구보고서 양식을 따릅니다. 주요 항목은 '○ ', 세부 항목은 '  - '로 시작합니다.\n"
    "- 마크다운 기호(**, ##, *, `, ---)는 절대 사용하지 않습니다. 평문으로만 작성합니다.\n"
    "- 문체는 '~하였다', '~로 나타났다', '~인 것으로 판단된다'를 사용합니다.\n"
    "- 표에 없는 수치나 사실은 만들어 내지 않습니다.\n\n"
    "【반드시 지킬 규칙】\n"
    "1. 제공된 JSON에 없는 숫자·처리명·p값·출처·원인을 만들어 내지 않습니다.\n"
    "2. 검정을 수행하지 않은 결과에 'p<0.05' 같은 표현을 붙이지 않습니다. "
    "yield_statistical_test.p_value가 null이면 유의성을 단정하지 않습니다.\n"
    "3. 통계적 유의성과 농업적·경제적 중요성을 구분해 서술합니다.\n"
    "4. 인과관계가 확인되지 않았으면 원인을 단정하지 않습니다('~로 추정된다').\n"
    "5. 입력 근거가 부족하면 '제공된 결과만으로는 판단할 수 없다'고 답합니다.\n"
    "6. 기준연도가 오래된 가격은 최신값처럼 표현하지 않고 기준연도를 함께 밝힙니다.\n"
    "7. 계산 결과와 서술이 어긋나면 계산 결과를 우선합니다.")

# AI 모델 목록 — 환경변수나 models.txt로 덮어쓸 수 있음(새 모델 나오면 코드 수정 불필요)
#   · 환경변수 예)  AI_MODELS_GEMINI="gemini-2.5-flash,gemini-2.5-pro"
#   · 또는 app.py 옆에 models.txt 파일:  gemini=gemini-2.5-flash,gemini-2.5-pro
_AI_PROVIDERS_BASE = {
    "Claude (Anthropic)": {
        "key": "claude",
        "models": ["claude-haiku-4-5-20251001", "claude-sonnet-5"],
        "key_hint": "console.anthropic.com 에서 발급 (sk-ant-...)"},
    "Gemini (Google)": {
        "key": "gemini",
        "models": ["gemini-3.6-flash", "gemini-3.5-flash"],
        "key_hint": "aistudio.google.com 에서 발급 (AIza...)"},
    "ChatGPT (OpenAI)": {
        "key": "openai",
        "models": ["gpt-5-mini", "gpt-5.1"],
        "key_hint": "platform.openai.com 에서 발급 (sk-...)"},
}

def _load_model_overrides():
    """환경변수·models.txt에서 모델 목록을 읽어 덮어씀"""
    import os
    over = {}
    for name, cfg in _AI_PROVIDERS_BASE.items():
        env = os.environ.get(f"AI_MODELS_{cfg['key'].upper()}")
        if env:
            over[cfg["key"]] = [m.strip() for m in env.split(",") if m.strip()]
    try:
        p = os.path.join(os.path.dirname(os.path.abspath(__file__)), "models.txt")
        if os.path.exists(p):
            with open(p, encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith("#") or "=" not in line:
                        continue
                    k, v = line.split("=", 1)
                    ms = [m.strip() for m in v.split(",") if m.strip()]
                    if ms: over[k.strip().lower()] = ms
    except Exception:
        pass
    return over

def build_ai_providers():
    over = _load_model_overrides()
    out = {}
    for name, cfg in _AI_PROVIDERS_BASE.items():
        models = over.get(cfg["key"], cfg["models"])
        out[name] = {**cfg, "models": list(models)}
    return out

_AI_PROVIDERS = build_ai_providers()

def test_ai_connection(provider, api_key, model):
    """API 연결을 짧게 확인. 반환: dict(ok, provider, model, message, sample).

    Streamlit 위젯 key인 ``ai_provider``는 위젯이 생성된 뒤 수정할 수 없으므로,
    연결 테스트에서는 세션 상태를 바꾸지 않고 provider를 ai_call에 직접 전달한다.
    """
    out = {"ok": False, "provider": provider.split()[0], "model": model or "(미지정)",
           "message": "", "sample": ""}
    if not api_key:
        out["message"] = "API 키가 비어 있습니다."
        return out
    if not model:
        out["message"] = "모델명을 지정해 주세요."
        return out
    # GPT-5 계열은 max_output_tokens 안에 보이지 않는 추론 토큰도 포함됩니다.
    # 20토큰처럼 너무 작게 잡으면 인증은 성공해도 추론 예산을 모두 써서
    # 실제 텍스트가 비어 있을 수 있으므로 연결 테스트에는 넉넉한 예산을 둡니다.
    r = ai_call("연결 테스트입니다. '확인'이라고만 답하세요.",
                api_key=api_key, model=model, max_tokens=512,
                system="한 단어로만 답하세요.", provider=provider)
    if isinstance(r, str) and r.startswith("⚠️"):
        low = r.lower()
        if "401" in r or "auth" in low or "unauthor" in low or "api key" in low:
            out["message"] = "인증 실패 — API 키를 다시 확인해 주세요."
        elif "404" in r or "not found" in low or "model" in low:
            out["message"] = f"모델 '{model}'을(를) 찾을 수 없습니다. 모델명을 확인해 주세요."
        elif "429" in r or "rate" in low or "quota" in low:
            out["message"] = "사용량 한도에 걸렸습니다. 잠시 후 다시 시도해 주세요."
        elif "timeout" in low or "timed out" in low:
            out["message"] = "응답 시간이 초과되었습니다. 네트워크를 확인해 주세요."
        elif "500" in r or "502" in r or "503" in r:
            out["message"] = "제공사 서버 오류입니다. 잠시 후 다시 시도해 주세요."
        else:
            out["message"] = r.replace("⚠️", "").strip()[:120]
        return out
    if not str(r).strip():
        out["message"] = "응답이 비어 있습니다(안전 필터 차단 가능성)."
        return out
    out["ok"] = True
    out["message"] = "정상"
    out["sample"] = str(r).strip()
    return out


def _ai_mask(text, api_key):
    """오류 메시지에 API 키가 섞여 나가지 않도록 제거"""
    t = str(text)
    if api_key:
        t = t.replace(str(api_key), "***")
        if len(str(api_key)) > 8:
            t = t.replace(str(api_key)[:8], "***")
    return t


def _ai_http(method_fn, *args, api_key=None, retries=2, **kw):
    """HTTP 호출 + 제한적 재시도(429·5xx만, 최대 2회). 무한 재시도 없음."""
    import time
    last = None
    for attempt in range(retries + 1):
        try:
            r = method_fn(*args, **kw)
        except Exception as ex:
            last = ("exception", f"{type(ex).__name__}")
            if attempt < retries:
                time.sleep(1.0 * (attempt + 1))
                continue
            return None, f"네트워크 오류({last[1]}) — 연결을 확인해 주세요."
        if r.status_code == 200:
            return r, None
        if r.status_code == 429 and attempt < retries:
            time.sleep(2.0 * (attempt + 1))      # 짧은 backoff
            continue
        if 500 <= r.status_code < 600 and attempt < retries:
            time.sleep(1.5 * (attempt + 1))
            continue
        return None, _ai_error_message(r.status_code,
                                       _ai_mask(getattr(r, "text", ""), api_key))
    return None, "요청이 반복 실패했습니다. 잠시 후 다시 시도해 주세요."


def _ai_error_message(code, body=""):
    """HTTP 상태코드를 사용자 친화적 메시지로"""
    b = str(body)[:160]
    if code in (401, 403):
        return f"⚠️ 인증 실패({code}) — API 키를 다시 확인해 주세요."
    if code == 404:
        return f"⚠️ 모델을 찾을 수 없습니다({code}) — 모델명을 확인해 주세요."
    if code == 429:
        return f"⚠️ 사용량 한도 초과({code}) — 잠시 후 다시 시도해 주세요."
    if code == 400:
        return f"⚠️ 요청 형식 오류({code}): {b}"
    if 500 <= code < 600:
        return f"⚠️ 제공사 서버 오류({code}) — 잠시 후 다시 시도해 주세요."
    return f"⚠️ 호출 실패({code}): {b}"


def call_claude(prompt, api_key, model, max_tokens=900, system=None, timeout=60):
    """Claude(Anthropic) 호출"""
    if not _HAS_ANTHROPIC:
        return "⚠️ anthropic 라이브러리가 없습니다. (pip install anthropic)"
    try:
        client = anthropic.Anthropic(api_key=api_key, timeout=timeout, max_retries=2)
        msg = client.messages.create(
            model=model, max_tokens=max_tokens, system=system or _AI_SYS,
            messages=[{"role": "user", "content": prompt}])
        parts = [getattr(b, "text", "") for b in getattr(msg, "content", [])
                 if getattr(b, "type", "") == "text"]
        text = "".join(parts).strip()
        if not text:
            return "⚠️ 응답이 비어 있습니다(안전 필터 차단 또는 토큰 부족 가능성)."
        return text
    except Exception as ex:
        msg = _ai_mask(f"{type(ex).__name__}: {ex}", api_key)
        low = msg.lower()
        if "authentication" in low or "401" in msg or "api key" in low:
            return "⚠️ 인증 실패 — API 키를 다시 확인해 주세요."
        if "not_found" in low or "404" in msg or "model" in low:
            return f"⚠️ 모델 '{model}'을(를) 찾을 수 없습니다."
        if "rate" in low or "429" in msg:
            return "⚠️ 사용량 한도 초과 — 잠시 후 다시 시도해 주세요."
        if "timeout" in low:
            return "⚠️ 응답 시간이 초과되었습니다."
        return f"⚠️ Claude 호출 오류: {msg[:150]}"


def list_gemini_models(api_key, timeout=20):
    """현재 키로 generateContent를 지원하는 Gemini 모델 목록을 조회한다."""
    if not _HAS_REQUESTS:
        return [], "requests 라이브러리가 없습니다."
    r, err = _ai_http(_requests.get,
                      "https://generativelanguage.googleapis.com/v1beta/models",
                      api_key=api_key, headers={"x-goog-api-key": api_key},
                      params={"pageSize": 1000}, timeout=timeout, retries=1)
    if err:
        return [], err
    try:
        js = r.json()
    except Exception:
        return [], "Gemini 모델 목록 응답을 해석하지 못했습니다."
    out = []
    for item in js.get("models") or []:
        actions = item.get("supportedGenerationMethods") or item.get("supportedActions") or []
        if "generateContent" not in actions:
            continue
        name = str(item.get("name", "")).replace("models/", "", 1)
        if name and name not in out:
            out.append(name)
    return out, None


def list_openai_models(api_key, timeout=20):
    """현재 키에 열려 있는 텍스트 생성용 OpenAI 모델 후보를 조회한다."""
    if not _HAS_REQUESTS:
        return [], "requests 라이브러리가 없습니다."
    r, err = _ai_http(
        _requests.get, "https://api.openai.com/v1/models", api_key=api_key,
        headers={"Authorization": f"Bearer {api_key}"}, timeout=timeout, retries=1)
    if err:
        return [], err
    try:
        data = r.json().get("data") or []
    except Exception:
        return [], "OpenAI 모델 목록 응답을 해석하지 못했습니다."
    excluded = ("audio", "realtime", "transcribe", "tts", "image", "search",
                "chatgpt", "codex", "embedding", "moderation")
    allowed_prefix = ("gpt-5", "gpt-4.1", "gpt-4o", "o3", "o4")
    models = []
    for item in data:
        mid = str((item or {}).get("id", ""))
        low = mid.lower()
        if mid.startswith(allowed_prefix) and not any(x in low for x in excluded):
            models.append(mid)
    # 날짜 고정 스냅샷보다 일반 별칭을 우선 표시한다.
    models = sorted(set(models), key=lambda x: (x.count("-"), len(x), x))
    return models, None


def list_claude_models(api_key, timeout=20):
    """현재 키에 열려 있는 Claude 모델 목록을 조회한다."""
    if not _HAS_REQUESTS:
        return [], "requests 라이브러리가 없습니다."
    r, err = _ai_http(
        _requests.get, "https://api.anthropic.com/v1/models", api_key=api_key,
        headers={"x-api-key": api_key, "anthropic-version": "2023-06-01"},
        params={"limit": 100}, timeout=timeout, retries=1)
    if err:
        return [], err
    try:
        data = r.json().get("data") or []
    except Exception:
        return [], "Claude 모델 목록 응답을 해석하지 못했습니다."
    models = [str((item or {}).get("id", "")) for item in data
              if str((item or {}).get("id", "")).startswith("claude-")]
    return list(dict.fromkeys(m for m in models if m)), None


def call_gemini(prompt, api_key, model, max_tokens=900, system=None, timeout=60):
    """Gemini(Google) 호출 — 인증키를 URL이 아닌 헤더로 전달"""
    if not _HAS_REQUESTS:
        return "⚠️ requests 라이브러리가 없습니다. (pip install requests)"
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    payload = {"contents": [{"parts": [{"text": prompt}]}],
               "generationConfig": {"maxOutputTokens": max_tokens}}
    if system:
        payload["systemInstruction"] = {"parts": [{"text": system}]}
    headers = {"x-goog-api-key": api_key, "Content-Type": "application/json"}
    r, err = _ai_http(_requests.post, url, api_key=api_key,
                      json=payload, headers=headers, timeout=timeout)
    if err:
        return err
    try:
        js = r.json()
    except Exception:
        return "⚠️ Gemini 응답을 해석하지 못했습니다."
    cands = js.get("candidates") or []
    if not cands:
        fb = (js.get("promptFeedback") or {}).get("blockReason")
        if fb:
            return f"⚠️ 안전 필터에 차단되었습니다(사유: {fb})."
        return "⚠️ Gemini 응답이 비어 있습니다."
    parts = ((cands[0] or {}).get("content") or {}).get("parts") or []
    text = "".join(p.get("text", "") for p in parts if isinstance(p, dict)).strip()
    if not text:
        fr = (cands[0] or {}).get("finishReason", "")
        return f"⚠️ Gemini 응답이 비어 있습니다{f'(사유: {fr})' if fr else ''}."
    return text


def call_openai(prompt, api_key, model, max_tokens=900, system=None, timeout=60):
    """OpenAI Responses API 호출. 결과 배열 구조가 달라도 텍스트를 안전하게 추출한다."""
    if not _HAS_REQUESTS:
        return "⚠️ requests 라이브러리가 없습니다. (pip install requests)"
    import re as _re
    requested_tokens = max(1, int(max_tokens))
    model_l = str(model).lower()
    is_gpt5 = model_l.startswith("gpt-5")
    # gpt-5.1, gpt-5.2 처럼 '점(.)'이 붙은 버전은 reasoning.effort로 minimal을 지원하지
    # 않거나(모델에 따라 오류) 권장하지 않는다 — OpenAI는 GPT-5.1부터 'none'을 새로 추가하고
    # 이를 기본값·권장값으로 안내한다. 점 없는 gpt-5(-mini/-nano 포함)는 계속 minimal을 쓴다.
    is_gpt51_plus = bool(_re.match(r"gpt-5\.\d", model_l))
    # Responses API의 max_output_tokens는 화면에 보이는 답변뿐 아니라 추론 토큰도 포함합니다.
    # GPT-5 계열에서 너무 작은 값은 status=incomplete / reason=max_output_tokens와
    # 빈 output_text를 만들 수 있으므로 최소 예산과 낮은 추론 강도를 적용합니다.
    output_budget = max(requested_tokens, 512) if is_gpt5 else requested_tokens
    payload = {"model": model, "input": prompt,
               "max_output_tokens": output_budget, "store": False}
    if is_gpt51_plus:
        payload["reasoning"] = {"effort": "none"}
    elif is_gpt5:
        payload["reasoning"] = {"effort": "minimal"}
    if system:
        payload["instructions"] = system
    r, err = _ai_http(
        _requests.post, "https://api.openai.com/v1/responses", api_key=api_key,
        headers={"Authorization": f"Bearer {api_key}",
                 "Content-Type": "application/json"},
        json=payload, timeout=timeout)
    if err:
        return err
    try:
        js = r.json()
    except Exception:
        return "⚠️ ChatGPT 응답을 해석하지 못했습니다."
    text = str(js.get("output_text") or "").strip()
    if not text:
        parts = []
        for item in js.get("output") or []:
            if not isinstance(item, dict):
                continue
            for content in item.get("content") or []:
                if not isinstance(content, dict):
                    continue
                if content.get("type") in ("output_text", "text"):
                    value = content.get("text", "")
                    if isinstance(value, dict):
                        value = value.get("value", "")
                    if value:
                        parts.append(str(value))
        text = "".join(parts).strip()
    if not text:
        status = js.get("status", "")
        incomplete = (js.get("incomplete_details") or {}).get("reason", "")
        suffix = incomplete or status
        if incomplete == "max_output_tokens":
            return ("⚠️ ChatGPT가 답변 토큰 한도에 도달했습니다. "
                    "연결 자체는 성공했지만 출력 예산이 부족했습니다. 다시 시도해 주세요.")
        return f"⚠️ ChatGPT 응답이 비어 있습니다{f'(사유: {suffix})' if suffix else ''}."
    return text


_AI_DISPATCH = {"Claude": call_claude, "Gemini": call_gemini, "ChatGPT": call_openai}


def ai_call(prompt, api_key=None, model=None, max_tokens=900, system=None, provider=None):
    """제공사별 함수로 분기.

    provider가 전달되면 해당 값을 우선 사용하고, 일반 분석 호출에서는
    사이드바에서 고른 제공사를 사용한다. 위젯 생성 뒤 session_state를 수정하지 않는다.
    """
    provider = provider or st.session_state.get("ai_provider", "Claude (Anthropic)")
    api_key = api_key or st.session_state.get("api_key")
    model = model or st.session_state.get("ai_model_g")
    system = system or _AI_SYS
    if not api_key:
        return "⚠️ 사이드바에 API 키를 입력하면 AI 해석을 사용할 수 있어요."
    if not model:
        return "⚠️ 모델명이 지정되지 않았습니다. 사이드바에서 모델을 선택해 주세요."
    fn = None
    for key, f in _AI_DISPATCH.items():
        if str(provider).startswith(key):
            fn = f
            break
    if fn is None:
        return "⚠️ 알 수 없는 AI 제공사입니다."
    return fn(prompt, api_key, model, max_tokens=max_tokens, system=system)


# ================================================================ 이미지/음성 데이터 입력

def _extract_ai_text_from_openai_response(js):
    txt = str((js or {}).get("output_text") or "").strip()
    if txt:
        return txt
    parts = []
    for item in (js or {}).get("output") or []:
        if not isinstance(item, dict):
            continue
        for c in item.get("content") or []:
            if isinstance(c, dict) and c.get("type") in ("output_text", "text"):
                v = c.get("text", "")
                if isinstance(v, dict): v = v.get("value", "")
                if v: parts.append(str(v))
    return "".join(parts).strip()


def ai_multimodal_text(binary, mime_type, prompt, kind="image"):
    """현재 사이드바에서 선택한 AI 제공사로 이미지/오디오를 읽어 텍스트를 반환."""
    import base64
    provider = st.session_state.get("ai_provider", "Claude (Anthropic)")
    api_key = st.session_state.get("api_key")
    model = st.session_state.get("ai_model_g")
    if not api_key or not model:
        return "⚠️ 먼저 사이드바의 '🤖 AI 기능 켜기'에서 API 키와 모델을 설정해 주세요."
    if not _HAS_REQUESTS:
        return "⚠️ requests 라이브러리가 필요합니다."
    b64 = base64.b64encode(binary).decode("ascii")

    # Claude: 현재 모델은 이미지 입력 지원. 오디오는 직접 입력 미지원이므로 안내.
    if str(provider).startswith("Claude"):
        if kind == "audio":
            return "⚠️ Claude 선택 상태에서는 음성 전사를 지원하지 않습니다. ChatGPT 또는 Gemini를 선택해 주세요."
        try:
            client = anthropic.Anthropic(api_key=api_key, timeout=90, max_retries=2)
            msg = client.messages.create(
                model=model, max_tokens=4000,
                messages=[{"role":"user","content":[
                    {"type":"image", "source":{"type":"base64", "media_type":mime_type, "data":b64}},
                    {"type":"text", "text":prompt},
                ]}])
            return "".join(getattr(x, "text", "") for x in msg.content if getattr(x, "type", "") == "text").strip()
        except Exception as ex:
            return f"⚠️ 이미지 인식 오류: {_ai_mask(str(ex), api_key)[:160]}"

    if str(provider).startswith("Gemini"):
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
        payload = {"contents":[{"parts":[
            {"inline_data":{"mime_type":mime_type, "data":b64}}, {"text":prompt}
        ]}], "generationConfig":{"maxOutputTokens":4000}}
        r, err = _ai_http(_requests.post, url, api_key=api_key,
                          headers={"x-goog-api-key":api_key, "Content-Type":"application/json"},
                          json=payload, timeout=100)
        if err: return err
        try:
            parts = (((r.json().get("candidates") or [])[0].get("content") or {}).get("parts") or [])
            return "".join(p.get("text", "") for p in parts if isinstance(p, dict)).strip()
        except Exception:
            return "⚠️ Gemini 멀티모달 응답을 읽지 못했습니다."

    # OpenAI: image는 Responses API, audio는 공식 Transcriptions API.
    if str(provider).startswith("ChatGPT"):
        if kind == "audio":
            try:
                files = {"file": ("voice.wav", binary, mime_type or "audio/wav")}
                data = {"model": "gpt-transcribe",
                        "prompt": "한국어 농업 시험 조사 데이터입니다. 처리구, 반복, 초장, 수량, 과장, 과폭 등 숫자와 단위를 정확히 전사하세요."}
                r = _requests.post("https://api.openai.com/v1/audio/transcriptions",
                                   headers={"Authorization": f"Bearer {api_key}"},
                                   files=files, data=data, timeout=100)
                if r.status_code != 200:
                    return _ai_error_message(r.status_code, _ai_mask(r.text, api_key))
                return str((r.json() or {}).get("text") or "").strip()
            except Exception as ex:
                return f"⚠️ 음성 전사 오류: {_ai_mask(str(ex), api_key)[:160]}"
        try:
            payload = {
                "model": model, "store": False, "max_output_tokens": 4000,
                "input":[{"role":"user","content":[
                    {"type":"input_text", "text":prompt},
                    {"type":"input_image", "image_url":f"data:{mime_type};base64,{b64}"},
                ]}]
            }
            r, err = _ai_http(_requests.post, "https://api.openai.com/v1/responses", api_key=api_key,
                              headers={"Authorization":f"Bearer {api_key}", "Content-Type":"application/json"},
                              json=payload, timeout=100)
            if err: return err
            return _extract_ai_text_from_openai_response(r.json())
        except Exception as ex:
            return f"⚠️ 이미지 인식 오류: {_ai_mask(str(ex), api_key)[:160]}"
    return "⚠️ 지원하지 않는 AI 제공사입니다."


def _json_from_ai_text(text):
    import json, re
    if not text or str(text).startswith("⚠️"):
        return None
    t = str(text).strip()
    t = re.sub(r"^```(?:json)?\\s*", "", t, flags=re.I)
    t = re.sub(r"\\s*```$", "", t)
    # 앞뒤 설명이 붙어도 첫 JSON object/array를 최대한 복원
    candidates = [t]
    for left, right in (("{", "}"), ("[", "]")):
        if left in t and right in t:
            candidates.append(t[t.find(left):t.rfind(right)+1])
    for c in candidates:
        try:
            return json.loads(c)
        except Exception:
            pass
    return None


def image_to_dataframe(binary, mime_type):
    prompt = """이 이미지는 연구/조사 데이터 표입니다. 표의 글자와 숫자를 그대로 읽어 구조화하세요.
반드시 JSON만 출력하세요. 형식:
{"columns":["열1","열2"],"rows":[[값,값],[값,값]],"warnings":["애매한 셀 설명"]}
규칙: 1) 보이지 않는 값을 추측하지 말고 null, 2) 소수점/음수/단위를 특히 정확히, 3) 병합 머리글은 의미가 보존되도록 한 줄 열 이름으로 합치기, 4) 표가 여러 개면 가장 큰 데이터 표 하나를 우선."""
    raw = ai_multimodal_text(binary, mime_type, prompt, kind="image")
    js = _json_from_ai_text(raw)
    if not isinstance(js, dict):
        return None, [raw if raw else "표를 구조화하지 못했습니다."]
    cols, rows = js.get("columns") or [], js.get("rows") or []
    if not cols or not isinstance(rows, list):
        return None, ["열 또는 행을 찾지 못했습니다."]
    fixed = []
    for r in rows:
        if isinstance(r, dict):
            fixed.append([r.get(c) for c in cols])
        elif isinstance(r, list):
            fixed.append((r + [None] * len(cols))[:len(cols)])
    return clean_columns(pd.DataFrame(fixed, columns=cols)), list(js.get("warnings") or [])


def voice_text_to_row(transcript, columns=None):
    cols = [str(c) for c in (columns or [])]
    prompt = f"""다음은 연구자가 음성으로 말한 한 행의 조사 데이터입니다.
음성: {transcript}
현재 데이터 열: {cols if cols else '없음'}
반드시 JSON만 출력하세요.
현재 열이 있으면 {{"row": {{"열이름": 값, ...}}, "warnings": []}} 형식으로 해당 열 이름을 그대로 사용하세요.
현재 열이 없으면 {{"row": {{"처리구":"A", "반복":1, ...}}, "warnings": []}} 형태로 의미 있는 열을 만드세요.
말하지 않은 값은 null로 두고 숫자는 가능하면 숫자형으로 반환하세요. 추측하지 마세요."""
    raw = ai_call(prompt, max_tokens=1200, system="데이터 입력 도우미입니다. JSON만 출력합니다.")
    js = _json_from_ai_text(raw)
    if not isinstance(js, dict) or not isinstance(js.get("row"), dict):
        return None, [raw if raw else "음성을 행 데이터로 변환하지 못했습니다."]
    row = js["row"]
    if cols:
        row = {c: row.get(c, None) for c in cols}
    return row, list(js.get("warnings") or [])

def ai_disclaimer():
    st.warning("⚠️ **AI가 만든 초안입니다.** 논문·보고서에 넣기 전에 반드시 연구자가 "
               "수치와 해석이 맞는지 확인하고 수정하세요. AI는 없는 인과관계를 서술하거나 "
               "수치를 잘못 인용할 수 있습니다.")

def _json_safe(obj):
    """NaN·NumPy 타입을 JSON으로 안전하게 변환"""
    import math
    if obj is None:
        return None
    if isinstance(obj, (np.integer,)):
        return int(obj)
    if isinstance(obj, (np.floating, float)):
        v = float(obj)
        return None if (math.isnan(v) or math.isinf(v)) else round(v, 6)
    if isinstance(obj, (np.bool_, bool)):
        return bool(obj)
    if isinstance(obj, (np.ndarray,)):
        return [_json_safe(x) for x in obj.tolist()]
    if isinstance(obj, pd.DataFrame):
        return [{str(k): _json_safe(v) for k, v in row.items()}
                for row in obj.to_dict(orient="records")]
    if isinstance(obj, pd.Series):
        return {str(k): _json_safe(v) for k, v in obj.items()}
    if isinstance(obj, dict):
        return {str(k): _json_safe(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple, set)):
        return [_json_safe(x) for x in obj]
    if pd.isna(obj) if np.isscalar(obj) else False:
        return None
    return obj if isinstance(obj, (str, int)) else str(obj)


def build_data_overview(df, max_cols=40):
    """AI에게 넘길 데이터 개요. 결측치·자료형을 '명시적으로' 적어 추측을 막는다.

    이전에는 describe() 결과만 넘겨서 AI가 결측치가 있는 열을 잘못 지목하는 일이
    있었다. 여기서는 열별 결측 개수를 문장으로 못 박아 전달한다.
    """
    if df is None or getattr(df, "empty", True):
        return "데이터가 없습니다."
    lines = [f"행 {len(df):,}개, 열 {len(df.columns)}개"]
    miss = df.isna().sum()
    has_miss = miss[miss > 0]
    lines.append("")
    lines.append("[열 정보]  형식 | 결측 | 고유값")
    for c in list(df.columns)[:max_cols]:
        kind = "숫자형" if pd.api.types.is_numeric_dtype(df[c]) else "문자형"
        lines.append(f"- {c} | {kind} | 결측 {int(miss[c])}개 | 고유 {int(df[c].nunique(dropna=True))}종")
    if len(df.columns) > max_cols:
        lines.append(f"- (이하 {len(df.columns)-max_cols}개 열 생략)")
    lines.append("")
    if has_miss.empty:
        lines.append("[결측치] 모든 열에 결측치가 없습니다. "
                     "어떤 열에도 '결측치가 있다'고 서술하지 마세요.")
    else:
        detail = ", ".join(f"'{c}' {int(v)}개" for c, v in has_miss.items())
        lines.append(f"[결측치] 결측치가 있는 열은 다음뿐입니다: {detail}. "
                     "여기에 없는 열은 결측치가 0개이므로 결측을 언급하지 마세요.")
    try:
        lines.append("")
        lines.append("[기술통계]")
        lines.append(df.describe().round(2).to_string())
    except Exception:
        pass
    return "\n".join(lines)


def build_group_profiles(df, question="", max_groups=15, max_numeric=12):
    """질문과 관련된 처리·품종별 평균/표준편차/n을 AI에 전달한다."""
    if df is None or df.empty:
        return {}
    nums = df.select_dtypes(include=np.number).columns.tolist()[:max_numeric]
    cats = [c for c in df.columns if c not in nums and 2 <= df[c].nunique(dropna=True) <= max_groups]
    q = str(question or "").lower()
    selected = []
    for c in cats:
        levels = [str(x) for x in df[c].dropna().unique()]
        if str(c).lower() in q or any(v.lower() in q for v in levels):
            selected.append(c)
    if not selected:
        selected = cats[:2]
    out = {}
    for c in selected:
        if not nums:
            continue
        g = df.groupby(c)[nums].agg(["mean", "std", "count"])
        out[str(c)] = _json_safe(g.reset_index())
    return out


def build_anova_context(**kw):
    """분산분석 결과를 AI에 넘길 구조화 JSON으로 정리"""
    ctx = {
        "analysis_type": kw.get("analysis_type", "분산분석"),
        "design": kw.get("design"),
        "treatment_column": kw.get("trt"),
        "block_column": kw.get("blk"),
        "response_variable": kw.get("yv"),
        "group_stats": kw.get("group_stats"),
        "anova_table": kw.get("anova_table"),
        "p_treatment": kw.get("p_treatment"),
        "p_block": kw.get("p_block"),
        "cv_percent": kw.get("cv"),
        "lsd": kw.get("lsd"),
        "mse": kw.get("mse"),
        "df_residual": kw.get("df_resid"),
        "posthoc_method": kw.get("posthoc"),
        "significance_letters": kw.get("letters"),
        "dunnett_table": kw.get("dunnett"),
        "assumption_tests": kw.get("assumptions"),
        "n_missing_excluded": kw.get("n_missing"),
        "cautions": kw.get("cautions", []),
    }
    return _json_safe({k: v for k, v in ctx.items() if v is not None})


def bootstrap_diff_ci(a, b, n_boot=2000, alpha=0.05, seed=0):
    """두 그룹 평균 차이(a−b)의 부트스트랩 신뢰구간. 소표본·비정규 자료용."""
    a = np.asarray(pd.to_numeric(pd.Series(a), errors="coerce").dropna(), dtype=float)
    b = np.asarray(pd.to_numeric(pd.Series(b), errors="coerce").dropna(), dtype=float)
    if len(a) < 2 or len(b) < 2:
        return None
    rng = np.random.default_rng(seed)
    diffs = np.empty(n_boot)
    for i in range(n_boot):
        diffs[i] = (rng.choice(a, len(a), replace=True).mean()
                    - rng.choice(b, len(b), replace=True).mean())
    lo, hi = np.percentile(diffs, [alpha / 2 * 100, (1 - alpha / 2) * 100])
    return {"diff": float(a.mean() - b.mean()), "low": float(lo), "high": float(hi),
            "n_boot": int(n_boot)}


def econ_metric_test(row_df, trt_col, value_col, control=None, blk_col=None):
    """반복별 소득·순수익 ANOVA와 모형 기반 대조구 비교.

    블록이 있으면 ANOVA와 Dunnett 모두 같은 RCBD 모형을 사용한다.
    부트스트랩 구간은 관측 평균 차이의 보조 정보로 별도 표시한다.
    """
    use_cols = [c for c in (trt_col, value_col, blk_col) if c and c in row_df.columns]
    d = row_df[use_cols].dropna()
    ok, _ = validate_anova_data(d, trt_col, value_col)
    if not ok:
        return None
    out = {"n_groups": int(d[trt_col].nunique()),
           "n_total": int(len(d)), "value": value_col,
           "block": blk_col if blk_col in d.columns else None}
    model = None
    try:
        f = safe_formula(value_col, [trt_col] + ([blk_col] if blk_col in d.columns else []))
        model = ols(f, data=d).fit()
        aov = sm.stats.anova_lm(model, typ=2)
        k = f"C({q_ref(trt_col)})"
        out["anova_p"] = float(aov.loc[k, "PR(>F)"]) if k in aov.index else None
        if blk_col in d.columns:
            bk = f"C({q_ref(blk_col)})"
            out["block_p"] = float(aov.loc[bk, "PR(>F)"]) if bk in aov.index else None
        out["model"] = f
    except Exception as ex:
        out["anova_p"] = None
        out["error"] = str(ex)[:120]

    str_levels = {str(x): x for x in d[trt_col].dropna().unique()}
    control_level = str_levels.get(str(control)) if control is not None else None
    if control_level is not None and model is not None:
        try:
            ph = posthoc_from_model(model, d, trt_col, "던넷(Dunnett)",
                                     control=control_level)
            tab = ph.get("table", pd.DataFrame()).copy()
            if not tab.empty:
                tab = tab.rename(columns={
                    "처리 평균(보정)": "평균(보정)",
                    "평균 차이": f"'{control}' 대비 차이",
                    "p(동시보정)": "p(보정)",
                    "95% 동시CI 하한": "95% 하한",
                    "95% 동시CI 상한": "95% 상한",
                })
                keep = [c for c in ["처리구", "평균(보정)", f"'{control}' 대비 차이",
                                     "p(보정)", "95% 하한", "95% 상한", "판정"]
                        if c in tab.columns]
                out["dunnett"] = tab[keep]
        except Exception as ex:
            out["dunnett_error"] = str(ex)[:120]

        # 비모수적 불확실성 참고: 원자료 평균차 부트스트랩
        g = d.assign(__trt=d[trt_col].astype(str)).groupby("__trt")[value_col]
        names = [n for n in g.groups.keys() if n != str(control)]
        boot = {}
        for n in names:
            r = bootstrap_diff_ci(g.get_group(n).values,
                                  g.get_group(str(control)).values)
            if r:
                boot[n] = r
        if boot:
            out["bootstrap"] = boot
    return out


def build_econ_context(**kw):
    """경제성 결과를 AI에 넘길 구조화 JSON으로 정리"""
    ctx = {
        "analysis_type": "경제성분석",
        "base_area": kw.get("base_area", "10a"),
        "control_treatment": kw.get("control"),
        "treatments": kw.get("treatments"),
        "price_assumptions": kw.get("prices"),
        "cost_columns_used": kw.get("cost_cols"),
        "cost_columns_excluded_for_duplication": kw.get("excluded_cols"),
        "sensitivity": kw.get("sensitivity"),
        "yield_statistical_test": kw.get("yield_test"),
        "income_statistical_test": kw.get("income_test"),
        "profit_statistical_test": kw.get("profit_test"),
        "cautions": kw.get("cautions", []),
    }
    return _json_safe({k: v for k, v in ctx.items() if v is not None})


# ---------------------------------------------------------------- 경제성 분석 길잡이(규칙 기반)
_ECON_MODE_PARTIAL = "📕 부분예산표 (손실적·이익적 요소)"
_ECON_MODE_INCOME = "📗 소득분석"
_ECON_MODE_MRR = "📘 신기술 경제성 (부분예산·한계수익률)"
_ECON_MODE_INVEST = "📙 시설·장기투자 경제성 (NPV·B/C·IRR)"


def recommend_economic_guide(goal, change, comparison, period, data_items=None):
    """초보자용 경제성 분석 길잡이의 순수 규칙 엔진.

    STEP 1~5 응답을 받아 현재 앱의 경제성 모듈 중 가장 적합한 것을 추천한다.
    AI/API를 쓰지 않으며, '잘 모르겠어요'가 포함되어도 나머지 답으로 판단한다.
    정책·공공사업 CBA는 기존 농가단위 모듈로 억지 연결하지 않는다.
    """
    data_items = list(data_items or [])
    values = [str(goal or ''), str(change or ''), str(comparison or ''), str(period or '')]
    unknown_count = sum('잘 모르' in v or '아직 모르' in v for v in values)
    if any('아직 거의 준비' in str(x) or '자료가 어떤' in str(x) for x in data_items):
        unknown_count += 1

    def has(text, *tokens):
        s = str(text or '')
        return any(t in s for t in tokens)

    scores = {
        _ECON_MODE_PARTIAL: 0.0,
        _ECON_MODE_INCOME: 0.0,
        _ECON_MODE_MRR: 0.0,
        _ECON_MODE_INVEST: 0.0,
    }
    reasons = []
    tags = []
    warnings = []

    # 정책·공공사업은 현재 농가단위 모듈 범위를 벗어난다.
    if has(goal, '정책', '사회적') or has(change, '사회적', '환경적', '공공사업'):
        return {
            'primary_mode': None,
            'title': '🏛️ 비용편익분석(CBA) — 현재 직접 계산 미지원',
            'confidence': '높음',
            'scores': scores,
            'reasons': [
                '정책·공공사업은 농가 개인의 수입·비용뿐 아니라 사회 전체의 편익·비용과 외부효과를 평가해야 합니다.',
                '현재 프로그램의 소득분석·부분예산·MRR은 농가 또는 기술대안 단위 분석이므로 범위가 다릅니다.',
            ],
            'needs': ['사회적 편익', '사회적 비용', '외부효과의 화폐가치', '분석기간', '사회적 할인율'],
            'together': ['재무성 분석과 경제성(CBA)을 구분', '비시장 편익·비용의 평가 근거 명시'],
            'tags': ['CBA'], 'warnings': [], 'ambiguous': False,
            'missing_reported': [], 'top_two': [],
        }

    # STEP 1: 연구 목적
    if has(goal, '기존 방식보다', '신품종', '신기술'):
        scores[_ECON_MODE_PARTIAL] += 7
        scores[_ECON_MODE_MRR] += 2
        reasons.append('기존 방식과 신기술의 차이를 평가하는 목적입니다.')
    elif has(goal, '현재 작목', '현재 처리', '수익성'):
        scores[_ECON_MODE_INCOME] += 8
        reasons.append('현재 한 해의 조수입·소득·순수익 자체가 핵심 질문입니다.')
    elif has(goal, '여러 대안', '가장 경제적', '무엇을 권'):
        scores[_ECON_MODE_MRR] += 8
        scores[_ECON_MODE_INCOME] += 2
        reasons.append('여러 대안 중 추가비용 대비 추가편익을 비교하려는 목적입니다.')
    elif has(goal, '시설', '농기계', '투자할 가치'):
        scores[_ECON_MODE_INVEST] += 10
        reasons.append('초기 투자비를 들여 여러 해 사용하는 자산의 투자 타당성이 핵심 질문입니다.')
    elif has(goal, '어느 가격', '어느 수량', '손해'):
        scores[_ECON_MODE_INCOME] += 8
        tags.append('손익분기점')
        reasons.append('현재 비용구조를 기준으로 손익이 0이 되는 가격·수량을 찾는 질문입니다.')
    elif has(goal, '가격', '수량', '유지', '위험'):
        scores[_ECON_MODE_INCOME] += 8
        tags.append('민감도 분석')
        reasons.append('기준 소득을 계산한 뒤 가격·수량 변동에 대한 위험을 확인하는 질문입니다.')
    elif has(goal, '여러 작형', '여러 품종', '경영성과'):
        scores[_ECON_MODE_INCOME] += 7
        scores[_ECON_MODE_PARTIAL] += 1
        reasons.append('같은 기간의 처리·작형별 경영성과를 동일 기준으로 비교하려는 목적입니다.')

    # STEP 2: 실제로 달라지는 것
    if has(change, '품종', '방제', '재배법', '재배기술'):
        scores[_ECON_MODE_PARTIAL] += 5
        scores[_ECON_MODE_MRR] += 1
        reasons.append('품종·방제·재배기술 변경은 기존 방식 대비 변화분 비교가 중요합니다.')
    elif has(change, '투입수준', '투입량', '비료량', '농약량', '노동량'):
        scores[_ECON_MODE_MRR] += 5
        scores[_ECON_MODE_PARTIAL] += 2
        reasons.append('투입수준에 따라 비용이 단계적으로 달라지는 구조입니다.')
    elif has(change, '시설', '농기계', '신규 투자'):
        scores[_ECON_MODE_INVEST] += 8
        reasons.append('시설·농기계의 신규 투자가 포함됩니다.')
    elif has(change, '판매가격', '상품수량', '상품률', '수량·가격'):
        scores[_ECON_MODE_INCOME] += 4
        tags.append('민감도 분석')
        reasons.append('가격·수량 변화가 수익성에 미치는 영향 확인이 필요합니다.')
    elif has(change, '특별한 변경 없음', '현재 경영성과'):
        scores[_ECON_MODE_INCOME] += 6
        reasons.append('특정 신기술의 변화분보다 현재 경영성과 자체를 평가하는 구조입니다.')

    # STEP 3: 비교 구조
    if has(comparison, '대조구 1개', '신기술 1', '신품종 1'):
        scores[_ECON_MODE_PARTIAL] += 6
        reasons.append('대조구와 신기술을 직접 비교하는 구조라 부분예산법과 잘 맞습니다.')
    elif has(comparison, '3개 이상', '비용이 다른', '여러 대안'):
        scores[_ECON_MODE_MRR] += 8
        reasons.append('비용이 다른 3개 이상 대안은 지배분석과 MRR로 단계적 채택 여부를 보기 좋습니다.')
    elif has(comparison, '여러 품종', '여러 작형', '한 해 성과'):
        scores[_ECON_MODE_INCOME] += 6
        reasons.append('여러 처리의 한 해 소득·순수익을 같은 기준으로 비교하는 구조입니다.')
    elif has(comparison, '비교대상 없음'):
        scores[_ECON_MODE_INCOME] += 3
        reasons.append('비교대상이 없으므로 우선 현재 수익성의 기준선을 만드는 것이 적합합니다.')

    # STEP 4: 분석기간
    if has(period, '한 작기', '1년'):
        scores[_ECON_MODE_PARTIAL] += 2
        scores[_ECON_MODE_INCOME] += 2
        scores[_ECON_MODE_MRR] += 2
        scores[_ECON_MODE_INVEST] -= 1
        reasons.append('경제효과를 한 작기·1년 기준으로 평가합니다.')
    elif has(period, '2년 이상'):
        scores[_ECON_MODE_INVEST] += 4
        warnings.append('여러 해 자료라도 매년 독립적인 재배기술 비교라면 연도별 부분예산/소득분석을 병행할 수 있습니다.')
        reasons.append('효과가 여러 해 지속되므로 시간가치를 확인할 필요가 있습니다.')
    elif has(period, '내용연수', '사용기간 전체'):
        scores[_ECON_MODE_INVEST] += 8
        reasons.append('시설·기계의 내용연수 전체를 보므로 할인현금흐름 분석이 필요합니다.')

    # STEP 5: 보유자료 — 추천을 뒤집기보다 실행가능성 판단에 가중치를 조금만 준다.
    items_text = ' | '.join(map(str, data_items))
    if has(items_text, '최초 투자비'):
        scores[_ECON_MODE_INVEST] += 2
    if has(items_text, '연도별 편익', '연간 편익'):
        scores[_ECON_MODE_INVEST] += 2
    if has(items_text, '할인율', '잔존가치'):
        scores[_ECON_MODE_INVEST] += 2
    if has(items_text, '달라지는 비용'):
        scores[_ECON_MODE_PARTIAL] += 1
        scores[_ECON_MODE_MRR] += 1
    if has(items_text, '항목별 경영비'):
        scores[_ECON_MODE_INCOME] += 1

    ranked = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    top_mode, top_score = ranked[0]
    second_mode, second_score = ranked[1]
    margin = top_score - second_score

    # 모든 응답이 거의 미정이면 억지 추천하지 않는다.
    ambiguous = bool(top_score < 5 or (unknown_count >= 3 and margin < 3))
    if ambiguous:
        primary_mode = None
        title = '🧭 추천 보류 — 두 가지만 더 정리하면 정확히 고를 수 있어요'
        confidence = '낮음'
    else:
        primary_mode = top_mode
        confidence = '높음' if unknown_count <= 1 and margin >= 3 else ('중간' if margin >= 1.5 else '낮음')
        title_map = {
            _ECON_MODE_PARTIAL: '📕 부분예산법',
            _ECON_MODE_INCOME: '📗 소득분석',
            _ECON_MODE_MRR: '📘 신기술 경제성(MRR)',
            _ECON_MODE_INVEST: '📙 시설·장기투자 분석',
        }
        title = title_map[top_mode]

    meta = {
        _ECON_MODE_PARTIAL: {
            'needs': ['대조구·신기술구 구분', '조사면적과 수량(또는 판매수입)', '실제 판매가격', '신기술 때문에 달라진 비용: 종묘·비료·농약·자재·노동·위탁·임차 등 변화분만'],
            'together': ['반복시험이면 수량·소득의 통계검정', '가격 변동이 크면 민감도 분석'],
            'data_groups': [
                ('처리구/대조구 구분', ['처리구/대조구 구분']),
                ('수량·생산량', ['수량·생산량']),
                ('판매가격/판매액', ['판매가격/판매액']),
                ('변화 비용', ['신기술로 달라지는 비용만', '항목별 경영비']),
            ],
        },
        _ECON_MODE_INCOME: {
            'needs': ['처리구/작형·반복(비교 시)', '조사면적과 생산량', '실제 판매가격과 부산물수입(있으면)', '경영비: 종자·종묘, 비료, 농약, 수도광열, 재료, 소농구, 감가상각, 수선, 임차, 위탁영농, 고용노동 등 실제 발생 항목', '순수익까지 볼 때: 자가노동시간, 자본용역비, 자가토지 용역비'],
            'together': ['손익분기점', '가격·수량 민감도', '반복자료가 있으면 소득·순수익 통계검정'],
            'data_groups': [
                ('수량·생산량', ['수량·생산량']),
                ('판매가격/판매액', ['판매가격/판매액']),
                ('항목별 경영비', ['항목별 경영비']),
            ],
        },
        _ECON_MODE_MRR: {
            'needs': ['비용이 다른 여러 처리구(보통 3개 이상)와 대조구', '조사면적·처리별 수량', '실제 판매가격', '처리 수준에 따라 달라지는 가변비용: 비료·농약·노동·자재·위탁비 등', '수량 조정률 및 최소수용 MRR 기준'],
            'together': ['부분예산', '지배분석', '최소수용 MRR', '가격·수량 민감도'],
            'data_groups': [
                ('처리구/대조구 구분', ['처리구/대조구 구분']),
                ('수량·생산량', ['수량·생산량']),
                ('판매가격/판매액', ['판매가격/판매액']),
                ('가변비용', ['신기술로 달라지는 비용만', '항목별 경영비']),
            ],
        },
        _ECON_MODE_INVEST: {
            'needs': ['최초 투자비(설치·구입·부대공사 포함)', '분석기간/내용연수', '연도별 또는 연간 추가수입·비용절감 편익', '연간 운영·유지·수선비와 예상 교체비', '할인율', '잔존가치(있으면)'],
            'together': ['NPV', '할인 B/C', 'IRR', '단순·할인 회수기간', '편익·비용 민감도'],
            'data_groups': [
                ('최초 투자비', ['최초 투자비']),
                ('연도별 편익·운영비', ['연도별 편익·운영비']),
                ('분석기간·할인율', ['분석기간·할인율·잔존가치']),
            ],
        },
    }
    chosen_meta = meta.get(top_mode, meta[_ECON_MODE_INCOME])
    present = set(map(str, data_items))
    missing = []
    for label, alternatives in chosen_meta['data_groups']:
        if not any(a in present for a in alternatives):
            missing.append(label)

    # 목적별 보조 분석 태그
    if top_mode == _ECON_MODE_PARTIAL and '반복(블록) 자료' in present:
        tags.append('통계검정 병행')
    if top_mode == _ECON_MODE_MRR:
        tags.extend(['지배분석', 'MRR'])
    if top_mode == _ECON_MODE_INVEST:
        tags.extend(['NPV', '할인 B/C', 'IRR'])
    tags = list(dict.fromkeys(tags))

    if ambiguous:
        reasons = reasons[-4:] if reasons else [
            '연구목적·비교대상·분석기간 중 아직 정해지지 않은 항목이 많습니다.',
            '“기존 방식과 다른 처리가 있는지”와 “효과가 1년인지 여러 해인지”만 정하면 대부분의 경우 분석법을 고를 수 있습니다.'
        ]
    else:
        reasons = list(dict.fromkeys(reasons))[-5:]

    return {
        'primary_mode': primary_mode,
        'title': title,
        'confidence': confidence,
        'scores': scores,
        'reasons': reasons,
        'needs': chosen_meta['needs'],
        'together': chosen_meta['together'],
        'tags': tags,
        'warnings': warnings,
        'ambiguous': ambiguous,
        'missing_reported': missing,
        'top_two': [(ranked[0][0], ranked[0][1]), (ranked[1][0], ranked[1][1])],
    }


def parse_ai_json(text):
    """AI 응답에서 {verified_facts, interpretation, limitations, recommendation}
    구조를 찾아 파싱. 실패하면 (None, 원문)을 반환해 그대로 보여줌."""
    import json as _json
    import re as _re
    if not text:
        return None, ""
    t = str(text).strip()
    t = _re.sub(r"^```(?:json)?\s*|\s*```$", "", t, flags=_re.S).strip()
    m = _re.search(r"\{.*\}", t, _re.S)
    if not m:
        return None, str(text)
    try:
        obj = _json.loads(m.group(0))
    except Exception:
        return None, str(text)
    if not isinstance(obj, dict):
        return None, str(text)
    keys = ("verified_facts", "interpretation", "limitations", "recommendation")
    if not any(k in obj for k in keys):
        return None, str(text)
    out = {}
    for k in keys:
        v = obj.get(k, [])
        if isinstance(v, str):
            v = [v]
        elif not isinstance(v, list):
            v = [str(v)]
        out[k] = [strip_md(str(x)).strip() for x in v if str(x).strip()]
    return out, str(text)


def render_ai_json(parsed):
    """파싱된 AI 응답을 보기 좋게 표시하고 보고서용 평문을 반환"""
    labels = [("verified_facts", "확인된 사실"), ("interpretation", "해석"),
              ("limitations", "한계"), ("recommendation", "권장 사항")]
    lines = []
    for key, title in labels:
        items = parsed.get(key) or []
        if not items:
            continue
        st.markdown(f"**{title}**")
        for it in items:
            st.markdown(f"- {it}")
        lines.append(f"○ {title}")
        lines += [f"  - {it}" for it in items]
    return "\n".join(lines)


def ai_interpret_advanced(slot, kind, table_df, extra="", context=None, capture_slot=None):
    """분석 결과를 3가지 스타일(보고서·고찰·현장지도)로 해석. 결과는 화면에 계속 남음."""
    key = st.session_state.get("api_key")
    out_key = f"__ai_out_{slot}"
    STYLES = {
        "1️⃣ 보고서용": (
            "농촌진흥청 시험연구보고서 '주요 연구결과' 항목에 그대로 넣을 수 있게 작성하세요.\n"
            "형식: 주요 항목은 '○ '로 시작, 세부 내용은 '  - '로 시작.\n"
            "표의 실제 수치를 반드시 인용하고, p-value·CV%·사후검정 문자를 근거로 제시하세요.\n"
            "분량: 6~10줄."),
        "2️⃣ 논문 고찰용 (Discussion)": (
            "학술논문의 고찰(Discussion) 초안으로 작성하세요.\n"
            "결과의 통계적 의미를 먼저 정리하고, 관찰된 경향이 나타난 농학적 원인을 추정하되 "
            "'~로 추정된다', '~때문으로 판단된다'처럼 단정하지 않는 표현을 쓰세요.\n"
            "마지막에 본 시험의 한계와 후속 연구 방향을 1~2문장 제시하세요.\n"
            "분량: 6~10문장의 서술형 문단(불릿 없이)."),
        "4️⃣ 구조화(검증·해석·한계·권장)": (
            "다음 JSON 형식으로만 답하세요. 다른 설명은 붙이지 마세요.\n"
            '{"verified_facts": [], "interpretation": [], "limitations": [], '
            '"recommendation": []}\n'
            "- verified_facts: 제공된 JSON에서 그대로 확인되는 사실(수치 포함)\n"
            "- interpretation: 그 사실이 농업적으로 무엇을 뜻하는지\n"
            "- limitations: 이 결과로 말할 수 없는 것, 검정하지 않은 부분\n"
            "- recommendation: 다음에 확인하거나 시도할 것\n"
            "각 항목은 문자열 배열이며, 마크다운 기호는 쓰지 마세요."),
        "3️⃣ 현장 지도용 (3줄 요약)": (
            "농가·현장 지도용으로 통계를 모르는 사람도 이해할 수 있게 작성하세요.\n"
            "형식: 정확히 3줄. 각 줄은 '○ '로 시작.\n"
            "1줄=어떤 처리가 가장 좋았는지, 2줄=그 차이가 믿을 만한지, 3줄=현장에서 어떻게 하면 되는지.\n"
            "전문용어(p값, 유의수준, 변이계수) 대신 쉬운 말로 바꿔 쓰세요."),
    }
    with st.expander("🤖 AI 해석 (보고서·고찰·현장지도)"):
        if not key:
            st.info("왼쪽 사이드바 **🤖 AI 기능 켜기**에 API 키를 넣으면 "
                    "이 결과를 3가지 형태의 문장으로 바꿔 드립니다.")
            return
        want = st.radio("어떤 형태로 만들까요?", list(STYLES.keys()), key="aim_" + slot)
        if st.button("✨ AI 해석 생성", key="aib_" + slot):
            ctx = ""
            if context:
                import json as _json
                try:
                    ctx = ("\n\n[분석 결과 JSON — 이 안의 값만 사용하세요]\n"
                           + _json.dumps(_json_safe(context), ensure_ascii=False, indent=1))
                except Exception:
                    ctx = "\n\n[분석 맥락]\n" + "\n".join(f"- {k}: {v}" for k, v in context.items() if v)
            with st.spinner("AI가 해석 중..."):
                prompt = (f"다음은 '{kind}' 분석 결과입니다.\n\n"
                          f"{table_df.to_string(index=False)}\n{ctx}\n\n"
                          f"{extra}\n\n{STYLES[want]}\n\n"
                          "반드시 한국어로 작성하고, 마크다운 기호(**, ##, *, `, ---)는 "
                          "어떤 경우에도 사용하지 마세요. 표에 없는 수치는 만들어 내지 마세요.")
                raw = ai_call(prompt, key, st.session_state.get("ai_model_g"), max_tokens=1600)
                if want.startswith("4️⃣"):
                    _parsed, _orig = parse_ai_json(raw)
                    st.session_state[out_key + "_json"] = _parsed
                    st.session_state[out_key] = strip_md(_orig) if _parsed is None else ""
                    if _parsed is None:
                        st.warning("⚠️ AI가 요청한 JSON 형식으로 답하지 않아 원문을 그대로 표시합니다.")
                else:
                    st.session_state[out_key + "_json"] = None
                    st.session_state[out_key] = strip_md(raw)
                log_action(f"AI 해석 생성({kind} / {want.split()[-1]})")
        _pj = st.session_state.get(out_key + "_json")
        saved = st.session_state.get(out_key)
        if _pj:
            st.markdown("###### 생성된 해석")
            _plain = render_ai_json(_pj)
            saved = _plain
            with st.expander("보고서용 평문 보기"):
                st.code(_plain, language=None)
        elif saved:
            st.markdown("###### 생성된 문장")
            st.code(saved, language=None)
        if saved:
            ai_disclaimer()
            cap = st.session_state.get(capture_slot) if capture_slot else None
            b1, b2 = st.columns(2)
            if cap:
                if b1.button("➕ 분석 결과 + AI 해석 함께 담기", key="aiadd_" + slot,
                             help="표·그림과 AI 해석이 보고서에서 같은 항목으로 이어 붙습니다."):
                    merged = merge_ai_into_capture(cap, kind, saved)
                    st.session_state.report_items.append(merged)
                    st.success("분석 결과와 해석을 함께 담았습니다! "
                               f"(현재 {len(st.session_state.report_items)}개)")
            else:
                if b1.button("➕ 보고서에 담기", key="aiadd_" + slot):
                    st.session_state.report_items.append(
                        {"heading": f"{kind} 해석", "text": saved,
                         "table": None, "image": None})
                    st.success("보고서에 담았습니다!")
            if b2.button("🗑️ 지우기", key="aidel_" + slot):
                st.session_state[out_key] = None
                st.session_state[out_key + "_json"] = None
                st.rerun()

def merge_ai_into_capture(capture, kind, ai_text):
    """담아둔 분석 결과(표·그림)와 AI 해석을 하나의 보고서 항목으로 합친다.

    예전에는 해석이 별도 항목으로 붙어 보고서에서 표와 떨어져 나왔다.
    """
    import copy as _copy
    item = _copy.copy(capture or {})
    blocks = list(item.get("blocks") or [])
    if not blocks:
        blocks = [b for b in [
            {"text": item.get("text")} if item.get("text") else None,
            {"caption": item.get("heading", ""), "table": item.get("table")}
            if item.get("table") is not None else None,
            {"caption": item.get("heading", ""), "image": item.get("image")}
            if item.get("image") else None,
        ] if b]
    # AI가 '○ ...' 로 시작하는 제목 줄을 만들면 우리 제목과 한 줄로 합친다.
    lines = [ln for ln in str(ai_text).rstrip().split("\n")]
    head = f"○ AI 해석({kind})"
    first = lines[0].strip() if lines else ""
    if first.startswith(("○", "◦", "●")):
        head += " - " + first.lstrip("○◦● ").strip()
        lines = lines[1:]
    body = "\n".join([head] + lines)
    ai_block = {"text": body, "ai": True}
    # 보고서 관행상 '결과 문장 → 표·그림' 순서이므로, 해석문은 표 앞에 넣는다.
    pos = next((i for i, b in enumerate(blocks)
                if b.get("table") is not None or b.get("image")), len(blocks))
    blocks.insert(pos, ai_block)
    return {"heading": item.get("heading") or f"{kind} 분석",
            "text": None, "table": None, "image": None, "blocks": blocks}


# 이전 이름 호환
def ai_interpret_button(slot, kind, table_df, extra="", capture_slot=None):
    return ai_interpret_advanced(slot, kind, table_df, extra,
                                 capture_slot=capture_slot)

# ---------------------------------------------------------------- 추천/해석
_BLOCK_KEYS = ["반복", "블록", "구역", "block", "rep", "blk"]
_TRT_KEYS = ["처리", "시험구", "구분", "품종", "계통", "약제", "농도", "시비", "수준"]
# 숫자지만 '조사한 값'이 아닌 열 — 측정항목 후보에서 제외한다.
_ID_KEYS = ["연도", "년도", "year", "일자", "날짜", "date", "번호", "no.", "id", "코드"]


def split_code_columns(df):
    """숫자로 적혀 있지만 사실은 처리구·반복 코드인 열을 찾아낸다.

    엑셀에서 반복을 1, 2, 3으로 적는 경우가 매우 흔한데, 이것을 측정값으로
    오인하면 난괴법이 완전임의배치로 분석되어 결론이 뒤집힐 수 있다.
    반환: (측정값 후보, 범주형 후보, 숫자→범주로 승격된 열)
    """
    num = df.select_dtypes(include=np.number).columns.tolist()
    promoted = []
    for c in list(num):
        s = pd.to_numeric(df[c], errors="coerce").dropna()
        if s.empty:
            continue
        name = str(c).lower()
        name_hit = any(k in name for k in _BLOCK_KEYS + _TRT_KEYS)
        # 연도·번호처럼 보이는 열은 측정값이 아니므로 이름이 맞으면 그대로 승격
        try:
            is_int = bool(np.allclose(s.to_numpy(dtype=float),
                                      np.round(s.to_numpy(dtype=float))))
        except (TypeError, ValueError):
            is_int = False
        nu = int(s.nunique())
        vmin, vmax = float(s.min()), float(s.max())
        # 코드값은 보통 1,2,3... 또는 0,1,2...처럼 빈틈없이 이어진다.
        # 이 조건이 없으면 '폭우일(7,8,10,12일)' 같은 실제 측정값까지 코드로 오인한다.
        contiguous = is_int and vmin in (0.0, 1.0) and (vmax - vmin + 1) == nu
        code_like = contiguous and 2 <= nu <= 15 and len(s) >= nu * 2
        if name_hit or code_like:
            num.remove(c)
            promoted.append(c)
    cat = [c for c in df.columns if c not in num]
    return num, cat, promoted


def detect_design(df):
    """데이터 구조를 보고 실험설계를 자동 판별.
    반환: dict(design, trt, blk, sub, ys, reason, confidence, promoted)"""
    num, cat, promoted = split_code_columns(df)
    res = {"design": "판별 불가", "trt": None, "blk": None, "sub": None,
           "ys": num, "reason": "", "confidence": "낮음", "promoted": promoted}
    if not num:
        res["reason"] = "숫자형 측정값 열이 없습니다."
        return res
    # 반복(블록) 후보: 이름 기반
    blk = next((c for c in cat if any(k in str(c).lower() for k in _BLOCK_KEYS)), None)
    # 처리구 후보: 반복이 아니면서 수준이 2~15개
    trt_cands = [c for c in cat if c != blk and 2 <= df[c].nunique() <= 15]
    # 응답자ID·이름 같은 열 제외
    trt_cands = [c for c in trt_cands if df[c].nunique() < len(df) * 0.9]
    if not trt_cands:
        res["reason"] = "처리구로 볼 만한 범주형 열이 없습니다."
        return res
    # 이름 우선순위로 처리구 선택
    trt = next((c for k in _TRT_KEYS for c in trt_cands if k in str(c)), trt_cands[0])
    others = [c for c in trt_cands if c != trt]
    res["trt"] = trt
    res["blk"] = blk
    # 균형 여부 확인
    def balanced(cols, strict=True):
        try:
            t = df.groupby(cols).size()
            if strict:
                return t.nunique() == 1
            return t.min() >= 1 and (t.max() - t.min()) <= 1
        except Exception:
            return False
    if blk and others:
        # 반복 + 요인 2개 → 요인배치/분할구. 단, 주 관심 요인을 '처리'로 우선
        res["sub"] = others[0]
        if balanced([trt, others[0], blk], strict=False):
            res["design"] = "난괴법 요인배치(2요인 × 반복)"
            res["reason"] = (f"반복('{blk}')이 있고 요인이 2개('{trt}', '{others[0]}')입니다. "
                             "→ 이원배치로 상호작용까지 보거나, 한 요인만 골라 난괴법으로 분석할 수 있습니다. "
                             "관수·경운처럼 큰 구역 요인이 있다면 **분할구법**을 쓰세요.")
            res["confidence"] = "높음"
        else:
            res["design"] = "이원배치(요인배치)"
            res["reason"] = f"두 요인('{trt}', '{others[0]}')이 있습니다."
            res["confidence"] = "중간"
    elif blk:
        if balanced([trt, blk], strict=False):
            res["design"] = "난괴법(RCBD)"
            n_rep = df[blk].nunique()
            res["reason"] = (f"처리구 '{trt}'({df[trt].nunique()}개)가 반복 '{blk}'"
                             f"({n_rep}반복)에 균형 있게 배치되어 있습니다.")
            res["confidence"] = "높음"
        else:
            res["design"] = "난괴법(RCBD, 불균형)"
            res["reason"] = f"반복 '{blk}'이 있으나 처리구별 반복 수가 고르지 않습니다."
            res["confidence"] = "중간"
    elif others:
        res["sub"] = others[0]
        res["design"] = "이원배치(요인배치)"
        res["reason"] = f"범주형 요인이 2개('{trt}', '{others[0]}') 있고 반복 열은 없습니다."
        res["confidence"] = "중간"
    else:
        res["design"] = "완전임의배치(CRD)"
        cnt = df.groupby(trt).size()
        res["reason"] = (f"처리구 '{trt}'({df[trt].nunique()}개)만 있고 반복 열이 없습니다. "
                         f"처리당 {cnt.min()}~{cnt.max()}개 관측치.")
        res["confidence"] = "높음" if cnt.min() >= 3 else "중간"
    # 연도·일련번호는 숫자지만 조사값이 아니므로 측정항목 후보에서 뺀다.
    ys = [c for c in num if not any(k in str(c).lower() for k in _ID_KEYS)]
    if not ys:
        ys = list(num)
    # 측정값 우선순위 정렬(수량류 먼저)
    res["ys"] = sorted(ys, key=lambda c: 0 if any(
        k in str(c) for k in ["수량", "수확", "생산량", "무게", "중"]) else 1)
    return res

def recommend_analysis(df):
    nc = df.select_dtypes(include=np.number).columns.tolist()
    cc = df.select_dtypes(exclude=np.number).columns.tolist()
    recs = []
    for c in cc:
        try:
            uniq = df[c].nunique()
            ng = int(uniq) if np.isscalar(uniq) else int(np.asarray(uniq).ravel()[0])
        except Exception:
            continue  # 중복 열 이름 등 비정상 구조는 건너뜀
        if 2 <= ng <= 15 and nc:
            recs.append((f"'{c}' 그룹이 {ng}개예요 → " +
                         ("**t-검정/ANOVA**" if ng == 2 else "**ANOVA + 사후검정(a,b,c)**") + " 이 적합합니다.", 3 if ng > 2 else 2))
    if len(cc) >= 2 and nc:
        recs.append(("범주형 변수가 2개 이상 → **이원배치 분산분석**으로 상호작용도 볼 수 있어요.", 2))
    if len(nc) >= 3:
        recs.append((f"숫자형 변수가 {len(nc)}개 → **PCA**로 특성을 압축·시각화할 수 있어요.", 1))
    if len(nc) >= 2:
        recs.append(("숫자형 변수가 여러 개 → **상관분석/히트맵**과 **회귀분석**이 가능합니다.", 2))
    seen, out = set(), []
    for msg, _ in sorted(recs, key=lambda x: -x[1]):
        if msg not in seen: seen.add(msg); out.append(msg)
    return out[:5] if out else ["데이터 구조상 뚜렷한 추천이 어려워요."]

def _josa(word, pair="이/가"):
    """받침 여부에 따라 조사 선택 (수비초가 / 청양이)"""
    a, b = pair.split("/")
    w = str(word).strip()
    if not w: return a
    ch = w[-1]
    if ch.isdigit():   # 숫자로 끝나면 읽는 소리로 판단 (1,3,6,7,8,0=받침 있음)
        return a if ch in "136078" else b
    if not ("가" <= ch <= "힣"):
        return a
    return a if (ord(ch) - 0xAC00) % 28 else b

def report_sentence_anova(gc, vc, pval, means, letters, ci=None, ph=""):
    """시험연구보고서 양식(○ / -)의 결과 문장 자동 작성"""
    order = means.sort_values("mean", ascending=False).index.tolist()
    top, low = order[0], order[-1]
    lines = [f"○ {gc}별 {vc} 분석 결과"]
    lines.append(f"  - {top}{_josa(top)} {means.loc[top,'mean']:.1f}로 가장 높았고, "
                 f"{low}{_josa(low)} {means.loc[low,'mean']:.1f}로 가장 낮았다.")
    if pval < 0.05:
        lines.append(f"  - 처리 간 유의한 차이가 인정되었다(p={pval:.4f}).")
        by = {}
        for g, l in letters.items(): by.setdefault(l, []).append(str(g))
        same = [v for v in by.values() if len(v) > 1]
        if same:
            _grp = ', '.join(same[0])
            lines.append(f"  - 다만 {_grp}{_josa(same[0][-1], '은/는')} 같은 문자군에 속하여 "
                         "통계적으로 동등한 수준이었다.")
    else:
        lines.append(f"  - 처리 간 유의한 차이는 인정되지 않았다(p={pval:.4f}).")
    if ci and not np.isnan(ci.get("CV", np.nan)):
        lines.append(f"  - 시험의 변이계수(CV)는 {ci['CV']:.1f}%로 "
                     f"{cv_grade(ci['CV'])} 수준이었다.")
    if ph:
        lines.append(f"  - 평균 간 비교는 {ph}(p<0.05)으로 실시하였다.")
    return "\n".join(lines)

def interpret_anova(pval, letters):
    if pval < 0.001: s = "처리구 간 **매우 뚜렷한 차이**가 있습니다 (p < 0.001)."
    elif pval < 0.05: s = f"처리구 간 **통계적으로 유의한 차이**가 있습니다 (p = {pval:.3f})."
    else: return f"처리구 간 유의한 차이가 **없습니다** (p = {pval:.3f} ≥ 0.05)."
    by = {}
    for g, l in letters.items(): by.setdefault(l, []).append(g)
    same = [v for v in by.values() if len(v) > 1]
    if same: s += f" 같은 문자를 가진 처리구({', '.join(map(str, same[0]))})끼리는 차이가 없습니다."
    return s

def interpret_corr(corr, sel):
    pairs = []
    for i in range(len(sel)):
        for j in range(i+1, len(sel)):
            pairs.append((abs(corr.iloc[i, j]), sel[i], sel[j], corr.iloc[i, j]))
    if not pairs: return ""
    pairs.sort(reverse=True); _, a, b, r = pairs[0]
    s = "강한" if abs(r) >= 0.7 else ("뚜렷한" if abs(r) >= 0.4 else "약한")
    d = "양(+)의" if r > 0 else "음(-)의"
    t = "한 변수가 커질수록 다른 변수도 커집니다." if r > 0 else "한 변수가 커질수록 다른 변수는 작아집니다."
    return f"가장 관계가 큰 변수는 **'{a}'와 '{b}'** 로, {s} {d} 상관입니다 (r = {r:.2f}). {t}"

EXPLAIN = {
"sd_se": """둘 다 '±' 뒤에 붙는 값이지만 **뜻이 완전히 다릅니다.**

| | 표준편차 (SD) | 표준오차 (SE) |
|---|---|---|
| 무엇을 재나요 | **개체들이** 평균에서 얼마나 흩어져 있나 | **평균값 자체가** 얼마나 믿을 만한가 |
| 반복을 늘리면 | 거의 그대로 (자연적인 변이라서) | **작아집니다** (√반복수로 나누므로) |
| 계산 | 자료의 흩어짐 | SD ÷ √반복수 |
| 언제 쓰나요 | "이 품종은 개체 간 편차가 크다"를 보일 때 | "처리 평균의 차이"를 보일 때 |

**쉽게 말하면**
- SD: *"이 처리구의 고추 무게는 개체마다 얼마나 들쭉날쭉한가?"*
- SE: *"이 처리구의 평균 무게를 얼마나 믿어도 되나?"*

**포장시험에서는 어떤 걸?**
- 처리구 간 비교 그래프(막대 + 오차막대) → **SE**를 더 많이 씁니다. 오차막대가 짧을수록 평균이 안정적이라는 뜻이라 처리 간 차이를 보기 좋습니다.
- 품종·계통의 균일도, 개체 변이를 설명할 때 → **SD**.

> ⚠️ SE는 반복수가 많을수록 무조건 작아지므로, **오차막대가 짧다고 처리 효과가 크다는 뜻은 아닙니다.**
> 유의차 판단은 오차막대가 아니라 **사후검정 문자(a, b, c)** 나 LSD로 하세요.
>
> 논문·보고서에는 `평균 ± SD` 인지 `평균 ± SE` 인지, 그리고 반복수(n)를 **표 각주에 반드시 밝혀야** 합니다.""",

"prep": """**전처리**는 분석 전에 데이터를 정리하는 단계입니다. 여기서 빠뜨린 문제는 이후 모든 결과를 왜곡시킵니다.

**꼭 확인할 것**
- **결측치**: 조사 누락·측정 실패로 비어 있는 칸
- **이상값**: 입력 실수(예: 15.0을 150으로)나 극단적으로 튀는 값
- **중복 행**: 같은 자료가 두 번 입력된 경우
- **자료형**: 숫자여야 하는데 문자로 읽힌 열(엑셀에서 '12.5 ' 처럼 공백이 섞이면 발생)

**순서 권장**: 자료형 확인 → 중복 제거 → 이상값 확인 → 결측치 처리""",

"outlier": """**이상값(outlier)**은 다른 값들과 유난히 동떨어진 값입니다. 평균과 분산을 크게 흔들어 분석 결과를 왜곡합니다.

**두 가지 탐지 방법**
- **IQR(사분위수) 방법**: 자료를 크기순으로 줄 세워 가운데 50% 구간(IQR)을 구한 뒤, 그 범위의 1.5배를 벗어나면 이상값으로 봅니다. 분포가 치우쳐 있어도 잘 작동합니다.
- **Z-점수 방법**: 평균에서 표준편차의 3배 이상 떨어지면 이상값으로 봅니다. 정규분포에 가까울 때 적합합니다.

**처리 방법 고르기**
- 측정 실수·입력 오류가 확실 → **해당 행 삭제**
- 실제 값이지만 너무 극단적 → **경계값으로 대체(윈저화)**
- 판단 보류 → **결측치로 변경** 후 따로 검토

⚠️ 이상값이 항상 오류인 것은 아닙니다. 특이한 개체가 실제로 존재할 수 있으니, 지우기 전에 원본 조사표를 꼭 확인하세요.""",

"derive": """**파생변수**는 기존 열을 조합해 새로운 열을 만드는 기능입니다. 분석에 필요한 지표가 원자료에 없을 때 사용합니다.

**세 가지 방식**
1. **두 열 사칙연산** — 예: `생체중 ÷ 초장` = 단위 길이당 무게, `수량 × 단가` = 조수입
2. **조건 열** — 조건을 만족하면 1, 아니면 0. 예: `일 최고기온 ≥ 33` → 폭염일 표시
3. **그룹별 집계** — 연도·처리구별로 합계·평균 등을 계산

**활용 예: 연간 폭염일수 구하기**
① 조건 열로 `기온 ≥ 33` 만들기 → ② 그룹별 집계에서 '연도별 **합계**' → 연도마다 폭염일이 며칠인지 나옵니다.
(냉해일수는 `기온 ≤ 0`, 강우일수는 `강수량 ≥ 30` 등으로 같은 방식)""",

"corr": """**상관분석**은 두 변수가 함께 변하는 정도를 하나의 숫자(r)로 나타냅니다.

**상관계수 r 읽는 법** (−1 ~ +1)
- **+**: 한쪽이 커지면 다른 쪽도 커짐 (예: 생체중↑ → 수량↑)
- **−**: 한쪽이 커지면 다른 쪽은 작아짐
- **0에 가까움**: 뚜렷한 관계 없음

|절댓값|해석|
|---|---|
|0.7 이상|강한 상관|
|0.4~0.7|뚜렷한 상관|
|0.2~0.4|약한 상관|
|0.2 미만|거의 없음|

**Pearson vs Spearman**
- **Pearson**: 직선 관계를 봅니다. 자료가 정규분포에 가까울 때 사용.
- **Spearman**: 순위로 바꿔서 계산합니다. 정규분포가 아니거나, 등급·순위 자료(1등급·2등급 등)일 때 사용.

**히트맵**은 여러 변수의 상관을 색으로 한눈에 보여줍니다. 붉을수록 양(+), 푸를수록 음(−)의 상관입니다.

⚠️ **상관은 인과가 아닙니다.** 두 변수가 같이 움직인다고 해서 하나가 다른 하나의 원인이라는 뜻은 아닙니다.""",

"anova": """**분산분석(ANOVA)**은 세 개 이상 처리구의 평균이 서로 다른지 검정하는 방법입니다.
(두 개만 비교할 때는 t-검정을 쓰지만, ANOVA로도 같은 결론이 나옵니다)

**왜 필요한가?** 처리구가 4개일 때 t-검정을 6번 반복하면, 실제로는 차이가 없는데도 우연히 "차이 있다"고 나올 확률이 크게 올라갑니다. ANOVA는 한 번에 검정해 이 문제를 피합니다.

**결과 읽는 법**
- **p < 0.05**: "적어도 한 처리구는 다르다" → 사후검정으로 어느 것이 다른지 확인
- **p ≥ 0.05**: 처리 간 차이가 뚜렷하지 않음

---
**📐 실험 설계 고르기 (중요)**

- **완전임의배치(CRD)**: 처리를 완전히 무작위로 배치. 온실처럼 환경이 균일할 때.
- **난괴법(RCBD)**: 포장을 몇 개 블록(반복)으로 나누고 각 블록 안에 모든 처리를 배치. **우리 포장시험의 표준**입니다.
  - 👉 반복(블록)을 두었다면 **'반복(블록) 열'을 반드시 지정**하세요. 지정하지 않으면 블록 간 토양·경사 차이가 오차에 섞여, 실제로는 있는 처리 효과를 놓칠 수 있습니다.
- **이원배치**: 두 요인을 동시에 봅니다(예: 품종 × 시비량).
  - **상호작용이 유의하다** = "한 요인의 효과가 다른 요인에 따라 달라진다"는 뜻. 예를 들어 A품종은 시비를 늘리면 수량이 늘지만 B품종은 오히려 줄어드는 경우입니다. 이때는 주효과만 보면 안 되고 조합별로 해석해야 합니다.

---
**✅ 가정 검정 (분석 전 자동 수행)**

ANOVA는 두 가지를 전제합니다.
- **정규성** (Shapiro-Wilk): 각 처리구 자료가 정규분포를 따르는가
- **등분산** (Levene): 처리구들의 흩어진 정도가 비슷한가

p ≥ 0.05면 가정을 만족합니다. 위배되면 **비모수검정(Kruskal-Wallis)** 을 쓰는 것이 안전합니다.

---
**🔤 사후검정과 유의성 문자(a, b, c)**

ANOVA는 "어딘가 다르다"까지만 알려줍니다. **어느 처리구끼리** 다른지는 사후검정으로 확인합니다.

- **같은 문자를 공유하면 차이 없음**, 문자가 완전히 다르면 차이 있음
- 예: 처리2(a), 처리1(ab), 대조구(b) → 처리2와 대조구는 차이 있지만, 처리1은 둘 중 어느 쪽과도 뚜렷한 차이가 없음

|방법|특징|
|---|---|
|**Tukey HSD**|국제 표준. 위양성을 잘 통제해 **논문 투고에 안전**|
|**던컨(DMRT)**|농업 논문 관행. 차이를 잘 잡아내지만 **위양성 위험이 큼**|
|**Bonferroni**|매우 엄격. 확실한 차이만 인정|

⚠️ 시비량·재식밀도처럼 **연속적인 수준**을 처리로 둔 경우에는 사후검정보다 **회귀분석**이 적절합니다.""",

"nonparam": """**비모수 검정**은 정규분포를 가정하지 않는 검정입니다.

**언제 쓰나요?**
- ANOVA의 정규성·등분산 가정이 깨졌을 때
- 표본이 매우 적을 때(처리당 5개 미만)
- 등급·순위처럼 간격이 일정하지 않은 자료(1=매우나쁨 ~ 5=매우좋음 등)
- 병해 발생 정도처럼 점수로 매긴 자료

**두 가지 방법 (자동 선택됩니다)**
- **Kruskal-Wallis**: 3개 이상 그룹 비교 → ANOVA의 비모수 버전
- **Mann-Whitney U**: 2개 그룹 비교 → t-검정의 비모수 버전

**결과 읽는 법**: p < 0.05면 그룹 간 차이가 있습니다. 평균 대신 **중앙값**으로 비교합니다(극단값의 영향을 덜 받기 때문).""",

"pca": """**주성분분석(PCA)**은 변수가 너무 많을 때, 정보를 최대한 유지하면서 **2개의 축으로 압축**해 그림 하나로 보여주는 방법입니다.

---
**🤔 왜 필요한가요?**

품종 10개에 대해 초장·엽수·생체중·과장·과경·당도·수량… 10가지 형질을 조사했다고 해봅시다.
"어떤 품종끼리 서로 비슷한가?"를 알고 싶은데, 형질이 10개면 그래프를 10차원으로 그려야 해서 눈으로 볼 수가 없습니다.
PCA는 이 10개 정보를 **가장 정보 손실이 적은 2개의 새 축**으로 요약해, 평면 위 산점도 하나로 보여줍니다.

**비유**: 사람을 여러 각도에서 찍을 수 있지만, 얼굴이 가장 잘 드러나는 각도 하나를 고르는 것과 비슷합니다. PCA는 데이터가 가장 잘 퍼져 보이는 각도를 수학적으로 찾아줍니다.

---
**📊 결과 읽는 법**

**① 설명분산비율 (가장 중요)**
- `PC1 45%, PC2 30% (누적 75%)` → 원래 정보의 75%를 이 그림 하나로 설명한다는 뜻
- **누적 70% 이상이면 신뢰할 만합니다.** 50% 미만이면 2차원 요약이 무리라는 뜻이니 해석에 주의하세요.

**② 산점도 (점들의 위치)**
- **가까운 점 = 서로 비슷한 개체/품종**
- 처리구별로 색을 나눴을 때 **무리가 뚜렷하게 갈리면**, 그 형질들로 처리구를 구분할 수 있다는 의미입니다.
- 반대로 색이 뒤섞여 있으면 처리 간 특성 차이가 크지 않다는 뜻입니다.

**③ 로딩표 (변수별 기여도)**
- 각 변수가 PC1·PC2를 만드는 데 얼마나 기여했는지 보여줍니다.
- **절댓값이 큰 변수**가 그 축의 의미를 결정합니다.
- 예: PC1에서 수량 0.52, 생체중 0.49, 초장 0.47처럼 크기 관련 형질이 모두 크면 → "PC1은 대체로 **식물체의 크기**를 나타내는 축"이라고 해석합니다.

---
**🌱 농업 연구에서 이렇게 씁니다**
- 여러 계통·품종을 형질 전체로 묶어 **유연관계 파악** (육종)
- 처리구가 여러 형질에서 **전체적으로 구분되는지** 확인
- 서로 비슷한(중복된) 형질을 찾아 **조사 항목 줄이기**

⚠️ 변수마다 단위가 달라도 괜찮습니다(자동 표준화됨). 다만 숫자형 변수가 **3개 이상** 필요하고, 표본이 너무 적으면(10개 미만) 결과가 불안정합니다.""",

"reg": """**회귀분석**은 한 변수(Y)를 다른 변수(X)로 **설명하거나 예측하는 식**을 만듭니다.

**ANOVA와 차이**: ANOVA는 "처리구별로 다른가?"(범주 비교), 회귀는 "X가 1 늘면 Y는 얼마나 변하나?"(수량적 관계)를 봅니다. 시비량처럼 **연속적인 수준**을 다룰 때는 회귀가 적합합니다.

**결과 읽는 법**
- **계수**: X가 1 증가할 때 Y의 변화량. 예: 계수 2.5 → 질소 1kg 증가 시 수량 2.5kg 증가
- **p-value**: 0.05 미만이면 그 변수가 의미 있게 기여함
- **R²(결정계수)**: 모델이 Y의 변동을 몇 % 설명하는지. 0.7이면 70% 설명

**VIF(다중공선성)** — 독립변수를 2개 이상 넣으면 자동 표시됩니다.
서로 너무 비슷한 변수(예: 초장과 생체중)를 함께 넣으면 계수가 뒤죽박죽이 됩니다.
**VIF 10 이상이면 경고**가 뜨니, 둘 중 하나를 빼세요.

**로지스틱 회귀**: Y가 **두 가지 값**일 때 사용합니다(발병/미발병, 합격/불합격 등).

⚠️ 관측 범위를 벗어난 예측은 위험합니다. 질소 0~30kg 자료로 만든 식을 60kg에 적용하면 안 됩니다.""",

"ml": """**머신러닝**은 데이터의 복잡한 패턴을 학습해 값을 예측하는 방법입니다.

**언제 쓰나요?**
- 변수가 많고 관계가 복잡해서 단순 회귀로 설명이 어려울 때
- **예측 자체**가 목적일 때(수량 예측, 등급 판정, 병해 발생 예측 등)
- 설문·센서·기상 등 **자료가 많을 때**

**알고리즘 고르기**
- **트리·앙상블**: 랜덤포레스트, Extra Trees, 그래디언트부스팅, 히스토그램 부스팅, AdaBoost, 의사결정나무
- **거리·경계 기반**: SVM(RBF), KNN — 변수 단위가 달라도 자동 표준화합니다.
- **회귀 전용 규제모형**: Ridge, Lasso, ElasticNet — 선형 관계와 다중공선성이 있을 때 유용합니다.
- **분류 전용 기본모형**: 로지스틱 회귀, GaussianNB

처음이라면 **랜덤포레스트**를 기준모형으로 먼저 돌리고, 다른 알고리즘과 테스트 성능을 비교하세요.

**결과 읽는 법**
- **R²**(회귀) / **정확도**(분류): 1에 가까울수록 잘 맞춤. 학습에 쓰지 않은 자료(테스트)로 평가한 값입니다.
- **변수 중요도**: 예측에 어떤 변수가 크게 기여했는지 순위

⚠️ **표본이 적으면 쓰지 마세요.** 포장시험처럼 반복이 3~4회뿐인 자료는 과적합(외운 것처럼 보이지만 새 자료에서는 틀림)이 일어납니다.
**처리 효과 검정은 반드시 분산분석**을 쓰고, 머신러닝은 참고용으로만 보세요.""",

"econ": """**경제성 분석**은 시험 결과를 '돈'으로 환산해, 그 처리가 실제로 농가에 이득인지 판단합니다.

**계산 체계** (농촌진흥청 농축산물 소득조사 기준)
- **총수입(조수입)** = 주산물가액(수량×단가) + **부산물가액**
- **경영비** = 생산에 투입된 경영비(종묘비·비료비·농약비·고용노력비·임차료·감가상각비 등). 현금지출만을 뜻하지는 않습니다.
- **생산비** = 경영비 + 자가노력비 + 유동자본용역비 + 고정자본용역비 + 자가토지 용역비
- **소득 = 총수입 − 경영비** → 경영비를 차감한 농업경영 성과
- **순수익 = 총수입 − 생산비** → 자기 노동·토지의 기회비용까지 뺀 순수 이익
- **소득률(%) = 소득 ÷ 총수입 × 100**

**지표 읽는 법**
- **소득률**: 고추는 대체로 50% 내외입니다(2024년 시설고추 56.3%).
- **단년도 총수입/생산비**: 1보다 크면 해당 연도의 입력 조건에서 총수입이 생산비보다 큼. 시설투자의 할인 B/C와는 다른 지표입니다.
- **손익분기수량**: (생산비 − 부산물가액 − 수량비례비) ÷ (단가 − 단위당 수량비례비)로 구합니다. 수확·선별·포장비처럼 수량에 비례하는 비용을 따로 지정하지 않으면 **(생산비 − 부산물가액) ÷ 단가**가 되어, "그해 들어간 비용을 회수하려면 몇 kg을 수확해야 하는가"를 뜻합니다. 실제 수량이 이보다 많아야 이익입니다.
- **가격 민감도**: 단가가 떨어져도 소득이 (+)로 유지되는 처리가 가격 위험에 강합니다.

**네 가지 분석 방식**
- **📕 부분예산표**: 기존 기술과 비교해 바뀌는 비용·수입만 계산
- **📗 소득분석**: 처리별 소득·순수익을 계산해 비교 (보고서·소득자료용)
- **📘 부분예산·MRR**: 지배분석 후 추가 비용 대비 순편익 증가율을 비교
- **📙 시설·장기투자**: 여러 해의 현금흐름을 할인해 NPV·할인 B/C·IRR·회수기간을 계산

MRR은 사용자가 정한 최소수용 기준과 자료 신뢰도·민감도를 함께 보고 판단합니다.""",

"survey": """**설문 분석**은 응답자 특성과 문항 응답을 함께 살펴봅니다.

**문항 유형별 분석 방법**
- **리커트 척도**(1~5점): 평균·표준편차, 긍정응답 비율, 신뢰도, 집단별 비교
- **객관식**(하나만 선택): 빈도·비율
- **다중응답**(모두 선택): 응답률(응답자 대비). 합계가 100%를 넘는 것이 정상입니다.
- **주관식**(자유 서술): 응답 목록, 주요 단어 빈도, AI 요약
- **교차분석**: 두 문항의 관계를 카이제곱 검정으로 확인

**크론바흐 알파(α) — 신뢰도**
여러 문항이 **같은 개념을 일관되게 측정하는지** 보는 지표입니다.

|α 값|해석|
|---|---|
|0.9 이상|매우 높음|
|0.8~0.9|높음|
|0.7~0.8|양호 (일반적 기준)|
|0.7 미만|낮음 — 문항 재검토 필요|

**'문항 제외 시 α'**를 보면, 어떤 문항을 뺐을 때 신뢰도가 크게 올라가는지 알 수 있습니다. 그 문항은 다른 문항들과 방향이 다르다는 뜻이니 재검토 대상입니다.

**🤖 자동 인식**을 쓰면 각 열의 값을 보고 문항 유형을 스스로 판별해 한 번에 분석합니다.""",

"report": """**자동 보고서**는 여러 분석 결과를 하나의 한글(hwpx) 문서로 만들어 줍니다.

**사용 순서**
1. 각 분석을 실행합니다.
2. 결과 아래 **'➕ 이 결과를 보고서에 담기'** 를 누릅니다.
3. 원하는 분석을 모두 담은 뒤, 이 화면에서 **'보고서 생성'** 을 누릅니다.

**보고서에 들어가는 것**: 소제목 · 해석 문장 · 결과표(`<표 1>` 캡션) · 그래프(`<그림 1>` 캡션, 가운데 정렬)

표 서식(글꼴·크기·음영·선 굵기·행 높이)은 사이드바 **⚙️ 한글 표 서식 설정**에서 미리 바꿀 수 있습니다.
**🕘 분석 이력**에는 언제 어떤 분석을 했는지 자동 기록되며, 이 기록도 보고서에 첨부할 수 있습니다.""",
}

# ---------------------------------------------------------------- 샘플 데이터
def make_sample(kind):
    if kind == "실험":
        # 2품종 × 4처리 × 4반복 = 32행 (일원·이원배치·난괴법 모두 시연 가능)
        rng = np.random.default_rng(42)
        trt = {"대조구": 0, "처리1": 13, "처리2": 18, "처리3": 5}
        var = {"청양": 0, "수비초": 8}
        blk = {"I": -4, "II": 0, "III": 3, "IV": 1}
        rows = []
        for v, ve in var.items():
            for t, te in trt.items():
                for b, be in blk.items():
                    rows.append({
                        "품종": v, "처리구": t, "반복": b,
                        "초장(cm)": round(95 + te*0.8 + ve*0.5 + be*0.4 + rng.normal(0, 3), 1),
                        "엽수(개)": round(8 + te*0.15 + ve*0.2 + rng.normal(0, 0.8)),
                        "생체중(g)": round(120 + te*2.2 + ve*1.5 + be + rng.normal(0, 6), 1),
                        "수량(kg/10a)": round(480 + te*6 + ve*4 + be*2 + rng.normal(0, 15)),
                    })
        return pd.DataFrame(rows)

    if kind == "경제성":
        return pd.DataFrame({
            "처리구": ["대조구", "처리1", "처리2", "처리3"],
            "수량": [250, 290, 310, 265], "단가": [15000]*4,
            "종자비": [90000]*4, "비료비": [350000, 400000, 450000, 380000],
            "농약비": [300000, 290000, 285000, 295000],
            "고용노력비": [600000, 640000, 660000, 620000],
            "재료비": [250000, 255000, 260000, 252000], "감가상각비": [200000]*4,
            "자가노동시간": [95, 105, 110, 98]})

    if kind == "설문":
        # 응답자별 '전반적 만족도(잠재요인)'를 두어 문항 간 상관이 생기도록 구성
        rng = np.random.default_rng(7); N = 60
        latent = rng.normal(0, 1, N)
        def item(bias):
            v = 3.8 + bias + latent*0.85 + rng.normal(0, 0.55, N)
            return np.clip(np.round(v), 1, 5).astype(int)
        opts = ["분산분석", "상관분석", "그래프 작성", "한글 표 생성", "머신러닝", "경제성 분석"]
        return pd.DataFrame({
            "응답자ID": [f"R{i+1:03d}" for i in range(N)],
            "성별": rng.choice(["남", "여"], N),
            "연령대": rng.choice(["20대", "30대", "40대", "50대 이상"], N),
            "소속": rng.choice(["농업기술원", "농업기술센터", "국립연구소"], N),
            "경력": rng.choice(["5년 미만", "5~10년", "10년 이상"], N),
            "주사용목적": rng.choice(["논문 작성", "시험 보고서", "현장 지도", "교육 자료"], N),
            "사용기능(다중)": [";".join(rng.choice(opts, size=int(rng.integers(1, 4)), replace=False))
                          for _ in range(N)],
            "Q1_사용편의성": item(0.30), "Q2_분석속도": item(-0.25), "Q3_결과신뢰도": item(0.10),
            "Q4_기능충분성": item(-0.15), "Q5_보고서품질": item(0.35), "Q6_재사용의향": item(-0.05),
            "개선의견": rng.choice([
                "통계 종류가 더 많았으면 좋겠습니다",
                "한글 표 서식이 편리해서 업무 시간이 줄었습니다",
                "그래프 색상을 더 다양하게 바꾸고 싶습니다",
                "사용법 설명이 자세해서 초보자도 쓰기 좋습니다",
                "엑셀 시트가 많을 때 처리가 편리했으면 합니다",
                "보고서 자동 생성 기능이 가장 유용했습니다", ""], N)})

    if kind == "반복측정":
        rng = np.random.default_rng(5); rows = []
        for i in range(1, 13):
            b0 = rng.normal(18, 3)
            for w, add in zip(["2주", "4주", "6주", "8주"], [0, 9, 17, 23]):
                rows.append({"개체번호": f"P{i:02d}", "조사시기": w,
                             "초장(cm)": round(b0 + add + rng.normal(0, 1.8), 1)})
        return pd.DataFrame(rows)

    if kind == "분할구":
        # 주구(관수) 2 × 세구(품종) 3 × 반복 3 = 18
        rng = np.random.default_rng(7); rows = []
        for rep in ["I", "II", "III"]:
            for mp, me in [("관수", 8), ("무관수", 0)]:
                for _cv, se in [("청양", 0), ("수비초", 5), ("칼미007", 9)]:
                    rows.append({"반복": rep, "관수방법": mp, "품종": _cv,
                                 "수량(kg/10a)": round(480 + me*6 + se*5 + rng.normal(0, 12), 1)})
        return pd.DataFrame(rows)

    if kind == "프로빗":
        rng = np.random.default_rng(9); rows = []
        for line, shift in [("감수성계통", 0.0), ("저항성계통", 0.55)]:
            for dose in [5, 10, 20, 40, 80, 160]:
                p = stats.norm.cdf((np.log10(dose) - (1.25 + shift)) * 2.3)
                rows.append({"계통": line, "농도(ppm)": dose, "공시충수": 30,
                             "사충수": int(np.clip(round(30*p + rng.normal(0, 1.2)), 0, 30))})
        return pd.DataFrame(rows)
    return pd.DataFrame()

# ================================================================ 세션
# 메뉴를 옮겨다녀도 각 화면의 선택 상태가 초기화되지 않도록 붙잡아 둔다.
# (스트림릿은 화면에 그려지지 않은 위젯의 상태를 자동으로 버린다)
# 화면(메뉴)이 바뀌면 스트림릿은 그려지지 않은 위젯의 상태를 버린다.
# 아래 키들은 각 메뉴 안에서만 그려지므로 값을 다시 써 넣어 유지한다.
# (사이드바 위젯은 매번 그려지므로 대상이 아니다)
_PINNED_DEFAULTS = {"ap_ph": "Tukey HSD", "ap_err": "표준편차(SD)", "ap_max": 8}
for _k, _v in _PINNED_DEFAULTS.items():
    st.session_state.setdefault(_k, _v)

# 전역 설정(글꼴·그래프·API키 등)과 데이터 그 자체 — 데이터가 바뀌어도 그대로 두는 키들.
_PIN_GLOBAL_EXACT = {
    "files", "df", "cur_key", "price_db", "report_items",
    "hdr_rows", "del_sel", "err_type", "round_n", "plot_color",
    "svy_type", "econ_mode", "stat_sub", "svy_chart",
    "menu_choice", "menu_main", "menu_support",
    "ap_ph", "ap_err", "ap_max",
    "_pin_store", "_pin_owner", "_pin_deny",
}
_PIN_GLOBAL_PREFIX = ("hwp_", "sup_", "fig_", "kamis_", "kosis_", "price_",
                      "pj_", "ai_", "api_", "dl_", "plan_", "report_",
                      "gen_report", "FormSubmitter", "$$", "uncaught")

# ★ 버튼·다운로드버튼 키는 st.session_state 로 값을 써 넣을 수 없다(스트림릿이 막는다).
#   여기 빠뜨리면 "Values for the widget with key '...' cannot be set using
#   st.session_state" 오류로 화면 전체가 멈추므로, 아래 목록 + 자동 학습(_pin_deny)으로
#   이중으로 막는다.
# st.file_uploader 도 session_state 로 값을 되돌릴 수 없다(ml_pf).
_PIN_BUTTON_EXACT = {
    "econ_selftest", "econ_test_run", "ml_predict1", "ml_predict2", "ml_dlpred",
    "pbd_fill", "pbd_clear", "fixnum", "sc", "price_up", "ml_pf",
    # 인증/이미지/카메라/음성 입력은 Streamlit이 값을 소유하는 비-settable 위젯이다.
    # 이 키를 _pin_sync가 session_state에 다시 쓰면 StreamlitValueAssignmentNotAllowedError가 난다.
    "auth_login_btn", "auth_signup_btn", "auth_reset_btn", "auth_logout_sidebar",
    "table_img_up", "table_cam", "img_parse_btn", "image_table_editor", "use_image_table",
    "voice_data_audio", "voice_parse_btn", "voice_rows_editor", "voice_append",
    "voice_new", "voice_clear",
    # 경제성 분석 길잡이 초기화 버튼은 버튼 상태를 session_state로 복원하면 안 된다.
    "econ_guide_reset", "econ_guide_home", "econ_switch_guide",
}
# 버튼뿐 아니라 st.data_editor 도 session_state 로 값을 써 넣을 수 없다.
# 이 앱의 해당 위젯 전부:
#   data_editor  → rank_*, pb_gain_*, pb_loss_*, price_editor
#   button       → __btn_*, btn_*, aib_*, aiadd_*, aidel_*, errai_*, list_models_*,
#                  p_*, pbd_*, ml_predict*, econ_selftest, econ_test_run,
#                  ai_conn_test, kamis_apply, kosis_*
#   download     → dl_*, ml_dlpred
# (price_*, ai_*, kamis_*, kosis_*, dl_* 는 이미 전역 목록에서 걸러진다)
_PIN_BUTTON_PREFIX = ("__btn_", "btn_", "aib_", "aiadd_", "aidel_", "errai_",
                      "list_models_", "rm_", "p_", "rank_",
                      "pb_gain_", "pb_loss_", "pbd_", "ml_predict", "ml_dl", "ms_plot_dl_",
                      "econ_entry_", "econ_g_",
                      "hwx_", "dcx_", "csv_", "xls_", "gai_",
                      "svyhwp_", "svyxls_")

# 데이터와 무관하지만 '메뉴 안에서만' 그려지는 위젯들 — 데이터별로 나눌 필요는 없어도
# 매 실행마다 붙잡아 두지 않으면 다른 메뉴에 다녀올 때 기본값으로 돌아간다.
_PIN_TOUCH_GLOBAL = ("svy_type", "econ_mode", "stat_sub", "svy_chart", "ai_mode",
                     "ap_ph", "ap_err", "ap_max")


def _pin_scoped_keys():
    """데이터(시트)마다 따로 기억해야 하는 화면 선택 키 목록.

    경제성분석의 열 선택처럼 **열 이름에 묶인** 값들이다. 다른 데이터로 옮겼을 때
    그대로 남아 있으면 없는 열을 가리켜 위젯 오류가 나므로, 데이터별로 보관했다가
    돌아왔을 때 되돌려 준다.
    """
    out = []
    for k in list(st.session_state.keys()):
        if not isinstance(k, str):
            continue
        if k in _PIN_GLOBAL_EXACT or k.startswith(_PIN_GLOBAL_PREFIX):
            continue
        if k in _PIN_BUTTON_EXACT or k.startswith(_PIN_BUTTON_PREFIX):
            continue
        if k in st.session_state.get("_pin_deny", ()):   # 실행 중에 배운 버튼 키
            continue
        out.append(k)
    return out


def _pin_restore_defaults():
    """상태를 지운 뒤 기본값을 다시 채운다.

    `st.slider("...", 1, 15, key="ap_max")` 처럼 **value= 없이 key= 로만** 만든 위젯은
    세션에 값이 없으면 최솟값(1)으로 떨어진다. 원클릭 보고서의 '한 번에 분석할
    조사항목 수'가 1이 되어 측정항목을 하나만 분석하던 원인이었다.
    """
    for k, v in _PINNED_DEFAULTS.items():
        st.session_state.setdefault(k, v)


def _pin_sync():
    """(1) 그려지지 않은 위젯의 값이 버려지지 않도록 매번 다시 써 넣고,
       (2) 분석할 데이터가 바뀌면 데이터별로 상태를 보관·복원한다.

    스트림릿은 '이번 실행에서 화면에 그려지지 않은 위젯'의 값을 버린다. 그래서
    경제성분석 → 설문조사 → 경제성분석 으로 메뉴를 옮기면 경제성 화면의 선택이
    전부 기본값으로 돌아가고, '단위 확인' 체크가 풀리면서 실행 버튼이 비활성화돼
    분석 결과까지 사라졌다. 시트를 바꿔 가며 쓰는 경우(경제성=②시트, 설문=③시트)도
    같은 문제가 생기므로 데이터별로 나눠서 보관한다.
    """
    for k in _PIN_TOUCH_GLOBAL:
        if k in st.session_state:
            try:
                st.session_state[k] = st.session_state[k]
            except Exception:
                pass
    cur = st.session_state.get("cur_key")
    store = st.session_state.setdefault("_pin_store", {})
    if "_pin_owner" not in st.session_state:
        st.session_state["_pin_owner"] = cur      # 첫 실행: 지금 상태를 현재 데이터 것으로 인정
    if st.session_state["_pin_owner"] != cur:
        prev = st.session_state["_pin_owner"]
        keys = _pin_scoped_keys()
        if prev is not None:
            store[prev] = {k: st.session_state[k] for k in keys}
        for k in keys:
            try:
                del st.session_state[k]
            except Exception:
                pass
        st.session_state["_pin_owner"] = cur
        for k, v in store.get(cur, {}).items():
            try:
                st.session_state[k] = v
            except Exception:
                pass
        _pin_restore_defaults()
        return
    for k in _pin_scoped_keys():
        try:
            st.session_state[k] = st.session_state[k]
        except Exception:
            pass
    _pin_restore_defaults()

if "files" not in st.session_state: st.session_state.files = {}
if "df" not in st.session_state: st.session_state.df = None
if "price_db" not in st.session_state: st.session_state["price_db"] = None
if "report_items" not in st.session_state: st.session_state.report_items = []

# Supabase가 설정되고 AUTH_REQUIRED=true이면 로그인한 사용자만 아래 앱을 렌더링합니다.
render_auth_gate()

# ================================================================ 사이드바
st.sidebar.title("📊 스마트 통계 에이전트")
st.sidebar.caption("실험 데이터 자동 통계 분석 시스템")

# 로그인 사용자의 소속을 사이드바에 표시한다.
_auth_u = _current_auth_user()
if _auth_u:
    _meta = _auth_u.get("user_metadata") or {}
    _who = _meta.get("name") or str(_auth_u.get("email", "")).split("@")[0]
    _org = _meta.get("organization") or "소속 미입력"
    st.sidebar.markdown(f"**👤 {_who}**  \n🏢 {_org}")
    if st.sidebar.button("로그아웃", width="stretch", key="auth_logout_sidebar"):
        _auth_logout()

with st.sidebar.expander("📂 데이터 불러오기", expanded=True):
    _input_mode = st.radio("입력 방식", ["📁 Excel/CSV", "📷 이미지/사진", "🎤 음성"],
                           horizontal=False, key="data_input_mode")

    if _input_mode == "📁 Excel/CSV":
        ups = st.file_uploader("Excel / CSV 업로드 (여러 개 가능)",
                               type=["xlsx", "xls", "csv"], accept_multiple_files=True)
        hdr_rows = st.radio("↳ 머리글(변수명) 행 수", [1, 2], horizontal=True, key="hdr_rows",
                            help="변수명이 두 줄로 되어 있으면 2를 선택하세요. 두 줄이 합쳐진 이름으로 만들어집니다.")
        st.caption("엑셀에 시트가 여러 개면 시트별로 나뉘어 들어옵니다. "
                   "값을 바꾸면 올려둔 파일을 **자동으로 다시 읽습니다**(새로고침 불필요).")
        # 머리글 행 수를 바꾸면 이미 올린 파일을 다시 읽어 화면에도 즉시 반영한다.
        _hdr_changed = st.session_state.get("__hdr_prev") not in (None, hdr_rows)
        st.session_state["__hdr_prev"] = hdr_rows
        if ups:
            head = [0, 1] if hdr_rows == 2 else 0
            for uf in ups:
                try:
                    if uf.name.endswith(".csv"):
                        d = None
                        for _enc in ("utf-8-sig", "cp949", "euc-kr", "utf-8"):
                            try:
                                uf.seek(0)
                                d = pd.read_csv(uf, header=head, encoding=_enc)
                                break
                            except (UnicodeDecodeError, LookupError):
                                continue
                            except Exception:
                                uf.seek(0)
                                d = pd.read_csv(uf, header=head, encoding_errors="replace")
                                break
                        if d is None:
                            st.error(f"'{uf.name}' 파일의 문자 인코딩을 읽지 못했습니다. "
                                     "엑셀에서 'CSV UTF-8'로 다시 저장해 보세요.")
                            continue
                        if hdr_rows == 2:
                            d.columns = [" ".join([str(x) for x in c if "Unnamed" not in str(x)]).strip()
                                         for c in d.columns]
                        # 병합된 두 줄 헤더를 이어붙이면, 블록 사이의 빈 구분용 열까지 앞 헤더를
                        # 그대로 물려받아 "가격" 같은 이름이 중복 생성될 수 있다. 데이터가 전혀
                        # 없는(전부 결측) 열은 실제 문항이 아니라 이 구분용 유령 열이므로 제거한다.
                        d = d.dropna(axis=1, how="all")
                        st.session_state.files[uf.name] = clean_columns(d)
                    else:
                        xls = pd.ExcelFile(uf)
                        for sh in xls.sheet_names:      # 시트별로 저장
                            d = pd.read_excel(xls, sheet_name=sh, header=head)
                            if hdr_rows == 2:
                                d.columns = [" ".join([str(x) for x in c if "Unnamed" not in str(x)]).strip()
                                             for c in d.columns]
                            # 위와 동일한 이유로, 병합헤더가 물려준 빈 구분용 열(전부 결측)은 제거
                            d = d.dropna(axis=1, how="all")
                            key = f"{uf.name} – {sh}" if len(xls.sheet_names) > 1 else uf.name
                            st.session_state.files[key] = clean_columns(d)
                except Exception as e:
                    st.error(f"{uf.name} 읽기 실패: {e}")
            if _hdr_changed:
                # 다시 읽은 결과를 현재 분석 화면에도 즉시 적용
                _cur = st.session_state.get("cur_key")
                if _cur in st.session_state.files:
                    st.session_state.df = st.session_state.files[_cur].copy()
                st.success(f"머리글 {hdr_rows}행 기준으로 다시 읽었습니다.")
        elif _hdr_changed:
            st.info("머리글 행 수를 바꿨습니다. 파일을 다시 올리면 새 기준으로 읽습니다.")

        st.markdown("---")
        st.markdown("**🧪 샘플 데이터 (연습용)**")
        st.caption("데이터가 없다면 아래 버튼을 눌러 체험해 보세요.")
        c6, c7 = st.columns(2)
        if c6.button("🔁 반복측정", width="stretch"):
            st.session_state.files["샘플_반복측정"] = make_sample("반복측정")
        if c7.button("🧪 프로빗", width="stretch"):
            st.session_state.files["샘플_프로빗"] = make_sample("프로빗")
        if st.button("🌾 분할구(Split-plot)", width="stretch"):
            st.session_state.files["샘플_분할구"] = make_sample("분할구")
        c1, c2, c3 = st.columns(3)
        if c1.button("🌱 실험", width="stretch"):
            st.session_state.files["샘플_실험데이터"] = make_sample("실험")
        if c2.button("💰 경제성", width="stretch"):
            st.session_state.files["샘플_경제성"] = make_sample("경제성")
        if c3.button("📋 설문", width="stretch"):
            st.session_state.files["샘플_설문"] = make_sample("설문")

    elif _input_mode == "📷 이미지/사진":
        st.caption("엑셀 화면 캡처·조사표 사진을 AI가 표 데이터로 바꿉니다. 분석 전 반드시 값을 확인하세요.")
        _img = st.file_uploader("표 이미지 업로드", type=["png", "jpg", "jpeg", "webp"], key="table_img_up")
        _cam = st.camera_input("또는 카메라로 촬영", key="table_cam")
        _src = _cam or _img
        if _src is not None:
            st.image(_src, caption="인식할 이미지", width="stretch")
            if st.button("✨ AI로 표 인식", width="stretch", key="img_parse_btn"):
                with st.spinner("표의 행·열과 숫자를 읽는 중..."):
                    _bytes = _src.getvalue()
                    _mime = getattr(_src, "type", None) or ("image/png" if str(getattr(_src, "name", "")).lower().endswith("png") else "image/jpeg")
                    _idf, _warn = image_to_dataframe(_bytes, _mime)
                if _idf is None:
                    st.error("표 인식에 실패했습니다.")
                    for _w in _warn[:3]: st.caption(str(_w)[:180])
                else:
                    st.session_state["image_table_preview"] = _idf
                    st.session_state["image_table_warn"] = _warn
        if st.session_state.get("image_table_preview") is not None:
            st.markdown("**✅ 인식 결과 확인/수정**")
            _edited_img = st.data_editor(st.session_state["image_table_preview"], num_rows="dynamic",
                                         width="stretch", key="image_table_editor", height=240)
            for _w in st.session_state.get("image_table_warn", [])[:3]:
                st.warning(f"확인 필요: {_w}")
            if st.button("📌 이 표를 분석 데이터로 사용", type="primary", width="stretch", key="use_image_table"):
                _key = "이미지_인식데이터"
                st.session_state.files[_key] = clean_columns(_edited_img.copy())
                st.session_state.cur_key = _key
                st.session_state.df = st.session_state.files[_key].copy()
                st.success("이미지에서 읽은 표를 분석 데이터로 적용했습니다.")
                st.rerun()
        if not st.session_state.get("api_key"):
            st.info("이미지 표 인식은 아래 '🤖 AI 기능 켜기'에서 API 키를 설정한 뒤 사용할 수 있습니다.")

    else:  # 음성
        st.caption("예: '처리구 A, 반복 1, 초장 72.3, 수량 615.4'처럼 한 행씩 말해 주세요.")
        _aud = st.audio_input("🎙️ 한 행 말하기", sample_rate=16000, key="voice_data_audio")
        if _aud is not None:
            st.audio(_aud)
            if st.button("📝 음성을 데이터 한 행으로 변환", width="stretch", key="voice_parse_btn"):
                with st.spinner("음성을 듣고 숫자와 변수명을 정리하는 중..."):
                    _tr = ai_multimodal_text(_aud.getvalue(), getattr(_aud, "type", None) or "audio/wav",
                                             "한국어 음성을 정확히 전사하세요.", kind="audio")
                if str(_tr).startswith("⚠️"):
                    st.error(_tr)
                else:
                    st.session_state["voice_transcript"] = _tr
                    _cols = list(st.session_state.df.columns) if isinstance(st.session_state.get("df"), pd.DataFrame) else []
                    _row, _warn = voice_text_to_row(_tr, _cols)
                    if _row is not None:
                        st.session_state.setdefault("voice_rows", [])
                        st.session_state["voice_rows"].append(_row)
                        st.session_state["voice_warn"] = _warn
        if st.session_state.get("voice_transcript"):
            st.caption("인식 문장: " + str(st.session_state["voice_transcript"]))
        if st.session_state.get("voice_rows"):
            _vdf = pd.DataFrame(st.session_state["voice_rows"])
            _ved = st.data_editor(_vdf, num_rows="dynamic", width="stretch", key="voice_rows_editor", height=220)
            for _w in st.session_state.get("voice_warn", [])[:3]: st.warning(f"확인 필요: {_w}")
            cva, cvb = st.columns(2)
            if cva.button("➕ 현재 데이터에 추가", width="stretch", key="voice_append"):
                if isinstance(st.session_state.get("df"), pd.DataFrame) and len(st.session_state.df.columns):
                    base = st.session_state.df.copy()
                    add = _ved.reindex(columns=base.columns)
                    st.session_state.df = pd.concat([base, add], ignore_index=True)
                    ck = st.session_state.get("cur_key") or "음성_추가데이터"
                    st.session_state.files[ck] = st.session_state.df.copy()
                    st.success(f"{len(add)}행을 현재 데이터에 추가했습니다.")
                else:
                    st.warning("먼저 기존 데이터를 불러오거나 '새 데이터로 사용'을 눌러 주세요.")
            if cvb.button("📌 새 데이터로 사용", width="stretch", key="voice_new"):
                key = "음성_입력데이터"
                st.session_state.files[key] = clean_columns(_ved.copy())
                st.session_state.cur_key = key
                st.session_state.df = st.session_state.files[key].copy()
                st.success("음성 입력 데이터를 새 분석 데이터로 적용했습니다.")
                st.rerun()
            if st.button("🗑️ 음성 입력 목록 비우기", width="stretch", key="voice_clear"):
                st.session_state["voice_rows"] = []
                st.session_state.pop("voice_transcript", None)
                st.rerun()
        if not st.session_state.get("api_key"):
            st.info("음성 인식은 ChatGPT 또는 Gemini API 키를 설정한 뒤 사용할 수 있습니다.")

# 데이터 선택 + 삭제
if st.session_state.files:
    names = list(st.session_state.files.keys())
    opts = names + (["🔗 모두 세로로 합치기"] if len(names) > 1 else [])
    _cur_for_select = st.session_state.get("cur_key")
    _sel_idx = opts.index(_cur_for_select) if _cur_for_select in opts else 0
    choice = st.sidebar.selectbox("📌 분석할 데이터 선택", opts, index=_sel_idx)
    # 선택이 바뀔 때만 새로 불러옴 (전처리 결과가 유지되도록)
    if choice != st.session_state.get("cur_key"):
        st.session_state.cur_key = choice
        if choice == "🔗 모두 세로로 합치기":
            try:
                st.session_state.df = pd.concat(list(st.session_state.files.values()), ignore_index=True)
            except Exception as e:
                st.sidebar.error(f"합치기 실패: {e}")
        else:
            st.session_state.df = st.session_state.files[choice].copy()

    if st.sidebar.button("↩️ 원본 데이터로 되돌리기", width="stretch",
                         help="전처리·파생변수 작업을 모두 취소하고 처음 불러온 상태로 복원합니다."):
        k = st.session_state.get("cur_key")
        if k in st.session_state.files:
            st.session_state.df = st.session_state.files[k].copy()
            st.sidebar.success("원본으로 되돌렸습니다.")
            st.rerun()

    with st.sidebar.expander("🗑️ 데이터 삭제"):
        dels = st.multiselect("삭제할 데이터 선택", names, key="del_sel")
        if dels and st.button("선택한 데이터 삭제", width="stretch"):
            for d_ in dels: st.session_state.files.pop(d_, None)
            st.session_state.df = None
            st.rerun()

# 서버에 한글 폰트가 없으면 그래프 글자가 전부 □로 나온다 — 미리 알려 준다.
if _KOREAN_FONT is None:
    st.sidebar.warning("⚠️ 한글 폰트를 찾지 못해 그래프 글자가 □로 나옵니다. "
                       "서버라면 `packages.txt`에 `fonts-nanum`을 넣고 다시 배포하세요.")

# 메뉴를 옮겨 다니거나 시트를 바꿔도 화면 선택·분석 결과가 그대로 남게 한다.
_pin_sync()
if st.session_state.files and st.sidebar.button(
        "🔄 이 데이터의 화면 선택 초기화", width="stretch",
        help="열 선택이 꼬였을 때, 지금 선택한 데이터의 화면 선택만 처음 상태로 되돌립니다."):
    for _k in _pin_scoped_keys():
        try:
            del st.session_state[_k]
        except Exception:
            pass
    st.session_state.get("_pin_store", {}).pop(st.session_state.get("cur_key"), None)
    st.rerun()

# 메뉴 — 주요 분석 흐름을 크게, AI/설명서는 보조 기능으로 작게 분리한다.
_MAIN_MENU_OPTIONS = [
    "⚡ 원클릭 보고서",
    "📊 통계분석",
    "💰 경제성분석",
    "📋 설문조사 분석",
    "📑 보고서",
]
_SUPPORT_MENU_OPTIONS = ["🧠 AI 도우미", "📖 사용설명서"]
if _is_admin_user():
    _SUPPORT_MENU_OPTIONS.append("👑 관리자")

# 현재 선택은 두 라디오 사이에서 하나만 유지한다.
_all_menu_options = _MAIN_MENU_OPTIONS + _SUPPORT_MENU_OPTIONS
if st.session_state.get("menu_choice") not in _all_menu_options:
    st.session_state["menu_choice"] = _MAIN_MENU_OPTIONS[0]

def _menu_from_main():
    value = st.session_state.get("menu_main")
    if value:
        st.session_state["menu_choice"] = value

def _menu_from_support():
    value = st.session_state.get("menu_support")
    if value:
        st.session_state["menu_choice"] = value

# key가 있는 container에는 st-key-* 클래스가 붙으므로 메뉴 영역만 안전하게 스타일링한다.
st.sidebar.markdown("""
<style>
[data-testid="stSidebar"] .st-key-main_menu_block [data-testid="stRadio"] div[role="radiogroup"] {
    gap: 0.22rem;
}
[data-testid="stSidebar"] .st-key-main_menu_block [data-testid="stRadio"] label {
    padding: 0.22rem 0.34rem;
    border-radius: 8px;
}
[data-testid="stSidebar"] .st-key-main_menu_block [data-testid="stRadio"] label p {
    font-size: 1.04rem !important;
    font-weight: 750 !important;
    line-height: 1.45 !important;
}
[data-testid="stSidebar"] .st-key-main_menu_block [data-testid="stRadio"] label:has(input:checked) {
    background: #EAF3FA;
}
[data-testid="stSidebar"] .st-key-support_menu_block [data-testid="stRadio"] div[role="radiogroup"] {
    gap: 0.06rem;
}
[data-testid="stSidebar"] .st-key-support_menu_block [data-testid="stRadio"] label {
    padding: 0.08rem 0.28rem;
    border-radius: 7px;
}
[data-testid="stSidebar"] .st-key-support_menu_block [data-testid="stRadio"] label p {
    font-size: 0.84rem !important;
    font-weight: 500 !important;
    color: #5F7285 !important;
    line-height: 1.3 !important;
}
[data-testid="stSidebar"] .st-key-support_menu_block [data-testid="stRadio"] label:has(input:checked) {
    background: #F3F7FA;
}
[data-testid="stSidebar"] .menu-support-title {
    margin: 0.75rem 0 0.10rem 0.15rem;
    font-size: 0.72rem;
    font-weight: 700;
    color: #8A9AAA;
    letter-spacing: 0.02em;
}
</style>
""", unsafe_allow_html=True)

st.sidebar.markdown("### 📁 주요 기능")
_current_menu = st.session_state.get("menu_choice")
with st.sidebar.container(key="main_menu_block"):
    _main_idx = (_MAIN_MENU_OPTIONS.index(_current_menu)
                 if _current_menu in _MAIN_MENU_OPTIONS else None)
    st.radio(
        "주요 기능", _MAIN_MENU_OPTIONS, index=_main_idx, key="menu_main",
        label_visibility="collapsed", on_change=_menu_from_main,
    )

st.sidebar.markdown('<div class="menu-support-title">보조 기능</div>', unsafe_allow_html=True)
_current_menu = st.session_state.get("menu_choice")
with st.sidebar.container(key="support_menu_block"):
    _support_idx = (_SUPPORT_MENU_OPTIONS.index(_current_menu)
                    if _current_menu in _SUPPORT_MENU_OPTIONS else None)
    st.radio(
        "보조 기능", _SUPPORT_MENU_OPTIONS, index=_support_idx, key="menu_support",
        label_visibility="collapsed", on_change=_menu_from_support,
    )

menu = st.session_state.get("menu_choice", _MAIN_MENU_OPTIONS[0])

with st.sidebar.expander("⚙️ 한글 표 서식 설정"):
    _HWP_FONTS = ["휴먼명조", "함초롬바탕", "함초롬돋움", "바탕", "신명조",
                  "맑은 고딕", "나눔명조", "나눔고딕", "Noto Sans KR", "돋움", "굴림", "직접 입력…"]
    st.selectbox("표·보고서 글씨체", _HWP_FONTS, key="hwp_font",
                 help="한글(hwpx)로 내려받는 표와 보고서 본문에 적용됩니다.")
    if st.session_state.get("hwp_font") == "직접 입력…":
        st.text_input("사용할 글꼴 이름", key="hwp_font_custom",
                      placeholder="예) KoPub바탕체 Medium",
                      help="한글의 글꼴 목록에 표시되는 이름을 그대로 입력하세요.")
    st.caption(f"현재 적용 글꼴: **{_selected_hwp_font()}**")
    c1, c2 = st.columns(2)
    with c1:
        st.selectbox("글자 크기(pt)", [8, 9, 10, 11, 12], index=2, key="hwp_size")
        st.color_picker("머리행 음영", "#D9D9D9", key="hwp_shade")
    with c2:
        st.selectbox("선 굵기", ["0.1 mm", "0.12 mm", "0.15 mm", "0.2 mm", "0.4 mm"], key="hwp_lw")
        st.color_picker("표 선 색", "#000000", key="hwp_line")
    st.checkbox("좌우 바깥 세로선 표시", value=False, key="hwp_sides",
                help="끄면 논문에서 흔히 쓰는 형태(양쪽 세로선 없음)가 됩니다.")
    st.slider("행 높이(mm)", 4.0, 15.0, 6.5, 0.5, key="hwp_rowh",
              help="값을 줄이면 표의 위아래 간격이 촘촘해집니다.")
    c3, c4 = st.columns(2)
    c3.slider("위첨자 크기(%)", 40, 90, 65, 5, key="sup_size")
    c4.slider("위첨자 올림(%)", 0, 70, 35, 5, key="sup_off",
              help="값이 클수록 유의성 문자(a,b,c)가 더 위로 올라갑니다.")
    st.slider("줄글 표 자간(%)", -30, 0, -14, 1, key="hwp_tight",
              help="부분예산표처럼 글이 긴 표에서 글자를 좁혀 한 줄에 담습니다. "
                   "0으로 두면 좁히지 않습니다.")
    st.color_picker("AI 해석 글자색", "#0000FF", key="hwp_aicolor",
                    help="AI가 만든 문장을 사람이 쓴 문장과 구분하기 위한 색입니다. "
                         "검정으로 바꾸면 구분 없이 나옵니다.")
    st.caption("한글에 설치된 글꼴이어야 정확히 표시됩니다.")

with st.sidebar.expander("📈 분석·그래프 설정"):
    st.radio("오차막대 기준", ["표준편차(SD)", "표준오차(SE)"], key="err_type",
             help="SD는 '개체들이 얼마나 흩어져 있나', SE는 '평균값이 얼마나 믿을 만한가'를 봅니다.")
    with st.expander("❓ 표준편차(SD)와 표준오차(SE), 뭐가 다른가요?"):
        st.markdown(EXPLAIN["sd_se"])
    st.selectbox("소수점 자릿수", [1, 2, 3, 4], index=2, key="round_n")
    st.selectbox("그래프 색상", ["파랑", "초록", "주황", "보라", "회색"], key="plot_color")
    c1, c2 = st.columns(2)
    c1.number_input("그래프 가로", 3.0, 16.0, 6.0, 0.5, key="fig_w")
    c2.number_input("그래프 세로", 2.0, 12.0, 4.0, 0.5, key="fig_h")
    st.checkbox("✨ 깔끔한 스타일 (그라데이션·값 표시)", value=True, key="fig_style",
                help="막대에 옅은→진한 색을 입히고 값을 표시하며, 전체 그래프의 축·간격을 통일합니다.")
    st.checkbox("⬛ 막대·원형 조각 검은 테두리", value=True, key="fig_border",
                help="그래프 전체 외곽선이 아니라 막대와 원형/도넛 조각의 경계선에만 검은색을 적용합니다.")
    st.checkbox("막대 위에 값 표시", value=True, key="fig_vlabel")
    st.checkbox("격자선 표시", value=False, key="fig_grid",
                help="'깔끔한 스타일'을 켜면 가로 격자선은 자동으로 들어갑니다.")
    st.checkbox("그래프 제목 표시", value=True, key="fig_title")

_PALETTE = {"파랑": "#6c8ebf", "초록": "#82b366", "주황": "#d79b00", "보라": "#9673a6", "회색": "#808080"}

# 옅은 색 → 진한 색 그라데이션. 값이 큰 막대일수록 진하게 칠해 한눈에 들어오게 한다.
_RAMP = {
    "파랑": ["#dce9f5", "#c2d9ee", "#a3c4e2", "#82acd3", "#6291c2", "#4576ab", "#2d5a8e", "#1f4569"],
    "초록": ["#e2eedd", "#cbe2c2", "#b0d2a3", "#93c083", "#76ad65", "#5c934c", "#457539", "#33582b"],
    "주황": ["#fdeadb", "#fbd7b9", "#f7bd8d", "#f2a061", "#e5833c", "#cd6a25", "#a95318", "#833f11"],
    "보라": ["#eae3f1", "#d9cce6", "#c3b0d7", "#ac93c7", "#9478b3", "#7a5e9a", "#5f487a", "#46345a"],
    "회색": ["#ececec", "#dadada", "#c2c2c2", "#a8a8a8", "#8d8d8d", "#727272", "#585858", "#3f3f3f"],
}


def pretty_on():
    return bool(st.session_state.get("fig_style", True))


def bar_colors(values=None, n=None):
    """막대 색을 옅은→진한 순으로 만든다.

    values 를 주면 **값이 큰 막대일수록 진하게** 칠한다(가장 좋은 처리가 눈에 띈다).
    '깔끔한 스타일'을 끄면 예전처럼 한 가지 색으로 돌아간다.
    """
    ramp = _RAMP.get(st.session_state.get("plot_color", "파랑"), _RAMP["파랑"])
    k = int(n if n is not None else (len(values) if values is not None else 1))
    if not pretty_on() or k <= 0:
        return [pcolor()] * max(k, 1)
    if k == 1:
        return [ramp[len(ramp) // 2 + 1]]
    lo, hi = 1, len(ramp) - 1                      # 너무 옅은 색은 빼서 인쇄해도 보이게
    picked = [ramp[round(lo + (hi - lo) * i / (k - 1))] for i in range(k)]
    if values is not None:
        try:
            import numpy as _np
            order = _np.argsort(_np.argsort(_np.asarray(values, dtype=float)))
            return [picked[int(r)] for r in order]   # 값이 클수록 진한 색
        except Exception:
            pass
    return picked


def bar_value_labels(ax, xs, values, errs=None, dec=None, offset=0.03):
    """막대 위에 값을 적는다(오차막대가 있으면 그 위로 올린다)."""
    if not (pretty_on() and st.session_state.get("fig_vlabel", True)):
        return 0.0
    import numpy as _np
    vals = _np.asarray(values, dtype=float)
    e = _np.zeros_like(vals) if errs is None else _np.nan_to_num(_np.asarray(errs, dtype=float))
    span = float(_np.nanmax(_np.abs(vals))) or 1.0
    d = rnd() if dec is None else dec
    finite = vals[_np.isfinite(vals)]
    if len(finite) and _np.allclose(finite, _np.round(finite)):
        d = 0                       # 도수(개수)처럼 정수뿐이면 '5.0' 대신 '5'
    elif span >= 100:
        d = min(d, 1)               # 큰 값은 소수점을 줄여야 글자가 겹치지 않는다
    for xi, v, ei in zip(xs, vals, e):
        if not _np.isfinite(v):
            continue
        ax.text(xi, v + ei + span * offset, f"{v:,.{d}f}", ha="center", va="bottom",
                fontsize=9, fontweight="bold", color="#33383d")
    return span * offset * 2.2                      # 글자가 차지한 높이(윗여백 확보용)



def bar_value_sig_labels(ax, xs, values, errs=None, sigs=None, dec=None, offset=0.03):
    """막대 위에 '평균값 + 유의성 문자'를 한 번에 표시한다.

    예: 105.7ᵃ
    값과 a/b/c를 별도의 ax.text로 두 번 찍으면 겹치기 쉬우므로 절대 분리하지 않는다.
    """
    if not (pretty_on() and st.session_state.get("fig_vlabel", True)):
        return 0.0
    import numpy as _np
    vals = _np.asarray(values, dtype=float)
    e = _np.zeros_like(vals) if errs is None else _np.nan_to_num(_np.asarray(errs, dtype=float))
    sigs = list(sigs or [""] * len(vals))
    span = float(_np.nanmax(_np.abs(vals))) or 1.0
    d = rnd() if dec is None else dec
    finite = vals[_np.isfinite(vals)]
    if len(finite) and _np.allclose(finite, _np.round(finite)):
        d = 0
    elif span >= 100:
        d = min(d, 1)
    ypad = span * offset
    for i, (xi, v, ei) in enumerate(zip(xs, vals, e)):
        if not _np.isfinite(v):
            continue
        sig = str(sigs[i] if i < len(sigs) else "").strip()
        # 배포 서버의 한글 폰트에는 Unicode 위첨자(ᵃ, ᵇ...)가 없는 경우가 있어 □로 깨진다.
        # Matplotlib mathtext의 영문 superscript를 사용하면 서버 폰트와 무관하게 안정적으로 보인다.
        _sig = "".join(ch for ch in sig if ch.isalpha() or ch == "*")
        label = f"{v:,.{d}f}" + (rf"$^{{{_sig}}}$" if _sig else "")
        ax.text(xi, v + ei + ypad, label,
                ha="center", va="bottom", fontsize=9,
                fontweight="bold", color="#33383d")
    return ypad * 2.6

def pcolor(): return _PALETTE.get(st.session_state.get("plot_color", "파랑"), "#6c8ebf")
def rnd(): return int(st.session_state.get("round_n", 3))
def figsize(w=None, h=None):
    return (w or float(st.session_state.get("fig_w", 6.0)),
            h or float(st.session_state.get("fig_h", 4.0)))
def deco(ax, title="", ylabel_top=True):
    """공통 그래프 마감: 전체 검은 프레임 없이 깔끔한 축/격자/제목만 적용."""
    ax.set_facecolor("white")
    ax.set_axisbelow(True)
    if pretty_on():
        ax.grid(axis="y", color="#DCE7F0", linewidth=.8, alpha=.95)
        ax.grid(axis="x", visible=False)
    elif st.session_state.get("fig_grid", False):
        ax.grid(alpha=.3, linestyle="--")

    # 전체 테두리는 사용하지 않는다. 논문형 좌·하단 축선만 옅게 유지한다.
    for side in ("top", "right"):
        ax.spines[side].set_visible(False)
    for side in ("left", "bottom"):
        ax.spines[side].set_visible(True)
        ax.spines[side].set_color("#AEBECD")
        ax.spines[side].set_linewidth(.8)

    ax.tick_params(colors="#4B5F73", length=3, labelsize=9, direction="out")
    ax.xaxis.label.set_color("#31485E")
    ax.yaxis.label.set_color("#31485E")
    _has_ylab = bool(ylabel_top and ax.get_ylabel())
    if pretty_on() and _has_ylab:
        ax.set_ylabel(ax.get_ylabel(), rotation=0, ha="left", va="bottom",
                      fontsize=9.5, color="#31485E")
        ax.yaxis.set_label_coords(-0.02, 1.025)
    if title and st.session_state.get("fig_title", True):
        ax.set_title(title, fontsize=11.5 if pretty_on() else None,
                     fontweight="bold" if pretty_on() else None,
                     color="#23394D" if pretty_on() else None,
                     pad=24 if (pretty_on() and _has_ylab) else 12)
    lg = ax.get_legend()
    if lg is not None:
        try:
            lg.get_frame().set_linewidth(0)
            lg.get_frame().set_facecolor("white")
        except Exception:
            pass
    return ax


# 위첨자 변환기는 sup_text/sup_display(파일 위쪽) 하나로 통일한다.
# (예전 _SUP 표는 a~h 까지만 있어서 처리구가 9개 이상이면 'i', 'j'가 '^i'로 그대로 보였다.)
sup_show = sup_text
sup_df = sup_display

def read_uploaded_text(f, limit=12000):
    """업로드 파일에서 텍스트 추출 (txt/csv/xlsx/hwpx/pdf/docx)"""
    name = f.name.lower()
    try:
        if name.endswith((".txt", ".md")):
            return f.getvalue().decode("utf-8", errors="ignore")[:limit]
        if name.endswith(".csv"):
            return pd.read_csv(f).to_string()[:limit]
        if name.endswith((".xlsx", ".xls")):
            return pd.read_excel(f).to_string()[:limit]
        if name.endswith(".hwpx"):
            import tempfile, os
            from hwpx import HwpxDocument
            t = tempfile.NamedTemporaryFile(delete=False, suffix=".hwpx"); t.write(f.getvalue()); t.close()
            d = HwpxDocument.open(t.name)
            txt = "\n".join(p.text for p in d.paragraphs if p.text)
            os.unlink(t.name)
            return txt[:limit]
        if name.endswith(".pdf"):
            try:
                from pypdf import PdfReader
            except Exception:
                return "⚠️ PDF를 읽으려면 pypdf가 필요합니다. (pip install pypdf)"
            r = PdfReader(io.BytesIO(f.getvalue()))
            return "\n".join((pg.extract_text() or "") for pg in r.pages)[:limit]
        if name.endswith(".docx"):
            try:
                import docx
            except Exception:
                return "⚠️ Word를 읽으려면 python-docx가 필요합니다. (pip install python-docx)"
            d = docx.Document(io.BytesIO(f.getvalue()))
            return "\n".join(p.text for p in d.paragraphs)[:limit]
        return "⚠️ 지원하지 않는 형식입니다. (txt, csv, xlsx, hwpx, pdf, docx)"
    except Exception as e:
        return f"⚠️ 파일을 읽는 중 오류: {e}"

# 리커트(동의/만족 정도)가 아니라 명목형 범주를 숫자 코드로 적은 열임을 강하게 시사하는 이름들.
# 값만 봐서는 "1,2,3"이 3점 리커트인지 '있다/없다/모르겠다' 같은 코드인지 구분할 수 없으므로
# 열 이름 힌트로 우선 걸러낸다 (economic_core의 반복/처리 열 판별과 같은 접근).
# '의향'은 넣지 않는다 — '재사용의향'·'추천의향'처럼 실제로는 거의 항상 리커트 문항이라,
# 이 힌트에 넣으면 리커트 문항이 명목 코드로 오분류된다(실제로 발생했던 문제).
_NOMINAL_CODE_HINTS = ["여부", "유무", "선택", "성별", "지역", "종류",
                       "품종", "구분", "방법", "코드", "처리구", "그룹"]


def _looks_like_nominal_code(colname):
    return any(k in str(colname) for k in _NOMINAL_CODE_HINTS)


def detect_question_types(df):
    """설문 문항의 유형을 자동으로 추정"""
    out = []
    n = len(df)
    for c in df.columns:
        ser = df[c].dropna()
        if ser.empty:
            out.append({"열 이름": c, "추정 유형": "빈 열", "근거": "-"}); continue
        nu = ser.nunique()
        if pd.api.types.is_numeric_dtype(ser):
            v = ser.astype(float)
            is_int = np.allclose(v, np.round(v))
            nominal_hint = _looks_like_nominal_code(c)
            if is_int and nu == 2:      # 이분형을 먼저 판정 (리커트로 오인 방지)
                out.append({"열 이름": c, "추정 유형": "이분형(예/아니오)",
                            "근거": f"값 2종({int(v.min())}, {int(v.max())})"})
            elif is_int and nominal_hint and nu <= 10:
                # 값 범위만 보면 리커트(3~10점)와 똑같아 보이지만, 열 이름이 '의향/여부/지역' 등
                # 명목 코드를 강하게 시사하므로 객관식(단일선택)으로 분류한다.
                out.append({"열 이름": c, "추정 유형": "객관식(단일선택)",
                            "근거": f"열 이름상 명목 코드로 추정, 보기 {nu}개"})
            elif is_int and 3 <= nu <= 10 and v.min() >= 0 and v.max() <= 10:
                out.append({"열 이름": c, "추정 유형": "리커트 척도",
                            "근거": f"{int(v.min())}~{int(v.max())}점 정수, 보기 {nu}개"})
            else:
                out.append({"열 이름": c, "추정 유형": "연속형 수치", "근거": f"평균 {v.mean():.1f}"})
            continue
        t = ser.astype(str).str.strip()
        avg_len = t.str.len().mean()
        if nu >= n * 0.9 and avg_len < 15:
            out.append({"열 이름": c, "추정 유형": "응답자 ID", "근거": "거의 모두 고유값"}); continue
        sep_found = None
        for sep in [";", ",", "/", "|"]:
            if t.str.contains(sep, regex=False).mean() > 0.3:
                sep_found = sep; break
        if sep_found:
            out.append({"열 이름": c, "추정 유형": "다중응답", "근거": f"구분기호 '{sep_found}'"})
        elif avg_len >= 15 or nu > n * 0.5:
            out.append({"열 이름": c, "추정 유형": "주관식(서술형)", "근거": f"평균 {avg_len:.0f}자"})
        else:
            out.append({"열 이름": c, "추정 유형": "객관식(단일선택)", "근거": f"보기 {nu}개"})
    return pd.DataFrame(out)

def guess_idx(cols, keys, default=0):
    """열 이름에 키워드가 있으면 그 위치를 기본 선택값으로"""
    for i, c in enumerate(cols):
        if any(k in str(c) for k in keys): return i
    return min(default, max(len(cols)-1, 0))


def reorder_by_rank(items, key, label="↕️ 표시 순서 바꾸기"):
    """멀티셀렉트는 고른 순서가 아니라 원래 열 순서대로 결과를 돌려주므로,
    표·그래프에 나오는 순서를 바꾸고 싶으면 여기서 순서 번호를 직접 매긴다."""
    if len(items) <= 1:
        return items
    with st.expander(label):
        rank_df = pd.DataFrame({"항목": items, "순서": range(1, len(items) + 1)})
        edited = st.data_editor(rank_df, key=f"rank_{key}_{hash(tuple(items))}", hide_index=True,
                                width="stretch", disabled=["항목"],
                                column_config={"순서": st.column_config.NumberColumn(min_value=1, step=1)})
    return edited.sort_values("순서", kind="stable")["항목"].tolist()

# 설문 그래프도 원클릭 보고서와 같은 블루 계열로 통일한다.
# 항목 수가 많아도 무지개색을 쓰지 않고 밝기 차이로만 구분해 전체 앱 분위기를 유지한다.
_SURVEY_COLORS = ["#DCE9F5", "#C2D9EE", "#A3C4E2", "#82ACD3", "#6291C2",
                  "#4576AB", "#2D5A8E", "#1F4569", "#7EA6C9", "#B5CEE3"]


def _survey_palette(k):
    if k <= 0:
        return []
    if k <= len(_SURVEY_COLORS):
        # 너무 옅은 색부터 시작하면 흰 배경에서 흐려 보이므로 중간 톤부터 순환
        base = _SURVEY_COLORS[2:] + _SURVEY_COLORS[:2]
        return base[:k]
    cmap = plt.get_cmap("Blues")
    return [cmap(0.35 + 0.55 * (i / max(k - 1, 1))) for i in range(k)]


def _autopct(vals, min_pct=5.0):
    def f(pct):
        n = int(round(pct/100.0*sum(vals)))
        return f"{pct:.1f}%\n({n}명)" if pct >= min_pct else ""
    return f

def pie_chart(counts, title, donut=False):
    """응답자 특성 → 원형/도넛 그래프 (보기 항목이 많으면 자동으로 범례로 전환)"""
    k = len(counts)
    many = k > 6
    fig, ax = plt.subplots(figsize=(figsize()[0] * (1.25 if many else 1.0), figsize()[1]))
    colors = _survey_palette(k)
    _pie_edge = "#111111" if st.session_state.get("fig_border", True) else "white"
    w = dict(width=0.45, edgecolor=_pie_edge, linewidth=0.7) if donut \
        else dict(edgecolor=_pie_edge, linewidth=0.7)
    # 조각이 완전히 붙어 보이지 않도록 아주 조금씩 띄운다.
    # (과한 explode는 보고서용 그래프에서 산만해 보이므로 2% 안팎만 적용)
    _explode = [0.008 if k > 6 else 0.012] * k
    wedges, *_ = ax.pie(counts.values,
                        labels=None if many else counts.index.astype(str),
                        autopct=_autopct(counts.values), startangle=90,
                        counterclock=False, colors=colors, wedgeprops=w,
                        explode=_explode,
                        pctdistance=0.72 if donut else 0.62,
                        textprops={"fontsize": 9, "color": "#2b2b2b"})
    if many:
        ax.legend(wedges, [f"{i} ({v}명)" for i, v in zip(counts.index.astype(str), counts.values)],
                  loc="center left", bbox_to_anchor=(1.0, 0.5), fontsize=8, frameon=False)
    if donut:
        ax.text(0, 0, f"n={int(counts.sum())}", ha="center", va="center",
                fontsize=12, fontweight="bold", color="#444")
    ax.set_title(title, fontsize=11, pad=10, fontweight="bold", color="#333")
    ax.axis("equal")
    plt.tight_layout()
    return fig

def break_even_qty(fixed_cost, variable_cost_total, qty, price):
    """손익분기수량 = 고정비 ÷ (판매단가 − 단위당 변동비).

    ★ 여기서 '변동비'는 **산출량(수량)에 비례하는** 비용만을 말한다. 10a 기준 작물
    예산에서는 종묘비·비료비·농약비·토지용역비·자가노력비처럼 대부분의 비용이
    '면적'에 대해 정해지며, 그해 수량이 줄어도 같이 줄지 않는다. 이런 비용까지
    변동비로 넣으면 단위당 변동비가 실제보다 훨씬 커지고 고정비는 거의 남지 않아
    손익분기수량이 터무니없이 작게 나온다(실제 318kg인데 23kg 같은 값).
    수량에 실제로 비례하는 비용(수확·선별·포장·운송비 등)이 없으면
    variable_cost_total = 0 이 되고, 이때 Q* = 고정비 ÷ 단가 가 된다.

    variable_cost_total은 해당 수량 qty를 생산하는 데 든 수량비례비 합계(단위당이 아님).
    단위당 수량비례비가 판매단가 이상이면(margin<=0) 아무리 팔아도 손익분기가 불가능하므로 NaN.
    """
    vc_unit = variable_cost_total / qty.where(qty != 0) if hasattr(qty, "where") \
        else (variable_cost_total / qty if qty else np.nan)
    margin = price - vc_unit
    if hasattr(margin, "where"):
        return fixed_cost / margin.where(margin > 0)
    return fixed_cost / margin if margin > 0 else np.nan


def crosstab_bar_label(v, pv, ymax):
    """교차분석 누적막대의 칸 안 글자를 정한다.

    칸이 넓으면 '9명\\n(30.0%)' 두 줄, 좁으면 '9명(30.0%)' 한 줄로 줄여서라도
    '명'과 '%'를 항상 남긴다. '명'·'%'를 통째로 빼고 숫자만 남기면 뭘 나타내는지
    알 수 없기 때문이다. 정말 작은 칸(전체의 5% 미만)만 아예 생략한다.

    반환: (문자열, 글자크기) 또는 표시할 수 없으면 None.
    """
    two_line, one_line = ymax * 0.11, ymax * 0.05
    if v >= two_line:
        return f"{int(v)}명\n({pv:.1f}%)", 7
    if v >= one_line:
        return f"{int(v)}명({pv:.1f}%)", 6.5
    return None


def likert_diverging(summ_counts, cats, title):
    """현대적인 블루/레드 다이버징 리커트 차트.

    부정은 부드러운 레드, 중립은 블루그레이, 긍정은 스마트 블루로 표시한다.
    """
    qs = list(summ_counts.keys())
    n_cat = len(cats)
    mid = n_cat // 2
    h = max(4.2, len(qs) * 0.62 + 1.8)
    w = max(8.8, figsize()[0] * 1.35)
    fig, ax = plt.subplots(figsize=(w, h))

    # 낮은 점수(부정) → 옅은~진한 레드 / 높은 점수(긍정) → 옅은~진한 블루
    neg_full = ["#F3D8D8", "#E8AAAA", "#C96767", "#A94D4D"]
    pos_full = ["#D6E7F4", "#9EC5E5", "#6291C2", "#2D5A8E"]
    neg = neg_full[max(0, len(neg_full)-mid):]
    pos = pos_full[max(0, len(pos_full)-mid):]
    neutral = [_SMART_CHART_NEUTRAL] if n_cat % 2 == 1 else []
    palette = neg + neutral + pos
    if len(palette) != n_cat:
        palette = [plt.get_cmap("RdBu")(0.18 + 0.64*i/max(n_cat-1,1)) for i in range(n_cat)]

    data = np.array([summ_counts[q] for q in qs], dtype=float)
    den = data.sum(axis=1, keepdims=True)
    den[den == 0] = 1
    pct = data / den * 100
    base = pct[:, :mid].sum(axis=1) + (pct[:, mid]/2 if n_cat % 2 == 1 else 0)
    starts = -base
    for i, cat in enumerate(cats):
        left_now = starts.copy()
        bars = ax.barh(qs, pct[:, i], left=left_now, color=palette[i], label=str(cat),
                       edgecolor="white", linewidth=.9, height=.58)
        for _b, _v, _l in zip(bars, pct[:, i], left_now):
            if _v >= 8.0:
                _is_dark = ((i < mid and i >= max(mid-1, 0)) or
                            (i >= mid + (1 if n_cat % 2 else 0) + max(len(pos)-2, 0)))
                # 홀수 척도의 중립 구간은 중심이 정확히 x=0이라 기준선과 글자가 겹친다.
                # 중립 라벨만 0선 오른쪽의 구간 안쪽으로 살짝 옮긴다.
                if n_cat % 2 == 1 and i == mid:
                    _half = float(_v) / 2.0
                    _x_text = min(max(3.5, float(_v) * 0.18), max(3.5, _half - 2.0))
                else:
                    _x_text = _l + _v/2
                ax.text(_x_text, _b.get_y()+_b.get_height()/2, f"{_v:.0f}%",
                        ha="center", va="center", fontsize=7.5, zorder=4,
                        color="white" if _is_dark else "#31485E", fontweight="bold")
        starts = starts + pct[:, i]

    ax.axvline(0, color="#71869A", lw=.65, zorder=1)
    ax.set_xlabel("응답 비율(%)")
    ax.invert_yaxis()
    ax.set_xlim(-105, 105)
    deco(ax, title, ylabel_top=False)
    ax.grid(axis="y", visible=False)
    ax.grid(axis="x", color="#E4EDF5", linewidth=.8)
    ax.legend(ncol=min(n_cat, 7), loc="upper center", bbox_to_anchor=(0.5, 1.14),
              fontsize=8, frameon=False, columnspacing=1.2, handlelength=1.4)
    fig.subplots_adjust(top=0.82, bottom=0.12, left=0.20 if len(qs) > 4 else 0.16, right=0.98)
    return fig

def build_stat_method_text(logs, extra=None):
    """분석 이력을 읽어 논문·보고서의 '통계처리' 문단을 자동 생성"""
    acts = " ".join(str(l.get("작업", "")) for l in logs)
    used = []
    if "난괴법" in acts or "블록" in acts or "ANOVA" in acts or "분산분석" in acts:
        used.append("분산분석(ANOVA)")
    if "반복측정" in acts: used.append("반복측정 분산분석")
    if "ANCOVA" in acts or "공분산" in acts: used.append("공분산분석(ANCOVA)")
    if "상관" in acts: used.append("상관분석")
    if "회귀" in acts: used.append("회귀분석")
    if "비모수" in acts or "Kruskal" in acts: used.append("비모수 검정")
    if "프로빗" in acts: used.append("프로빗 분석")
    if "PCA" in acts or "주성분" in acts: used.append("주성분분석")
    if "교차분석" in acts or "카이제곱" in acts: used.append("카이제곱 검정")
    if "크론바흐" in acts or "리커트" in acts or "설문" in acts: used.append("신뢰도 분석")

    ph = None
    for name, label in [("Tukey", "Tukey의 HSD 검정"), ("던컨", "던컨의 다중검정(DMRT)"),
                        ("Duncan", "던컨의 다중검정(DMRT)"), ("Bonferroni", "Bonferroni 보정")]:
        if name in acts: ph = label; break

    s = "모든 자료의 통계분석은 Python의 statsmodels, scipy 라이브러리를 이용하여 수행하였다."
    if used:
        s += " 분석 방법으로는 " + ", ".join(dict.fromkeys(used)) + "을(를) 적용하였다."
    if extra and extra.get("design"):
        s += f" 시험은 {extra['design']}으로 배치하였다."
    if ph:
        s += f" 처리 평균 간 비교는 {ph}(p<0.05)으로 실시하였다."
    if extra and extra.get("cv"):
        s += f" 시험의 변이계수(CV)는 {extra['cv']}%였다."
    s += " 유의수준은 5%로 하였다."
    return s

def build_abstract(items, meta=None):
    """보고서에 담긴 분석 결과를 읽어 '적요(요약)' 초안을 자동 작성"""
    meta = meta or {}
    lines = []
    title = meta.get("title")
    purpose = meta.get("purpose")
    design = meta.get("design")
    if purpose:
        lines.append(f"본 시험은 {purpose}"
                     + ("" if str(purpose).rstrip().endswith(("다.", "다", ".")) else "를 위하여 수행하였다."))
    elif title:
        lines.append(f"본 시험은 '{title}'을(를) 목적으로 수행하였다.")
    if design:
        lines.append(f"시험은 {design}으로 배치하였다.")

    # 담긴 분석에서 유의한 결과 추출
    findings, tables_seen = [], 0
    for it in items:
        blocks = it.get("blocks") or [{"text": it.get("text"), "table": it.get("table")}]
        for b in blocks:
            tb = b.get("table")
            if tb is None or not hasattr(tb, "columns"):
                continue
            cols = [str(c) for c in tb.columns]
            tables_seen += 1
            # 분산분석 요약형(측정 항목/유의성)
            if "측정 항목" in cols and "유의성" in cols:
                sig = tb[tb["유의성"].astype(str).str.contains("유의")]
                for _, r in sig.iterrows():
                    _pv = r.get("p-value")
                    _ptag = ""
                    if pd.notna(_pv):
                        try:
                            _pf = float(_pv)
                            _ptag = " (p<0.001)" if _pf < 0.001 else f" (p={_pf:.4f})"
                        except Exception:
                            _ptag = ""
                    findings.append(f"{r['측정 항목']}은 '{r.get('최고 처리구','')}'에서 "
                                    f"{r.get('최고 평균','')}로 가장 높아" + "||" + _ptag)
            # 평균+유의성형
            elif "유의성" in cols and "평균" in cols:
                try:
                    top = tb.sort_values("평균", ascending=False).iloc[0]
                    gcol = cols[0]
                    if str(top.get("유의성", "")).strip():
                        findings.append(f"{gcol} 중 '{top[gcol]}'의 평균이 {top['평균']}로 가장 높아||")
                except Exception:
                    pass
            # 경제성형
            elif "소득" in cols and "소득률(%)" in cols:
                try:
                    top = tb.sort_values("소득", ascending=False).iloc[0]
                    findings.append(f"'{top[cols[0]]}'의 소득이 "
                                    f"{int(top['소득']):,}원/10a(소득률 {float(top['소득률(%)']):.1f}%)으로 가장 높아||")
                except Exception:
                    pass
            # 증수형
            elif "증수율(%)" in cols and "소득증가액" in cols:
                try:
                    top = tb.sort_values("소득증가액", ascending=False).iloc[0]
                    if float(top["증수율(%)"]) > 0:
                        findings.append(f"'{top[cols[0]]}'은 대조구 대비 {top['증수율(%)']}% 증수되어 "
                                        f"소득증가액 {int(top['소득증가액']):,}원/10a을 나타내||")
                except Exception:
                    pass
    if findings:
        uniq = list(dict.fromkeys(findings))[:4]
        lines.append("주요 결과는 다음과 같다.")
        for u in uniq:
            _body, _sep, _tag = u.partition("||")
            lines.append("  - " + _body.replace("가장 높아", "가장 높았다")
                                       .replace("나타내", "나타내었다") + _tag)
    else:
        lines.append("분석 결과를 보고서에 담으면 주요 결과가 자동으로 요약됩니다.")
    if meta.get("cv"):
        lines.append(f"시험의 변이계수(CV)는 {meta['cv']}%로 시험 정밀도는 양호하였다.")
    lines.append("이상의 결과를 종합할 때, 본 시험에서 얻어진 결과는 "
                 "현장 적용 및 후속 연구의 기초 자료로 활용될 수 있을 것으로 판단된다.")
    return "\n".join(ln if ln.startswith("  -") else f"◦ {ln}" for ln in lines)

def set_df(new_df, memo=""):
    """데이터를 바꾸기 전에 현재 상태를 되돌리기 스택에 저장"""
    hist = st.session_state.setdefault("undo_stack", [])
    cur = st.session_state.get("df")
    if cur is not None:
        hist.append({"df": cur.copy(), "memo": memo})
        if len(hist) > 10:      # 최근 10단계만 유지
            hist.pop(0)
    st.session_state.df = new_df

def undo_df():
    hist = st.session_state.get("undo_stack", [])
    if hist:
        last = hist.pop()
        st.session_state.df = last["df"]
        return last["memo"]
    return None

def keep_running(key, label, **kw):
    """버튼을 누른 뒤 화면이 다시 그려져도 결과를 유지하되, 자료·유효성 변경 시 해제한다."""
    flag, sig_key = f"__ran_{key}", f"__sig_{key}"
    _d = st.session_state.get("df")
    try:
        cur_sig = dataframe_signature(_d) if _d is not None else None
    except Exception:
        cur_sig = (_d.shape, tuple(map(str, _d.columns))) if _d is not None else None
    # 검증 오류로 버튼이 비활성화되면 과거 실행 상태도 반드시 해제한다.
    # 입력이 '잠깐' 유효하지 않은 것일 수도 있으므로(메뉴를 옮겼다 와서 위젯이 초기화된
    # 직후 등) 실행 기록 자체는 지우지 않는다. 이번 화면에서만 결과를 감춘다.
    if kw.get("disabled", False):
        st.button(label, key=f"__btn_{key}", **kw)
        return False
    if st.button(label, key=f"__btn_{key}", **kw):
        st.session_state[flag] = True
        st.session_state[sig_key] = cur_sig
    # 자료가 바뀌면 결과를 감추되 기록은 남긴다 — 원래 자료로 돌아오면 다시 보인다.
    return bool(st.session_state.get(flag)) and st.session_state.get(sig_key) == cur_sig

def fmt_num(df, cols=None, dec=0):
    out = df.copy()
    if cols is None:
        cols = [c for c in out.columns if pd.api.types.is_numeric_dtype(out[c])]
    for c in cols:
        if c in out.columns and pd.api.types.is_numeric_dtype(out[c]):
            out[c] = out[c].map(lambda v: "-" if pd.isna(v) else f"{v:,.{dec}f}")
    return out

def show_money(df, money_cols, dec=0):
    """경제성 금액 표시: 내부 원값은 유지하고 ROUND_HALF_UP으로만 표시한다."""
    out = df.copy()
    for c in [c for c in money_cols if c in out.columns]:
        out[c] = out[c].map(
            lambda v: "-" if pd.isna(v) else f"{round_half_up(v, dec):,.{dec}f}")
    return out

# 공식 조사자료 기반 기준단가 (기준연도 명시 / 매년 갱신 필요)
_PRICE_DEFAULTS = [
    # 항목, 단가, 단위, 기준연도, 출처
    ("농업노임(남)", 153520, "원/일", "2025년",
     "통계청 농가판매·구입가격조사(KOSIS, 국가승인 306001)"),
    ("농업노임(여)", 121392, "원/일", "2025년",
     "통계청 농가판매·구입가격조사(KOSIS, 국가승인 306001)"),
    ("농업노임(남, 시간)", 19190, "원/시간", "2025년",
     "남자 일당 153,520원 ÷ 8시간 (소득조사 환산 기준)"),
    ("농업노임(여, 시간)", 15174, "원/시간", "2025년",
     "여자 일당 121,392원 ÷ 8시간 (소득조사 환산 기준)"),
    ("요소비료(20kg)", 17900, "원/포", "2026년",
     "농협 무기질비료 판매가 (보조금 적용 실구매가 16,250원)"),
    ("무기질비료(톤)", 871000, "원/톤", "2026년",
     "농협 무기질비료 평균 판매가격(전년 825,000원, +5.6%)"),
    ("토지용역비(논)", 275, "원/㎡", "2024년",
     "농지임차료실태조사(농식품부, 국가승인 114062) — 10a=1,000㎡"),
    ("토지용역비(밭)", 260, "원/㎡", "2024년",
     "농지임차료실태조사(농식품부) — 10a 환산 시 260,000원"),
    ("토지용역비(과수원)", 342, "원/㎡", "2024년",
     "농지임차료실태조사(농식품부)"),
    ("자본이자율", 5.0, "%", "관행",
     "농촌진흥청 소득조사 적용 이자율 — 소득자료집 원문 확인 권장"),
    ("농기계 임차료(경운기/일)", 0, "원/일", "-",
     "시군 농기계임대사업소 개별 고시 — 지역별로 다름"),
    ("위탁영농비", 0, "원/10a", "-",
     "농촌진흥청 소득자료집 작목별 경영비 항목 참고"),
]

def default_price_db():
    """⑫ 메타데이터를 모두 포함한 기준단가 표"""
    import datetime as _dt
    today = _dt.date.today().isoformat()
    return pd.DataFrame(
        [{"항목": a, "단가": b, "단위": c, "기준연도": d, "출처": e,
          "조회방식": "기본값", "갱신일": today, "환산식": "", "사용자수정": False}
         for a, b, c, d, e in _PRICE_DEFAULTS])


def price_db_warnings(db, current_year=None):
    """기준연도가 오래됐거나 단가가 비어 있는 항목을 찾아 경고 문구 목록 반환"""
    import datetime as _dt
    import re as _re
    if db is None or getattr(db, "empty", True):
        return []
    cy = current_year or _dt.date.today().year
    old_items, zero_items = [], []
    for _, r in db.iterrows():
        name = str(r.get("항목", "")).strip()
        if not name:
            continue
        try:
            val = float(r.get("단가", 0) or 0)
        except (TypeError, ValueError):
            val = 0.0
        if val <= 0:
            zero_items.append(name)
        yr_txt = str(r.get("기준연도", ""))
        m = _re.search(r"(19|20)\d{2}", yr_txt)
        if m and (cy - int(m.group(0))) >= 2:
            old_items.append(f"{name}({m.group(0)}년)")
    msgs = []
    if old_items:
        msgs.append("현재 분석에는 " + ", ".join(old_items[:5])
                    + (" 등" if len(old_items) > 5 else "")
                    + " 기준 단가가 사용되었습니다. 최신 지역 단가와 차이가 있을 수 있습니다.")
    if zero_items:
        msgs.append("단가가 0이거나 비어 있는 항목: " + ", ".join(zero_items[:6])
                    + (" 등" if len(zero_items) > 6 else "")
                    + " — 값을 채우기 전에는 계산에 자동 반영하지 않습니다.")
    return msgs

def _kamis_parse_unit(unit_text):
    """'20kg' → (20.0, '20kg') 처럼 kg 환산계수를 추출. 환산 불가면 (None, 원문)"""
    import re as _re
    if not unit_text:
        return None, ""
    t = str(unit_text).strip()
    m = _re.match(r"^\s*([\d.]+)?\s*(kg|g|톤|개|포|망|상자|단)\s*$", t, _re.I)
    if not m:
        return None, t
    num = float(m.group(1)) if m.group(1) else 1.0
    unit = m.group(2).lower()
    if unit == "kg":
        return num, t
    if unit == "g":
        return num / 1000.0, t
    if unit == "톤":
        return num * 1000.0, t
    return None, t          # 개·포·망 등은 kg 환산 불가


def _kamis_ssl_context(verify=True, legacy=True):
    """KAMIS(구형 국내 서버) 호환 SSLContext.

    OpenSSL 3.x 는 기본적으로
      - RFC5746 미지원 서버(레거시 재협상)
      - SECLEVEL 2 미만의 약한 키/암호군
      - TLS 1.0/1.1
    을 모두 거부한다. 국내 공공기관 서버는 이 중 하나에 걸리는 경우가 많아
    핸드셰이크 단계에서 SSLError 가 난다. 아래 컨텍스트는 검증(인증서 확인)은
    그대로 유지한 채 호환성 옵션만 풀어 준다.
    """
    import ssl as _ssl
    ctx = _ssl.create_default_context()
    if legacy:
        # OP_LEGACY_SERVER_CONNECT (0x4) : RFC5746 미지원 서버 허용
        ctx.options |= getattr(_ssl, "OP_LEGACY_SERVER_CONNECT", 0x4)
        try:
            import warnings as _w
            with _w.catch_warnings():
                _w.simplefilter("ignore", DeprecationWarning)
                ctx.minimum_version = _ssl.TLSVersion.TLSv1
        except Exception:
            pass
        for _cipher in ("DEFAULT@SECLEVEL=1", "ALL:@SECLEVEL=1"):
            try:
                ctx.set_ciphers(_cipher)
                break
            except Exception:
                continue
    if not verify:
        ctx.check_hostname = False
        ctx.verify_mode = _ssl.CERT_NONE
    return ctx


def _kamis_session(ssl_context=None):
    """지정한 SSLContext 를 http/https 양쪽에 적용한 requests 세션."""
    from requests.adapters import HTTPAdapter

    class _KamisAdapter(HTTPAdapter):
        def __init__(self, ssl_context=None, **kw):
            self._kamis_ctx = ssl_context
            super().__init__(**kw)

        def init_poolmanager(self, *a, **kw):
            if self._kamis_ctx is not None:
                kw["ssl_context"] = self._kamis_ctx
            return super().init_poolmanager(*a, **kw)

        def proxy_manager_for(self, *a, **kw):
            if self._kamis_ctx is not None:
                kw["ssl_context"] = self._kamis_ctx
            return super().proxy_manager_for(*a, **kw)

    sess = _requests.Session()
    adapter = _KamisAdapter(ssl_context=ssl_context, max_retries=0)
    sess.mount("https://", adapter)
    sess.mount("http://", adapter)
    return sess


def kamis_request(url, params, timeout=20):
    """KAMIS 공식 Open-API 호출 (다단계 폴백).

    공식 도메인은 ``www.kamis.or.kr`` 이고 공식 예제는 http/https 를 모두 쓴다.
    KAMIS 서버는 구형 TLS 설정을 쓰는 경우가 있어 OpenSSL 3.x 환경
    (Streamlit Cloud 등)에서 기본 설정만으로는 핸드셰이크가 실패할 수 있다.
    아래 순서대로 시도하고, 성공한 방식을 ``r._smart_transport`` 에 기록한다.

      1) https + 기본 설정                → 'https'
      2) http  (리다이렉트 시 레거시 컨텍스트) → 'http-fallback'
      3) https + 레거시 호환 컨텍스트(검증 유지) → 'https-legacy'
      4) https + 레거시 + 인증서 검증 생략   → 'https-insecure' (경고 표시)

    4단계는 인증서 검증을 끄므로 최후 수단이며, 성공해도 UI 에 경고를 띄운다.
    """
    if not _HAS_REQUESTS:
        raise RuntimeError("requests 라이브러리가 없습니다. pip install requests")

    raw_url = str(url).strip()
    if raw_url.startswith("http://"):
        https_url = "https://" + raw_url[len("http://"):]
    elif raw_url.startswith("https://"):
        https_url = raw_url
    else:
        https_url = "https://" + raw_url.lstrip("/")
    http_url = "http://" + https_url[len("https://"):]

    headers = {
        'User-Agent': ('Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                       'AppleWebKit/537.36 (KHTML, like Gecko) '
                       'Chrome/131.0 Safari/537.36'),
        'Accept': 'application/json, application/xml, text/xml, */*',
        'Connection': 'close',
    }

    attempts = []          # (라벨, 호출가능객체) 순서대로 시도
    attempts.append(("https", lambda: _requests.get(
        https_url, params=params, timeout=timeout, headers=headers)))

    def _legacy_http():
        sess = _kamis_session(_kamis_ssl_context(verify=True, legacy=True))
        try:
            return sess.get(http_url, params=params, timeout=timeout,
                            headers=headers)
        finally:
            try: sess.close()
            except Exception: pass
    attempts.append(("http-fallback", _legacy_http))

    def _legacy_https():
        sess = _kamis_session(_kamis_ssl_context(verify=True, legacy=True))
        try:
            return sess.get(https_url, params=params, timeout=timeout,
                            headers=headers)
        finally:
            try: sess.close()
            except Exception: pass
    attempts.append(("https-legacy", _legacy_https))

    def _insecure_https():
        try:
            import urllib3 as _u3
            _u3.disable_warnings(_u3.exceptions.InsecureRequestWarning)
        except Exception:
            pass
        sess = _kamis_session(_kamis_ssl_context(verify=False, legacy=True))
        try:
            return sess.get(https_url, params=params, timeout=timeout,
                            headers=headers, verify=False)
        finally:
            try: sess.close()
            except Exception: pass
    attempts.append(("https-insecure", _insecure_https))

    errors = []
    for label, call in attempts:
        try:
            r = call()
        except Exception as ex:
            errors.append(f"{label}: {_kamis_err_text(ex)}")
            continue
        try:
            r._smart_transport = label
        except Exception:
            pass
        return r

    raise RuntimeError(
        "KAMIS 서버에 연결하지 못했습니다. 시도한 방식과 원인:\n- "
        + "\n- ".join(errors))


def _kamis_err_text(ex, limit=200):
    """SSLError 등의 실제 원인(가장 안쪽 예외)까지 풀어서 문자열로 만든다."""
    parts, seen, cur = [], set(), ex
    while cur is not None and id(cur) not in seen:
        seen.add(id(cur))
        txt = str(cur).strip()
        if txt and txt not in parts:
            parts.append(txt)
        cur = getattr(cur, "__cause__", None) or getattr(cur, "__context__", None)
    # urllib3 는 원인을 reason 속성에 담아 두기도 한다
    reason = getattr(ex, "reason", None)
    if reason is not None and str(reason).strip() not in parts:
        parts.append(str(reason).strip())
    msg = " / ".join(parts) if parts else type(ex).__name__
    # 가장 유용한 [SSL: XXX] 코드가 있으면 앞으로 끌어올린다
    import re as _re
    m = _re.search(r"\[SSL:[^\]]+\][^\)\'\"]*", msg)
    head = f"{type(ex).__name__}"
    if m:
        return f"{head} {m.group(0).strip()} | {msg[:limit]}"
    return f"{head}: {msg[:limit]}"


KAMIS_ITEM_CATEGORY_MAP = {
    "111": "100", "211": "200", "225": "200", "231": "200", "245": "200",
    "258": "200", "312": "300", "411": "400",
}


def _normalize_kamis_date(year, regday):
    import re as _re
    y = str(year or "").strip()
    d = str(regday or "").strip().replace("/", "-").replace(".", "-")
    d = _re.sub(r"-+", "-", d).strip("-")
    if not d:
        return y
    if len(d.split("-")) == 2 and y:
        mm, dd = d.split("-", 1)
        return f"{y}-{int(mm):02d}-{int(dd):02d}"
    if len(d.split("-")) == 3:
        yy, mm, dd = d.split("-", 2)
        return f"{int(yy):04d}-{int(mm):02d}-{int(dd):02d}"
    return f"{y}-{d}".strip("-")


def kamis_fetch(cert_key, cert_id, item_code, kind_code, rank_code,
                days=7, country_code="", product_cls="02", category_code=None,
                market_type="도매", convert_kg=False, county_code=None):
    """KAMIS 기간별 품목가격 조회. 반환: (DataFrame|None, 오류메시지|None).

    county_code는 이전 버전 호출과의 호환용이며 실제 요청에는 공식 p_countrycode를 사용한다.
    """
    if county_code and not country_code:
        country_code = county_code
    missing = [n for n, v in [("인증키", cert_key), ("아이디", cert_id),
                              ("품목코드", item_code)] if not str(v or "").strip()]
    if missing:
        return None, f"{', '.join(missing)}을(를) 입력해 주세요."
    try:
        days = int(days)
        if not (1 <= days <= 365):
            return None, "조회 기간은 1~365일 사이여야 합니다."
    except (TypeError, ValueError):
        return None, "조회 기간이 올바르지 않습니다."
    item_code = str(item_code).strip()
    category_code = str(category_code or KAMIS_ITEM_CATEGORY_MAP.get(item_code, "")).strip()
    if not category_code:
        return None, "품목에 맞는 부류코드를 지정해 주세요(예: 채소 200, 특용작물 300)."
    if market_type not in ("도매", "소매"):
        return None, "시장 유형은 '도매' 또는 '소매'여야 합니다."

    import datetime as _dt
    end = _dt.date.today()
    start = end - _dt.timedelta(days=days)
    # KAMIS의 일별 품목별 도·소매 가격 API는 periodProductList 하나를 사용하고
    # p_productclscode로 도매(02)/소매(01)를 구분한다.
    action = "periodProductList"
    product_cls = "02" if market_type == "도매" else "01"
    url = "https://www.kamis.or.kr/service/price/xml.do"
    params = {"action": action,
              "p_cert_key": str(cert_key).strip(), "p_cert_id": str(cert_id).strip(),
              "p_returntype": "json", "p_startday": start.isoformat(),
              "p_endday": end.isoformat(), "p_productclscode": str(product_cls),
              "p_itemcategorycode": category_code, "p_itemcode": item_code,
              "p_kindcode": str(kind_code or ""), "p_productrankcode": str(rank_code or ""),
              "p_countrycode": str(country_code or ""),
              "p_convert_kg_yn": "Y" if convert_kg else "N"}
    try:
        r = kamis_request(url, params)
    except RuntimeError as ex:
        return None, str(ex)
    except Exception as ex:
        return None, f"KAMIS 연결 오류 - {_kamis_err_text(ex)}"
    if getattr(r, "status_code", 0) != 200:
        return None, f"KAMIS 서버 응답 오류(HTTP {getattr(r, 'status_code', '?')})"
    # KAMIS는 xml.do 주소에서 p_returntype=json을 써도 환경/오류 종류에 따라
    # XML을 돌려주는 경우가 있어 JSON과 XML을 모두 받을 수 있게 한다.
    js = None
    try:
        js = r.json()
    except Exception:
        try:
            import xml.etree.ElementTree as _ETK
            _root = _ETK.fromstring(getattr(r, "text", "") or "")
            def _xml_dict(el):
                children = list(el)
                if not children:
                    return (el.text or "").strip()
                out = {}
                for ch in children:
                    val = _xml_dict(ch)
                    if ch.tag in out:
                        if not isinstance(out[ch.tag], list): out[ch.tag] = [out[ch.tag]]
                        out[ch.tag].append(val)
                    else:
                        out[ch.tag] = val
                return out
            js = _xml_dict(_root)
        except Exception:
            _ct = str(getattr(r, "headers", {}).get("Content-Type", ""))[:80]
            return None, ("KAMIS 응답 형식을 해석하지 못했습니다. "
                          f"HTTP {getattr(r, 'status_code', '?')}, Content-Type={_ct or '미상'}")

    if isinstance(js, dict):
        condition = js.get("condition") or {}
        code = (condition.get("code") if isinstance(condition, dict) else None)
        code = code or js.get("error_code") or js.get("errCode")
        if code and str(code) not in ("000", "0"):
            msg = condition.get("message") if isinstance(condition, dict) else ""
            return None, f"KAMIS 오류({code}): {msg or '인증 정보 또는 조회 조건을 확인해 주세요.'}"
        data = js.get("data", js)
    else:
        data = js
    if isinstance(data, dict):
        _dcode = (data.get("error_code") or data.get("errCode") or
                  data.get("result_code") or data.get("resultCode"))
        if _dcode and str(_dcode) not in ("000", "0"):
            _dmsg = data.get("error_msg") or data.get("message") or data.get("result_msg") or ""
            return None, f"KAMIS 오류({_dcode}): {_dmsg or '인증 정보 또는 조회 조건을 확인해 주세요.'}"
    items = data.get("item") if isinstance(data, dict) else (data if isinstance(data, list) else None)
    if not items:
        return None, "조회 결과가 없습니다. 품목·품종·등급·지역·기간을 확인해 주세요."
    if isinstance(items, dict):
        items = [items]
    rows = []
    for it in items:
        if not isinstance(it, dict):
            continue
        date_s = _normalize_kamis_date(it.get("yyyy"), it.get("regday"))
        unit_raw = it.get("unit") or it.get("unitname") or ("1kg" if convert_kg else "")
        kg_factor, unit_text = _kamis_parse_unit(unit_raw)
        price_raw = str(it.get("price", "")).replace(",", "").strip()
        try:
            price = float(price_raw) if price_raw not in ("", "-") else float("nan")
        except ValueError:
            price = float("nan")
        rows.append({"기준일": date_s, "시장유형": market_type,
                     "품목": it.get("itemname", ""), "품종": it.get("kindname", ""),
                     "시장/지역": it.get("countyname", it.get("marketname", "")),
                     "가격": price, "단위": unit_text or "(단위 미상)",
                     "kg환산계수": kg_factor, "부류코드": category_code,
                     "품목코드": item_code, "출처": "KAMIS 농산물유통정보"})
    if not rows:
        return None, "조회 결과를 표로 만들지 못했습니다."
    _dfk = pd.DataFrame(rows)
    try:
        _dfk.attrs["kamis_transport"] = getattr(r, "_smart_transport", "https")
    except Exception:
        pass
    return _dfk, None


@st.cache_data(show_spinner=False, ttl=3600, max_entries=30)
def kosis_fetch_url(url):
    """KOSIS 오픈API URL을 그대로 호출해 표로 변환 (KOSIS 사이트에서 복사한 URL 사용)"""
    if not _HAS_REQUESTS:
        return None, "requests 라이브러리가 없습니다."
    if "kosis.kr" not in url:
        return None, "KOSIS 주소가 아닙니다. kosis.kr 로 시작하는 URL을 넣어 주세요."
    try:
        u = url.strip().replace("http://", "https://")
        if "format=" not in u:
            u += ("&" if "?" in u else "?") + "format=json&jsonVD=Y"
        r = _requests.get(u, timeout=30)
        if r.status_code != 200:
            return None, f"KOSIS 오류({r.status_code})"
        js = r.json()
        if isinstance(js, dict) and js.get("err"):
            return None, f"KOSIS 응답 오류: {js.get('errMsg', js.get('err'))}"
        if not isinstance(js, list) or not js:
            return None, "조회 결과가 없습니다. 통계표·시점 설정을 확인해 주세요."
        rows = []
        for it in js:
            rows.append({
                "시점": it.get("PRD_DE", ""),
                "항목": it.get("ITM_NM", ""),
                "분류1": it.get("C1_NM", ""),
                "분류2": it.get("C2_NM", ""),
                "값": it.get("DT", ""),
                "단위": it.get("UNIT_NM", ""),
                "통계표": it.get("TBL_NM", "")})
        return pd.DataFrame(rows), None
    except Exception as ex:
        return None, f"조회 실패: {str(ex)[:80]}"

def kosis_build_url(api_key, org_id, tbl_id, itm_id="ALL", obj_l1="ALL",
                    prd_se="Y", count=5):
    """파라미터로 KOSIS 요청 URL 만들기"""
    return ("https://kosis.kr/openapi/Param/statisticsParameterData.do"
            f"?method=getList&apiKey={api_key}&itmId={itm_id}&objL1={obj_l1}"
            f"&format=json&jsonVD=Y&prdSe={prd_se}&newEstPrdCnt={int(count)}"
            f"&orgId={org_id}&tblId={tbl_id}")

def get_price(item, default=0.0):
    """기준단가 DB에서 값을 읽음(없으면 default)"""
    db = st.session_state.get("price_db")
    if db is None or db.empty: return default
    row = db[db["항목"].astype(str) == item]
    if row.empty: return default
    try:
        v = float(row.iloc[0]["단가"])
        return v if v > 0 else default
    except Exception:
        return default

def run_autopilot_engine(df, ph="Tukey HSD", err_type=None, max_items=8,
                         trt_override=None, blk_override=None):
    """원클릭 오토파일럿: 정제 → 설계인지 → 분석 → 그래프 → 문장 → 보고서까지 한 번에.
    반환: dict(ok, blocks, summary, abstract, design, msgs)"""
    msgs = []
    prog = st.progress(0.0, text="1/5 데이터 점검 중...")

    # ---------- 1단계: 숫자 정제 ----------
    with st.spinner("1/5 데이터를 점검하고 정리하는 중..."):
        work = clean_columns(df.copy())
        # 처리구·반복으로 쓸 열은 숫자 변환에서 제외 (이름이 숫자로 바뀌는 것 방지)
        _keep_text = set()
        _pre = detect_design(work)
        for _k in (trt_override, blk_override, _pre.get("trt"), _pre.get("blk"), _pre.get("sub")):
            if _k: _keep_text.add(_k)
        fixed = []
        for c, ratio in find_numeric_like(work).items():
            if c in _keep_text:
                continue
            conv = to_numeric_clean(work[c])
            if conv.notna().mean() >= 0.6:
                work[c] = conv
                fixed.append(c)
        if fixed:
            msgs.append(f"문자로 읽힌 열을 숫자로 변환했습니다: {', '.join(fixed)}")
        n_before = len(work)
        work = work.dropna(how="all")
        if len(work) < n_before:
            msgs.append(f"완전히 빈 행 {n_before - len(work)}개를 제외했습니다.")
    prog.progress(0.2, text="2/5 실험설계 인지 중...")

    # ---------- 2단계: 설계 자동 인지 ----------
    with st.spinner("2/5 처리구·반복·측정항목을 찾는 중..."):
        dsg = detect_design(work)
        if dsg.get("promoted"):
            msgs.append("숫자로 적혀 있지만 처리구·반복 코드로 보이는 열을 그룹으로 인식했습니다: "
                        + ", ".join(map(str, dsg["promoted"]))
                        + " (실제 측정값이라면 위 ⚙️ 설정에서 처리구·반복 열을 직접 지정하세요)")
        if not dsg.get("trt"):
            prog.empty()
            return {"ok": False, "msgs": msgs + ["처리구로 볼 만한 열을 찾지 못했습니다. "
                                                 "'분산분석' 화면에서 직접 선택해 주세요."]}
        trt = trt_override if trt_override in work.columns else dsg["trt"]
        blk = blk_override if blk_override in work.columns else dsg["blk"]
        if blk == trt: blk = None
        # 처리구·반복·부요인 및 '그룹처럼 보이는' 열은 측정항목에서 제외
        exclude = {trt, blk, dsg.get("sub")}
        orig_cat = set(df.columns) - set(df.select_dtypes(include=np.number).columns)
        ys = []
        for c in dsg["ys"]:
            if c in exclude or c is None:
                continue
            if c in orig_cat:          # 원래 문자였던 열(처리구 등)은 측정값이 아님
                continue
            uniq = work[c].nunique(dropna=True)
            if uniq <= 2 and len(work) > 6:   # 값이 2종류뿐이면 구분용 열일 가능성
                continue
            ys.append(c)
        ys = ys[:max_items]
        if not ys:
            prog.empty()
            return {"ok": False, "msgs": msgs + ["분석할 숫자형 측정항목이 없습니다."]}
        msgs.append(f"실험설계: {dsg['design']} (확신도 {dsg['confidence']}) — "
                    f"처리구 '{trt}'" + (f", 반복 '{blk}'" if blk else ", 반복 없음"))
    prog.progress(0.4, text=f"3/5 {len(ys)}개 항목 분석 중...")

    # ---------- 3단계: 항목별 분산분석 + 사후검정 ----------
    use_se = (err_type or st.session_state.get("err_type", "표준편차(SD)")).startswith("표준오차")
    blocks, summary_rows, sentences = [], [], []
    with st.spinner(f"3/5 {len(ys)}개 항목을 분석하고 그래프를 만드는 중..."):
        for k, yv in enumerate(ys):
            cols_need = [trt, yv] + ([blk] if blk else [])
            data = work[cols_need].dropna()
            ok_d, vmsg = validate_anova_data(data, trt, yv)
            if not ok_d:
                msgs.append(f"[{yv}] 건너뜀 — {vmsg[0] if vmsg else '자료 부족'}")
                continue
            try:
                formula = safe_formula(yv, [trt] + ([blk] if blk else []))
                model = ols(formula, data=data).fit()
                aov = sm.stats.anova_lm(model, typ=2)
                tkey = f"C({q_ref(trt)})"
                pval = float(aov.loc[tkey, "PR(>F)"]) if tkey in aov.index else float(aov["PR(>F)"].iloc[0])
                ci = calc_cv_lsd(model, data, trt, yv)
                _is_dunnett = str(ph).startswith(("던넷", "Dunnett"))
                _ctrl = st.session_state.get("dunnett_ctrl") if _is_dunnett else None
                _phres = posthoc_from_model(model, data, trt, ph, control=_ctrl)
                ns = _phres["not_sig"]
                means = data.groupby(trt)[yv].agg(["mean", "std", "count"])
                order = means.sort_values("mean", ascending=False).index.tolist()
                # 던넷(Dunnett)은 각 처리를 대조구와만 비교하므로, 전체 처리 쌍을 요구하는
                # 유의성 문자(a,b,c)를 만들 수 없다. 지금은 원클릭 사후검정 선택지에
                # 던넷이 없지만, 나중에 추가되더라도 여기서 막아 잘못된 문자가
                # 보고서에 실리지 않게 한다.
                letters = ({} if _is_dunnett
                           else (compact_letter_display(order, ns) if pval < .05 else {}))

                # 논문 표 형식: 평균값에 유의성 문자를 위첨자로 붙임 (예: 607.6^a)
                _dec = rnd()
                _se_v = (means["std"] / np.sqrt(means["count"]))
                res = pd.DataFrame({
                    trt: [str(g) for g in means.index],
                    f"{yv}": [f"{means.loc[g, 'mean']:.{_dec}f}"
                              + (f"^{letters[g]}" if letters.get(g) else "")
                              for g in means.index],
                    "표준편차": [round(float(means.loc[g, "std"]), _dec) for g in means.index],
                    "표준오차": [round(float(_se_v.loc[g]), _dec) for g in means.index],
                    "n": [int(means.loc[g, "count"]) for g in means.index]})
                res = res.set_index(trt).loc[[str(o) for o in order]].reset_index()

                # ---------- 4단계 일부: 그래프 ----------
                err = (means["std"] / np.sqrt(means["count"])) if use_se else means["std"]
                fig, ax = plt.subplots(figsize=figsize())
                mm, ee = means.loc[order], err.loc[order]
                _xs = [str(o) for o in order]
                ax.bar(_xs, mm["mean"], yerr=ee, capsize=4,
                       color=bar_colors(values=mm["mean"].tolist()),
                       edgecolor="none", width=.62 if pretty_on() else .8,
                       error_kw={"ecolor": "#5a6067", "elinewidth": 1.1})
                top = float(mm["mean"].max())
                _pad = bar_value_sig_labels(
                    ax, range(len(_xs)), mm["mean"].tolist(), ee.tolist(),
                    [letters.get(g, "") for g in order], dec=rnd())
                ax.set_ylabel(yv)
                ax.margins(y=.16 if pretty_on() else .05)
                deco(ax, f"{trt}별 {yv} (평균±{'표준오차' if use_se else '표준편차'})")
                plt.tight_layout()
                png = fig_to_png(fig, show=False)

                sent = report_sentence_anova(trt, yv, pval, means, letters, ci, ph)
                sentences.append(sent)
                blocks.append({"text": sent})
                blocks.append({"caption": f"{trt}별 {yv}", "table": res, "image": png})

                summary_rows.append({
                    "측정 항목": yv,
                    "최고 처리구": order[0],
                    "최고 평균": round(float(means.loc[order[0], "mean"]), rnd()),
                    "최저 처리구": order[-1],
                    "p-value": round(pval, 4),
                    "유의성": "유의(*)" if pval < .05 else "n.s.",
                    "CV(%)": round(ci["CV"], 1) if not np.isnan(ci["CV"]) else "-",
                    "LSD(0.05)": round(ci["LSD"], 2) if not np.isnan(ci["LSD"]) else "-"})
            except Exception as ex:
                msgs.append(f"[{yv}] 분석 실패: {str(ex)[:60]}")
            prog.progress(0.4 + 0.3 * (k + 1) / len(ys),
                          text=f"3/5 분석 중... ({k+1}/{len(ys)})")

    if not summary_rows:
        prog.empty()
        return {"ok": False, "msgs": msgs + ["분석 가능한 항목이 없었습니다."]}

    # ---------- 4단계: 요약표 · 통계처리 문구 ----------
    prog.progress(0.75, text="4/5 요약표와 문장을 정리하는 중...")
    with st.spinner("4/5 종합 요약을 만드는 중..."):
        summary = pd.DataFrame(summary_rows)
        n_sig = int((summary["유의성"] == "유의(*)").sum())
        head_txt = (f"◦ {dsg['design']}으로 배치된 시험 자료를 분석하였다.\n"
                    f"◦ 총 {len(summary)}개 측정항목 중 {n_sig}개 항목에서 "
                    "처리 간 유의한 차이가 인정되었다.")
        stat_line = build_stat_method_text(
            st.session_state.get("log", []),
            {"design": dsg["design"] if blk else None})
        head_blocks = [{"text": head_txt},
                       {"caption": "측정항목별 분석 종합", "table": summary}]

    # ---------- 5단계: 적요 + 보고서 파일 ----------
    prog.progress(0.9, text="5/5 보고서를 만드는 중...")
    with st.spinner("5/5 적요와 보고서 문서를 만드는 중..."):
        items_for_abs = [{"heading": "분석 종합", "blocks": head_blocks}]
        abstract = build_abstract(items_for_abs,
                                  {"design": dsg["design"] if blk else None})
        report_items = [
            {"heading": "적요(要約)", "text": abstract, "table": None, "image": None},
            {"heading": "분석 종합", "blocks": head_blocks + blocks},
            {"heading": "통계 처리", "text": stat_line, "table": None, "image": None}]
        try:
            hwpx = build_report_hwpx(report_items, "시험 통계 분석 보고서(초안)")
        except Exception as ex:
            hwpx = None; msgs.append(f"한글 보고서 생성 실패: {str(ex)[:60]}")
        docx_data = None
        if _HAS_DOCX:
            try:
                docx_data = build_report_docx(report_items, "시험 통계 분석 보고서(초안)")
            except Exception as ex:
                msgs.append(f"워드 보고서 생성 실패: {str(ex)[:60]}")
    prog.progress(1.0, text="완료!")
    prog.empty()
    return {"ok": True, "design": dsg, "summary": summary, "blocks": blocks,
            "head_blocks": head_blocks, "abstract": abstract, "stat_line": stat_line,
            "sentences": sentences, "report_items": report_items,
            "hwpx": hwpx, "docx": docx_data, "msgs": msgs, "ys": ys, "trt": trt, "blk": blk}

def two_way_sensitivity(base_qty, base_price, mgmt_cost, byproduct=0,
                        yield_variable_cost=0, q_range=(-20, 20), p_range=(-20, 20), step=10):
    """수량×단가 위험 매트릭스. 수량비례비용은 수량 변화에 함께 연동한다."""
    q_rates = list(range(int(q_range[0]), int(q_range[1]) + 1, int(step)))
    p_rates = list(range(int(p_range[0]), int(p_range[1]) + 1, int(step)))
    base_yvc = max(float(yield_variable_cost or 0), 0.0)
    fixed_like_mgmt = max(float(mgmt_cost) - base_yvc, 0.0)
    mat = np.zeros((len(q_rates), len(p_rates)))
    for i, qr in enumerate(q_rates):
        q_factor = 1 + qr/100
        for j, pr_ in enumerate(p_rates):
            revenue = base_qty * q_factor * base_price * (1 + pr_/100) + byproduct
            adjusted_cost = fixed_like_mgmt + base_yvc * q_factor
            mat[i, j] = revenue - adjusted_cost
    return pd.DataFrame(mat,
                        index=[f"{r:+d}%" for r in q_rates],
                        columns=[f"{r:+d}%" for r in p_rates])

def plot_sensitivity_heatmap(mat, title="수량·단가가 동시에 변할 때의 소득 변화"):
    """RdYlGn 히트맵 — 적자(빨강) ~ 흑자(초록)"""
    import seaborn as sns
    h = max(4.0, 0.55 * len(mat) + 1.6)
    w = max(5.5, 0.95 * len(mat.columns) + 2.2)
    fig, ax = plt.subplots(figsize=(w, h))
    disp = mat / 10000.0     # 만원 단위로 표시
    from matplotlib.colors import LinearSegmentedColormap
    _econ_cmap = LinearSegmentedColormap.from_list(
        "smart_econ_div", ["#C96767", "#F5E2E2", "#FFFFFF", "#D9EAF7", "#3D6F9F"])
    sns.heatmap(disp, annot=True, fmt=".0f", cmap=_econ_cmap, center=0,
                linewidths=.8, linecolor="white", ax=ax,
                cbar_kws={"label": "소득 (만원/10a)", "shrink": .82})
    ax.set_xlabel("단가 변동률"); ax.set_ylabel("수량 변동률")
    deco(ax, title + "  (단위: 만원/10a)", ylabel_top=False)
    ax.grid(False)
    for _sp in ax.spines.values():
        _sp.set_visible(False)
    try:
        _cb = ax.collections[0].colorbar
        if _cb is not None:
            _cb.outline.set_visible(False)
    except Exception:
        pass
    plt.tight_layout(pad=1.2)
    return fig

LABOR_HINTS = ("노동시간", "노력시간", "작업시간", "소요시간", "노동(시간", "시간")


def looks_like_hours(name):
    """'자가노동시간'처럼 값이 '원'이 아니라 '시간'인 열인지 이름으로 추정한다."""
    t = str(name).replace(" ", "")
    if any(k in t for k in ("시간당", "원/시간", "임률", "노임")):
        return False
    return any(k in t for k in LABOR_HINTS)


def partial_budget_from_data(data, trt_col, control, treated, qty_col, price,
                             cost_cols=(), area_a=10.0,
                             hour_cols=(), wage_per_hour=0.0):
    """올린 데이터에서 '대조구 → 신기술구'로 바뀐 것만 뽑아 부분예산표 두 장을 만든다.

    - 손실적 요소(A) = 늘어난 비용 + 줄어든 수익
    - 이익적 요소(B) = 늘어난 수익 + 줄어든 비용
    반복이 여러 개면 처리구 평균을 쓰고, 자료 기준면적을 10a로 환산한다.
    산출근거는 검산기가 읽을 수 있게 숫자와 사칙연산만 넣는다.

    hour_cols: 값이 '원'이 아니라 '시간'인 열(자가노동시간 등).
               늘어난 시간 × wage_per_hour 로 금액을 만들어 넣는다.
               (환산하지 않으면 '10시간'이 '10원'으로 들어가 버린다)
    """
    hour_cols = [c for c in (hour_cols or []) if c and c != qty_col]
    cost_cols = [c for c in (cost_cols or [])
                 if c and c != qty_col and c not in hour_cols]
    wage = float(wage_per_hour or 0.0)
    if hour_cols and wage <= 0:
        raise ValueError("노동시간 열을 금액으로 바꾸려면 시간당 노임을 입력해야 합니다.")
    if control == treated:
        raise ValueError("대조구와 신기술 처리구가 같습니다.")
    area = float(area_a)
    if not np.isfinite(area) or area <= 0:
        raise ValueError("자료 기준면적은 0보다 커야 합니다.")
    price = float(price)
    if not np.isfinite(price) or price < 0:
        raise ValueError("단가는 0 이상이어야 합니다.")
    need = [trt_col, qty_col] + cost_cols + hour_cols
    missing = [c for c in need if c not in data.columns]
    if missing:
        raise ValueError("자료에 없는 열: " + ", ".join(map(str, missing)))
    d = data[list(dict.fromkeys(need))].copy()
    d[trt_col] = d[trt_col].astype(str).str.strip()
    for c in [qty_col] + cost_cols + hour_cols:
        d[c] = pd.to_numeric(d[c], errors="coerce")
    d = d.dropna(subset=[qty_col])
    got = set(d[trt_col])
    for g in (control, treated):
        if str(g) not in got:
            raise ValueError(f"'{g}' 처리구가 자료에 없습니다.")
    m = d.groupby(trt_col).mean(numeric_only=True)
    factor = 10.0 / area

    loss, gain, detail = [], [], []

    def _row(name, basis, amount):
        return {"항목": name, "산출근거": basis, "금액(원)": float(round(amount))}

    q_c = float(m.loc[str(control), qty_col]) * factor
    q_t = float(m.loc[str(treated), qty_col]) * factor
    dq = q_t - q_c
    detail.append({"구분": "수량(10a)", "대조구": round(q_c, 1), "신기술구": round(q_t, 1),
                   "차이": round(dq, 1)})
    if abs(dq) > 1e-9 and price > 0:
        # 산출근거에 적힌 숫자와 금액이 정확히 맞아떨어져야 검산기를 통과한다
        dq_r, price_r = round(abs(dq), 1), round(price)
        basis = f"{price_r:,.0f} * {dq_r:,.1f}"
        if dq > 0:
            gain.append(_row("판매수익 증가", basis, dq_r * price_r))
        else:
            loss.append(_row("판매수익 감소(수량 감소)", basis, dq_r * price_r))

    for c in cost_cols:
        if c not in m.columns or pd.isna(m.loc[str(control), c]) or pd.isna(m.loc[str(treated), c]):
            continue
        c_c = round(float(m.loc[str(control), c]) * factor)
        c_t = round(float(m.loc[str(treated), c]) * factor)
        dc = c_t - c_c
        detail.append({"구분": f"{c}(10a)", "대조구": c_c, "신기술구": c_t, "차이": dc})
        if dc == 0:
            continue
        basis = f"{max(c_t, c_c):,.0f} - {min(c_t, c_c):,.0f}"
        if dc > 0:
            loss.append(_row(f"{c} 증가", basis, dc))
        else:
            gain.append(_row(f"{c} 절감", basis, abs(dc)))

    # 노동시간 열: 시간 차이 × 시간당 노임 → 금액
    for c in hour_cols:
        if c not in m.columns or pd.isna(m.loc[str(control), c]) or pd.isna(m.loc[str(treated), c]):
            continue
        h_c = round(float(m.loc[str(control), c]) * factor, 1)
        h_t = round(float(m.loc[str(treated), c]) * factor, 1)
        dh = round(h_t - h_c, 1)
        detail.append({"구분": f"{c}(10a, 시간)", "대조구": h_c, "신기술구": h_t, "차이": dh})
        if abs(dh) < 1e-9:
            continue
        w = round(wage)
        basis = f"{w:,.0f} * {abs(dh):,.1f}"
        amt = w * abs(dh)
        label = f"{abs(dh):g}시간"   # 표 칸이 좁으므로 짧게
        if dh > 0:
            loss.append(_row(f"자가노동비 증가({label})", basis, amt))
        else:
            gain.append(_row(f"자가노동비 절감({label})", basis, amt))

    cols = ["항목", "산출근거", "금액(원)"]
    empty = pd.DataFrame({"항목": [""], "산출근거": [""], "금액(원)": [None]})
    loss_df = pd.DataFrame(loss, columns=cols) if loss else empty.copy()
    gain_df = pd.DataFrame(gain, columns=cols) if gain else empty.copy()
    return loss_df, gain_df, pd.DataFrame(detail)


def money_table(df, dec_overrides=None):
    """경제성 표를 보고서·다운로드용으로 정리한다: 숫자 열을 반올림하고
    천 단위 콤마를 넣은 문자열로 바꾼다 (예: 3750000.333... → '3,750,000').

    소수 자릿수는 열 이름으로 추정한다:
    - '(%)'로 끝나면 1자리 (예: 소득률(%))
    - 'B/C'는 2자리
    - '손익분기수량'처럼 '수량'이 들어간 열은 1자리
    - 그 밖의 숫자 열(원 단위 금액)은 0자리
    dec_overrides={"열이름": 자릿수} 로 개별 지정할 수 있다.
    """
    out = df.copy()
    dec_overrides = dec_overrides or {}
    for c in out.columns:
        if not pd.api.types.is_numeric_dtype(out[c]):
            continue
        name = str(c)
        if name in dec_overrides:
            dec = dec_overrides[name]
        elif name.endswith("(%)"):
            dec = 1
        elif name in ("B/C", "단년도 총수입/생산비", "할인 B/C"):
            dec = 2
        elif "수량" in name:
            dec = 1
        else:
            dec = 0
        out[c] = out[c].map(
            lambda v, dec=dec: "-" if pd.isna(v) else f"{round_half_up(v, dec):,.{dec}f}")
    return out


def log_action(what):
    import datetime
    st.session_state.setdefault("log", []).append(
        {"시각": datetime.datetime.now().strftime("%H:%M:%S"), "작업": what})
    _record_usage(what)

with st.sidebar.expander("🤖 AI 기능 켜기 (API 키 입력)"):
    st.caption("여기는 **설정 칸**입니다. 키를 넣으면 각 분석 결과 아래에 "
               "**🤖 AI 해석** 버튼이 생기고, 통계분석 탭의 **🧠 AI도우미**도 쓸 수 있어요. "
               "키가 없어도 다른 기능은 모두 정상 작동합니다.")
    provider = st.selectbox("AI 제공사", list(_AI_PROVIDERS.keys()), key="ai_provider")
    st.caption(f"🔑 {_AI_PROVIDERS[provider]['key_hint']}")
    st.text_input(f"{provider.split()[0]} API 키", type="password", key="api_key")
    _models = list(_AI_PROVIDERS[provider]["models"])
    if st.session_state.get("api_key"):
        if provider.startswith("Gemini"):
            _list_fn, _live_key, _btn_label = (
                list_gemini_models, "gemini_models_live", "🔄 사용 가능한 Gemini 모델 조회")
        elif provider.startswith("ChatGPT"):
            _list_fn, _live_key, _btn_label = (
                list_openai_models, "openai_models_live", "🔄 사용 가능한 OpenAI 모델 조회")
        else:
            _list_fn, _live_key, _btn_label = (
                list_claude_models, "claude_models_live", "🔄 사용 가능한 Claude 모델 조회")
        if st.button(_btn_label, width="stretch", key=f"list_models_{_live_key}"):
            with st.spinner("제공사에서 모델 목록을 확인하는 중..."):
                _live_models, _live_err = _list_fn(st.session_state.get("api_key"))
            if _live_err:
                st.warning(_live_err)
            elif _live_models:
                st.session_state[_live_key] = _live_models
                st.success(f"사용 가능한 모델 {len(_live_models)}개를 확인했습니다.")
            else:
                st.warning("사용 가능한 텍스트 모델을 찾지 못했습니다. 직접 입력을 이용해 주세요.")
        if st.session_state.get(_live_key):
            _models = list(st.session_state[_live_key])

    _labels = {}
    if len(_models) >= 1:
        _labels[_models[0]] = f"{_models[0]} (저렴·빠름)"
    if len(_models) >= 2:
        _labels[_models[1]] = f"{_models[1]} (정교함)"
    _opts = list(_models) + ["✏️ 직접 입력"]
    if not _models:
        st.caption("설정된 모델 목록이 없습니다. 모델명을 직접 입력해 주세요.")
        _sel = st.text_input("모델명 직접 입력", value="",
                             placeholder="예) claude-sonnet-5")
    else:
        _sel = st.selectbox("모델", _opts, format_func=lambda m: _labels.get(m, m))
        if _sel == "✏️ 직접 입력":
            _sel = st.text_input("모델명 직접 입력", value=_models[0],
                                 help="새 모델이 나왔을 때 여기에 이름을 넣으면 바로 쓸 수 있습니다.")
    st.session_state["ai_model_g"] = _sel
    ai_model = _sel
    if st.session_state.get("api_key"):
        st.success(f"✅ {provider.split()[0]} 활성화됨 — 분석 결과 아래 'AI 해석'을 눌러보세요.")
        if st.button("🔌 API 연결 테스트", width="stretch", key="ai_conn_test"):
            with st.spinner("연결을 확인하는 중..."):
                _r = test_ai_connection(provider, st.session_state.get("api_key"), _sel)
            if _r["ok"]:
                st.success(f"✅ {_r['provider']} / {_r['model']} 연결 성공")
                if _r.get("sample"):
                    st.caption("응답 예시: " + _r["sample"][:60])
            else:
                st.error(f"❌ {_r['provider']} / {_r['model']} — {_r['message']}")
    else:
        st.info("키를 넣으면 AI 해석이 켜집니다.")

if not _HAS_DOCX:
    st.sidebar.caption("💡 워드(docx) 저장을 쓰려면: pip install python-docx")

st.sidebar.markdown("---")
st.sidebar.caption("스마트 통계 에이전트 얏호(*/ω＼*)")

# ================================================================ 떠 있는 AI 도우미
_ATT_LIMIT = 12000          # 첨부 전체에서 AI 에게 넘길 글자 수 상한


def _read_attachment(f, budget=6000):
    """올린 파일을 AI 가 읽을 수 있는 **글자**로 바꾼다.

    ai_call() 은 글자만 주고받으므로(제공사가 여러 곳이라 그림 전송 방식이 제각각),
    표·문서는 내용을 뽑아 넘기고 그림은 넘길 수 없음을 분명히 알린다.
    """
    name = getattr(f, "name", "첨부파일")
    ext = ("." + name.rsplit(".", 1)[-1].lower()) if "." in name else ""
    try:
        raw = f.getvalue()
    except Exception:
        raw = f.read()
    head = f"\n\n===== 첨부: {name} ({len(raw):,} bytes) =====\n"

    try:
        if ext in (".csv", ".tsv", ".txt", ".md", ".json"):
            if ext in (".csv", ".tsv"):
                d = pd.read_csv(io.BytesIO(raw), sep=None, engine="python")
                return head + _df_digest(d, budget)
            for enc in ("utf-8", "cp949", "utf-8-sig"):
                try:
                    return head + raw.decode(enc)[:budget]
                except UnicodeDecodeError:
                    continue
            return head + raw.decode("utf-8", "replace")[:budget]

        if ext in (".xlsx", ".xls", ".xlsm"):
            xls = pd.ExcelFile(io.BytesIO(raw))
            per = max(600, budget // max(1, len(xls.sheet_names)))
            out = [head + f"(시트 {len(xls.sheet_names)}개: {', '.join(xls.sheet_names)})"]
            for sh in xls.sheet_names:
                out.append(f"\n--- 시트: {sh} ---\n"
                           + _df_digest(pd.read_excel(xls, sh), per))
            return "".join(out)

        if ext == ".pdf":
            from pypdf import PdfReader
            rd = PdfReader(io.BytesIO(raw))
            txt = []
            for pg in rd.pages[:30]:
                txt.append(pg.extract_text() or "")
                if sum(map(len, txt)) > budget:
                    break
            body = "\n".join(txt).strip()
            return head + (f"(총 {len(rd.pages)}쪽)\n" + body[:budget] if body
                           else "⚠️ 글자를 뽑을 수 없는 PDF입니다(스캔본일 수 있음).")

        if ext == ".docx":
            import docx as _dx
            d = _dx.Document(io.BytesIO(raw))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for t in d.tables:
                for r in t.rows:
                    parts.append(" | ".join(c.text.strip() for c in r.cells))
            return head + "\n".join(parts)[:budget]

        if ext == ".hwpx":
            import zipfile
            from lxml import etree as _et
            parts = []
            with zipfile.ZipFile(io.BytesIO(raw)) as z:
                for n in sorted(x for x in z.namelist()
                                if x.startswith("Contents/section") and x.endswith(".xml")):
                    root = _et.fromstring(z.read(n))
                    parts += [t.text for t in root.iter() if t.tag.endswith("}t") and t.text]
            return head + ("\n".join(parts)[:budget] or "⚠️ 내용을 읽지 못했습니다.")

        if ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"):
            return head + ("⚠️ 그림 파일은 아직 AI에게 전달할 수 없습니다. "
                           "그림 속 표는 엑셀·CSV로, 문서는 PDF로 올려 주세요.")

        if ext == ".hwp":
            return head + ("⚠️ 옛 한글(.hwp)은 읽을 수 없습니다. "
                           "한글에서 **다른 이름으로 저장 → hwpx 또는 PDF**로 바꿔 올려 주세요.")

        return head + f"⚠️ 지원하지 않는 형식입니다({ext or '확장자 없음'})."
    except Exception as e:
        return head + f"⚠️ 읽기 실패: {type(e).__name__}: {e}"


def _df_digest(d, budget):
    """표를 '크기 + 열 목록 + 앞부분 + 기술통계'로 요약한다(통째로 넘기면 너무 길다)."""
    lines = [f"크기: {len(d):,}행 × {len(d.columns)}열",
             f"열: {', '.join(map(str, d.columns))}", "", "[앞부분]",
             d.head(20).to_string(index=False)]
    try:
        num = d.select_dtypes(include=np.number)
        if len(num.columns):
            lines += ["", "[기술통계]", num.describe().round(2).to_string()]
    except Exception:
        pass
    return "\n".join(lines)[:budget]


def _ai_panel(df, menu_name):
    """AI 질문 상자 본체. 떠 있는 창과 사이드바가 같은 내용을 공유한다."""
    if not st.session_state.get("api_key"):
        st.caption("사이드바 **🤖 AI 기능 켜기**에서 API 키를 넣으면 사용할 수 있어요.")
        return
    st.caption(f"지금 화면: **{menu_name}**"
               + (f" · 데이터: **{st.session_state.get('cur_key')}**"
                  if df is not None else " · (데이터 없음)"))
    for role, txt in st.session_state.get("gai_hist", [])[-6:]:
        with st.chat_message("user" if role == "q" else "assistant"):
            st.markdown(txt)
    q = st.text_area("궁금한 점", key="gai_q", height=80, label_visibility="collapsed",
                     placeholder="예) 지금 이 표에서 어떤 처리가 가장 좋아? / 이 화면 어떻게 쓰는 거야?")
    ups = st.file_uploader(
        "📎 파일 첨부 (선택)", key="gai_files", accept_multiple_files=True,
        type=["csv", "tsv", "txt", "md", "json", "xlsx", "xls", "xlsm",
              "pdf", "docx", "hwpx"],
        help="엑셀·CSV·PDF·워드·한글(hwpx)의 내용을 읽어 함께 물어봅니다. "
             "그림 파일과 옛 한글(.hwp)은 아직 읽을 수 없어요.")
    if ups:
        st.caption("첨부: " + ", ".join(getattr(f, "name", "?") for f in ups))
    b1, b2 = st.columns([3, 1])
    if b1.button("물어보기", key="gai_send", type="primary", width="stretch") and q.strip():
        try:
            ctx = ""
            if df is not None:
                # build_group_profiles() 는 문자열이 아니라 dict 를 돌려준다.
                # 그대로 + 로 이으면 TypeError 가 나므로 f-string 으로 문자열화한다.
                _prof = build_group_profiles(df, question=q)
                ctx = (f"{build_data_overview(df)}\n\n[처리·품종별 요약]\n{_prof}")
            att = ""
            if ups:
                per = max(1500, _ATT_LIMIT // len(ups))
                att = "".join(_read_attachment(f, per) for f in ups)[:_ATT_LIMIT]
            with st.spinner("AI가 생각하는 중..."):
                ans = ai_call(
                    "당신은 농업연구사를 돕는 통계 전문가이자 이 앱의 사용 안내자입니다. "
                    f"사용자는 지금 '{menu_name}' 화면을 보고 있습니다. "
                    "아래 실제 데이터 요약만 근거로 한국어로 쉽고 정확하게 답하세요. "
                    "입력에 없는 수치나 유의성은 추측하지 마세요. "
                    "앱 사용법을 묻는다면 화면 이름을 들어 안내하세요.\n\n"
                    + (ctx or "(현재 선택된 데이터가 없습니다.)")
                    + (f"\n\n[사용자가 올린 파일]{att}" if att else "")
                    + f"\n\n질문: {q}", max_tokens=900)
        except Exception as _ex:
            ans = f"⚠️ AI 호출 실패: {type(_ex).__name__}: {_ex}"
        _qlog = q + (("\n\n📎 " + ", ".join(getattr(f, "name", "?") for f in ups))
                     if ups else "")
        st.session_state["gai_hist"] = (st.session_state.get("gai_hist", [])
                                        + [("q", _qlog), ("a", ans)])[-12:]
        st.rerun()
    if b2.button("지우기", key="gai_clear", width="stretch"):
        st.session_state["gai_hist"] = []
        st.rerun()
    st.caption("⚠️ AI 답변은 초안입니다. 수치와 해석은 연구자가 확인해 주세요.")


def _dock_supported():
    """오른쪽 아래에 띄우려면 st.container(key=...) 가 필요하다(스트림릿 1.48 이상).

    옛 버전에서는 TypeError 가 나므로 미리 확인해서 사이드바로 대신 내보낸다.
    (예전에는 이 실패를 그냥 삼켜서 버튼이 아무 데도 안 보였다.)
    """
    try:
        import inspect as _isp
        return ("key" in _isp.signature(st.container).parameters
                and hasattr(st, "popover") and hasattr(st, "chat_message"))
    except Exception:
        return False


def floating_ai(df, menu_name):
    """어느 화면에서든 쓸 수 있는 AI 질문 상자.

    기본은 오른쪽 아래에 떠 있는 버튼이고, 스트림릿 버전이 낮아 그렇게 만들 수
    없으면 **사이드바 맨 위**에 같은 상자를 대신 넣는다. 어느 쪽이든 사라지지 않는다.
    """
    if not _dock_supported():
        with st.sidebar:
            with st.expander("💬 AI에게 물어보기", expanded=False):
                _ai_panel(df, menu_name)
                st.caption("ℹ️ 스트림릿 1.48 이상이면 화면 오른쪽 아래에 떠 있는 창으로 "
                           "쓸 수 있어요. (`pip install -U streamlit`)")
        return

    st.markdown("""<style>
      div.st-key-gai_dock { position: fixed !important; right: 1.5rem; bottom: 1.5rem;
          z-index: 9999; width: auto !important; }
      div.st-key-gai_dock button { border-radius: 2.2rem !important;
          padding: .95rem 1.9rem !important; font-weight: 700 !important;
          font-size: 1.12rem !important; line-height: 1.25 !important;
          box-shadow: 0 6px 22px rgba(0,0,0,.30) !important; }
      div.st-key-gai_dock button:hover { transform: translateY(-2px); }
      div.st-key-gai_dock button p { font-size: 1.12rem !important;
          font-weight: 700 !important; margin: 0 !important; }
      @media (max-width: 640px) { div.st-key-gai_dock { right: .6rem; bottom: .6rem; }
          div.st-key-gai_dock button { padding: .8rem 1.4rem !important;
              font-size: 1rem !important; } }
    </style>""", unsafe_allow_html=True)
    with st.container(key="gai_dock"):
        with st.popover("💬 AI에게 물어보기"):
            st.markdown("###### 💬 AI 도우미")
            _ai_panel(df, menu_name)


df = st.session_state.df
if df is not None and len(df.columns) != len(set(map(str, df.columns))):
    df = clean_columns(df)
    st.session_state.df = df
    st.info("ℹ️ 열 이름이 중복되어 자동으로 구분했습니다(예: 값, 값_2). "
            "원본 엑셀의 머리글을 확인해 보세요.")

try:
    floating_ai(df, menu)
except Exception as _gex:
    # 조용히 사라지면 원인을 알 수 없다 — 사이드바에 최소한의 대체 창을 남긴다.
    try:
        with st.sidebar.expander("💬 AI에게 물어보기", expanded=False):
            st.caption(f"떠 있는 창을 만들지 못했습니다 ({type(_gex).__name__}). "
                       "여기서 이용해 주세요.")
            _ai_panel(df, menu)
    except Exception:
        pass


# ================================================================ 공통 가드
if df is None and menu not in ("📑 보고서", "📖 사용설명서"):
    st.title("실험 데이터 자동 통계 분석")
    st.caption("엑셀만 올리면 통계분석부터 한글 보고서까지 한 번에")
    st.info("👈 왼쪽 **📂 데이터 불러오기**에서 파일을 올리거나, 아래 샘플 데이터로 바로 체험해 보세요.")

    st.markdown("#### 🚀 샘플 데이터로 바로 시작하기")
    s1, s2, s3 = st.columns(3)
    if s1.button("🌱 실험 데이터 체험\n\n처리구별 생육·수량", width="stretch"):
        st.session_state.files["샘플_실험데이터"] = make_sample("실험")
        st.rerun()
    if s2.button("💰 경제성 데이터 체험\n\n수량·단가·경영비", width="stretch"):
        st.session_state.files["샘플_경제성"] = make_sample("경제성")
        st.rerun()
    if s3.button("📋 설문 데이터 체험\n\n응답자 60명 만족도", width="stretch"):
        st.session_state.files["샘플_설문"] = make_sample("설문")
        st.rerun()

    st.divider()
    st.markdown("#### 📌 이렇게 사용하세요")
    w1, w2, w3, w4 = st.columns(4)
    for col, (n_, t_, d_) in zip([w1, w2, w3, w4], [
        ("1️⃣", "데이터 올리기", "엑셀·CSV를 올립니다.\n여러 파일·여러 시트도 OK"),
        ("2️⃣", "전처리·확인", "결측치와 이상값을 정리하고\n파생변수를 만듭니다"),
        ("3️⃣", "분석 실행", "AI가 추천한 분석을 클릭.\n유의성 문자(a,b,c) 자동"),
        ("4️⃣", "보고서 생성", "결과를 담아 한글(hwpx)\n보고서로 자동 완성")]):
        col.markdown(f"### {n_}\n**{t_}**\n\n{d_}")

    st.divider()
    st.markdown("#### 🧰 주요 기능")
    f1, f2, f3 = st.columns(3)
    f1.markdown("""
**📈 통계 분석**
- 일원·이원배치 분산분석
- 사후검정 (Tukey / 던컨 / Bonferroni)
- 정규성·등분산 가정 검정
- 비모수 검정 (Kruskal-Wallis 등)
- 상관분석 · 회귀분석 · PCA
""")
    f2.markdown("""
**🤖 자동화 · AI**
- 데이터 구조 분석 → 통계 방법 추천
- 결과를 쉬운 말로 자동 해석
- 머신러닝 예측 + 변수 중요도
- AI 자연어 질의응답 (선택)
- 연구계획서 기반 분석 추천 (선택)
""")
    f3.markdown("""
**📄 한글 문서 자동화**
- 결과표를 한글(hwpx)로 바로 저장
- 표 서식(글꼴·음영·선) 자유 설정
- 그래프 자동 삽입 + 중앙 정렬
- `<표 1>` `<그림 1>` 캡션 자동
- 여러 분석을 하나의 보고서로
""")
    st.divider()
    st.caption("💡 통계를 잘 몰라도 괜찮아요. 각 분석마다 'ℹ️ 이 분석이 뭔가요?' 설명이 준비되어 있습니다.")
    st.stop()

if df is not None:
    num_cols = df.select_dtypes(include=np.number).columns.tolist()
    cat_cols = df.select_dtypes(exclude=np.number).columns.tolist()
else:
    num_cols, cat_cols = [], []

# ================================================================ 통계분석
# ================================================================ 원클릭 오토파일럿
if menu == "⚡ 원클릭 보고서":
    st.title("⚡ 원클릭 보고서")
    st.caption("데이터만 올리면 **정제 → 설계 인지 → 분석 → 그래프 → 문장 → 보고서**까지 "
               "한 번에 만들어 드립니다.")
    st.info("💡 결과물은 **초안**입니다. 연구 목적과 고찰은 연구자가 확인·보완해 주세요. "
            "세부 조정이 필요하면 '📊 통계분석' 메뉴에서 직접 분석할 수 있습니다.")

    if df is None:
        st.warning("먼저 왼쪽에서 데이터를 올리거나 🧪 샘플 데이터를 눌러 주세요.")
    else:
        c1, c2, c3 = st.columns(3)
        c1.metric("행", f"{len(df):,}"); c2.metric("열", len(df.columns))
        c3.metric("숫자형 항목", len(df.select_dtypes(include=np.number).columns))
        _pre = detect_design(df)
        if _pre.get("trt"):
            st.caption(f"🔬 자동 인지: **{_pre['design']}** — 처리구 '{_pre['trt']}'"
                       + (f", 반복 '{_pre['blk']}'" if _pre.get("blk") else ", 반복 없음"))
        with st.expander("⚙️ 설정 (필요할 때만)"):
            _allc = df.columns.tolist()
            _bopts = ["(자동)"] + _allc
            # index= 와 key= 를 함께 쓰면 스트림릿이 경고를 띄우므로,
            # 기본값은 session_state에만 넣고 위젯은 key로만 만든다.
            if st.session_state.get("ap_trt") not in _allc:
                st.session_state["ap_trt"] = (_pre["trt"] if _pre.get("trt") in _allc
                                              else (_allc[0] if _allc else None))
            if st.session_state.get("ap_blk") not in _bopts:
                st.session_state["ap_blk"] = "(자동)"
            st.selectbox("처리구 열 (자동 인지 결과를 바꾸려면)", _allc, key="ap_trt")
            st.selectbox("반복(블록) 열", _bopts, key="ap_blk")
            ap_ph = st.selectbox("사후검정", ["Tukey HSD", "던컨(Duncan)", "Bonferroni"], key="ap_ph")
            ap_err = st.radio("오차막대", ["표준편차(SD)", "표준오차(SE)"], horizontal=True, key="ap_err",
                              help="SD=자료가 흩어진 정도, SE=평균의 정확도. 아래 설명을 참고하세요.")
            with st.expander("❓ 표준편차(SD)와 표준오차(SE), 뭐가 다른가요?"):
                st.markdown(EXPLAIN["sd_se"])
            ap_max = st.slider("한 번에 분석할 조사항목 수", 1, 15, key="ap_max",
                               help="초장·엽수·수량처럼 숫자로 조사한 항목이 여러 개일 때, "
                                    "앞에서부터 몇 개까지 자동 분석할지 정합니다. "
                                    "8이면 조사항목이 12개라도 8개만 분석해 보고서가 너무 길어지지 않습니다.")
            st.caption("💡 '조사항목'은 초장(cm)·수량(kg/10a)처럼 **측정한 숫자 열**을 뜻합니다. "
                       "수량 관련 항목이 먼저 분석됩니다.")

        if st.button("🚀 원클릭 분석 시작", type="primary", width="stretch"):
            _bsel = st.session_state.get("ap_blk", "(자동)")
            res = run_autopilot_engine(df, ph=st.session_state.get("ap_ph", "Tukey HSD"),
                                       err_type=st.session_state.get("ap_err"),
                                       max_items=int(st.session_state.get("ap_max", 8)),
                                       trt_override=st.session_state.get("ap_trt"),
                                       blk_override=(None if _bsel == "(자동)" else _bsel))
            st.session_state["autopilot"] = res
            if res.get("ok"):
                st.markdown(
                    "<div style='text-align:center;font-size:2.1rem;letter-spacing:.35rem;"
                    "padding:.5rem 0'>🌱 🌾 🌶️ 🍃 🌾 🌱</div>"
                    "<div style='text-align:center;color:#4b7d3a;font-weight:600;"
                    "padding-bottom:.6rem'>분석이 잘 마무리되었습니다</div>",
                    unsafe_allow_html=True)
                log_action(f"원클릭 보고서 생성({len(res['summary'])}개 항목)")

        ap = st.session_state.get("autopilot")
        if ap and not ap.get("ok"):
            for m in ap.get("msgs", []): st.warning(m)
        elif ap and ap.get("ok"):
            st.success(f"✅ 완료! {len(ap['summary'])}개 항목을 분석해 보고서를 만들었습니다.")
            d1, d2 = st.columns(2)
            if ap.get("hwpx"):
                d1.download_button("📘 보고서 내려받기 — 한글(hwpx)", ap["hwpx"],
                                   "통계분석_초안.hwpx", type="primary", width="stretch")
            if ap.get("docx"):
                d2.download_button("📝 보고서 내려받기 — 워드(docx)", ap["docx"],
                                   "통계분석_초안.docx", type="primary", width="stretch")
            elif not _HAS_DOCX:
                d2.caption("워드 저장: pip install python-docx 후 사용 가능")
            try:
                _xl = make_xlsx_multi(
                    [{"caption": "분석 종합", "table": ap["summary"]}]
                    + [b for b in ap["blocks"] if b.get("table") is not None],
                    "원클릭 분석 결과")
                if _xl:
                    st.download_button("📈 엑셀(xlsx) — 스마트 블루 항목별 시트 + 편집 가능한 그래프",
                                       _xl, "통계분석_초안.xlsx", width="stretch",
                                       key="xls_autopilot",
                                       help="조사항목마다 시트가 하나씩 만들어지고, 각 시트에 "
                                            "엑셀 기본 차트가 들어갑니다. 막대 색·글꼴·축 범위·"
                                            "차트 종류를 원하는 대로 바꿀 수 있습니다.")
            except Exception as _e:
                st.caption(f"엑셀 파일 생성 실패 ({type(_e).__name__})")
            if st.button("➕ 이 결과를 '📑 보고서'에도 담기", width="stretch"):
                st.session_state.report_items.extend(ap["report_items"])
                st.success("보고서 메뉴에 담았습니다! 다른 분석과 합쳐서 편집할 수 있어요.")

            st.markdown("---")
            for m in ap.get("msgs", []): st.caption("• " + m)

            st.markdown("### 📋 분석 종합")
            smart_table(ap["summary"], width="stretch", hide_index=True)

            st.markdown("### 📄 적요(초안)")
            st.code(ap["abstract"], language=None)

            st.markdown("### 📈 항목별 결과")
            _tables = [b for b in ap["blocks"] if b.get("table") is not None]
            _texts = [b for b in ap["blocks"] if b.get("text")]
            for i, blk in enumerate(_tables):
                with st.expander(f"{blk['caption']}", expanded=(i == 0)):
                    if i < len(_texts):
                        st.code(_texts[i]["text"], language=None)
                    smart_table(sup_display(blk["table"]), width="stretch", hide_index=True)
                    if blk.get("image"):
                        st.image(blk["image"], width=640)

            st.markdown("### 🧾 통계 처리 문구")
            st.code(ap["stat_line"], language=None)

elif menu == "📊 통계분석":
    st.title("실험 데이터 자동 통계 분석")
    st.markdown("### 🤖 AI 통계 방법 추천")
    for rec in recommend_analysis(df): st.success(rec)
    st.divider()

    _SUB = ["📋 데이터", "🧹 전처리", "🧮 파생변수", "🔗 상관분석", "📈 분산분석",
            "🧪 비모수검정", "🧬 PCA", "📉 회귀분석", "🤖 머신러닝"]
    sub = st.radio("분석 선택", _SUB, horizontal=True, key="stat_sub",
                   label_visibility="collapsed")
    st.markdown("---")

    class _Show:
        """선택된 화면만 그리도록 (숨은 화면은 계산하지 않아 훨씬 빠릅니다)"""
        def __init__(self, name): self.on = (sub == name)
        def __enter__(self): return self
        def __exit__(self, *a): return False
    
    tab_data = _Show("📋 데이터"); tab_prep = _Show("🧹 전처리")
    tab_derive = _Show("🧮 파생변수"); tab_corr = _Show("🔗 상관분석")
    tab_anova = _Show("📈 분산분석"); tab_np = _Show("🧪 비모수검정")
    tab_pca = _Show("🧬 PCA"); tab_reg = _Show("📉 회귀분석")
    tab_ml = _Show("🤖 머신러닝")

    if tab_data.on:
        st.subheader("데이터 미리보기")
        smart_table(df, width="stretch")
        c1, c2, c3 = st.columns(3)
        c1.metric("행 개수", df.shape[0]); c2.metric("열 개수", df.shape[1])
        c3.metric("결측치 개수", int(df.isna().sum().sum()))

        if len(df) > 20000:
            st.warning(f"⚠️ 행이 {len(df):,}개로 많습니다. 분석·그래프가 느려질 수 있어요. "
                       "필요한 기간·처리만 걸러서 사용하시길 권장합니다.")
        # ---------- 데이터 검진 ----------
        st.markdown("#### 🩺 데이터 검진")
        issues = []
        numlike = find_numeric_like(df)
        if numlike:
            issues.append(("warn",
                "숫자로 보이지만 문자로 읽힌 열이 있습니다: "
                f"{', '.join(f'{k}({v}%)' for k, v in numlike.items())}"))
        n_miss = int(df.isna().sum().sum())
        if n_miss:
            issues.append(("warn", f"결측치 {n_miss}개 — '전처리' 탭에서 처리할 수 있습니다."))
        n_dup = int(df.duplicated().sum())
        if n_dup:
            issues.append(("warn", f"완전히 같은 행이 {n_dup}개 있습니다(중복 입력 가능성)."))
        numc = df.select_dtypes(include=np.number).columns
        neg = [c for c in numc if (df[c] < 0).any()]
        if neg:
            issues.append(("info", f"음수가 포함된 열: {', '.join(neg)} — 입력 오류가 아닌지 확인하세요."))
        const = [c for c in df.columns if df[c].nunique(dropna=True) <= 1]
        if const:
            issues.append(("info", f"값이 모두 같은 열: {', '.join(const)} — 분석에서 제외됩니다."))
        # 실험설계 추정 (숫자로 적힌 반복·처리 코드도 함께 인식)
        _ys_c, catc, _promoted_c = split_code_columns(df)
        if _promoted_c:
            issues.append(("info",
                f"숫자로 적혀 있지만 처리구·반복 코드로 보이는 열: {', '.join(map(str, _promoted_c))} "
                "— 측정값이 아니라 그룹 구분으로 인식했습니다. 실제 측정값이면 분산분석에서 직접 골라 주세요."))
        design_msg = None
        for cand in catc:
            if any(k in str(cand).lower() for k in _BLOCK_KEYS):
                other = [c for c in catc if c != cand]
                if other:
                    tab = df.groupby([other[0], cand]).size()
                    balanced = tab.nunique() == 1
                    design_msg = (f"'{other[0]}' × '{cand}' 구조가 감지되었습니다 → "
                                  + ("**난괴법(RCBD)**으로 보입니다. 분산분석에서 반복(블록) 열로 "
                                     f"'{cand}'을 꼭 지정하세요." if balanced
                                     else "반복 수가 고르지 않습니다(불균형). 결측을 확인하세요."))
                break
        if design_msg:
            st.success("🔬 " + design_msg)
        if not issues:
            st.success("✅ 특별한 문제가 발견되지 않았습니다.")
        else:
            for kind, msg in issues:
                (st.warning if kind == "warn" else st.info)(msg)
        if numlike:
            st.markdown("###### 🔧 숫자로 자동 변환")
            fixcols = st.multiselect("변환할 열", list(numlike.keys()),
                                     default=list(numlike.keys()), key="fixnum")
            st.caption("쉼표(1,200)·단위(120kg)·공백이 섞인 값에서 숫자만 뽑아냅니다. "
                       "변환할 수 없는 값은 결측치가 됩니다.")
            if fixcols and st.button("숫자로 변환하기"):
                d2 = df.copy()
                report = []
                for c in fixcols:
                    conv = to_numeric_clean(d2[c])
                    report.append({"열": c, "변환 성공": int(conv.notna().sum()),
                                   "변환 실패(결측)": int(conv.isna().sum() - d2[c].isna().sum())})
                    d2[c] = conv
                set_df(d2, "숫자 변환")
                smart_table(pd.DataFrame(report), width="stretch")
                log_action(f"숫자 변환: {', '.join(fixcols)}")
                st.success("변환했습니다!"); st.rerun()

        st.write("**기술통계**"); smart_table(df.describe(), width="stretch")

    # ---------- 전처리 ----------
    if tab_prep.on:
        st.subheader("데이터 전처리")
        with st.expander("ℹ️ 전처리가 뭔가요?"):
            st.markdown(EXPLAIN["prep"])
        st.markdown("###### 📋 현재 데이터 (작업 결과가 즉시 반영됩니다)")
        pc1, pc2, pc3 = st.columns(3)
        pc1.metric("행", f"{len(df):,}"); pc2.metric("열", len(df.columns))
        pc3.metric("결측치", f"{int(df.isna().sum().sum()):,}")
        smart_table(df.head(20), width="stretch")
        st.caption(f"열 목록: {', '.join(map(str, df.columns))}")
        _stack = st.session_state.get("undo_stack", [])
        u1, u2 = st.columns([1, 3])
        if u1.button(f"↩️ 실행취소 ({len(_stack)})", width="stretch",
                     disabled=not _stack, help="바로 전 전처리 작업을 되돌립니다."):
            memo = undo_df()
            log_action(f"실행취소: {memo}")
            st.success(f"'{memo}' 작업을 되돌렸습니다."); st.rerun()
        if _stack:
            u2.caption(f"되돌릴 수 있는 작업: {' → '.join(h['memo'] for h in _stack[-3:])}"
                       + (" (최근 3개)" if len(_stack) > 3 else ""))
        else:
            u2.caption("전처리를 실행하면 되돌리기가 활성화됩니다. (최근 10단계까지)")
        pmode = st.radio("작업 선택",
                         ["결측치 처리", "이상값 처리", "중복 행 제거", "자료형 변환", "열 삭제/이름변경", "표준화·정규화"])

        if pmode == "결측치 처리":
            miss = df[df.isna().any(axis=1)]
            if len(miss) == 0:
                st.success("결측치가 없습니다.")
            else:
                st.warning(f"결측치가 있는 행: 총 {len(miss)}개")
                smart_table(miss.style.highlight_null(color="#FFF3B0"), width="stretch")
                st.caption(f"행 번호: {list(miss.index)}")
            m = st.radio("처리 방법", ["행 삭제", "평균 대체", "중앙값 대체", "0으로 대체"], horizontal=True)
            if st.button("적용", key="p_miss"):
                d = st.session_state.df.copy(); aff = list(miss.index)
                if m == "행 삭제": d = d.dropna()
                elif m == "평균 대체":
                    for c in d.select_dtypes(include=np.number).columns: d[c] = d[c].fillna(d[c].mean())
                elif m == "중앙값 대체":
                    for c in d.select_dtypes(include=np.number).columns: d[c] = d[c].fillna(d[c].median())
                else: d = d.fillna(0)
                set_df(d, "전처리")
                st.success(f"{len(aff)}개 행 처리 완료. 처리된 행 번호: {aff}")
                smart_table(d, width="stretch")

        elif pmode == "이상값 처리":
            with st.expander("ℹ️ 이상값이 뭔가요?"):
                st.markdown(EXPLAIN["outlier"])
            if not num_cols:
                st.warning("숫자형 변수가 필요합니다.")
            else:
                c1, c2 = st.columns(2)
                ocols = c1.multiselect("검사할 열", num_cols, default=num_cols)
                omethod = c2.radio("탐지 방법", ["IQR (1.5배)", "Z-점수 (±3)"])
                if ocols:
                    mask = pd.Series(False, index=df.index)
                    info = []
                    for c in ocols:
                        s = df[c]
                        if omethod.startswith("IQR"):
                            q1, q3 = s.quantile(.25), s.quantile(.75); iqr = q3 - q1
                            lo, hi = q1 - 1.5*iqr, q3 + 1.5*iqr
                        else:
                            mu, sd = s.mean(), s.std()
                            lo, hi = mu - 3*sd, mu + 3*sd
                        m_ = (s < lo) | (s > hi)
                        mask |= m_.fillna(False)
                        info.append({"변수": c, "하한": round(lo, 2), "상한": round(hi, 2), "이상값 수": int(m_.sum())})
                    smart_table(pd.DataFrame(info), width="stretch")
                    out_rows = df[mask]
                    if len(out_rows) == 0:
                        st.success("이상값이 없습니다.")
                    else:
                        st.warning(f"이상값이 포함된 행: {len(out_rows)}개 (행 번호: {list(out_rows.index)})")
                        smart_table(out_rows, width="stretch")
                    fig, ax = plt.subplots(figsize=(min(12, 1.5*len(ocols)+2), 4))
                    df[ocols].plot(kind="box", ax=ax); ax.set_title("상자그림(이상값 확인)")
                    plt.xticks(rotation=30); show_plot(fig); plt.close(fig)
                    act = st.radio("처리 방법", ["해당 행 삭제", "경계값으로 대체(윈저화)", "결측치로 변경"], horizontal=True)
                    if st.button("적용", key="p_out"):
                        d = st.session_state.df.copy()
                        if act == "해당 행 삭제":
                            d = d[~mask]
                        else:
                            for c in ocols:
                                s = d[c]
                                if omethod.startswith("IQR"):
                                    q1, q3 = s.quantile(.25), s.quantile(.75); iqr = q3-q1
                                    lo, hi = q1-1.5*iqr, q3+1.5*iqr
                                else:
                                    mu, sd = s.mean(), s.std(); lo, hi = mu-3*sd, mu+3*sd
                                if act.startswith("경계값"): d[c] = s.clip(lo, hi)
                                else: d[c] = s.where((s >= lo) & (s <= hi))
                        set_df(d, "전처리")
                        st.success(f"이상값 처리 완료. (영향 행: {list(out_rows.index)})")
                        smart_table(d, width="stretch")

        elif pmode == "중복 행 제거":
            dup = df[df.duplicated(keep=False)]
            st.write(f"중복 행: {len(dup)}개")
            if len(dup): smart_table(dup, width="stretch")
            if st.button("중복 제거", key="p_dup"):
                set_df(st.session_state.df.drop_duplicates().reset_index(drop=True), "중복 제거")
                st.success("중복 행을 제거했습니다."); smart_table(st.session_state.df, width="stretch")

        elif pmode == "자료형 변환":
            c1, c2 = st.columns(2)
            col = c1.selectbox("열 선택", df.columns.tolist(), key="t_col")
            to = c2.radio("변환", ["숫자형으로", "문자형으로"], horizontal=True)
            st.caption(f"현재 자료형: {df[col].dtype}")
            if st.button("변환", key="p_type"):
                d = st.session_state.df.copy()
                if to == "숫자형으로":
                    d[col] = pd.to_numeric(d[col].astype(str).str.replace(",", "").str.strip(), errors="coerce")
                    st.info(f"변환 실패(결측 처리)된 값: {int(d[col].isna().sum() - df[col].isna().sum())}개")
                else:
                    d[col] = d[col].astype(str)
                set_df(d, "자료형 변환"); st.success("변환 완료"); smart_table(d.head(), width="stretch")

        elif pmode == "열 삭제/이름변경":
            c1, c2 = st.columns(2)
            with c1:
                drops = st.multiselect("삭제할 열", df.columns.tolist())
                if drops and st.button("열 삭제", key="p_drop"):
                    set_df(st.session_state.df.drop(columns=drops), "열 삭제")
                    st.success(f"{len(drops)}개 열 삭제"); st.rerun()
            with c2:
                oldc = st.selectbox("이름 바꿀 열", df.columns.tolist(), key="rn")
                newc = st.text_input("새 이름", value=str(oldc))
                if st.button("이름 변경", key="p_rn"):
                    st.session_state.df = st.session_state.df.rename(columns={oldc: newc})
                    st.success("변경 완료"); st.rerun()

        else:  # 표준화·정규화
            sc = st.multiselect("변환할 열", num_cols, default=num_cols, key="sc")
            how = st.radio("방법", ["표준화(Z-점수)", "정규화(0~1)"], horizontal=True)
            st.caption("단위가 다른 변수들을 비교하거나 머신러닝에 넣을 때 사용합니다. 새 열로 추가됩니다.")
            if sc and st.button("적용", key="p_sc"):
                d = st.session_state.df.copy()
                arr = (StandardScaler() if how.startswith("표준화") else MinMaxScaler()).fit_transform(d[sc])
                suffix = "_표준화" if how.startswith("표준화") else "_정규화"
                for i, c in enumerate(sc): d[c + suffix] = arr[:, i].round(4)
                st.session_state.df = d; st.success("완료"); smart_table(d.head(), width="stretch")

    # ---------- 파생변수 ----------
    if tab_derive.on:
        st.subheader("파생변수 생성")
        with st.expander("ℹ️ 이 기능이 뭔가요?"): st.markdown(EXPLAIN["derive"])
        st.markdown("###### 📋 현재 데이터 (새 열이 추가되면 바로 보입니다)")
        dc1, dc2 = st.columns(2)
        dc1.metric("행", f"{len(df):,}"); dc2.metric("열", len(df.columns))
        smart_table(df.head(20), width="stretch")
        st.caption(f"열 목록: {', '.join(map(str, df.columns))}")
        kind = st.radio("만들 방식", ["두 열 사칙연산", "조건 열 (예: 기온≥33)", "그룹별 집계"])
        if kind == "두 열 사칙연산":
            if not num_cols: st.warning("숫자형 변수가 필요합니다.")
            else:
                c1, c2, c3 = st.columns(3)
                a = c1.selectbox("열 A", num_cols, key="d_a")
                op = c2.selectbox("연산", ["+", "-", "×", "÷"], key="d_op")
                b = c3.selectbox("열 B", num_cols, key="d_b")
                nm = st.text_input("새 열 이름", value=f"{a}_{op}_{b}")
                if st.button("새 열 만들기"):
                    d = st.session_state.df.copy()
                    d[nm] = {"+": d[a]+d[b], "-": d[a]-d[b], "×": d[a]*d[b],
                             "÷": d[a]/d[b].replace(0, np.nan)}[op]
                    st.session_state.df = d; st.success(f"'{nm}' 생성"); smart_table(d.head(), width="stretch")
        elif kind.startswith("조건"):
            if not num_cols: st.warning("숫자형 변수가 필요합니다.")
            else:
                c1, c2, c3 = st.columns(3)
                col = c1.selectbox("기준 열", num_cols, key="d_col")
                cond = c2.selectbox("조건", ["≥", ">", "≤", "<", "="], key="d_cond")
                thr = c3.number_input("임계값", value=float(round(df[col].mean(), 1)))
                nm = st.text_input("새 열 이름", value=f"{col}_{cond}{thr}")
                st.caption("조건 만족 시 1, 아니면 0. 이후 '그룹별 집계 → 합계'로 '해당일 수'를 구할 수 있어요.")
                if st.button("조건 열 만들기"):
                    d = st.session_state.df.copy(); s = d[col]
                    flag = {"≥": s >= thr, ">": s > thr, "≤": s <= thr, "<": s < thr, "=": s == thr}[cond]
                    d[nm] = flag.astype(int); st.session_state.df = d
                    st.success(f"'{nm}' 생성 (1의 개수: {int(d[nm].sum())})"); smart_table(d.head(), width="stretch")
        else:
            if not num_cols: st.warning("숫자형 값 열이 필요합니다.")
            else:
                c1, c2, c3 = st.columns(3)
                g = c1.selectbox("그룹 열 (예: 연도)", df.columns.tolist(), key="agg_g")
                v = c2.selectbox("값 열", num_cols, key="agg_v")
                f = c3.selectbox("집계", ["합계", "평균", "최대", "최소", "개수"], key="agg_f")
                if keep_running("agg", "집계표 만들기"):
                    fmap = {"합계": "sum", "평균": "mean", "최대": "max", "최소": "min", "개수": "count"}
                    agg = df.groupby(g)[v].agg(fmap[f]).round(3).reset_index()
                    agg.columns = [g, f"{v}_{f}"]
                    smart_table(agg, width="stretch")
                    dl_table(agg, f"{g}별 {v} {f}", "aggregate1", "aggregate")
                    st.session_state.files[f"집계_{g}_{v}_{f}"] = agg
                    st.info("사이드바 '분석할 데이터 선택'에서 이 집계표를 고를 수 있어요.")

    # ---------- 상관 ----------
    if tab_corr.on:
        st.subheader("상관분석 & 히트맵")
        with st.expander("ℹ️ 이 분석이 뭔가요?"): st.markdown(EXPLAIN["corr"])
        if len(num_cols) < 2: st.warning("숫자형 변수가 2개 이상 필요합니다.")
        else:
            c1, c2 = st.columns([3, 1])
            sel = c1.multiselect("분석할 변수 선택", num_cols, default=num_cols)
            cmethod = c2.selectbox("상관계수", ["Pearson(선형)", "Spearman(순위)"],
                                   help="정규분포가 아니거나 순위·등급 자료면 Spearman을 쓰세요.")
            if len(sel) >= 2:
                corr = df[sel].corr(method="pearson" if cmethod.startswith("Pearson") else "spearman")
                smart_table(corr.round(3), width="stretch")
                # ---- 유의성 별표 표기 (논문 관행) ----
                _meth = stats.pearsonr if cmethod.startswith("Pearson") else stats.spearmanr
                star_tbl = pd.DataFrame(index=corr.index, columns=corr.columns, dtype=object)
                n_pairs = 0
                for a in sel:
                    for b in sel:
                        if a == b:
                            star_tbl.loc[a, b] = "1"
                            continue
                        sub = df[[a, b]].dropna()
                        if len(sub) < 3:
                            star_tbl.loc[a, b] = "-"
                            continue
                        try:
                            r_, p_ = _meth(sub[a], sub[b])
                            mark = "***" if p_ < .001 else "**" if p_ < .01 else "*" if p_ < .05 else ""
                            star_tbl.loc[a, b] = f"{r_:.3f}{mark}"
                            if mark and a < b: n_pairs += 1
                        except Exception:
                            star_tbl.loc[a, b] = "-"
                st.markdown("###### 📋 상관계수표 (유의성 별표 포함)")
                star_out = star_tbl.reset_index().rename(columns={"index": "변수"})
                smart_table(star_out, width="stretch")
                st.caption("* p<0.05, ** p<0.01, *** p<0.001 "
                           f"／ 유의한 상관을 보인 변수쌍 {n_pairs}개")
                dl_table(star_out, f"{'Pearson' if cmethod.startswith('Pearson') else 'Spearman'} 상관분석표",
                         "corrstar", "상관분석표")
                txt = interpret_corr(corr, sel); st.info("💡 " + txt)
                fig, ax = plt.subplots(figsize=(1.2*len(sel), 1.0*len(sel)))
                from matplotlib.colors import LinearSegmentedColormap
                _corr_cmap = LinearSegmentedColormap.from_list(
                    "smart_corr", ["#C96767", "#F5E2E2", "#FFFFFF", "#D9EAF7", "#3D6F9F"])
                sns.heatmap(corr, annot=True, fmt=".2f", cmap=_corr_cmap, center=0, ax=ax,
                            linewidths=.6, linecolor="white", square=True,
                            cbar_kws={"shrink": .82})
                deco(ax, "상관계수 히트맵", ylabel_top=False); ax.grid(False)
                for _sp in ax.spines.values():
                    _sp.set_visible(False)
                try:
                    _cb = ax.collections[0].colorbar
                    if _cb is not None:
                        _cb.outline.set_visible(False)
                except Exception:
                    pass
                png = fig_to_png(fig)
                st.download_button("🖼️ 히트맵 다운로드", png, "heatmap.png", "image/png")
                out = corr.round(3).reset_index().rename(columns={"index": "변수"})
                dl_table(out, "상관분석 결과표", "corr2", "corr")
                log_action(f"상관분석: {len(sel)}개 변수")
                report_capture("cap_corr", "상관분석", txt, out, png)
        report_button("cap_corr")

    # ---------- 분산분석 ----------
    if tab_anova.on:
        st.subheader("분산분석(ANOVA)")
        with st.expander("ℹ️ 이 분석이 뭔가요?"): st.markdown(EXPLAIN["anova"])
        mode = st.radio("분석 방식", ["일원배치 (요인 1개)", "이원배치 (요인 2개 + 상호작용)",
                                   "🌾 분할구법 (Split-plot)",
                                   "🔁 반복측정 (같은 개체 시기별 조사)", "🎚️ 공분산분석(ANCOVA)",
                                   "📊 여러 형질 한 표에 (요약표)"])
        if mode.startswith("📊"):
            st.caption("여러 측정 항목을 한 번에 분석해, 논문 양식처럼 **하나의 표**로 만듭니다. "
                       "각 수치 옆에 유의성 문자(a, b, c)가 위첨자로 붙습니다.")
            if not cat_cols or not num_cols:
                st.warning("그룹(범주형)과 측정(숫자형) 변수가 필요합니다.")
            else:
                c1, c2 = st.columns(2)
                gc = c1.selectbox("처리구(그룹)", cat_cols, key="ms_g")
                ph = c2.selectbox("사후검정", ["Tukey HSD", "던컨(Duncan)", "Bonferroni"], key="ms_ph")
                traits = st.multiselect("측정 항목(형질) 선택 — 여러 개", num_cols,
                                        default=num_cols, key="ms_t")
                c3, c4 = st.columns(2)
                dec = c3.selectbox("소수점 자릿수", [0, 1, 2, 3], index=1, key="ms_dec")
                show_ns = c4.checkbox("유의차 없으면 문자 생략", value=True,
                                      help="ANOVA에서 p≥0.05인 항목은 문자를 붙이지 않습니다.")
                if traits and keep_running("summary", "요약표 만들기"):
                    rows, notes = [], []
                    groups_order = None
                    result = {}
                    # 여러 형질 모드도 표만 만들고 끝내지 않고, 항목별 그래프를 함께 만든다.
                    # 각 그래프에는 평균값 + SD/SE 오차막대 + 유의성 문자(a,b,c)를 표시한다.
                    plot_records = []
                    for tr in traits:
                        data = df[[gc, tr]].dropna()
                        if data[gc].nunique() < 2: continue
                        try:
                            model = ols(safe_formula(tr, [gc]), data=data).fit()
                            pval = sm.stats.anova_lm(model, typ=2)["PR(>F)"].iloc[0]
                        except Exception:
                            pval = np.nan
                        means = data.groupby(gc)[tr].mean()
                        if groups_order is None:
                            groups_order = list(df[gc].dropna().unique())
                        letters = {}
                        if not (show_ns and (np.isnan(pval) or pval >= .05)):
                            try:
                                _phres = posthoc_from_model(model, data, gc, ph)
                                ns = _phres["not_sig"]
                                order = means.sort_values(ascending=False).index.tolist()
                                letters = compact_letter_display(order, ns)
                            except Exception:
                                letters = {}
                        col = {}
                        for g in groups_order:
                            if g in means.index:
                                v = f"{means[g]:.{dec}f}"
                                if letters.get(g): v += "^" + letters[g]
                                col[g] = v
                            else:
                                col[g] = "-"
                        result[tr] = col
                        notes.append({"항목": tr, "p-value": ("-" if np.isnan(pval) else round(pval, 4)),
                                      "유의성": "-" if np.isnan(pval) else ("**" if pval < .01 else "*" if pval < .05 else "n.s.")})

                        # ---- 여러 형질 모드: 항목별 화면 그래프 ----
                        try:
                            _stats = data.groupby(gc)[tr].agg(["mean", "std", "count"])
                            _order = [g for g in groups_order if g in _stats.index]
                            _use_se = st.session_state.get("err_type", "표준편차(SD)").startswith("표준오차")
                            _err = (_stats["std"] / np.sqrt(_stats["count"])) if _use_se else _stats["std"]
                            _elabel = "표준오차" if _use_se else "표준편차"
                            _m = _stats.loc[_order]
                            _e = _err.loc[_order].fillna(0)
                            _fig, _ax = plt.subplots(figsize=(min(6.2, max(4.8, len(_order) * 0.90)),
                                                             min(4.0, figsize()[1])))
                            _ax.bar(_order, _m["mean"], yerr=_e, capsize=4,
                                    color=bar_colors(values=_m["mean"].tolist()),
                                    edgecolor="none", width=.62 if pretty_on() else .8,
                                    error_kw={"ecolor": "#5a6067", "elinewidth": 1.1})
                            bar_value_sig_labels(
                                _ax, range(len(_order)), _m["mean"].tolist(), _e.tolist(),
                                [letters.get(g, "") for g in _order], dec=dec)
                            _ax.margins(y=.20 if pretty_on() else .10)
                            _ax.set_ylabel(tr); _ax.set_xlabel(gc)
                            deco(_ax, f"{gc}별 {tr} (평균±{_elabel}, {ph})")
                            _png = fig_to_png(_fig, show=False)
                            plot_records.append({"trait": tr, "png": _png, "error_label": _elabel})
                        except Exception:
                            # 표 계산은 정상인데 특정 그래프만 실패한 경우 전체 요약분석을 중단하지 않는다.
                            pass
                    summary = pd.DataFrame(result)
                    summary.index.name = gc
                    summary = summary.reset_index()
                    st.markdown("#### 처리구별 형질 요약표")
                    smart_table(sup_df(summary), width="stretch")
                    st.markdown("#### 항목별 분산분석 유의성")
                    ndf = pd.DataFrame(notes)
                    smart_table(ndf, width="stretch")

                    if plot_records:
                        st.markdown("#### 📊 항목별 그래프")
                        st.caption("막대는 처리 평균, 오차막대는 선택한 SD/SE이며, 숫자 뒤 a·b·c는 사후검정 유의성 그룹입니다. "
                                   "ANOVA가 유의하지 않은 항목은 문자를 생략합니다.")
                        _cols = st.columns(2)
                        for _i, _pr in enumerate(plot_records):
                            with _cols[_i % 2]:
                                st.markdown(f"**{_pr['trait']}**")
                                st.image(_pr["png"], width="stretch")
                                st.download_button(
                                    f"🖼️ {_pr['trait']} 그래프 PNG", _pr["png"],
                                    f"anova_multi_{_i+1}.png", "image/png",
                                    key=f"ms_plot_dl_{_i}", width="stretch")

                    txt = (f"{gc}에 따라 {len(traits)}개 형질을 {ph}로 분석했습니다. "
                           "같은 문자를 가진 처리구끼리는 통계적 차이가 없습니다. "
                           f"(유의: {(ndf['유의성'] != 'n.s.').sum()}개 항목)")
                    st.info("💡 " + txt)
                    dl_table(summary, f"{gc}별 생산력 검정 결과 ({ph})", "summary_table3", "summary_table")
                    log_action(f"요약표 생성: {gc} × {len(traits)}개 형질 ({ph})")

                    _ms_blocks = [
                        {"caption": "처리구별 형질 요약표", "table": summary},
                        {"caption": "항목별 분산분석 유의성", "table": ndf},
                    ]
                    _ms_blocks += [
                        {"caption": f"{_pr['trait']} 처리구별 평균±{_pr['error_label']}",
                         "image": _pr["png"]}
                        for _pr in plot_records
                    ]
                    report_capture("cap_ms", f"{gc}별 형질 요약표", text=txt, blocks=_ms_blocks)
                    ai_interpret_button("ms", f"{gc}별 여러 형질 요약표", summary,
                                        "각 수치 옆 a,b,c는 처리 간 유의성 그룹입니다.",
                                        capture_slot="cap_ms")
                report_button("cap_ms")
        elif mode.startswith("일원배치"):
            if not cat_cols or not num_cols: st.warning("그룹(범주형)과 측정(숫자형) 변수가 필요합니다.")
            else:
                c1, c2, c3 = st.columns(3)
                gc = c1.selectbox("처리구(그룹)", cat_cols, key="aov_g")
                vc = c2.selectbox("측정값", num_cols, key="aov_v")
                ph = c3.selectbox("사후검정", ["Tukey HSD", "던컨(Duncan)", "Bonferroni",
                                            "던넷(Dunnett, 대조구 대비)"], key="aov_ph")
                if ph.startswith("던넷"):
                    _lvls = df[gc].dropna().astype(str).unique().tolist()
                    _ci = guess_idx(_lvls, ["대조", "관행", "무처리", "control", "CK"])
                    st.selectbox("대조구(비교 기준)", _lvls, index=_ci, key="dunnett_ctrl")
                    st.caption("던넷 검정은 **모든 처리를 대조구와만** 비교합니다. "
                               "신품종 vs 대비품종처럼 기준이 뚜렷할 때 검정력이 가장 높습니다.")
                blk_opts = ["(없음 · 완전임의배치)"] + [c for c in df.columns if c not in (gc, vc)]
                blk = st.selectbox("반복(블록) 열 — 난괴법이면 반드시 선택", blk_opts,
                                   index=guess_idx(blk_opts, ["반복", "블록", "구역", "block", "rep"]),
                                   key="aov_b")
                if ph == "던컨(Duncan)":
                    st.caption("⚠️ 던컨(DMRT)은 검정력이 높지만 위양성(제1종 오류)을 통제하지 못합니다. "
                               "논문 투고 시에는 Tukey HSD가 더 안전합니다.")
                if blk.startswith("(없음"):
                    st.caption("💡 포장시험에서 반복(블록)을 두었다면 반드시 지정하세요. "
                               "지정하지 않으면 블록 간 변이가 오차에 섞여 처리 효과를 놓칠 수 있습니다.")
                if keep_running("anova", "ANOVA 분석 실행"):
                    use_blk = not blk.startswith("(없음")
                    cols_need = [gc, vc] + ([blk] if use_blk else [])
                    data = df[cols_need].dropna()
                    # ---- 사전 검증 (분석 가능한 자료인지) ----
                    ok_to_run, msgs = validate_anova_data(data, gc, vc)
                    for m in msgs: st.warning(m)
                    if not ok_to_run:
                        st.stop()
                    st.markdown("#### 1) 가정 검정")
                    nrows, nok = [], True
                    for g in data[gc].unique():
                        v = data[data[gc] == g][vc]
                        if len(v) >= 3:
                            w, p = stats.shapiro(v)
                            nrows.append({"그룹": g, "W": round(w, 3), "p": round(p, 3),
                                          "정규성": "만족" if p >= .05 else "위배"})
                            if p < .05: nok = False
                    lp = np.nan
                    try:
                        samples = [data[data[gc] == g][vc] for g in data[gc].unique()]
                        samples = [s for s in samples if len(s) >= 2]
                        if len(samples) >= 2:
                            ls, lp = stats.levene(*samples)
                    except Exception:
                        lp = np.nan
                    if nrows:
                        smart_table(pd.DataFrame(nrows), width="stretch")
                    else:
                        st.caption("각 처리구의 반복이 3개 미만이라 정규성 검정을 생략했습니다.")
                    if not np.isnan(lp):
                        st.write(f"등분산(Levene): 통계량={ls:.3f}, p={lp:.3f} → {'만족' if lp >= .05 else '위배'}")
                    else:
                        st.caption("반복이 부족해 등분산 검정을 생략했습니다.")
                    if nok and (np.isnan(lp) or lp >= .05):
                        st.success("가정을 모두 만족합니다. ANOVA 결과를 신뢰할 수 있어요.")
                    else:
                        st.warning("가정이 일부 위배되었습니다. 아래 **비모수 검정 결과**를 함께 확인하세요.")
                        # ---- 비모수 자동 전환 (원클릭) ----
                        try:
                            _grps = [g[vc].values for _, g in data.groupby(gc)]
                            if len(_grps) == 2:
                                _st, _p = stats.mannwhitneyu(*_grps)
                                _nm = "Mann-Whitney U 검정"
                            else:
                                _st, _p = stats.kruskal(*_grps)
                                _nm = "Kruskal-Wallis 검정"
                            _med = data.groupby(gc)[vc].median().round(rnd())
                            with st.expander(f"🧪 비모수 대안: {_nm} 결과 보기", expanded=True):
                                k1, k2 = st.columns(2)
                                k1.metric("검정 통계량", f"{_st:.3f}")
                                k2.metric("p-value", f"{_p:.4f}",
                                          "유의함" if _p < .05 else "유의하지 않음")
                                smart_table(pd.DataFrame({gc: _med.index.astype(str),
                                                           f"{vc} 중앙값": _med.values}),
                                             width="stretch")
                                st.caption("정규성·등분산 가정을 쓰지 않는 방법입니다. "
                                           "평균 대신 **중앙값**으로 비교합니다. "
                                           "아래 ANOVA 결과와 결론이 다르면 비모수 결과를 우선하세요.")
                        except Exception:
                            st.caption("비모수 검정을 자동 수행하지 못했습니다. '비모수검정' 탭을 이용하세요.")
                    st.markdown("#### 2) 분산분석 결과")
                    formula = safe_formula(vc, [gc] + ([blk] if use_blk else []))
                    model = ols(formula, data=data).fit()
                    aov = sm.stats.anova_lm(model, typ=2)
                    smart_table(aov.round(4), width="stretch")
                    st.caption("설계: " + ("**난괴법(RCBD)** — 반복(블록) 효과를 모형에 포함했습니다."
                                          if use_blk else "**완전임의배치(CRD)** — 블록 없음"))
                    tkey = f"C({q_ref(gc)})"
                    pval = aov.loc[tkey, "PR(>F)"] if tkey in aov.index else aov["PR(>F)"].iloc[0]
                    if use_blk:
                        bkey = f"C({q_ref(blk)})"
                        if bkey in aov.index:
                            bp = aov.loc[bkey, "PR(>F)"]
                            st.caption(f"블록('{blk}') 효과 p = {bp:.4f} → "
                                       + ("블록 간 차이가 있어 난괴법이 적절했습니다."
                                          if bp < .05 else "블록 간 차이는 뚜렷하지 않았습니다."))
                    _ctrl_for_ph = st.session_state.get("dunnett_ctrl") if ph.startswith("던넷") else None
                    _phres = posthoc_from_model(model, data, gc, ph, control=_ctrl_for_ph)
                    ns = _phres["not_sig"]
                    means = data.groupby(gc)[vc].agg(["mean", "std", "count"])
                    order = means.sort_values("mean", ascending=False).index.tolist()
                    # 전체 ANOVA가 유의하지 않으면 모든 처리에 'a'를 붙이지 않는다.
                    # 화면 메시지(유의차 없음)와 그래프/표가 서로 모순되어 보이는 것을 방지한다.
                    letters = ({} if float(pval) >= 0.05 or ph.startswith("던넷")
                               else compact_letter_display(order, ns))
                    # ---- CV% · LSD (시험연구보고서 필수 지표) ----
                    st.markdown("#### 3) 시험 정밀도 지표")
                    ci = calc_cv_lsd(model, data, gc, vc)
                    m1, m2, m3 = st.columns(3)
                    m1.metric("CV (변이계수)", f"{ci['CV']:.1f} %", cv_grade(ci["CV"]))
                    m2.metric("LSD (p<0.05)", f"{ci['LSD']:.2f}")
                    m3.metric("오차평균제곱(MSE)", f"{ci['MSE']:.2f}")
                    st.caption("CV%는 시험의 정밀도를 나타냅니다(포장시험 10~20% 양호, 20% 초과 시 재검토). "
                               f"두 처리 평균의 차이가 LSD({ci['LSD']:.2f})보다 크면 유의한 차이로 봅니다.")
                    if ci["CV"] > 30:
                        st.warning("⚠️ CV%가 30%를 넘습니다. 포장 불균일·조사 오차·이상값을 점검해 보세요.")
                    if ph.startswith("던넷"):
                        _c0 = st.session_state.get("dunnett_ctrl", "")
                        txt = (f"[{ph}] " + ("처리구 간 유의한 차이가 있습니다"
                                             if pval < .05 else "처리구 간 유의한 차이가 없습니다")
                               + f" (p = {pval:.4f}). 대조구 '{_c0}'와의 개별 비교는 아래 표를 보세요.")
                    else:
                        txt = f"[{ph}] " + interpret_anova(pval, letters)
                    st.info("💡 " + txt)
                    # ---- 던넷: 적합모형 기반 대조구 대비 전용 표 ----
                    if ph.startswith("던넷"):
                        _c = _phres.get("control")
                        dn_df = _phres.get("table", pd.DataFrame()).copy()
                        if not dn_df.empty:
                            for _cnum in ["대조구 평균(보정)", "처리 평균(보정)", "평균 차이",
                                          "t 통계량", "p(동시보정)", "95% 동시CI 하한", "95% 동시CI 상한"]:
                                if _cnum in dn_df.columns:
                                    dn_df[_cnum] = pd.to_numeric(dn_df[_cnum], errors="coerce").round(rnd())
                            st.markdown(f"###### 던넷 검정: '{_c}' 대비 모형 기반 비교")
                            smart_table(dn_df, width="stretch", hide_index=True)
                            st.caption("ANOVA와 같은 적합모형의 잔차·블록 보정을 사용한 동시비교입니다. "
                                       "95% 동시신뢰구간이 0을 포함하지 않으면 대조구와 유의한 차이가 있습니다.")
                            _sig = dn_df[dn_df["판정"].astype(str).str.startswith("유의")]["처리구"].tolist()
                            st.caption(f"대조구 '{_c}'와 유의한 차이를 보인 처리: "
                                       + (", ".join(map(str, _sig)) if _sig else "없음"))
                            dl_table(dn_df, f"{_c} 대비 모형 기반 던넷 검정", "dunnett1", "dunnett")
                        else:
                            st.warning("던넷 비교표를 만들 수 없습니다. 대조구와 처리 수준을 확인하세요.")
                    if ph.startswith("던넷"):
                        letters = {}   # 던넷은 문자(a,b,c) 표기를 쓰지 않음
                        st.info("ℹ️ 던넷 검정은 대조구와의 비교만 수행하므로 "
                                "유의성 문자(a, b, c)는 표기하지 않습니다. 위 표를 사용하세요.")
                    res = means.copy(); res["유의성"] = [letters.get(g, "") for g in res.index]
                    res = res.rename(columns={"mean": "평균", "std": "표준편차", "count": "n"}).round(rnd()).reset_index()
                    # 논문용 '평균±표준오차' 결합 표기 컬럼 추가
                    _se = (means["std"] / np.sqrt(means["count"])).round(rnd())
                    res["평균±SE"] = [f"{means.loc[g,'mean']:.{rnd()}f}±{_se.loc[g]:.{rnd()}f}"
                                     + (letters.get(g, "") and f"^{letters.get(g,'')}")
                                     for g in res[gc]]
                    smart_table(sup_display(res), width="stretch")
                    use_se = st.session_state.get("err_type", "표준편차(SD)").startswith("표준오차")
                    err = (means["std"] / np.sqrt(means["count"])) if use_se else means["std"]
                    elabel = "표준오차" if use_se else "표준편차"
                    fig, ax = plt.subplots(figsize=(min(6.6, max(5.0, len(order)*0.92)), min(4.2, figsize()[1])))
                    m = means.loc[order]; e_ = err.loc[order]
                    ax.bar(order, m["mean"], yerr=e_, capsize=4,
                           color=bar_colors(values=m["mean"].tolist()),
                           edgecolor="none", width=.62 if pretty_on() else .8,
                           error_kw={"ecolor": "#5a6067", "elinewidth": 1.1})
                    bar_value_sig_labels(
                        ax, range(len(order)), m["mean"].tolist(), e_.tolist(),
                        [letters.get(g, "") for g in order], dec=rnd())
                    ax.margins(y=.20 if pretty_on() else .07)
                    ax.set_ylabel(vc); ax.set_xlabel(gc)
                    deco(ax, f"{gc}별 {vc} (평균±{elabel}, {ph})")
                    png = fig_to_png(fig)
                    st.download_button("🖼️ 그래프 다운로드", png, "anova.png", "image/png")
                    # ---- 표 각주 자동 생성 ----
                    _phname = {"Tukey HSD": "Tukey의 HSD 검정", "던컨(Duncan)": "던컨의 다중검정(DMRT)",
                               "Bonferroni": "Bonferroni 보정 t-검정"}.get(ph, ph)
                    if ph.startswith("던넷"):
                        # ⑭ 던넷은 문자를 쓰지 않으므로 각주도 다르게
                        _c1 = st.session_state.get("dunnett_ctrl", "대조구")
                        footnote = (f"* 던넷 검정으로 대조구 '{_c1}'와 각 처리를 비교하였음"
                                    "(다중비교 보정 p값, 5% 수준).\n"
                                    f"* CV(%) = {ci['CV']:.1f}, 평균±{elabel} "
                                    f"(n = {int(means['count'].min())})")
                    else:
                        footnote = (f"* 같은 열의 다른 문자는 {_phname}으로 5% 수준에서 "
                                    "유의차가 있음을 나타냄.\n"
                                    f"* CV(%) = {ci['CV']:.1f}, LSD(0.05) = {ci['LSD']:.2f}, "
                                    f"평균±{elabel} (n = {int(means['count'].min())})")
                    st.markdown("###### 📋 표 각주 (복사해서 표 아래에 붙이세요)")
                    st.code(footnote, language=None)
                    dl_table(res, f"{gc}별 {vc} 분산분석 ({ph})", "anova4", "anova")
                    log_action(f"일원배치 ANOVA: {gc} × {vc} ({ph})")
                    _anova_ctx = build_anova_context(
                        design=("난괴법(RCBD)" if use_blk else "완전임의배치(CRD)"),
                        trt=gc, blk=(blk if use_blk else None), yv=vc,
                        group_stats=res, anova_table=aov.reset_index(),
                        p_treatment=pval,
                        p_block=(float(aov.loc[f"C({q_ref(blk)})", "PR(>F)"])
                                 if use_blk and f"C({q_ref(blk)})" in aov.index else None),
                        cv=ci.get("CV"), lsd=ci.get("LSD"), mse=ci.get("MSE"),
                        df_resid=ci.get("dfe"), posthoc=ph,
                        letters=(None if ph.startswith("던넷") else letters),
                        dunnett=(dn_df if (ph.startswith("던넷") and "dn_df" in dir()) else None),
                        assumptions={"normality": nrows, "levene_p": (None if np.isnan(lp) else lp)},
                        n_missing=int(len(df) - len(data)),
                        cautions=(["던넷 검정은 대조구 대비 비교만 수행하며 문자(a,b,c)를 쓰지 않음"]
                                  if ph.startswith("던넷") else []))
                    ai_interpret_advanced("anova", f"{gc}별 {vc} 분산분석({ph})", res,
                                          "유의성 문자가 같으면 처리 간 차이가 없다는 의미입니다.",
                                          context=_anova_ctx, capture_slot="cap_anova")
                    _rep_txt = report_sentence_anova(gc, vc, pval, means, letters, ci, ph)
                    st.markdown("###### 📋 보고서용 결과 문장")
                    st.code(_rep_txt, language=None)
                    report_capture("cap_anova", f"{gc}별 {vc} 분산분석", None,
                                   blocks=[{"text": _rep_txt},
                                           {"caption": f"{gc}별 {vc} 분산분석 ({ph})", "table": res,
                                            "image": png},
                                           {"text": footnote}])
                report_button("cap_anova")
        elif mode.startswith("이원배치"):
            if len(cat_cols) < 2 or not num_cols: st.warning("범주형 변수 2개와 측정값 1개가 필요합니다.")
            else:
                c1, c2, c3 = st.columns(3)
                f1 = c1.selectbox("요인 A", cat_cols, key="tw_a")
                f2 = c2.selectbox("요인 B", [c for c in cat_cols if c != f1], key="tw_b")
                yv = c3.selectbox("측정값", num_cols, key="tw_y")
                if keep_running("twoway", "이원배치 ANOVA 실행"):
                    data = df[[f1, f2, yv]].dropna()
                    model = ols(safe_formula(yv, [f1, f2], interactions=[(f1, f2)]),
                                data=data).fit()
                    aov = sm.stats.anova_lm(model, typ=2); out = aov.round(4)
                    smart_table(out, width="stretch")
                    terms = {f"C({q_ref(f1)})": f"요인A({f1})",
                             f"C({q_ref(f2)})": f"요인B({f2})",
                             f"C({q_ref(f1)}):C({q_ref(f2)})": "상호작용"}
                    txt = " / ".join(f"**{lab}**: {'유의' if aov.loc[k,'PR(>F)']<.05 else '비유의'}(p={aov.loc[k,'PR(>F)']:.3f})"
                                     for k, lab in terms.items() if k in aov.index)
                    st.info("💡 " + txt)
                    fig, ax = plt.subplots(figsize=figsize())
                    for lv in data[f2].unique():
                        s = data[data[f2] == lv].groupby(f1)[yv].mean()
                        ax.plot(s.index, s.values, marker="o", label=f"{f2}={lv}")
                    ax.set_xlabel(f1); ax.set_ylabel(yv); deco(ax, "상호작용 그래프"); ax.legend()
                    png = fig_to_png(fig)
                    out2 = out.reset_index().rename(columns={"index": "요인"})
                    dl_table(out2, f"{yv} 이원배치 분산분석", "twoway5", "twoway")
                    report_capture("cap_tw", f"{yv} 이원배치 분산분석", txt, out2, png)
                report_button("cap_tw")

        # ---------- 반복측정 ANOVA ----------
        # ---------- 분할구법 (Split-plot) ----------
        elif mode.startswith("🌾"):
            with st.expander("ℹ️ 분할구법이란?"):
                st.markdown("""
**관수·경운·재배법처럼 작은 구역에 나누기 어려운 요인**이 있을 때 쓰는 설계입니다.

- **주구(主區, Main plot)**: 큰 구역에 배치하는 요인 (예: 관수 방법, 경운 방법, 재배 양식)
- **세구(細區, Sub plot)**: 주구를 쪼개서 배치하는 요인 (예: 품종, 시비량)

**왜 따로 분석해야 하나요?**
주구와 세구는 **오차의 크기가 다릅니다.** 주구는 큰 구역이라 오차가 크고, 세구는 작아서 오차가 작습니다.
일반 이원배치로 분석하면 주구 효과가 **실제보다 과대평가**됩니다.
분할구 분석은 **주구오차(반복×주구)** 와 **세구오차**를 분리해 각각 올바른 검정을 합니다.

**필요한 열**: 반복(블록) · 주구 요인 · 세구 요인 · 측정값
""")
            if len(cat_cols) < 3 or not num_cols:
                st.warning("반복·주구·세구 3개의 범주형 열과 측정값 1개가 필요합니다.")
            else:
                allc = df.columns.tolist()
                c1, c2 = st.columns(2)
                rep_c = c1.selectbox("반복(블록) 열", cat_cols,
                                     index=guess_idx(cat_cols, ["반복", "블록", "block", "rep"]), key="sp_r")
                yv = c2.selectbox("측정값", num_cols, key="sp_y")
                c3, c4 = st.columns(2)
                main_c = c3.selectbox("주구(큰 구역) 요인", [c for c in cat_cols if c != rep_c],
                                      index=guess_idx([c for c in cat_cols if c != rep_c],
                                                      ["관수", "경운", "재배", "처리"]), key="sp_m")
                sub_opts = [c for c in cat_cols if c not in (rep_c, main_c)]
                sub_c = c4.selectbox("세구(작은 구역) 요인", sub_opts,
                                     index=guess_idx(sub_opts, ["품종", "계통", "시비"]), key="sp_s")
                if keep_running("splitplot", "분할구 분산분석 실행"):
                    data = df[[rep_c, main_c, sub_c, yv]].dropna()
                    ok_sp, msgs = validate_anova_data(data, main_c, yv)
                    for m in msgs: st.warning(m)
                    if not ok_sp: st.stop()
                    try:
                        f_full = safe_formula(yv, [rep_c, main_c, sub_c],
                                              interactions=[(rep_c, main_c),
                                                            (main_c, sub_c)])
                        mfull = ols(f_full, data=data).fit()
                        a = sm.stats.anova_lm(mfull, typ=2)
                        k_rep, k_main = f"C({q_ref(rep_c)})", f"C({q_ref(main_c)})"
                        k_erra = f"C({q_ref(rep_c)}):C({q_ref(main_c)})"
                        k_sub = f"C({q_ref(sub_c)})"
                        k_int = f"C({q_ref(main_c)}):C({q_ref(sub_c)})"
                        ms = lambda k: a.loc[k, "sum_sq"] / a.loc[k, "df"]
                        ms_errb = a.loc["Residual", "sum_sq"] / a.loc["Residual", "df"]
                        # 주구는 주구오차로, 세구·상호작용은 세구오차로 검정
                        F_main = ms(k_main) / ms(k_erra)
                        p_main = 1 - stats.f.cdf(F_main, a.loc[k_main, "df"], a.loc[k_erra, "df"])
                        F_sub = ms(k_sub) / ms_errb
                        p_sub = 1 - stats.f.cdf(F_sub, a.loc[k_sub, "df"], a.loc["Residual", "df"])
                        F_int = ms(k_int) / ms_errb
                        p_int = 1 - stats.f.cdf(F_int, a.loc[k_int, "df"], a.loc["Residual", "df"])
                        rows = [
                            {"요인": f"반복({rep_c})", "자유도": int(a.loc[k_rep, "df"]),
                             "제곱합": round(a.loc[k_rep, "sum_sq"], 3), "평균제곱": round(ms(k_rep), 3),
                             "F": "-", "p": "-"},
                            {"요인": f"주구: {main_c}", "자유도": int(a.loc[k_main, "df"]),
                             "제곱합": round(a.loc[k_main, "sum_sq"], 3), "평균제곱": round(ms(k_main), 3),
                             "F": f"{F_main:.3f}", "p": f"{p_main:.4f}"},
                            {"요인": "주구오차(Ea)", "자유도": int(a.loc[k_erra, "df"]),
                             "제곱합": round(a.loc[k_erra, "sum_sq"], 3), "평균제곱": round(ms(k_erra), 3),
                             "F": "-", "p": "-"},
                            {"요인": f"세구: {sub_c}", "자유도": int(a.loc[k_sub, "df"]),
                             "제곱합": round(a.loc[k_sub, "sum_sq"], 3), "평균제곱": round(ms(k_sub), 3),
                             "F": f"{F_sub:.3f}", "p": f"{p_sub:.4f}"},
                            {"요인": f"{main_c}×{sub_c}", "자유도": int(a.loc[k_int, "df"]),
                             "제곱합": round(a.loc[k_int, "sum_sq"], 3), "평균제곱": round(ms(k_int), 3),
                             "F": f"{F_int:.3f}", "p": f"{p_int:.4f}"},
                            {"요인": "세구오차(Eb)", "자유도": int(a.loc["Residual", "df"]),
                             "제곱합": round(a.loc["Residual", "sum_sq"], 3), "평균제곱": round(ms_errb, 3),
                             "F": "-", "p": "-"},
                        ]
                        sp_tbl = pd.DataFrame(rows)
                        st.markdown("#### 분할구 분산분석표")
                        smart_table(sp_tbl, width="stretch")
                        st.caption("주구는 **주구오차(Ea)**로, 세구와 상호작용은 **세구오차(Eb)**로 검정합니다. "
                                   "일반 이원배치로 분석하면 주구 효과가 과대평가됩니다.")
                        cva = np.sqrt(ms(k_erra)) / data[yv].mean() * 100
                        cvb = np.sqrt(ms_errb) / data[yv].mean() * 100
                        m1, m2 = st.columns(2)
                        m1.metric("CV(a) 주구", f"{cva:.1f} %", cv_grade(cva))
                        m2.metric("CV(b) 세구", f"{cvb:.1f} %", cv_grade(cvb))
                        parts = []
                        parts.append(f"주구인 '{main_c}'의 효과는 " +
                                     (f"유의하였다(p={p_main:.4f})." if p_main < .05 else f"유의하지 않았다(p={p_main:.4f})."))
                        parts.append(f"세구인 '{sub_c}'의 효과는 " +
                                     (f"유의하였다(p={p_sub:.4f})." if p_sub < .05 else f"유의하지 않았다(p={p_sub:.4f})."))
                        parts.append("두 요인의 상호작용은 " +
                                     (f"유의하여 조합별 해석이 필요하다(p={p_int:.4f})." if p_int < .05
                                      else f"유의하지 않았다(p={p_int:.4f})."))
                        txt = " ".join(parts)
                        st.info("💡 " + txt)
                        piv = data.pivot_table(index=main_c, columns=sub_c, values=yv, aggfunc="mean").round(rnd())
                        st.markdown("#### 주구 × 세구 평균")
                        smart_table(piv.reset_index(), width="stretch")
                        fig, ax = plt.subplots(figsize=figsize())
                        for sname in piv.columns:
                            ax.plot(piv.index.astype(str), piv[sname], marker="o", label=str(sname))
                        ax.set_xlabel(main_c); ax.set_ylabel(yv)
                        ax.legend(title=sub_c, fontsize=8); deco(ax, f"{main_c} × {sub_c}")
                        plt.tight_layout(); png = fig_to_png(fig)
                        st.download_button("🖼️ 그래프 다운로드", png, "splitplot.png", "image/png")
                        dl_table(sp_tbl, f"{main_c}(주구) × {sub_c}(세구) 분할구 분산분석", "sp1", "splitplot")
                        log_action(f"분할구 분산분석: {main_c} × {sub_c}")
                        ai_interpret_button("sp", f"{main_c}(주구)×{sub_c}(세구) 분할구 분산분석", sp_tbl,
                                            "주구는 주구오차로, 세구는 세구오차로 검정한 결과입니다.", capture_slot="cap_sp")
                        report_capture("cap_sp", f"{main_c}×{sub_c} 분할구 분산분석", None,
                                       blocks=[{"text": txt},
                                               {"caption": "분할구 분산분석표", "table": sp_tbl},
                                               {"caption": "주구×세구 평균", "table": piv.reset_index(),
                                                "image": png}])
                    except Exception as ex:
                        st.error(f"분할구 분석 실패: {ex}\n\n반복·주구·세구가 모두 갖춰진 균형 자료인지 확인해 주세요.")
                report_button("cap_sp")

        elif mode.startswith("🔁"):
            with st.expander("ℹ️ 반복측정 분산분석이란?"):
                st.markdown("""
**같은 개체를 여러 시기에 반복해서 조사한 자료**에 사용합니다.
(예: 같은 고추 개체의 초장을 정식 후 2·4·6·8주에 계속 측정)

**왜 일반 ANOVA를 쓰면 안 되나요?**
일반 ANOVA는 모든 관측치가 서로 **독립**이라고 가정합니다. 그런데 같은 개체를 반복 측정하면,
원래 잘 자라던 개체는 모든 시기에 계속 크게 나옵니다. 즉 관측치들이 서로 얽혀 있습니다.
이를 무시하면 **오차를 실제보다 작게 평가**해서, 차이가 없는데도 있다고 판단하기 쉽습니다.
반복측정 ANOVA는 '개체마다 원래 다른 정도'를 따로 분리해 이 문제를 해결합니다.

**자료 형태 (긴 형식)** — 한 행 = 한 개체의 한 시점

| 개체번호 | 조사시기 | 초장(cm) |
|---|---|---|
| P01 | 2주 | 18.2 |
| P01 | 4주 | 27.5 |
| P02 | 2주 | 16.9 |

**필요한 열**: 개체 번호 · 조사 시기 · 측정값
**조건**: 모든 개체가 **같은 시기에 빠짐없이** 측정되어 있어야 합니다(균형 자료).

**결과 읽는 법**: p < 0.05면 시기에 따라 측정값이 유의하게 변했다는 뜻입니다.

⚠️ 시기가 3개 이상이면 **구형성(sphericity)** 가정이 필요합니다. 결과가 p≈0.05 근처로 애매하면 해석에 주의하세요.
""")
            if len(df.columns) < 3:
                st.warning("개체·시기·측정값 3개 열이 필요합니다.")
            else:
                allc = df.columns.tolist()
                c1, c2, c3 = st.columns(3)
                subj = c1.selectbox("개체(반복 단위) 열", allc,
                                    index=guess_idx(allc, ["개체", "번호", "ID", "포기", "주"]), key="rm_s")
                within = c2.selectbox("조사 시기 열", allc,
                                      index=guess_idx(allc, ["시기", "일자", "주차", "조사", "date"]), key="rm_w")
                yv = c3.selectbox("측정값 열", num_cols, key="rm_y")
                if keep_running("rm", "반복측정 ANOVA 실행"):
                    data = df[[subj, within, yv]].dropna()
                    cnt = data.groupby([subj, within]).size()
                    _balance = validate_repeated_measure_balance(data, subj, within)
                    if _balance["duplicate_subjects"]:
                        st.error("⚠️ 같은 개체·같은 시기의 중복 자료가 있습니다: "
                                 + ", ".join(map(str, _balance["duplicate_subjects"][:8]))
                                 + ". 중복을 확인·정리한 뒤 다시 분석하세요.")
                    if _balance["bad_subjects"]:
                        st.error("⚠️ 개체마다 조사 시기 구성이 다릅니다. 반복측정 ANOVA는 "
                                 "**모든 개체가 동일한 시기 집합**을 가져야 합니다.\n\n"
                                 f"문제 개체({len(_balance['bad_subjects'])}개): "
                                 + ", ".join(map(str, _balance["bad_subjects"][:8]))
                                 + (" ..." if len(_balance["bad_subjects"]) > 8 else "")
                                 + f"\n\n전체 시기: {', '.join(map(str, _balance['expected_times']))}")
                    if _balance["ok"]:
                        try:
                            from statsmodels.stats.anova import AnovaRM
                            aovrm = AnovaRM(data, yv, subj, within=[within],
                                            aggregate_func="mean" if (cnt > 1).any() else None).fit()
                            tbl = aovrm.anova_table.round(4)
                            st.markdown("#### 반복측정 분산분석표")
                            smart_table(tbl, width="stretch")
                            p = float(tbl["Pr > F"].iloc[0])
                            txt = ("조사 시기에 따라 " +
                                   (f"측정값이 **유의하게 변화**했습니다 (p = {p:.4f} < 0.05)."
                                    if p < .05 else f"유의한 변화가 없었습니다 (p = {p:.4f})."))
                            st.info("💡 " + txt)
                            st.caption("※ 구형성(sphericity) 가정이 필요합니다. 시기 수가 3개 이상이고 "
                                       "결과가 경계값(p≈0.05)이면 해석에 주의하세요.")
                            g = data.groupby(within)[yv].agg(["mean", "std", "count"])
                            use_se = st.session_state.get("err_type", "표준편차(SD)").startswith("표준오차")
                            err = (g["std"]/np.sqrt(g["count"])) if use_se else g["std"]
                            res = g.rename(columns={"mean": "평균", "std": "표준편차", "count": "n"}).round(rnd()).reset_index()
                            smart_table(res, width="stretch")
                            fig, ax = plt.subplots(figsize=figsize())
                            ax.errorbar(g.index.astype(str), g["mean"], yerr=err, marker="o",
                                        capsize=4, color=pcolor(), lw=2)
                            ax.set_xlabel(within); ax.set_ylabel(yv)
                            deco(ax, f"시기별 {yv} 변화 (평균±{'표준오차' if use_se else '표준편차'})")
                            plt.tight_layout(); png = fig_to_png(fig)
                            st.download_button("🖼️ 그래프 다운로드", png, "rm.png", "image/png")
                            dl_table(res, f"시기별 {yv} 반복측정 분석", "rm6", "rm")
                            log_action(f"반복측정 ANOVA: {within} × {yv}")
                            ai_interpret_button("rm", f"{yv} 반복측정 분산분석", res, "시기에 따른 변화를 나타냅니다.", capture_slot="cap_rm")
                            report_capture("cap_rm", f"{yv} 반복측정 분산분석", txt, res, png)
                        except Exception as ex:
                            st.error(f"분석 실패: {ex}\n\n자료가 균형적인지(모든 개체 × 모든 시기) 확인해 주세요.")
                report_button("cap_rm")

        # ---------- ANCOVA ----------
        elif mode.startswith("🎚️"):
            with st.expander("ℹ️ 공분산분석(ANCOVA)이란?"):
                st.markdown("""
**출발선이 달랐던 것을 보정하고, 순수한 처리 효과만** 보는 방법입니다.

**이런 상황에서 씁니다**
시비 시험을 했는데, 하필 처리2 구역의 묘가 정식 당시부터 조금 더 컸다고 해봅시다.
나중에 처리2의 수량이 높게 나왔을 때, 이게 **비료 효과인지 원래 묘가 좋아서인지** 구분이 안 됩니다.
ANCOVA는 '정식 당시 묘 크기'를 **공변량**으로 넣어 그 영향을 통계적으로 걷어냅니다.

**공변량(covariate)이란?**
- 처리를 하기 **전부터** 존재하던 연속형 변수
- 예: 정식 시 묘 크기, 토양 유기물 함량, 초기 경경, 시험 전 수량
- ⚠️ 처리의 **결과로 생긴 변수**를 공변량으로 넣으면 안 됩니다(처리 효과까지 지워버림)

**결과 읽는 법**
- **원평균**: 보정 전, 실제 관측된 평균
- **보정평균(adjusted mean)**: 모든 처리구의 공변량이 **똑같았다면** 나왔을 평균 → 이 값으로 비교합니다
- 두 값의 차이가 클수록 초기 조건 차이가 컸다는 뜻입니다

**자동 점검**: 실행하면 '처리 × 공변량' 상호작용을 먼저 검사합니다.
이것이 유의하면(p<0.05) 처리구마다 공변량의 영향이 다르다는 뜻이라, ANCOVA 해석에 주의가 필요합니다.

**장점**: 공변량이 종속변수와 관련이 클수록 오차가 줄어 **검정력이 올라갑니다.**
""")
            if not cat_cols or len(num_cols) < 2:
                st.warning("그룹(범주형) 1개와 숫자형 2개(종속변수·공변량)가 필요합니다.")
            else:
                c1, c2, c3 = st.columns(3)
                gc = c1.selectbox("처리구(그룹)", cat_cols, key="an_g")
                yv = c2.selectbox("종속변수", num_cols, key="an_y")
                cov = c3.selectbox("공변량(보정할 변수)", [c for c in num_cols if c != yv], key="an_c")
                if keep_running("ancova", "ANCOVA 실행"):
                    data = df[[gc, yv, cov]].dropna()
                    m_int = ols(f"{q_ref(yv)} ~ C({q_ref(gc)}) * {q_ref(cov)}", data=data).fit()
                    a_int = sm.stats.anova_lm(m_int, typ=2)
                    ikey = [i for i in a_int.index if ":" in i]
                    p_int = a_int.loc[ikey[0], "PR(>F)"] if ikey else np.nan
                    if not np.isnan(p_int) and p_int < .05:
                        st.warning(f"⚠️ 처리×공변량 상호작용이 유의합니다 (p={p_int:.4f}). "
                                   "회귀기울기 동일 가정이 깨져 ANCOVA 해석에 주의가 필요합니다.")
                    else:
                        st.success(f"회귀기울기 동일 가정을 만족합니다 (상호작용 p={p_int:.3f}).")
                    model = ols(safe_formula(yv, [gc], covars=[cov]), data=data).fit()
                    aov = sm.stats.anova_lm(model, typ=2)
                    st.markdown("#### 공분산분석표")
                    smart_table(aov.round(4), width="stretch")
                    tkey = f"C({q_ref(gc)})"
                    p = aov.loc[tkey, "PR(>F)"] if tkey in aov.index else aov["PR(>F)"].iloc[0]
                    grand = data[cov].mean()
                    pred = data[[gc]].drop_duplicates().assign(**{cov: grand})
                    pred["보정평균"] = model.predict(pred).round(rnd())
                    raw = data.groupby(gc)[yv].mean().round(rnd()).rename("원평균").reset_index()
                    res = raw.merge(pred[[gc, "보정평균"]], on=gc)
                    st.markdown("#### 원평균 vs 보정평균")
                    smart_table(res, width="stretch")
                    st.caption(f"보정평균은 모든 처리구의 '{cov}'가 전체 평균({grand:.2f})으로 같다고 가정했을 때의 값입니다.")
                    txt = (f"'{cov}'를 보정한 결과 처리구 간 " +
                           (f"유의한 차이가 있습니다 (p = {p:.4f})." if p < .05
                            else f"유의한 차이가 없습니다 (p = {p:.4f})."))
                    st.info("💡 " + txt)
                    fig, ax = plt.subplots(figsize=figsize())
                    for lv in data[gc].unique():
                        sub = data[data[gc] == lv]
                        ax.scatter(sub[cov], sub[yv], alpha=.7, label=str(lv))
                    ax.set_xlabel(cov); ax.set_ylabel(yv); ax.legend(fontsize=8)
                    deco(ax, f"{cov} 보정 전 관계")
                    plt.tight_layout(); png = fig_to_png(fig)
                    dl_table(res, f"{yv} 공분산분석 보정평균", "ancova7", "ancova")
                    log_action(f"ANCOVA: {gc} × {yv} (공변량 {cov})")
                    report_capture("cap_ac", f"{yv} 공분산분석", txt, res, png)
                report_button("cap_ac")

    # ---------- 비모수 ----------
    if tab_np.on:
        st.subheader("비모수 검정")
        with st.expander("ℹ️ 이 분석이 뭔가요?"): st.markdown(EXPLAIN["nonparam"])
        if not cat_cols or not num_cols: st.warning("그룹(범주형)과 측정(숫자형) 변수가 필요합니다.")
        else:
            c1, c2 = st.columns(2)
            g = c1.selectbox("그룹 변수", cat_cols, key="np_g")
            v = c2.selectbox("측정값", num_cols, key="np_v")
            if keep_running("nonparam", "비모수 검정 실행"):
                data = df[[g, v]].dropna()
                grp = [data[data[g] == x][v] for x in data[g].unique()]
                if len(grp) >= 3:
                    h, p = stats.kruskal(*grp)
                    st.write(f"**Kruskal-Wallis** (그룹 {len(grp)}개)")
                    c1, c2 = st.columns(2); c1.metric("H", f"{h:.3f}"); c2.metric("p", f"{p:.4f}")
                    txt = "Kruskal-Wallis 결과 " + ("그룹 간 유의한 차이가 있습니다." if p < .05 else "차이가 없습니다.")
                elif len(grp) == 2:
                    u, p = stats.mannwhitneyu(grp[0], grp[1])
                    st.write("**Mann-Whitney U**")
                    c1, c2 = st.columns(2); c1.metric("U", f"{u:.1f}"); c2.metric("p", f"{p:.4f}")
                    txt = "Mann-Whitney U 결과 " + ("두 그룹 간 유의한 차이가 있습니다." if p < .05 else "차이가 없습니다.")
                else: txt = ""
                st.info("💡 " + txt)
                med = data.groupby(g)[v].median().round(3).reset_index(); med.columns = [g, "중앙값"]
                smart_table(med, width="stretch")
                dl_table(med, f"{g}별 {v} 중앙값(비모수)", "np8", "np")
                report_capture("cap_np", f"{g}별 {v} 비모수 검정", txt, med, None)
            report_button("cap_np")

    # ---------- PCA ----------
    if tab_pca.on:
        st.subheader("주성분분석 (PCA)")
        with st.expander("ℹ️ 이 분석이 뭔가요?"): st.markdown(EXPLAIN["pca"])
        if len(num_cols) < 3: st.warning("숫자형 변수가 3개 이상 필요합니다.")
        else:
            feats = st.multiselect("분석할 변수", num_cols, default=num_cols, key="pca_f")
            cby = st.selectbox("색상 구분(선택)", ["(없음)"] + cat_cols, key="pca_c")
            if len(feats) >= 3 and keep_running("pca", "PCA 실행"):
                data = df[feats].dropna()
                Xs = StandardScaler().fit_transform(data)
                pca = PCA(n_components=2).fit(Xs); sc_ = pca.transform(Xs); evr = pca.explained_variance_ratio_
                txt = f"주성분 2개가 원본 정보의 {evr.sum()*100:.1f}%를 설명합니다 (PC1 {evr[0]*100:.1f}%, PC2 {evr[1]*100:.1f}%)."
                st.info("💡 " + txt)
                fig, ax = plt.subplots(figsize=figsize(h=float(st.session_state.get("fig_h",4.0))+1))
                if cby != "(없음)":
                    cats = df.loc[data.index, cby]
                    for lv in cats.unique():
                        m_ = (cats == lv).values
                        ax.scatter(sc_[m_, 0], sc_[m_, 1], label=str(lv), alpha=.7)
                    ax.legend()
                else:
                    ax.scatter(sc_[:, 0], sc_[:, 1], alpha=.7, color="#6c8ebf")
                ax.set_xlabel(f"PC1 ({evr[0]*100:.1f}%)"); ax.set_ylabel(f"PC2 ({evr[1]*100:.1f}%)")
                deco(ax, "PCA 산점도"); ax.axhline(0, color="gray", lw=.5); ax.axvline(0, color="gray", lw=.5)
                png = fig_to_png(fig)
                st.download_button("🖼️ PCA 그래프", png, "pca.png", "image/png")
                load = pd.DataFrame(pca.components_.T, columns=["PC1", "PC2"], index=feats).round(3).reset_index().rename(columns={"index": "변수"})
                smart_table(load, width="stretch")
                dl_table(load, "PCA 로딩 결과", "pca9", "pca")
                report_capture("cap_pca", "주성분분석(PCA)", txt, load, png)
            report_button("cap_pca")

    # ---------- 회귀 ----------
    if tab_reg.on:
        st.subheader("회귀분석")
        with st.expander("ℹ️ 이 분석이 뭔가요?"): st.markdown(EXPLAIN["reg"])
        rt = st.radio("분석 종류", ["단순/다중 회귀분석", "로지스틱 회귀분석", "🧪 프로빗 분석 (LC50/LD50)"])
        if rt.startswith("단순"):
            if len(num_cols) < 2: st.warning("숫자형 변수가 2개 이상 필요합니다.")
            else:
                y = st.selectbox("종속변수 (Y)", num_cols, key="lin_y")
                xs = st.multiselect("독립변수 (X)", [c for c in num_cols if c != y], key="lin_x")
                if xs and keep_running("reg", "회귀분석 실행"):
                    data = df[[y]+xs].dropna()
                    model = sm.OLS(data[y], sm.add_constant(data[xs])).fit()
                    txt = f"이 모델은 '{y}'의 변동을 약 {model.rsquared*100:.1f}% 설명합니다 (R²={model.rsquared:.3f})."
                    st.info("💡 " + txt); st.text(model.summary())
                    coef = pd.DataFrame({"변수": model.params.index, "계수": model.params.values.round(4),
                                         "p-value": model.pvalues.values.round(4)})
                    if len(xs) >= 2:
                        try:
                            from statsmodels.stats.outliers_influence import variance_inflation_factor
                            Xv = sm.add_constant(data[xs])
                            vif = pd.DataFrame({"변수": xs,
                                "VIF": [round(variance_inflation_factor(Xv.values, i+1), 2) for i in range(len(xs))]})
                            st.markdown("**다중공선성 진단 (VIF)**")
                            smart_table(vif, width="stretch")
                            hi = vif[vif["VIF"] >= 10]["변수"].tolist()
                            if hi:
                                st.warning(f"⚠️ VIF가 10 이상인 변수: {', '.join(hi)} — "
                                           "변수들이 서로 너무 비슷해 계수 해석이 불안정합니다. 일부를 빼는 것이 좋습니다.")
                            else:
                                st.caption("VIF가 모두 10 미만이라 다중공선성 문제는 크지 않습니다.")
                        except Exception:
                            pass
                    png = None
                    if len(xs) == 1:
                        fig, ax = plt.subplots(figsize=figsize())
                        ax.scatter(data[xs[0]], data[y], alpha=.6, color=pcolor())
                        xl = np.linspace(data[xs[0]].min(), data[xs[0]].max(), 100)
                        ax.plot(xl, model.params.iloc[0]+model.params.iloc[1]*xl, color="red")
                        ax.set_xlabel(xs[0]); ax.set_ylabel(y); deco(ax, f"단순회귀: R²={model.rsquared:.3f}")
                        png = fig_to_png(fig)
                    # ---- 잔차 진단 (논문 부록·모형 타당성 확인용) ----
                    with st.expander("🔍 잔차 진단 (회귀모형이 적절한지 확인)"):
                        st.caption("**잔차**는 실제값과 예측값의 차이입니다. 회귀분석이 타당하려면 "
                                   "잔차가 특정 패턴 없이 0을 중심으로 고르게 흩어져야 합니다.")
                        resid = model.resid; fitted = model.fittedvalues
                        fg, axs = plt.subplots(1, 2, figsize=(figsize()[0]*2, figsize()[1]))
                        axs[0].scatter(fitted, resid, alpha=.6, color=pcolor())
                        axs[0].axhline(0, color="red", ls="--", lw=1)
                        axs[0].set_xlabel("예측값"); axs[0].set_ylabel("잔차")
                        deco(axs[0], "잔차 vs 예측값")
                        stats.probplot(resid, dist="norm", plot=axs[1])
                        axs[1].set_title("정규 Q-Q 도표", fontsize=11)
                        plt.tight_layout(); show_plot(fg, max_width=920); plt.close(fg)
                        try:
                            _w, _p = stats.shapiro(resid)
                            st.caption(f"잔차 정규성(Shapiro-Wilk) p = {_p:.4f} → "
                                       + ("만족 ✅ 모형이 적절합니다."
                                          if _p >= .05 else
                                          "위배 ⚠️ 변수 변환이나 다른 모형을 고려해 보세요."))
                        except Exception:
                            pass
                        st.caption("왼쪽 그림에서 깔때기·곡선 모양이 보이면 등분산·선형성 가정이 "
                                   "깨진 것입니다. 오른쪽 점들이 직선에 가까울수록 정규성이 좋습니다.")
                    dl_table(coef, f"{y} 회귀분석 결과", "reg10", "reg")
                    report_capture("cap_reg", f"{y} 회귀분석", txt, coef, png)
                report_button("cap_reg")
        elif rt.startswith("로지스틱"):
            if not num_cols: st.warning("숫자형 독립변수가 필요합니다.")
            else:
                y = st.selectbox("종속변수 (Y, 2범주)", df.columns.tolist(), key="log_y")
                xs = st.multiselect("독립변수 (X)", [c for c in num_cols if c != y], key="log_x")
                if xs and keep_running("logit", "로지스틱 회귀 실행"):
                    data = df[[y]+xs].dropna()
                    if data[y].nunique() != 2: st.error("종속변수는 2개의 범주여야 합니다.")
                    else:
                        yb = pd.factorize(data[y])[0]
                        clf = LogisticRegression(max_iter=1000).fit(data[xs], yb)
                        st.metric("정확도(훈련)", f"{accuracy_score(yb, clf.predict(data[xs])):.3f}")
                        coef = pd.DataFrame({"변수": xs, "계수": clf.coef_[0].round(4)})
                        smart_table(coef, width="stretch")
                        dl_table(coef, f"{y} 로지스틱 회귀", "logit11", "logit")

        # ---------- 프로빗 (LC50 / LD50) ----------
        else:
            with st.expander("ℹ️ 프로빗 분석이란?"):
                st.markdown("""
농도를 높일수록 얼마나 더 죽는지(**농도-사충률 관계**)를 분석해
**LC50**(반수치사농도) 또는 **LD50**(반수치사량)을 구하는 표준 방법입니다.

**LC50이란?** 시험 개체의 **절반(50%)이 죽는 농도**입니다.
- LC50이 **작을수록** 낮은 농도에서도 효과가 나타남 = 약효가 강하거나 해충이 민감함
- LC50이 **클수록** 많이 써야 효과가 남 = 저항성이 생겼을 가능성

**어디에 쓰나요?**
- 살충제·살균제 **감수성 검정**
- 지역별·계통별 **저항성 발달 여부** 비교
- 천연물·친환경 자재의 **살충 효과 평가**

**자료 형태** — 한 행 = 한 농도 처리

| 계통 | 농도(ppm) | 공시충수 | 사충수 |
|---|---|---|---|
| 감수성계통 | 10 | 30 | 9 |
| 감수성계통 | 20 | 30 | 16 |

**필요한 열**: 농도 · 공시충수(총 개체수) · 사충수(죽은 개체수)
**권장 조건**: 농도 **5수준 이상**, 각 농도당 30마리 내외, 사충률이 0%와 100% 사이에 골고루 분포

**결과 읽는 법**
- **LC50 / LC90**: 50%, 90%가 죽는 농도
- **95% 신뢰구간**: 두 계통의 신뢰구간이 **서로 겹치지 않으면** 감수성이 통계적으로 다르다고 봅니다
- **저항성비(RR)** = 저항성계통 LC50 ÷ 감수성계통 LC50 → 10 이상이면 저항성이 상당히 발달한 것으로 해석합니다
- **기울기**: 클수록 농도가 조금만 올라가도 사충률이 급격히 증가

⚠️ 농도는 0보다 커야 합니다(로그 변환을 사용). 무처리구(농도 0)는 제외하고, 자연사충률이 높으면 Abbott 보정을 먼저 하세요.
""")
            if len(num_cols) < 3:
                st.warning("농도·총개체수·사충수 3개의 숫자형 열이 필요합니다.")
            else:
                c1, c2, c3 = st.columns(3)
                dose = c1.selectbox("농도(처리량) 열", num_cols,
                                    index=guess_idx(num_cols, ["농도", "처리량", "dose", "conc"]), key="pb_d")
                ntot = c2.selectbox("총 개체수(공시충수) 열", num_cols,
                                    index=guess_idx(num_cols, ["공시", "총개체", "총충", "n"], 1), key="pb_n")
                nres = c3.selectbox("사충수(반응 개체수) 열", num_cols,
                                    index=guess_idx(num_cols, ["사충", "사망", "반응", "death"], 2), key="pb_r")
                grp_opts = ["(없음)"] + cat_cols
                gsel = st.selectbox("계통·약제별 구분 열 (선택)", grp_opts, key="pb_g")
                if keep_running("probit", "프로빗 분석 실행"):
                    need = [dose, ntot, nres] + ([gsel] if gsel != "(없음)" else [])
                    d0 = df[need].dropna().copy()
                    # ---------- 입력 검증 ----------
                    _err = []
                    _bad_dead = d0[pd.to_numeric(d0[nres], errors="coerce")
                                   > pd.to_numeric(d0[ntot], errors="coerce")]
                    if len(_bad_dead):
                        _err.append(f"사충수가 공시충수보다 큰 행이 {len(_bad_dead)}개 있습니다. "
                                    "입력을 확인해 주세요.")
                    if (pd.to_numeric(d0[nres], errors="coerce") < 0).any():
                        _err.append("사충수에 음수가 있습니다.")
                    if (pd.to_numeric(d0[ntot], errors="coerce") <= 0).any():
                        _err.append("공시충수가 0 이하인 행이 있습니다.")
                    if (pd.to_numeric(d0[dose], errors="coerce") <= 0).any():
                        _err.append("농도가 0 이하인 행이 있습니다(로그 변환 불가). "
                                    "무처리구(농도 0)는 제외해 주세요.")
                    if _err:
                        for _m in _err: st.error("⚠️ " + _m)
                        st.stop()
                    d0 = d0[(d0[dose] > 0) & (d0[ntot] > 0)]
                    if d0.empty:
                        st.error("농도가 0보다 크고 총 개체수가 있는 자료가 필요합니다.")
                    else:
                        groups = d0[gsel].unique() if gsel != "(없음)" else ["전체"]
                        rows, curves = [], {}
                        for g in groups:
                            sub = d0 if gsel == "(없음)" else d0[d0[gsel] == g]
                            sub = sub.copy()
                            sub["logd"] = np.log10(sub[dose].astype(float))
                            sub["dead"] = sub[nres].astype(float)
                            sub["alive"] = (sub[ntot].astype(float) - sub["dead"]).clip(lower=0)
                            _nlev = sub[dose].nunique()
                            if _nlev < 3:
                                rows.append({"구분": g, "LC50": "계산불가", "LC90": "-",
                                             "기울기": "-",
                                             "비고": f"고유 농도 {_nlev}수준(3수준 이상 필요)"})
                                continue
                            if sub["dead"].sum() == 0:
                                rows.append({"구분": g, "LC50": "계산불가", "LC90": "-",
                                             "기울기": "-", "비고": "사충 반응이 전혀 없음"})
                                continue
                            if (sub["alive"] <= 0).all():
                                rows.append({"구분": g, "LC50": "계산불가", "LC90": "-",
                                             "기울기": "-", "비고": "모든 농도에서 100% 사충"})
                                continue
                            try:
                                X = sm.add_constant(sub[["logd"]])
                                gm = sm.GLM(sub[["dead", "alive"]], X,
                                            family=sm.families.Binomial(
                                                link=sm.families.links.Probit())).fit()
                                b0, b1 = float(gm.params.iloc[0]), float(gm.params.iloc[1])
                                if abs(b1) < 1e-9:
                                    raise ValueError("기울기가 0에 가까움")
                                lc50 = 10 ** (-b0 / b1)
                                lc90 = 10 ** ((stats.norm.ppf(0.90) - b0) / b1)
                                # LC50 신뢰구간 — 절편·기울기의 공분산을 반영 (델타법)
                                _cov = np.asarray(gm.cov_params())
                                v00, v01, v11 = float(_cov[0, 0]), float(_cov[0, 1]), float(_cov[1, 1])
                                m = -b0 / b1
                                # Var(m) = (1/b1²)·V00 + (2·b0/b1³)·V01 + (b0²/b1⁴)·V11
                                var_m = (v00 / b1**2) + (2 * b0 * v01 / b1**3) + (b0**2 * v11 / b1**4)
                                se_m = float(np.sqrt(var_m)) if var_m > 0 else float("nan")
                                if np.isnan(se_m):
                                    lo = hi = float("nan")
                                else:
                                    lo, hi = 10 ** (m - 1.96*se_m), 10 ** (m + 1.96*se_m)
                                rows.append({"구분": g, "LC50": round(lc50, 3),
                                             "95% 하한": round(lo, 3), "95% 상한": round(hi, 3),
                                             "LC90": round(lc90, 3), "기울기": round(b1, 3),
                                             "n(농도수)": len(sub)})
                                curves[g] = (sub, b0, b1, lc50)
                            except Exception as ex:
                                rows.append({"구분": g, "LC50": "계산불가", "LC90": "-",
                                             "기울기": "-", "비고": str(ex)[:30]})
                        res = pd.DataFrame(rows)
                        st.markdown("#### 프로빗 분석 결과")
                        smart_table(res, width="stretch")
                        st.caption("LC50이 작을수록 낮은 농도에서 효과가 나타남을 의미합니다. "
                                   "95% 신뢰구간이 서로 겹치지 않으면 두 계통·약제의 감수성이 다르다고 봅니다.")
                        if curves:
                            fig, ax = plt.subplots(figsize=figsize())
                            for g, (sub, b0, b1, lc50) in curves.items():
                                obs = sub["dead"] / (sub["dead"] + sub["alive"]) * 100
                                ax.scatter(sub[dose], obs, alpha=.7, label=f"{g} (관측)")
                                xs = np.linspace(sub[dose].min()*0.8, sub[dose].max()*1.2, 200)
                                ys = stats.norm.cdf(b0 + b1*np.log10(xs)) * 100
                                ax.plot(xs, ys, lw=2, label=f"{g} LC50={lc50:.2f}")
                            ax.axhline(50, color="gray", ls="--", lw=.8)
                            ax.set_xscale("log"); ax.set_xlabel(f"{dose} (로그 눈금)")
                            ax.set_ylabel("사충률(%)"); ax.set_ylim(-5, 105)
                            ax.legend(fontsize=7); deco(ax, "농도-사충률 곡선")
                            plt.tight_layout(); png = fig_to_png(fig)
                            st.download_button("🖼️ 그래프 다운로드", png, "probit.png", "image/png")
                        else:
                            png = None
                        ok = res[res["LC50"] != "계산불가"]
                        txt = ("프로빗 분석 결과 LC50은 " +
                               ", ".join(f"{r['구분']} {r['LC50']}" for _, r in ok.iterrows()) +
                               " 입니다." if len(ok) else "LC50을 계산할 수 있는 자료가 없습니다.")
                        st.info("💡 " + txt)
                        dl_table(res, "프로빗 분석 (LC50/LD50)", "probit12", "probit")
                        log_action("프로빗 분석(LC50) 실행")
                        report_capture("cap_pr", "프로빗 분석(LC50)", txt, res, png)
                        ai_interpret_button("pr", "프로빗 분석(LC50/LD50)", res, "LC50이 작을수록 약효가 강하거나 해충이 민감합니다.", capture_slot="cap_pr")
                report_button("cap_pr")

    # ---------- 머신러닝 ----------
    if tab_ml.on:
        st.subheader("머신러닝")
        with st.expander("ℹ️ 이 분석이 뭔가요?"): st.markdown(EXPLAIN["ml"])
        if len(num_cols) < 2: st.warning("숫자형 변수가 2개 이상 필요합니다.")
        else:
            allc = df.columns.tolist()
            tgt = st.selectbox("예측 대상(Y)", allc,
                               index=(allc.index(num_cols[-1]) if num_cols else 0), key="ml_y")
            fts = st.multiselect("입력 변수(X)", [c for c in num_cols if c != tgt], key="ml_x")
            y_is_num = tgt in num_cols
            n_uni = int(df[tgt].nunique())
            auto_cls = (not y_is_num) or n_uni <= 10
            task = st.radio("문제 유형", ["회귀(연속값 예측)", "분류(범주 예측)"],
                            index=1 if auto_cls else 0)
            if not y_is_num:
                st.info(f"'{tgt}'은 문자(범주)형이므로 **분류**만 가능합니다.")
            elif n_uni <= 10:
                st.caption(f"'{tgt}'의 값이 {n_uni}종류뿐이라 분류가 자연스럽습니다.")
            _reg_algos = [
                "랜덤포레스트", "Extra Trees", "그래디언트부스팅", "히스토그램 부스팅",
                "AdaBoost", "의사결정나무", "SVM(RBF)", "KNN", "Ridge", "Lasso", "ElasticNet"]
            _clf_algos = [
                "랜덤포레스트", "Extra Trees", "그래디언트부스팅", "히스토그램 부스팅",
                "AdaBoost", "의사결정나무", "SVM(RBF)", "KNN", "로지스틱 회귀", "GaussianNB"]
            _is_reg_choice = task.startswith("회귀")
            _algo_options = _reg_algos if _is_reg_choice else _clf_algos
            if st.session_state.get("ml_algo") not in _algo_options:
                st.session_state["ml_algo"] = _algo_options[0]
            algo = st.selectbox("알고리즘", _algo_options, key="ml_algo")
            _algo_help = {
                "랜덤포레스트": "여러 나무의 결과를 평균/투표합니다. 비선형 관계에 강하고 기본 선택으로 무난합니다.",
                "Extra Trees": "랜덤포레스트보다 분할을 더 무작위화한 앙상블입니다. 빠르고 변수 관계가 복잡할 때 유용합니다.",
                "그래디언트부스팅": "앞선 모형의 오차를 순차적으로 보완합니다. 중소형 표형 데이터에서 강한 편입니다.",
                "히스토그램 부스팅": "연속값을 구간화해 빠르게 부스팅합니다. 관측치가 많은 표형 자료에 유리합니다.",
                "AdaBoost": "틀린 관측치에 더 가중치를 주며 약한 모형을 결합합니다.",
                "의사결정나무": "규칙을 나무 형태로 나눠 설명하기 쉽지만 과적합에 주의해야 합니다.",
                "SVM(RBF)": "비선형 경계를 학습합니다. 변수 단위의 영향을 줄이기 위해 자동 표준화합니다.",
                "KNN": "비슷한 관측치 주변의 값을 이용합니다. 자동 표준화하며 표본이 너무 적으면 불안정할 수 있습니다.",
                "Ridge": "다중공선성이 있는 여러 연속형 변수의 회귀에 유용한 L2 규제 선형모형입니다.",
                "Lasso": "덜 중요한 변수의 계수를 0으로 줄일 수 있는 L1 규제 회귀입니다.",
                "ElasticNet": "Ridge와 Lasso를 함께 사용하는 규제 회귀입니다.",
                "로지스틱 회귀": "분류확률을 추정하는 기본 선형 분류모형입니다. 자동 표준화합니다.",
                "GaussianNB": "각 변수의 분포를 이용하는 매우 빠른 확률 분류모형입니다.",
            }
            st.caption("💡 " + _algo_help.get(algo, ""))

            if fts and keep_running("ml", "모델 학습"):
                is_reg = task.startswith("회귀")
                if is_reg and not y_is_num:
                    st.error(f"⚠️ '{tgt}'은 문자(범주)형이라 회귀 예측을 할 수 없습니다. "
                             "문제 유형을 '분류(범주 예측)'로 바꾸거나 숫자형 열을 선택하세요.")
                    st.stop()
                data = df[[tgt] + fts].dropna().copy()
                if len(data) < 10:
                    st.error("⚠️ 결측치를 제외한 학습자료가 10개 미만입니다. 머신러닝 결과를 신뢰하기 어렵습니다.")
                    st.stop()
                if len(data) < 30:
                    st.warning(f"⚠️ 사용 가능한 자료가 {len(data)}개뿐입니다. 머신러닝 결과는 탐색적으로만 해석하세요.")
                X = data[fts].astype(float)

                if is_reg:
                    y = pd.to_numeric(data[tgt], errors="coerce")
                    Xtr, Xte, ytr, yte = train_test_split(X, y, test_size=.3, random_state=0)
                    _kn = max(1, min(5, len(Xtr)))
                    mreg = {
                        "랜덤포레스트": RandomForestRegressor(n_estimators=300, random_state=0),
                        "Extra Trees": ExtraTreesRegressor(n_estimators=300, random_state=0),
                        "그래디언트부스팅": GradientBoostingRegressor(random_state=0),
                        "히스토그램 부스팅": HistGradientBoostingRegressor(random_state=0),
                        "AdaBoost": AdaBoostRegressor(n_estimators=200, random_state=0),
                        "의사결정나무": DecisionTreeRegressor(random_state=0),
                        "SVM(RBF)": Pipeline([("scale", StandardScaler()), ("model", SVR(kernel="rbf"))]),
                        "KNN": Pipeline([("scale", StandardScaler()), ("model", KNeighborsRegressor(n_neighbors=_kn))]),
                        "Ridge": Pipeline([("scale", StandardScaler()), ("model", Ridge(alpha=1.0))]),
                        "Lasso": Pipeline([("scale", StandardScaler()), ("model", Lasso(alpha=0.01, max_iter=5000))]),
                        "ElasticNet": Pipeline([("scale", StandardScaler()), ("model", ElasticNet(alpha=0.01, l1_ratio=0.5, max_iter=5000))]),
                    }
                    model = mreg[algo].fit(Xtr, ytr)
                    pred = model.predict(Xte)
                    s_ = r2_score(yte, pred)
                    st.metric("R² (테스트)", f"{s_:.3f}")
                    txt = f"[{algo}] 테스트 R²는 {s_:.3f}입니다."
                    _saved_classes = None
                else:
                    y_codes, y_levels = pd.factorize(data[tgt])
                    y = pd.Series(y_codes, index=data.index)
                    Xtr, Xte, ytr, yte = train_test_split(X, y, test_size=.3, random_state=0)
                    _kn = max(1, min(5, len(Xtr)))
                    mclf = {
                        "랜덤포레스트": RandomForestClassifier(n_estimators=300, random_state=0),
                        "Extra Trees": ExtraTreesClassifier(n_estimators=300, random_state=0),
                        "그래디언트부스팅": GradientBoostingClassifier(random_state=0),
                        "히스토그램 부스팅": HistGradientBoostingClassifier(random_state=0),
                        "AdaBoost": AdaBoostClassifier(n_estimators=200, random_state=0),
                        "의사결정나무": DecisionTreeClassifier(random_state=0),
                        "SVM(RBF)": Pipeline([("scale", StandardScaler()), ("model", SVC(kernel="rbf", probability=True, random_state=0))]),
                        "KNN": Pipeline([("scale", StandardScaler()), ("model", KNeighborsClassifier(n_neighbors=_kn))]),
                        "로지스틱 회귀": Pipeline([("scale", StandardScaler()), ("model", LogisticRegression(max_iter=3000, random_state=0))]),
                        "GaussianNB": GaussianNB(),
                    }
                    model = mclf[algo].fit(Xtr, ytr)
                    pred = model.predict(Xte)
                    s_ = accuracy_score(yte, pred)
                    st.metric("정확도 (테스트)", f"{s_:.3f}")
                    txt = f"[{algo}] 테스트 정확도는 {s_:.3f}입니다."
                    _saved_classes = list(y_levels)

                # 모든 알고리즘에서 가능한 한 변수 중요도/영향도를 제공한다.
                # 트리계열은 내장 중요도, 선형계열은 계수 절댓값, 그 외는 테스트셋 permutation importance를 사용한다.
                png, imp = None, None
                _base_model = model.named_steps.get("model") if isinstance(model, Pipeline) else model
                _imp_vals = None
                _imp_method = None
                if hasattr(_base_model, "feature_importances_"):
                    _imp_vals = np.asarray(_base_model.feature_importances_, dtype=float)
                    _imp_method = "모형 내장 중요도"
                elif hasattr(_base_model, "coef_"):
                    _coef = np.asarray(_base_model.coef_, dtype=float)
                    _imp_vals = np.abs(_coef).mean(axis=0) if _coef.ndim > 1 else np.abs(_coef)
                    _imp_method = "표준화 계수 절댓값"
                else:
                    try:
                        _perm = permutation_importance(model, Xte, yte, n_repeats=12, random_state=0,
                                                       scoring=("r2" if is_reg else "accuracy"))
                        _imp_vals = np.clip(np.asarray(_perm.importances_mean, dtype=float), 0, None)
                        _imp_method = "순열 중요도"
                    except Exception:
                        _imp_vals = None

                if _imp_vals is not None and len(_imp_vals) == len(fts):
                    imp = (pd.DataFrame({"변수": fts, "중요도": np.round(_imp_vals, 4)})
                           .sort_values("중요도", ascending=False).reset_index(drop=True))
                    if len(imp) and float(imp["중요도"].max()) > 0:
                        st.info(f"💡 [{algo}] {_imp_method} 기준 가장 영향이 큰 변수는 **'{imp.iloc[0]['변수']}'** 입니다.")
                    else:
                        st.info(f"💡 [{algo}] {_imp_method}를 계산했지만 변수 간 중요도 차이가 거의 없었습니다.")
                    fig, ax = plt.subplots(figsize=figsize())
                    _vals = imp["중요도"].astype(float).to_numpy()
                    _bars = ax.barh(imp["변수"], _vals, color=bar_colors(values=_vals.tolist()))
                    ax.invert_yaxis()
                    _vmax = float(np.nanmax(_vals)) if len(_vals) else 0.0
                    _pad = (_vmax * 0.025) if _vmax > 0 else 0.01
                    for _bar, _v in zip(_bars, _vals):
                        ax.text(float(_v) + _pad, _bar.get_y() + _bar.get_height() / 2,
                                f"{float(_v):.3f}", va="center", ha="left", fontsize=9,
                                fontweight="bold", color="#33383d", clip_on=False)
                    if _vmax > 0: ax.set_xlim(0, _vmax * 1.18)
                    deco(ax, f"변수 중요도 ({algo})")
                    png = fig_to_png(fig)
                    st.caption(f"※ 중요도 산출 방식: {_imp_method}. 알고리즘 간 중요도 값의 절대크기를 직접 비교하지 마세요.")
                    dl_table(imp, f"{tgt} 예측 변수 중요도", "ml13", "ml")
                else:
                    st.info("💡 " + txt)

                report_capture("cap_ml", f"{tgt} 예측 머신러닝({algo})", txt, imp, png)
                st.session_state["ml_model"] = {
                    "model": model, "feats": fts, "target": tgt, "is_reg": is_reg,
                    "algo": algo, "classes": _saved_classes,
                    "ranges": {f: (float(data[f].min()), float(data[f].max()), float(data[f].mean())) for f in fts}}
            report_button("cap_ml")

            # ---- 새 데이터로 실제 예측 ----
            saved = st.session_state.get("ml_model")
            if saved:
                st.divider()
                st.markdown("#### 🔮 새 데이터로 예측하기")
                st.caption(f"학습된 [{saved['algo']}] 모델로 '{saved['target']}'을(를) 예측합니다. "
                           "아래에 값을 입력하세요.")
                pmode = st.radio("입력 방식", ["직접 입력", "엑셀/CSV 업로드"], horizontal=True, key="ml_pmode")
                if pmode == "직접 입력":
                    vals = {}
                    pcols = st.columns(min(3, len(saved["feats"])))
                    for i, f in enumerate(saved["feats"]):
                        lo, hi, mean = saved["ranges"][f]
                        with pcols[i % len(pcols)]:
                            vals[f] = st.number_input(f, value=round(mean, 2),
                                                      help=f"학습 데이터 범위: {lo:.1f} ~ {hi:.1f}",
                                                      key=f"pv_{f}")
                    if st.button("예측 실행", key="ml_predict1"):
                        Xnew = pd.DataFrame([vals])[saved["feats"]]
                        pred = saved["model"].predict(Xnew)[0]
                        oob = [f for f in saved["feats"]
                               if not (saved["ranges"][f][0] <= vals[f] <= saved["ranges"][f][1])]
                        if saved["is_reg"]:
                            st.success(f"### 예측 결과: {saved['target']} ≈ **{pred:.2f}**")
                        else:
                            label = saved["classes"][int(pred)] if saved["classes"] else pred
                            st.success(f"### 예측 결과: {saved['target']} = **{label}**")
                            m = saved["model"]
                            if hasattr(m, "predict_proba"):
                                proba = m.predict_proba(Xnew)[0]
                                pr = pd.DataFrame({"분류": saved["classes"],
                                                   "확률(%)": (proba*100).round(1)}).sort_values("확률(%)", ascending=False)
                                smart_table(pr, width="stretch")
                        if oob:
                            st.warning(f"⚠️ {', '.join(oob)} 값이 학습 데이터 범위를 벗어났습니다. "
                                       "범위 밖 예측은 신뢰도가 떨어질 수 있어요.")
                        log_action(f"머신러닝 예측: {saved['target']}")
                else:
                    st.caption(f"예측할 파일에 **{', '.join(saved['feats'])}** 열이 있어야 합니다.")
                    pf = st.file_uploader("예측할 데이터 (xlsx/csv)", type=["xlsx", "csv"], key="ml_pf")
                    if pf is not None:
                        newdf = pd.read_csv(pf) if pf.name.endswith(".csv") else pd.read_excel(pf)
                        miss = [f for f in saved["feats"] if f not in newdf.columns]
                        if miss:
                            st.error(f"필요한 열이 없습니다: {', '.join(miss)}")
                        elif st.button("일괄 예측 실행", key="ml_predict2"):
                            Xnew = newdf[saved["feats"]].dropna()
                            preds = saved["model"].predict(Xnew)
                            out = newdf.loc[Xnew.index].copy()
                            if saved["is_reg"]:
                                out[f"{saved['target']}_예측"] = np.round(preds, 2)
                            else:
                                out[f"{saved['target']}_예측"] = [saved["classes"][int(p)]
                                                                if saved["classes"] else p for p in preds]
                            st.success(f"{len(out)}건 예측 완료!")
                            smart_table(out, width="stretch")
                            csv = out.to_csv(index=False).encode("utf-8-sig")
                            st.download_button("📥 예측 결과 CSV 다운로드", csv, "예측결과.csv", key="ml_dlpred")
                            log_action(f"머신러닝 일괄 예측: {len(out)}건")

    # ---------- AI 도우미 ----------

# ================================================================ AI 도우미
elif menu == "🧠 AI 도우미":
    st.subheader("🧠 AI 도우미 (생성형 AI)")
    if not st.session_state.get("api_key"):
        st.warning("사이드바 **🤖 AI 기능 켜기**에서 API 키를 입력하세요. "
                   "(Claude·Gemini·ChatGPT 중 선택 가능)")
    amode = st.radio("기능", ["결과를 자연어로 질문", "📝 데이터 자동 요약(초록 초안)",
                             "연구계획서 기반 통계 추천"], key="ai_mode")
    summary = build_data_overview(df)
    # 처리구·품종별 실제 평균을 함께 전달해 AI가 전체 describe()만 보고 추측하지 않게 한다.
    _general_profiles = build_group_profiles(df)

    if amode.startswith("📝"):
        st.caption("현재 데이터의 특징을 AI가 살펴보고, 연구 초록(Abstract) 초안이나 핵심 요약을 만들어 줍니다.")
        purpose = st.text_input("연구 목적/배경 (한 줄, 선택)",
                                placeholder="예) 고추 신품종의 생육·수량 특성 비교")
        want = st.radio("형태", ["핵심 요약 (불릿)", "결과 요약 문단"], horizontal=True)
        if st.button("AI 요약 생성"):
            with st.spinner("AI가 데이터를 살펴보는 중..."):
                fmt = {"핵심 요약 (불릿)": "핵심 발견을 불릿 5개 이내로",
                       "연구 초록 초안": "학술논문 초록 형식(목적·방법·결과·결론)으로 200자 내외",
                       "결과 요약 문단": "결과를 서술한 한 문단으로"}[want]
                st.markdown(ai_call(
                    f"연구 목적: {purpose or '(미기재)'}\n\n{summary}\n\n"
                    f"[처리·품종별 요약]\n{_general_profiles}\n\n"
                    f"위 데이터를 {fmt} 한국어로 정리해 주세요. 데이터에 근거한 내용만 쓰고, "
                    "통계 검정을 따로 하진 않았으니 단정적 유의성 주장은 피하세요.",
                    st.session_state.get("api_key"), st.session_state.get("ai_model_g"), max_tokens=1000))
                ai_disclaimer()
                log_action("AI 데이터 요약 생성")
    elif amode.startswith("결과"):
        with st.expander("💬 이렇게 물어보세요 (예시)"):
            st.markdown("""
- 처리구별 수량 차이를 보고서에 쓸 문장으로 정리해줘
- 이 데이터로 어떤 분석을 하면 좋을지 알려줘
- 유의성 문자 a, b, c가 무슨 뜻인지 쉽게 설명해줘
- 상관계수 0.93이면 어느 정도로 강한 관계인지 설명해줘
- 이 결과를 비전공자인 상사에게 보고할 때 어떻게 말하면 좋을까?
- 처리2가 가장 좋은 이유를 데이터 근거로 설명해줘
""")
        q = st.text_area("궁금한 점", placeholder="예) 처리구별 수량 차이를 쉽게 설명해줘")
        if st.button("AI에게 물어보기"):
            with st.spinner("AI가 분석 중..."):
                _question_profiles = build_group_profiles(df, question=q)
                st.markdown(ai_call(
                    "당신은 농업연구사를 돕는 통계 전문가입니다. 아래 실제 데이터 요약만 참고해 "
                    "질문에 쉽고 정확하게 한국어로 답해주세요. 입력에 없는 수치나 유의성을 "
                    f"추측하지 마세요.\n\n{summary}\n\n"
                    f"[질문 관련 처리·품종별 요약]\n{_question_profiles}\n\n질문: {q}",
                    st.session_state.get("api_key"), st.session_state.get("ai_model_g")))
                ai_disclaimer()
    else:
        with st.expander("💬 이렇게 입력하세요 (예시)"):
            st.markdown("""
연구계획서의 **시험 목적·처리 내용·반복 수·조사 항목**을 붙여넣으면 가장 정확합니다.

> 예시) 고추 '청양' 품종을 대상으로 질소 시비량 4수준(0, 10, 20, 30kg/10a)을 난괴법 3반복으로 배치하여
> 초장, 착과수, 상품수량, 당도를 조사하고자 함. 시험은 2개 지역(영양, 안동)에서 동시 수행함.
""")
        plan_file = st.file_uploader("📎 연구계획서 파일 첨부 (선택)",
                                     type=["txt", "md", "csv", "xlsx", "hwpx", "pdf", "docx"],
                                     key="plan_file")
        file_text = ""
        if plan_file is not None:
            file_text = read_uploaded_text(plan_file)
            if file_text.startswith("⚠️"):
                st.warning(file_text); file_text = ""
            else:
                st.success(f"'{plan_file.name}'에서 {len(file_text):,}자를 읽었습니다.")
                with st.expander("읽어온 내용 확인"):
                    st.text(file_text[:2000] + ("..." if len(file_text) > 2000 else ""))
        plan = st.text_area("연구계획서 내용 (직접 입력하거나 위 파일 첨부)",
                            value=file_text, height=180,
                            placeholder="시험 목적, 처리 내용, 반복 수, 조사 항목 등")
        if st.button("통계 방법 추천받기"):
            with st.spinner("AI가 연구계획을 분석 중..."):
                st.markdown(ai_call("당신은 농업 실험설계·통계 전문가입니다. 아래 연구계획서와 데이터 구조를 보고 "
                                    "가장 적합한 통계 분석 방법(분산분석 종류, 사후검정, 상관/회귀, 비모수 등)을 "
                                    f"이유와 함께 한국어로 단계별 추천해주세요.\n\n[데이터]\n{summary}\n\n[연구계획서]\n{plan}",
                                    st.session_state.get("api_key"), st.session_state.get("ai_model_g"), max_tokens=1300))
                ai_disclaimer()
                log_action("AI 연구계획서 기반 통계 추천")
    st.caption("※ AI 응답은 참고용이며, 호출 시 사용량만큼 소액 비용이 발생할 수 있어요.")

# ================================================================ 경제성 분석
elif menu == "💰 경제성분석":
    st.title("💰 경제성 분석")
    st.info("**경제성 분석이 처음이면 길잡이로 시작하고, 분석방법을 알고 있다면 바로 분석할 수 있습니다.** 필요한 자료와 기준단가는 접어서 확인할 수 있습니다.")
    st.caption("🧭 경제성 분석 UX v3.6 · 단계형 길잡이 + 간소화 화면 · 2026-08-18 적용")

    # ---------------------------------------------------------------- 경제성 분석 시작 화면 / 길잡이
    # 기능을 한꺼번에 펼치지 않고, 처음에는 "길잡이"와 "바로 분석" 두 갈래만 보여준다.
    # 길잡이는 STEP 1~5를 한 단계씩 진행해 초보자도 현재 질문에만 집중할 수 있게 한다.
    def _render_econ_material_guide(expanded=False):
        with st.expander("📚 경제성 분석에 어떤 자료를 준비해야 하나요?", expanded=expanded):
            st.caption("모든 분석에 모든 비용이 필요한 것은 아닙니다. 연구목적에 맞는 자료만 준비하면 됩니다.")
            _mt1, _mt2, _mt3, _mt4 = st.tabs(["공통자료", "비용 항목", "분석별 자료", "빈 서식"])
            with _mt1:
                _base_need = pd.DataFrame([
                    ["처리구·품종·기술명", "무엇과 무엇을 비교했는지", "대조구, 신품종A, 신품종B", "비교 연구면 필수"],
                    ["반복(블록)", "같은 처리를 몇 번 반복했는지", "1, 2, 3 또는 Block1~3", "반복시험이면 권장"],
                    ["조사면적", "원자료가 몇 a 또는 ㎡ 기준인지", "5a, 10a, 1,000㎡", "10a 환산에 필요"],
                    ["생산량", "실제로 생산·수확한 양", "kg/조사구, kg/10a", "대부분의 분석에 필수"],
                    ["판매가격", "농가가 실제로 받는 가격", "원/kg, 원/상자", "수입 계산에 필수"],
                    ["부산물 수입", "주산물 외 판매수입이 있으면 기록", "부산물 판매액", "해당 시"],
                    ["상품등급별 수량·가격", "등급별 가격차가 크면 분리", "특·상·보통 수량과 단가", "원예작물에서 권장"],
                ], columns=["자료", "무엇을 뜻하나요?", "예시", "언제 필요한가요?"])
                smart_table(_base_need, width="stretch", hide_index=True)
                st.info("💡 한 줄은 보통 **처리구 × 반복 한 조사구**로 적는 것이 가장 안전합니다. "
                        "5a 자료라면 억지로 10a로 바꾸지 말고 5a 값을 그대로 넣은 뒤 기준면적을 5a로 지정하세요.")

            with _mt2:
                st.markdown("**소득분석에서는 실제로 발생한 비용만 기록합니다. 없는 항목은 만들 필요가 없습니다.**")
                _cost_need = pd.DataFrame([
                    ["종자·종묘비", "종자, 묘, 접목묘 등", "구입액 또는 사용량×단가"],
                    ["비료비", "기비·추비·액비·엽면시비 등", "비료 종류별 사용액"],
                    ["농약비", "살균제·살충제·제초제·생물농약 등", "약제별 사용액"],
                    ["수도·광열비", "전기, 유류, 난방, 관수", "전기료·경유·등유·가스 등"],
                    ["기타재료비", "멀칭필름, 상토, 트레이, 유인끈, 지주대, 포장재 등", "소모성 자재 사용액"],
                    ["소농구비", "내용연수가 짧은 소형 농기구", "가위, 호미, 소형도구 등"],
                    ["대농구상각비", "여러 해 쓰는 농기계의 연간 감가상각액", "트랙터, 관리기, 방제기 등"],
                    ["영농시설상각비", "여러 해 쓰는 시설의 연간 감가상각액", "하우스, 관수시설, 건조시설 등"],
                    ["수선비", "농기계·시설 수리 및 유지", "부품·수리비"],
                    ["임차료", "실제로 지불한 토지·시설·농기계 임차비", "농지·시설·장비 임차료"],
                    ["위탁영농비", "작업을 외부에 맡긴 비용", "경운·정지·수확·방제 위탁료"],
                    ["고용노동비", "외부 인력에게 실제 지급한 임금", "정식·유인·수확 인건비"],
                    ["조성비 상각", "과수·다년생 초기 조성비의 연간 배분", "과원 조성비의 연간 상각액"],
                ], columns=["비용 항목", "쉽게 말하면", "무엇을 적나요?"])
                smart_table(_cost_need, width="stretch", hide_index=True)
                st.markdown("**순수익(생산비)까지 보려면 아래 경제적 비용도 추가합니다.**")
                _full_cost = pd.DataFrame([
                    ["자가노동시간", "본인·가족이 일한 시간", "시간으로 입력 → 프로그램이 시간당 노임을 곱함"],
                    ["유동자본용역비", "재배기간 동안 투입자금이 묶인 비용", "이자율·재포기간으로 자동 계산"],
                    ["고정자본용역비", "농기계·시설 자본 사용의 기회비용", "부분현재가·작목부담률·이자율로 계산"],
                    ["토지용역비", "자가토지를 다른 용도로 쓸 수 있었던 가치", "자가토지 기회비용; 실제 임차료와 중복 금지"],
                ], columns=["추가 항목", "의미", "입력 방법"])
                smart_table(_full_cost, width="stretch", hide_index=True)
                st.warning("⚠️ `경영비합계`, `비용합계`, `소득`, `순수익` 같은 계산결과 열을 다시 비용으로 선택하지 마세요. "
                           "자가노동은 **원(비용)이 아니라 시간**으로 기록합니다.")

            with _mt3:
                _by_analysis = pd.DataFrame([
                    ["📕 부분예산법", "대조구·신기술구, 수량/수입, 가격, **신기술 때문에 달라지는 비용만**", "관행과 똑같이 발생한 비용은 조사하지 않아도 됨"],
                    ["📗 소득분석", "수량, 가격, 조사면적, 실제 발생한 항목별 경영비", "순수익까지 보면 자가노동·자본·토지도 추가"],
                    ["📘 지배분석·MRR", "비용이 다른 여러 처리, 수량, 가격, 처리별 가변비용, 대조구", "처리 수준에 따라 실제로 달라지는 비용을 구분"],
                    ["📙 시설·장기투자", "최초투자비, 내용연수, 연간 편익, 운영·유지비, 할인율, 잔존가치", "여러 해의 현금흐름을 평가"],
                    ["손익분기·민감도", "기준 수량·가격·비용", "수확·선별·포장처럼 수량에 따라 변하는 비용을 구분하면 더 현실적"],
                ], columns=["분석", "최소한 준비할 자료", "특히 주의할 점"])
                smart_table(_by_analysis, width="stretch", hide_index=True)
                st.markdown("**연구질문별 빠른 찾기**")
                st.markdown("- 현재 작목 10a당 수익성 → **소득분석**\n"
                            "- 신품종·신기술 vs 기존 방식 → **부분예산법**\n"
                            "- 비용이 다른 여러 기술 중 최적 대안 → **지배분석·MRR**\n"
                            "- 시설·농기계 투자 → **NPV·할인 B/C·IRR**\n"
                            "- 어느 가격·수량부터 손해인지 → **손익분기점**\n"
                            "- 가격·수량 변동에도 버티는지 → **민감도 분석**")

            with _mt4:
                _tmpl1 = pd.DataFrame(columns=["처리구", "반복", "조사면적(a)", "수량(kg)", "판매단가(원/kg)",
                                                "종자종묘비", "비료비", "농약비", "수도광열비", "기타재료비",
                                                "소농구비", "대농구상각비", "영농시설상각비", "수선비", "임차료",
                                                "위탁영농비", "고용노동비", "자가노동시간"])
                _tmpl2 = pd.DataFrame(columns=["처리구", "반복", "조사면적(a)", "수량(kg)", "판매단가(원/kg)",
                                                "신기술추가비용", "절감비용", "추가노동시간"])
                _tc1, _tc2 = st.columns(2)
                _tc1.download_button("📥 소득분석 빈 서식 CSV", _tmpl1.to_csv(index=False).encode("utf-8-sig"),
                                     "경제성_소득분석_빈서식.csv", mime="text/csv",
                                     key="p_econ_guide_template_income", width="stretch")
                _tc2.download_button("📥 신기술 비교 빈 서식 CSV", _tmpl2.to_csv(index=False).encode("utf-8-sig"),
                                     "경제성_신기술비교_빈서식.csv", mime="text/csv",
                                     key="p_econ_guide_template_partial", width="stretch")

    _econ_entry = st.session_state.get("econ_entry_mode")
    if _econ_entry not in ("guide", "direct"):
        st.markdown("### 어떻게 시작할까요?")
        _ec1, _ec2 = st.columns(2)
        with _ec1:
            with st.container(border=True):
                st.markdown("### 🧭 경제성 분석이 처음이에요")
                st.caption("연구내용에 맞는 분석법을 STEP 1~5로 한 단계씩 찾아드립니다.")
                if st.button("길잡이 시작", type="primary", width="stretch", key="econ_entry_guide"):
                    st.session_state["econ_entry_mode"] = "guide"
                    st.session_state["econ_guide_step"] = 1
                    st.rerun()
        with _ec2:
            with st.container(border=True):
                st.markdown("### 📊 분석방법을 알고 있어요")
                st.caption("부분예산 · 소득분석 · MRR · 시설투자 중 원하는 분석으로 바로 갑니다.")
                if st.button("바로 분석하기", width="stretch", key="econ_entry_direct"):
                    st.session_state["econ_entry_mode"] = "direct"
                    st.rerun()
        _render_econ_material_guide(expanded=False)
        st.stop()

    if _econ_entry == "guide":
        _top1, _top2 = st.columns([5, 1])
        with _top1:
            st.markdown("## 🧭 경제성 분석 길잡이")
            st.caption("한 번에 한 질문만 답하면 됩니다. 모르는 항목은 '잘 모르겠어요'를 선택해도 됩니다. AI/API 키는 필요하지 않습니다.")
        with _top2:
            if st.button("처음으로", width="stretch", key="econ_guide_home"):
                st.session_state["econ_entry_mode"] = None
                st.rerun()

        _step = int(st.session_state.get("econ_guide_step", 1))
        _step = min(max(_step, 1), 5)
        st.caption(f"STEP {_step} / 5")
        st.progress(_step / 5)

        if _step == 1:
            st.markdown("### STEP 1. 이번 연구에서 무엇을 확인하고 싶나요?")
            _g_goal = st.radio(
                "가장 가까운 상황을 하나 고르세요.",
                ["🧪 새로운 품종·기술이 기존 방식보다 경제적인지 알고 싶어요",
                 "🌱 현재 작목·처리의 한 해 수익성이 궁금해요",
                 "🏆 여러 기술·처리 중 가장 경제적인 대안을 고르고 싶어요",
                 "🏗️ 시설·농기계에 투자할 가치가 있는지 알고 싶어요",
                 "💸 어느 가격·수량부터 손해인지 알고 싶어요",
                 "📉 가격·수량이 변해도 경제성이 유지되는지 알고 싶어요",
                 "🔄 여러 품종·작형·처리의 한 해 경영성과를 비교하고 싶어요",
                 "🏛️ 정책·공공사업의 사회적 효과를 평가하고 싶어요",
                 "🤔 잘 모르겠어요 — 쉬운 질문으로 찾아주세요"],
                key="econ_guide_goal_v3", index=None)
            if st.button("다음 →", type="primary", width="stretch", key="econ_g_next1",
                         disabled=_g_goal is None):
                st.session_state["econ_guide_step"] = 2; st.rerun()

        elif _step == 2:
            st.markdown("### STEP 2. 이번 연구에서 실제로 무엇이 달라지나요?")
            _g_change = st.radio(
                "가장 큰 변화 하나를 고르세요.",
                ["품종·방제·재배법 등 기술이 바뀌어요",
                 "비료량·농약량·노동량 등 투입수준과 비용이 달라져요",
                 "시설·농기계를 새로 구입하거나 설치해요",
                 "판매가격·상품수량·상품률 등이 달라져요",
                 "특별한 기술변경 없이 현재 경영성과를 평가해요",
                 "사회적·환경적 편익과 비용을 평가해요",
                 "잘 모르겠어요"], key="econ_guide_change_v3", index=None)
            _b1, _b2 = st.columns(2)
            if _b1.button("← 이전", width="stretch", key="econ_g_prev2"):
                st.session_state["econ_guide_step"] = 1; st.rerun()
            if _b2.button("다음 →", type="primary", width="stretch", key="econ_g_next2",
                          disabled=_g_change is None):
                st.session_state["econ_guide_step"] = 3; st.rerun()

        elif _step == 3:
            st.markdown("### STEP 3. 무엇과 무엇을 비교하나요?")
            _g_compare = st.radio(
                "비교 구조를 고르세요.",
                ["대조구 1개와 신기술·신품종 1~2개를 비교해요",
                 "비용이 서로 다른 3개 이상 대안을 비교해요",
                 "여러 품종·작형·처리의 한 해 성과를 나란히 비교해요",
                 "비교대상 없이 현재 상태 하나만 평가해요",
                 "잘 모르겠어요"], key="econ_guide_compare_v3", index=None)
            _b1, _b2 = st.columns(2)
            if _b1.button("← 이전", width="stretch", key="econ_g_prev3"):
                st.session_state["econ_guide_step"] = 2; st.rerun()
            if _b2.button("다음 →", type="primary", width="stretch", key="econ_g_next3",
                          disabled=_g_compare is None):
                st.session_state["econ_guide_step"] = 4; st.rerun()

        elif _step == 4:
            st.markdown("### STEP 4. 경제효과를 어느 기간까지 보나요?")
            _g_period = st.radio(
                "분석기간을 고르세요.",
                ["한 작기 또는 1년 안에서 평가해요",
                 "같은 기술의 효과와 비용이 2년 이상 이어져요",
                 "시설·농기계의 내용연수(사용기간) 전체를 평가해요",
                 "잘 모르겠어요"], key="econ_guide_period_v3", index=None)
            _b1, _b2 = st.columns(2)
            if _b1.button("← 이전", width="stretch", key="econ_g_prev4"):
                st.session_state["econ_guide_step"] = 3; st.rerun()
            if _b2.button("다음 →", type="primary", width="stretch", key="econ_g_next4",
                          disabled=_g_period is None):
                st.session_state["econ_guide_step"] = 5; st.rerun()

        else:
            st.markdown("### STEP 5. 지금 어떤 자료가 준비되어 있나요?")
            _g_df = st.session_state.get("df")
            _cols = list(map(str, _g_df.columns)) if isinstance(_g_df, pd.DataFrame) and not _g_df.empty else []
            def _g_find(keys):
                return [c for c in _cols if any(k.lower() in c.lower() for k in keys)]
            _cand_trt = _g_find(["처리", "품종", "계통", "작형", "구분", "시험구"])
            _cand_qty = _g_find(["수량", "생산량", "수확량", "수확중", "상품수량"])
            _cand_price = _g_find(["단가", "가격", "판매가", "판매액"])
            _cand_cost = _g_find(["비용", "비료", "농약", "노력", "노동", "노임", "자재", "임차", "연료", "종묘"])
            _cand_rep = _g_find(["반복", "블록", "rep", "block", "blk"])
            _auto_data = []
            if _cand_trt: _auto_data.append("처리구/대조구 구분")
            if _cand_qty: _auto_data.append("수량·생산량")
            if _cand_price: _auto_data.append("판매가격/판매액")
            if _cand_cost: _auto_data.append("항목별 경영비")
            if _cand_rep: _auto_data.append("반복(블록) 자료")
            _g_data = st.multiselect(
                "가지고 있는 자료를 모두 선택하세요. 업로드된 데이터에서 찾은 항목은 자동 체크됩니다.",
                ["처리구/대조구 구분", "수량·생산량", "판매가격/판매액", "항목별 경영비",
                 "신기술로 달라지는 비용만", "반복(블록) 자료", "최초 투자비",
                 "연도별 편익·운영비", "분석기간·할인율·잔존가치",
                 "아직 거의 준비되지 않았거나 자료가 어떤 것인지 잘 모르겠어요"],
                default=_auto_data, key="econ_guide_data_v3")
            _render_econ_material_guide(expanded=False)

            _g_goal = st.session_state.get("econ_guide_goal_v3")
            _g_change = st.session_state.get("econ_guide_change_v3")
            _g_compare = st.session_state.get("econ_guide_compare_v3")
            _g_period = st.session_state.get("econ_guide_period_v3")
            _g_result = recommend_economic_guide(_g_goal, _g_change, _g_compare, _g_period, _g_data)
            st.divider()
            st.markdown("### 🎯 추천 결과")

            if _g_result.get("ambiguous"):
                st.warning(_g_result.get("title", "추천을 조금 더 좁혀야 합니다."))
                for _r in _g_result.get("reasons", []): st.markdown(f"- {_r}")
                st.info("💡 **기존 방식과 다른 처리·기술이 있는지**, 그리고 **효과가 1년인지 여러 해인지**를 다시 확인하면 추천을 좁힐 수 있습니다.")
            elif _g_result.get("primary_mode") is None:
                st.warning("현재 농가단위 분석 모듈로 억지 연결하지 않았습니다. 정책·공공사업의 사회적 효과는 별도의 비용편익분석(CBA)이 필요합니다.")
            else:
                _r1, _r2 = st.columns([3, 1])
                _r1.markdown(f"### {_g_result['title']}")
                _r2.metric("추천 확신도", _g_result.get("confidence", "-"))
                st.markdown("**추천 이유**")
                for _r in _g_result.get("reasons", []): st.markdown(f"- {_r}")
                _gc1, _gc2 = st.columns(2)
                with _gc1:
                    st.markdown("**필요한 자료**")
                    for _x in _g_result.get("needs", []): st.markdown(f"- {_x}")
                with _gc2:
                    st.markdown("**함께 보면 좋은 분석**")
                    for _x in _g_result.get("together", []): st.markdown(f"- {_x}")
                _missing = _g_result.get("missing_reported") or []
                if _missing:
                    st.warning("추가로 준비하면 좋은 자료: **" + ", ".join(_missing) + "**")
                if _cols:
                    with st.expander("📋 업로드한 데이터 열 후보", expanded=False):
                        smart_table(pd.DataFrame([
                            {"필요 항목":"처리구/비교대상", "현재 데이터 후보":", ".join(_cand_trt[:3]) if _cand_trt else "찾지 못함"},
                            {"필요 항목":"수량", "현재 데이터 후보":", ".join(_cand_qty[:3]) if _cand_qty else "찾지 못함"},
                            {"필요 항목":"단가/가격", "현재 데이터 후보":", ".join(_cand_price[:3]) if _cand_price else "찾지 못함"},
                            {"필요 항목":"비용", "현재 데이터 후보":", ".join(_cand_cost[:4]) if _cand_cost else "찾지 못함"},
                            {"필요 항목":"반복(선택)", "현재 데이터 후보":", ".join(_cand_rep[:2]) if _cand_rep else "찾지 못함"},
                        ]), width="stretch", hide_index=True)
                if st.button("🚀 추천 분석 바로 시작하기", type="primary", width="stretch", key="p_econ_guide_start_v4"):
                    st.session_state["econ_mode"] = _g_result["primary_mode"]
                    st.session_state["econ_entry_mode"] = "direct"
                    st.session_state["_econ_guide_applied"] = {"title": _g_result["title"], "tags": _g_result.get("tags") or []}
                    st.rerun()

            if st.button("← STEP 4로 돌아가기", width="stretch", key="econ_g_prev5"):
                st.session_state["econ_guide_step"] = 4; st.rerun()

        st.stop()

    # direct mode: 핵심 분석 화면만 보여주고, 도움말/설정은 접어서 필요할 때만 연다.
    _direct_head1, _direct_head2 = st.columns([5, 1])
    with _direct_head1:
        st.caption("분석방법을 선택하고 필요한 값만 입력하세요. 자료 준비 방법과 기준단가는 필요할 때만 열어볼 수 있습니다.")
    with _direct_head2:
        if st.button("🧭 길잡이", width="stretch", key="econ_switch_guide"):
            st.session_state["econ_entry_mode"] = "guide"
            st.session_state["econ_guide_step"] = 1
            st.rerun()
    _render_econ_material_guide(expanded=False)

    _guide_applied = st.session_state.pop("_econ_guide_applied", None)
    if _guide_applied:
        if isinstance(_guide_applied, dict):
            _focus = " · ".join(_guide_applied.get("tags") or [])
            _msg = f"✅ 길잡이가 추천한 **{_guide_applied.get('title', '분석')}**을 아래에서 자동 선택했습니다."
            if _focus:
                _msg += f" 결과에서 **{_focus}**도 함께 확인하세요."
            st.success(_msg)
        else:
            st.success("✅ 길잡이가 추천한 분석을 아래에서 자동 선택했습니다. 필요한 입력값을 확인하고 분석을 진행하세요.")

    st.markdown("### 분석방법")
    emode = st.radio("분석 방식",
                     ["📕 부분예산표 (손실적·이익적 요소)",
                      "📗 소득분석",
                      "📘 신기술 경제성 (부분예산·한계수익률)",
                      "📙 시설·장기투자 경제성 (NPV·B/C·IRR)"],
                     horizontal=True, key="econ_mode", label_visibility="collapsed")
    st.caption("선택한 분석에 필요한 입력과 결과만 아래에 표시됩니다.")

    with st.expander("🧪 경제성 계산 자가진단 (개발·점검용)", expanded=False):
        st.caption("면적 환산, 입력 검증, 유동·고정자본 산식, 부분예산, 장기투자 NPV/B·C까지 즉시 점검합니다.")
        if st.button("자가진단 실행", key="econ_selftest"):
            _self = run_economic_self_test()
            smart_table(_self, width="stretch", hide_index=True)
            _passed = int(_self["결과"].eq("PASS").sum())
            if _passed == len(_self):
                st.success(f"✅ {_passed}/{len(_self)}개 핵심 테스트 통과")
            else:
                st.error(f"❌ {_passed}/{len(_self)}개 통과 — FAIL 항목을 확인하세요.")
    with st.expander("💾 기준단가 관리 (노임·자재비·임차료 등)", expanded=False):
        st.caption("여기에 저장한 값이 아래 분석의 **기본값으로 자동 입력**됩니다. "
                   "기관에서 쓰는 공식 단가를 한 번 넣어두면 매번 입력할 필요가 없어요.")
        if st.session_state.get("price_db") is None:
            st.session_state["price_db"] = default_price_db()
        else:   # 옛 형식(메타데이터 없는 표)이면 빠진 열을 채움
            _pdb = st.session_state["price_db"]
            import datetime as _dt
            for _c, _dv in [("조회방식", "사용자 입력"),
                            ("갱신일", _dt.date.today().isoformat()),
                            ("환산식", ""), ("사용자수정", False)]:
                if _c not in _pdb.columns:
                    _pdb[_c] = _dv
            st.session_state["price_db"] = _pdb
        st.session_state["price_db"] = st.data_editor(
            st.session_state["price_db"], num_rows="dynamic", width="stretch", key="price_editor")
        for _w in price_db_warnings(st.session_state["price_db"]):
            st.warning("⚠️ " + _w)
        pcol1, pcol2 = st.columns(2)
        pcol1.download_button("📥 기준단가 내려받기(CSV)",
                              st.session_state["price_db"].to_csv(index=False).encode("utf-8-sig"),
                              "기준단가.csv", width="stretch")
        _up = pcol2.file_uploader("📤 기준단가 불러오기(CSV)", type=["csv"], key="price_up")
        if _up is not None:
            try:
                st.session_state["price_db"] = pd.read_csv(_up)
                st.success("기준단가를 불러왔습니다."); st.rerun()
            except Exception as ex:
                st.error(f"불러오기 실패: {ex}")
        st.markdown("---")
        st.markdown("###### 🌐 KAMIS 농산물 가격 자동 조회 (선택)")
        st.caption("KAMIS 오픈API 인증키가 있으면 최근 가격을 자동으로 불러올 수 있습니다. "
                   "발급: kamis.or.kr → 고객센터 → Open-API 이용안내 (무료)")
        kc1, kc2 = st.columns(2)
        k_key = kc1.text_input("KAMIS 인증키 (cert_key)", type="password", key="kamis_key")
        k_id = kc2.text_input("KAMIS 아이디 (cert_id)", key="kamis_id")
        kc3, kc4, kc5 = st.columns(3)
        _ITEMS = {"건고추": ("312", "300", "01"), "풋고추": ("225", "200", "01"),
                  "마늘": ("258", "200", "00"), "양파": ("245", "200", "00"),
                  "배추": ("211", "200", "00"), "무": ("231", "200", "00"),
                  "사과": ("411", "400", "05"), "쌀": ("111", "100", "01")}
        k_item = kc3.selectbox("품목", list(_ITEMS.keys()), key="kamis_item")
        k_rank = kc4.selectbox("등급", ["04 (상품)", "05 (중품)"], key="kamis_rank")
        k_days = kc5.number_input("최근 며칠", 1, 30, 7, key="kamis_days")
        km1, km2 = st.columns(2)
        k_market = km1.radio("가격 유형", ["도매", "소매"], horizontal=True, key="kamis_market")
        _COUNTRIES = {"전체": "", "서울": "1101", "부산": "2100", "대구": "2200",
                      "광주": "2401", "대전": "2501", "안동": "3714", "포항": "3711"}
        k_country_name = km2.selectbox("지역", list(_COUNTRIES), key="kamis_country")
        if st.button("📡 KAMIS 가격 조회", width="stretch"):
            if not (k_key and k_id):
                st.warning("인증키와 아이디를 모두 입력하세요.")
            else:
                code, category, kind = _ITEMS[k_item]
                with st.spinner("KAMIS에서 가격을 불러오는 중..."):
                    dfk, err = kamis_fetch(
                        k_key, k_id, code, kind, k_rank.split()[0], int(k_days),
                        country_code=_COUNTRIES[k_country_name], category_code=category,
                        market_type=k_market, convert_kg=False)
                if err:
                    st.error(err)
                    st.caption("품목·등급 조합에 따라 자료가 없을 수 있습니다. "
                               "KAMIS 홈페이지에서 코드를 확인해 주세요.")
                else:
                    _tr = getattr(dfk, "attrs", {}).get("kamis_transport")
                    if _tr == "http-fallback":
                        st.warning("⚠️ HTTPS 연결이 실패하여 공식 HTTP 호환 endpoint로 조회했습니다. "
                                   "KAMIS 서버의 SSL 상태가 정상화되면 자동으로 HTTPS를 다시 사용합니다.")
                    elif _tr == "https-legacy":
                        st.info("ℹ️ KAMIS 서버가 구형 TLS 설정을 사용하고 있어 호환 모드로 조회했습니다. "
                                "인증서 검증은 정상적으로 수행되었습니다.")
                    elif _tr == "https-insecure":
                        st.warning("⚠️ KAMIS 서버 인증서를 검증할 수 없어 검증을 생략하고 조회했습니다. "
                                   "가격 자료 확인 용도로만 사용하고, 인증키가 노출될 수 있는 환경에서는 "
                                   "수동 입력을 권장합니다.")
                    st.session_state["kamis_result"] = dfk
        if st.session_state.get("kamis_result") is not None:
            _kres = st.session_state["kamis_result"]
            smart_table(_kres, width="stretch", hide_index=True)
            _valid = _kres.dropna(subset=["가격"])
            if len(_valid):
                st.markdown("###### 📥 조회 결과를 기준단가에 반영")
                _rowsel = st.selectbox(
                    "반영할 행", list(range(len(_valid))),
                    format_func=lambda i: (f"{_valid.iloc[i]['기준일']} · "
                                           f"{_valid.iloc[i]['품목']} {_valid.iloc[i]['품종']} · "
                                           f"{_valid.iloc[i]['가격']:,.0f}원 / "
                                           f"{_valid.iloc[i]['단위']}"),
                    key="kamis_rowsel")
                _row = _valid.iloc[_rowsel]
                _kgf = _row["kg환산계수"]
                if pd.isna(_kgf) or not _kgf:
                    st.error(f"❌ 단위 '{_row['단위']}'는 kg으로 환산할 수 없습니다. "
                             "경제성 분석 단가로 자동 적용하지 않습니다. "
                             "포장 단위를 확인한 뒤 직접 입력해 주세요.")
                else:
                    _per_kg = float(_row["가격"]) / float(_kgf)
                    st.info(f"환산: {_row['가격']:,.0f}원 ÷ {_kgf:g}kg = "
                            f"**{_per_kg:,.0f}원/kg** (근거: 단위 '{_row['단위']}')")
                    _iname = st.text_input("기준단가 항목 이름",
                                           value=f"{_row['품목']} {_row['품종']}".strip(),
                                           key="kamis_itemname")
                    _existing = st.session_state["price_db"]["항목"].astype(str).eq(_iname).any()
                    _apply_mode = st.radio(
                        "같은 이름이 있을 때", ["기존값 유지", "새 값으로 교체", "새 이름으로 추가"],
                        horizontal=True, key="kamis_apply_mode",
                        disabled=not _existing) if _existing else "새 이름으로 추가"
                    if st.button("➕ 기준단가에 반영", key="kamis_apply"):
                        if _existing and _apply_mode == "기존값 유지":
                            st.info("기존값을 유지했습니다.")
                            st.stop()
                        import datetime as _dt
                        _db = st.session_state["price_db"].copy()
                        _new = {"항목": _iname, "단가": round(_per_kg, 2), "단위": "원/kg",
                                "기준연도": str(_row["기준일"])[:4], "기준일": str(_row["기준일"]), "출처": str(_row["출처"]),
                                "조회방식": "KAMIS",
                                "갱신일": _dt.date.today().isoformat(),
                                "환산식": f"{_row['가격']:,.0f}원 ÷ {_kgf:g}kg",
                                "사용자수정": False}
                        if "기준일" not in _db.columns:
                            _db["기준일"] = ""
                        _final_name = _iname
                        if _existing and _apply_mode == "새 이름으로 추가":
                            _final_name = f"{_iname} ({_row['기준일']})"
                            _new["항목"] = _final_name
                        _hit = _db.index[_db["항목"].astype(str) == _final_name]
                        if len(_hit) and _apply_mode == "새 값으로 교체":
                            st.warning(f"'{_final_name}' 항목을 "
                                       f"기존 {_db.loc[_hit[0], '단가']} → 새 값 {_new['단가']}로 "
                                       "덮어씁니다.")
                            for k, v in _new.items():
                                if k in _db.columns:
                                    _db.loc[_hit[0], k] = v
                        else:
                            _db = pd.concat([_db, pd.DataFrame([_new])], ignore_index=True)
                        st.session_state["price_db"] = _db
                        log_action(f"KAMIS 단가 반영: {_final_name}")
                        st.success("기준단가에 반영했습니다."); st.rerun()
            st.caption("조회 결과를 그대로 덮어쓰지 않고, 위에서 확인 후 반영합니다.")

        st.markdown("---")
        st.markdown("###### 📊 KOSIS 통계 자동 조회 (농업노임·가격지수)")
        st.caption("통계청 국가통계포털(KOSIS) 인증키가 있으면 농촌 일용노임·농가구입가격지수를 "
                   "자동으로 받아올 수 있습니다. 발급: kosis.kr/openapi → 회원가입 → 활용신청(자동승인, 무료)")
        with st.expander("📖 사용법 (처음이라면 꼭 읽어보세요)"):
            st.markdown("""
**가장 쉬운 방법 — KOSIS에서 주소 복사해 오기**

1. KOSIS(kosis.kr)에서 원하는 통계표를 찾습니다.
   예) 통계청 → 농업 → **농가판매및구입가격조사** → 농촌임료금
2. 표 화면에서 **[OpenAPI]** 버튼을 누릅니다.
3. 인증키를 선택하면 **요청 주소(URL)**가 만들어집니다. 그 주소를 통째로 복사하세요.
4. 아래 칸에 붙여넣고 '조회'를 누르면 값이 표로 나옵니다.

**직접 입력하는 방법**
- 기관코드(orgId): 통계청은 **101**
- 통계표ID(tblId): KOSIS 통계표 주소에 있는 `DT_...` 형태의 값
- 시점(prdSe): Y(연간), Q(분기), M(월간)
""")
        _kk1, _kk2 = st.columns([3, 1])
        kosis_key = _kk1.text_input("KOSIS 인증키 (직접 입력 방식일 때만 필요)",
                                    type="password", key="kosis_key")
        kosis_mode = _kk2.radio("방식", ["주소 붙여넣기", "직접 입력"], key="kosis_mode")

        if kosis_mode == "주소 붙여넣기":
            _kurl = st.text_input("KOSIS OpenAPI 주소",
                                  placeholder="https://kosis.kr/openapi/Param/statisticsParameterData.do?method=getList&apiKey=...",
                                  key="kosis_url")
            if st.button("📡 KOSIS 조회", width="stretch", key="kosis_go1"):
                if not _kurl.strip():
                    st.warning("주소를 붙여넣어 주세요.")
                else:
                    with st.spinner("KOSIS에서 자료를 불러오는 중..."):
                        dfk, err = kosis_fetch_url(_kurl.strip())
                    if err: st.error(err)
                    else:
                        st.session_state["kosis_result"] = dfk
                        log_action("KOSIS 통계 조회")
        else:
            _c1, _c2, _c3, _c4 = st.columns(4)
            _org = _c1.text_input("기관코드", value="101", key="kosis_org")
            _tbl = _c2.text_input("통계표ID", placeholder="DT_...", key="kosis_tbl")
            _prd = _c3.selectbox("시점", ["Y", "Q", "M"], key="kosis_prd")
            _cnt = _c4.number_input("최근 몇 개", 1, 20, 5, key="kosis_cnt")
            if st.button("📡 KOSIS 조회", width="stretch", key="kosis_go2"):
                if not (kosis_key and _tbl.strip()):
                    st.warning("인증키와 통계표ID를 입력하세요.")
                else:
                    _u = kosis_build_url(kosis_key, _org.strip(), _tbl.strip(),
                                         prd_se=_prd, count=int(_cnt))
                    with st.spinner("KOSIS에서 자료를 불러오는 중..."):
                        dfk, err = kosis_fetch_url(_u)
                    if err: st.error(err)
                    else:
                        st.session_state["kosis_result"] = dfk
                        log_action("KOSIS 통계 조회")

        if st.session_state.get("kosis_result") is not None:
            _kr = st.session_state["kosis_result"]
            smart_table(_kr, width="stretch", hide_index=True)
            st.caption("조회된 값을 위 기준단가 표의 '단가' 칸에 입력해 사용하세요. "
                       "기준연도도 함께 적어두면 나중에 갱신할 때 편합니다.")
            if st.button("🗑️ 조회 결과 지우기", key="kosis_clear"):
                st.session_state["kosis_result"] = None; st.rerun()

        st.markdown("""
**⚠️ 기본값의 기준연도를 반드시 확인하세요.** 표의 '기준연도' 칸을 보고, 최신 자료가 있으면 갱신해 주세요.

**자동으로 받아올 수 있는 자료**
- **농산물 가격** : KAMIS 오픈API (위에서 조회 가능, 무료·일별 갱신)
  또는 공공데이터포털 aT '지역별 품목별 도·소매 가격정보'(승인 대기 없음)
- **농촌 일용노임 / 농가구입가격지수** : KOSIS 공유서비스 오픈API (kosis.kr/openapi, 무료·분기 갱신)
  → 통계청 기관코드 101, '농가판매 및 구입가격조사'
- **농지 임차료** : 농지공간포털·KOSIS·공공데이터포털 (연 1회, 7월 공표)
- **농기계 임대정보** : 공공데이터포털 표준데이터(15017325, 월 갱신)

**사람이 직접 확인해 입력해야 하는 자료**
- **소득조사 산정계수**(자가노력비 평가노임, 자본용역비 이자율, 감가상각 내용연수)
  → 농촌진흥청 「농축산물 소득자료집」 부록, 또는 농산업경영과(063-238-1197) 문의
- **농협 비료·농약 실판매가** (연 1회, 1월경 공표 / 보조금 적용 실구매가 기준 권장)
- **위탁영농비 표준단가** → 소득자료집 작목별 경영비 항목 참고
- **지역별(경북) 노임** → 시군 조사 또는 지역 농협 확인

**공식 자료 출처**
- 농업노임·소득자료 : 농촌진흥청 「농축산물 소득자료집」 (농사로 경영자료실)
- 농산물 가격 : KAMIS 농산물유통정보 / 통계청 농가판매가격조사
- 비료·자재 가격 : 농협 자재가격 정보
- 농기계 임차료 : 지역 농기계은행 임대료표
- 농지 임차료 : 통계청 농지임차료 조사

⚠️ 이 프로그램은 외부 자료를 자동으로 받아오지 않습니다. 기관 공식 자료를 직접 확인해
입력하셔야 정확한 결과가 나옵니다.
""")

    # ---------------- 부분예산표 (손실적/이익적 요소) ----------------
    if emode.startswith("📕"):
        st.caption("신기술을 도입했을 때 **늘어나는 비용(손실적 요소 A)**과 **늘어나는 수익(이익적 요소 B)**을 "
                   "정리해 추정수익액(B−A)을 계산합니다. 시험연구보고서에 그대로 쓰는 형식입니다.")
        with st.expander("ℹ️ 어떻게 쓰나요?", expanded=False):
            st.markdown("""
- **손실적 요소(A)**: 신기술 때문에 **더 들어가는 비용** + 줄어드는 수익
  - 예) 추가 인건비, 상품성 저하 손실, 추가 농약·비료대, 운송비, 부대경비
- **이익적 요소(B)**: 신기술 때문에 **늘어나는 수익** + 줄어드는 비용
  - 예) 판매수익 증가, 절감된 노동비
- **추정수익액 = B − A** → 양수면 도입할 가치가 있습니다.
""")
        with st.expander("🧮 이 숫자가 어떻게 나온 건가요? (예시로 따라가기)", expanded=False):
            st.markdown("""
부분예산은 **바뀐 것만** 계산합니다. 관행과 신기술이 똑같이 쓰는 비용(종묘비 등)은
양쪽에서 지워지므로 **아예 넣지 않습니다.** 그래서 표가 짧습니다.

**① 자료가 이렇다고 합시다** (모두 10a 기준, 반복이 있으면 평균)

| | 관행 | 신기술 | 차이 |
|---|---|---|---|
| 수량 | 295 kg | 335 kg | **+40 kg** |
| 고용노력비 | 600,000원 | 640,000원 | **+40,000원** |
| 농약비 | 300,000원 | 290,000원 | **−10,000원** |

**② 늘어난 건 A(손실적), 줄어든 건 B(이익적)로 보냅니다**

| 바뀐 것 | 계산 | 어디로 |
|---|---|---|
| 수량 +40 kg | 10,000원 × 40 = **400,000원** | 이익적 요소(B) |
| 고용노력비 +40,000원 | 640,000 − 600,000 = **40,000원** | 손실적 요소(A) |
| 농약비 −10,000원 | 300,000 − 290,000 = **10,000원** | 이익적 요소(B) |

> 비용이 **줄면 그만큼 번 것**이므로 이익 쪽(B)으로 갑니다. 헷갈리기 쉬운 부분입니다.

**③ 합쳐서 추정수익액을 냅니다**

```
계(A) = 40,000원
계(B) = 400,000 + 10,000 = 410,000원
추정수익액(B − A) = 410,000 − 40,000 = 370,000원/10a
```

즉 **"이 신기술을 쓰면 10a당 37만원 더 남는다"** 는 뜻입니다.

---

**자주 하는 실수**

| 실수 | 왜 안 되나 |
|---|---|
| 총수입·총경영비를 통째로 넣음 | 부분예산은 **차이만** 봅니다. 안 바뀐 비용은 넣지 마세요. |
| 비용 절감을 A(손실)에 넣음 | 절감은 **이익(B)** 입니다. |
| 30a 포장 값을 그대로 씀 | 10a 기준으로 나눠서 넣거나, 자동채움에서 기준면적을 30으로 지정하세요. |
| 자가노동을 뺌 | 노동이 늘었으면 **노임 × 늘어난 시간**을 A에 넣어야 실제 이득이 보입니다. |
| 노동'시간'을 비용 열에 넣음 | 10시간이 10원이 됩니다. 자동채움의 **노동시간 열**에 넣어 노임을 곱하세요. |
""")
        with st.expander("📊 올린 데이터에서 자동으로 채우기 (대조구 ↔ 신기술구 비교)",
                         expanded=False):
            st.caption("처리구별 수량·비용이 들어 있는 자료를 올렸다면, **대조구 대비 달라진 부분만** "
                       "뽑아서 아래 표를 자동으로 채워 드립니다. 채운 뒤 손으로 고쳐도 됩니다.")
            if df is None:
                st.info("왼쪽에서 데이터를 먼저 올리면 이 기능을 쓸 수 있습니다. "
                        "(데이터 없이 직접 입력해도 됩니다.)")
            elif not cat_cols or not num_cols:
                st.info("처리구(문자) 열과 수량·비용(숫자) 열이 모두 있어야 합니다.")
            else:
                a1, a2 = st.columns(2)
                _pbg = a1.selectbox("처리구 열", cat_cols, key="pbd_g")
                _lv = df[_pbg].dropna().astype(str).unique().tolist()
                _pbq = a2.selectbox("수량 열 (주산물)", num_cols,
                                    index=guess_idx(num_cols, ["수량", "수확량", "생산량", "상품수량"]),
                                    key="pbd_q")
                b1, b2 = st.columns(2)
                _pbc = b1.selectbox("대조구 (관행)", _lv,
                                    index=guess_idx(_lv, ["대조", "관행", "무처리", "control", "CK"]),
                                    key="pbd_c")
                _tv = [v for v in _lv if v != _pbc] or _lv
                _pbt = b2.selectbox("신기술 처리구", _tv, key="pbd_t")
                c1_, c2_ = st.columns(2)
                _pbp = c1_.number_input("단가 (원/수량 1단위)", 0, 100000000, 0, 100,
                                        key="pbd_price",
                                        help="수량 열이 kg이면 원/kg, 상자면 원/상자를 넣으세요.")
                _pba = c2_.number_input("자료 기준면적 (a)", 0.1, 1000.0, 10.0, 0.1,
                                        key="pbd_area",
                                        help="자료가 10a 기준이면 10, 1a(=100㎡) 기준이면 1. "
                                             "10a 기준으로 환산해 계산합니다.")
                _numopt = [c for c in num_cols if c != _pbq]
                _hour_guess = [c for c in _numopt if looks_like_hours(c)]
                _pbhour = st.multiselect(
                    "노동시간 열 (시간 단위 — 노임을 곱해 금액으로 바꿉니다)",
                    _numopt, default=_hour_guess, key="pbd_hour",
                    help="자가노동시간처럼 '원'이 아니라 '시간'으로 적힌 열입니다. "
                         "여기에 넣지 않으면 10시간이 10원으로 계산됩니다.")
                _pbwage = 0
                if _pbhour:
                    _pbwage = st.number_input(
                        "시간당 노임 (원/시간)", 0, 1000000,
                        int(get_price("농업노임(남, 시간)", 19190)), 10, key="pbd_wage",
                        help="기준단가 관리에 넣어 둔 값을 기본으로 씁니다. "
                             "여자 노임(15,174원) 등으로 바꿔도 됩니다.")
                _pbcost = st.multiselect(
                    "비용 열 (원 단위 — 신기술 때문에 달라지는 비용만)",
                    [c for c in _numopt if c not in _pbhour], key="pbd_cost",
                    help="인건비·자재비·농약비처럼 처리구마다 값이 다른 비용 열을 고릅니다. "
                         "'합계' 같은 계산 결과 열은 넣지 마세요.")
                _odd = [c for c in _pbcost
                        if looks_like_hours(c)
                        or any(k in str(c).replace(" ", "")
                               for k in ("단가", "가격", "판매가", "수량", "면적", "%"))]
                if _odd:
                    st.warning("⚠️ 비용 열로 보기 어려운 열이 섞여 있습니다: "
                               + ", ".join(map(str, _odd))
                               + " — 시간 단위면 위 '노동시간 열'로 옮기고, "
                                 "단가·수량이면 빼 주세요.")
                if st.button("📥 이 조건으로 아래 표 채우기", key="pbd_fill", width="stretch"):
                    try:
                        _L, _G, _D = partial_budget_from_data(
                            df, _pbg, _pbc, _pbt, _pbq, _pbp,
                            cost_cols=_pbcost, area_a=_pba,
                            hour_cols=_pbhour, wage_per_hour=_pbwage)
                        st.session_state["pb_auto"] = {"loss": _L, "gain": _G, "detail": _D,
                                                       "control": _pbc, "treated": _pbt}
                        st.session_state["pb_fill_n"] = st.session_state.get("pb_fill_n", 0) + 1
                        st.session_state["pb_demo"] = False
                        st.success(f"'{_pbc}' 대비 '{_pbt}'의 차이를 아래 표에 채웠습니다.")
                        log_action(f"부분예산표 데이터 자동채움: {_pbc} → {_pbt}")
                    except Exception as _ex:
                        st.error(f"❌ {_ex}")
                if st.session_state.get("pb_auto"):
                    st.markdown("###### 🔎 10a 환산 비교 (자동 채움 근거)")
                    smart_table(st.session_state["pb_auto"]["detail"],
                                 width="stretch", hide_index=True)
                    if st.button("↩️ 자동 채움 지우고 직접 입력", key="pbd_clear"):
                        st.session_state.pop("pb_auto", None)
                        st.session_state["pb_fill_n"] = st.session_state.get("pb_fill_n", 0) + 1
                        st.rerun()

        _pb_auto = st.session_state.get("pb_auto")
        if _pb_auto:
            st.info(f"📊 **'{_pb_auto['control']}' 대비 '{_pb_auto['treated']}'** 자료에서 자동으로 채운 표입니다. "
                    "빠진 항목(운송비·부대경비 등)은 직접 추가하세요.")
        else:
            st.warning("✏️ 지금은 **직접 입력 모드**입니다. 아래 칸에 적은 항목·금액으로만 계산합니다. "
                       "위 **📊 올린 데이터에서 자동으로 채우기**를 쓰면 자료에서 바로 만들 수 있어요.")
        _pb_example = st.checkbox("예시 값 채워보기 (연습용)", value=False, key="pb_demo",
                                  disabled=bool(_pb_auto),
                                  help="켜면 작성 방법을 보여주는 예시 숫자가 들어갑니다. "
                                       "실제 분석에서는 끄고 직접 입력하세요.")
        _pb_key = f"{int(_pb_example)}_{st.session_state.get('pb_fill_n', 0)}"
        st.markdown("##### ① 손실적 요소(A) — 늘어나는 비용")
        if _pb_auto:
            loss_default = _pb_auto["loss"].copy()
        elif _pb_example:
            loss_default = pd.DataFrame({
                "항목": ["인건비", "손실비", "농약·비료대", "운송비", "부대경비(수수료 등)"],
                "산출근거": ["130,000원 × 10명", "290kg × 0.2(상품율) × 28,000원", "", "", ""],
                "금액(원)": [1300000, 1624000, 200000, 725000, 435000]})
        else:
            loss_default = pd.DataFrame({
                "항목": ["", "", "", ""],
                "산출근거": ["", "", "", ""],
                "금액(원)": [None, None, None, None]})
        loss_df = st.data_editor(loss_default, num_rows="dynamic", width="stretch",
                                 key=f"pb_loss_{_pb_key}")
        st.markdown("##### ② 이익적 요소(B) — 늘어나는 수익")
        if _pb_auto:
            gain_default = _pb_auto["gain"].copy()
        elif _pb_example:
            gain_default = pd.DataFrame({
                "항목": ["판매수익 증가"],
                "산출근거": ["52,000원 × 145상자"],
                "금액(원)": [7540000]})
        else:
            gain_default = pd.DataFrame({
                "항목": ["", ""], "산출근거": ["", ""], "금액(원)": [None, None]})
        gain_df = st.data_editor(gain_default, num_rows="dynamic", width="stretch",
                                 key=f"pb_gain_{_pb_key}")
        _pb_crop = st.text_input("작목명 (보고서 문장에 사용, 선택)", key="pb_crop",
                                 placeholder="예) 풋고추").strip()

        if keep_running("pbtable", "부분예산표 만들기"):
            st.session_state.pop("cap_pbtbl", None)
            L, _loss_errors, _loss_check = validate_manual_budget_table(
                loss_df, "손실적 요소")
            G, _gain_errors, _gain_check = validate_manual_budget_table(
                gain_df, "이익적 요소")
            _manual_errors = _loss_errors + _gain_errors
            if L.empty and G.empty:
                _manual_errors.append("손실적 요소와 이익적 요소가 모두 비어 있습니다.")
            if _manual_errors:
                for _err in _manual_errors:
                    st.error("❌ " + _err)
                st.warning("입력 오류를 수정하기 전에는 계산·보고서·AI 해석을 만들지 않습니다.")
            else:
                _checks = pd.concat([_loss_check, _gain_check], ignore_index=True)
                if not _checks.empty:
                    _check_show = _checks.copy()
                    for _c in ["산출근거 계산값", "입력 금액", "차이"]:
                        if _c in _check_show.columns:
                            _check_show[_c] = _check_show[_c].map(
                                lambda v: "-" if pd.isna(v) else f"{round_half_up(v):,}")
                    st.markdown("##### 🔎 산출근거 자동 검산")
                    smart_table(_check_show, width="stretch", hide_index=True)
                    _mismatch = _checks["판정"].eq("확인 필요")
                    if _mismatch.any():
                        st.warning(f"⚠️ 산출근거와 입력 금액이 다른 항목이 {_mismatch.sum()}개 있습니다. "
                                   "금액 또는 산출근거를 확인하세요.")
                    else:
                        st.success("✅ 입력 금액과 산출근거 검산을 통과했습니다.")

                tot_a = float(L["금액(원)"].sum())
                tot_b = float(G["금액(원)"].sum())
                profit = tot_b - tot_a
                tot_a_show, tot_b_show = round_half_up(tot_a), round_half_up(tot_b)
                profit_show = round_half_up(profit)
                c1, c2, c3 = st.columns(3)
                c1.metric("손실적 요소(A)", f"{tot_a_show:,} 원")
                c2.metric("이익적 요소(B)", f"{tot_b_show:,} 원")
                c3.metric("추정수익액(B−A)", f"{profit_show:,} 원",
                          "입력 조건에서 양수" if profit > 0 else "재검토 필요")

                left = ["○ 증가되는 비용 :"]
                for _, r in L.iterrows():
                    base = f" ({r['산출근거']})" if str(r.get("산출근거", "")).strip() else ""
                    left.append(f"  - {r['항목']}{base} = {round_half_up(r['금액(원)']):,}원")
                left.append(f"  - 계(A) : {tot_a_show:,}원")
                right = ["○ 증가되는 이익 :"]
                for _, r in G.iterrows():
                    base = f" ({r['산출근거']})" if str(r.get("산출근거", "")).strip() else ""
                    right.append(f"  - {r['항목']}{base} = {round_half_up(r['금액(원)']):,}원")
                right.append(f"  - 계(B) : {tot_b_show:,}원")
                n = max(len(left), len(right))
                left += [""] * (n - len(left)); right += [""] * (n - len(right))
                pb_tbl = pd.DataFrame({"손실적 요소(A)": left, "이익적 요소(B)": right})
                st.markdown("##### ③ 부분예산표 (보고서용)")
                smart_table(pb_tbl, width="stretch", hide_index=True)
                concl = f"○ 추정수익액(B-A) : {tot_b_show:,} - {tot_a_show:,} = {profit_show:,}원"
                st.code(concl, language=None)

                _crop_txt = f"{_pb_crop} " if _pb_crop else ""
                txt = (f"○ {_crop_txt}신기술 도입에 따른 부분예산 분석 결과\n"
                       f"  - 증가되는 비용은 {tot_a_show:,}원, 증가되는 이익은 {tot_b_show:,}원이었다.\n"
                       f"  - 입력한 가격·수량·비용 조건에서 추정수익액은 {profit_show:,}원/10a으로 산출되었다.\n"
                       + ("  - 추정수익액이 양수이므로 경제성 검토 대상이 될 수 있으나, "
                          "가격·수량 변동과 현장 적용성을 함께 확인해야 한다."
                          if profit > 0 else
                          "  - 추정수익액이 0 이하이므로 현재 조건에서는 도입을 재검토할 필요가 있다."))
                st.markdown("###### 📋 보고서용 문장")
                st.code(txt, language=None)
                out_tbl = pd.concat([pb_tbl,
                                     pd.DataFrame({"손실적 요소(A)": [concl],
                                                   "이익적 요소(B)": [""]})],
                                    ignore_index=True)
                dl_table(out_tbl, "부분예산 분석표", "pbtbl", "부분예산표")
                log_action("부분예산표(손실적·이익적) 작성")
                report_capture("cap_pbtbl", "경제성 분석 (부분예산)", None,
                               blocks=[{"text": txt},
                                       {"caption": "부분예산 분석표", "table": out_tbl,
                                        "plain": True},
                                       {"caption": "산출근거 검산표", "table": _check_show}])
                ai_interpret_button("pbtable", "부분예산(손실적·이익적 요소)", pb_tbl,
                                    f"검증된 추정수익액은 {profit_show:,}원입니다. "
                                    "수치를 다시 계산하지 말고 제공된 값만 해석하세요. "
                                    + (f"작목은 {_pb_crop}입니다. " if _pb_crop else
                                       "작목명은 제공되지 않았으므로 특정 작목을 언급하지 마세요. "),
                                    capture_slot="cap_pbtbl")
        report_button("cap_pbtbl")


    # ---------------- 소득분석 ----------------
    elif emode.startswith("📗"):
        crop = st.selectbox("작목 유형 (농촌진흥청 소득조사 분류)",
                            ["식량작물", "노지채소", "시설채소", "노지과수", "시설과수", "특용·약용작물", "직접 지정"],
                            index=1, key="e_crop",
                            help="작목 유형에 따라 경영비 주요 항목과 분석 관점이 달라집니다.")
        _CROP_INFO = {
            "식량작물": ("벼·보리·콩·감자 등. 기계화율이 높아 위탁영농비·광열동력비 비중이 큽니다.",
                     ["종자비", "비료비", "농약비", "광열동력비", "위탁영농비", "제재료비", "감가상각비"]),
            "노지채소": ("고추·마늘·양파·배추·무 등. 노동집약적이라 고용노력비·종묘비 비중이 큽니다.",
                     ["종자비", "종묘비", "비료비", "농약비", "멀칭·피복재비", "지주·유인비", "고용노력비", "제재료비", "감가상각비"]),
            "시설채소": ("시설고추·토마토·오이·딸기 등. 시설 감가상각비·난방비·전기료가 큰 비중을 차지합니다.",
                     ["종묘비", "비료비", "농약비", "난방비", "전기료", "양액·배지비", "고용노력비", "시설감가상각비", "제재료비"]),
            "노지과수": ("사과·배·포도·감귤 등. **다년생**이라 과수원 조성비(묘목·유목기 관리)를 반영해야 합니다.",
                     ["비료비", "농약비", "봉지·피복재비", "전정·유인 노력비", "수분수·방화곤충비", "고용노력비", "조성비상각", "감가상각비"]),
            "시설과수": ("하우스 감귤·포도 등. 시설비 + 다년생 조성비가 모두 들어갑니다.",
                     ["비료비", "농약비", "난방비", "전기료", "봉지·피복재비", "전정·유인 노력비", "고용노력비", "시설감가상각비", "조성비상각"]),
            "특용·약용작물": ("인삼·약용작물 등. 다년근이 많아 조성비·검사수수료·피복자재비가 특징입니다.",
                        ["종묘비", "비료비", "농약비", "해가림·피복재비", "검사수수료", "고용노력비", "조성비상각", "감가상각비"]),
            "직접 지정": ("아래에서 경영비 항목을 직접 선택하세요.", []),
        }
        info, suggested = _CROP_INFO[crop]
        st.caption(f"💡 {info}")
        is_perennial = crop in ("노지과수", "시설과수", "특용·약용작물")
        is_facility = crop in ("시설채소", "시설과수")

        with st.expander("ℹ️ 계산 체계 (농촌진흥청 농축산물 소득조사 기준)"):
            st.markdown("""
- **총수입(조수입)** = 주산물가액(수량×단가) + **부산물가액**
- **경영비** = 종묘비·비료비·농약비·광열동력비·수리비·제재료비·소농구비·상각비·임차료·위탁영농비·고용노력비 등 **직접 지출 비용**
- **생산비** = 경영비 + 자가노력비 + 유동자본용역비 + 고정자본용역비 + 토지용역비
- **유동자본용역비** = (경영비 − 감가상각성 비용) × 이자율 × 1/2 × 재포기간(월/12)
- **고정자본용역비** = 고정자산 부분현재가 × 해당 작목 부담률 × 이자율
- **소득 = 총수입 − 경영비** ／ **순수익 = 총수입 − 생산비**
- **소득률(%) = (소득 ÷ 총수입) × 100** ／ 순수익률(%) = (순수익 ÷ 총수입) × 100
- 모든 지표는 **10a(1,000㎡) 기준**으로 환산합니다.

**작목별 추가 고려**
- 🌳 **과수·약용(다년생)**: 심은 뒤 수확까지 수년이 걸리므로 **과수원 조성비**(묘목비 + 유목기 관리비)를 내용연수로 나눠 매년 상각합니다.
- 🏠 **시설재배**: 하우스·난방기 등 **시설 감가상각비**와 **난방비·전기료**가 경영비의 큰 부분입니다.
""")
        with st.expander("📋 어떤 엑셀을 올려야 하나요? (여기부터 보세요)", expanded=False):
            st.markdown("""
**한 줄 = 한 조사구(처리구 × 반복)** 로 적으면 됩니다. 열 이름은 자유롭게 쓰셔도 되고,
화면에서 어떤 열이 수량인지 단가인지 골라 주면 됩니다.

**규칙은 딱 세 개입니다.**
1. **모든 값은 10a(1,000㎡) 기준**으로 적습니다. (30a 포장이면 3으로 나눠서 적기)
2. **비용은 원 단위 숫자**만 적습니다. 쉼표·'원' 글자가 있어도 자동 변환되지만, 빈칸은 0으로 채우세요.
3. **합계·소득같은 계산 결과 열은 넣지 마세요.** 프로그램이 다시 계산해서 이중으로 잡힙니다.
""")
            _samp_inc = pd.DataFrame({
                "처리": ["관행", "관행", "신기술", "신기술"],
                "반복": [1, 2, 1, 2],
                "수량(kg/10a)": [295, 305, 345, 352],
                "단가(원/kg)": [28000, 28000, 28000, 28000],
                "종묘비": [180000, 180000, 180000, 180000],
                "비료비": [210000, 210000, 240000, 240000],
                "농약비": [400000, 400000, 250000, 250000],
                "고용노력비": [600000, 600000, 640000, 640000],
                "자가노동시간": [95, 98, 105, 104]})
            smart_table(_samp_inc, width="stretch", hide_index=True)
            st.markdown("""
- **수량 열** → `수량(kg/10a)`, **단가 열** → `단가(원/kg)` 로 고릅니다.
- **경영비 열** → `종묘비`·`비료비`·`농약비`·`고용노력비` 를 모두 고릅니다.
- **자가노동시간** 은 자가노력비(= 시간 × 농촌임료금) 계산에 쓰입니다. 없으면 비워도 됩니다.
- 단가가 처리구마다 같다면 그냥 같은 값을 반복해 적으면 됩니다.
- 반복(1, 2, …)이 있으면 **처리구 평균**으로 묶어 계산합니다.
""")
            st.download_button("📥 이 서식 그대로 내려받기 (CSV)",
                               _samp_inc.to_csv(index=False).encode("utf-8-sig"),
                               "소득분석_입력서식.csv", width="stretch",
                               key="dl_inc_form")
            st.caption("내려받아서 우리 시험 숫자로 바꾼 뒤, 왼쪽 사이드바에 다시 올리면 됩니다.")

        if suggested:
            st.caption(f"이 작목의 대표 경영비 항목: {', '.join(suggested)}")

        allc = df.columns.tolist()
        # 콤마·단위가 섞여 문자로 읽힌 열도 숫자 후보에 포함 (분석 시 자동 변환)
        _numlike = list(find_numeric_like(df).keys())
        num_cols = list(dict.fromkeys(num_cols + _numlike))
        if _numlike:
            st.caption("💡 숫자로 보이는 문자 열도 선택할 수 있습니다(분석 시 자동 변환): "
                       f"{', '.join(_numlike)}")
        c1, c2, c3 = st.columns(3)
        trt = c1.selectbox("처리구 열", allc, index=guess_idx(allc, ["처리", "품종", "시험구", "구분"]), key="e_t")
        yq = c2.selectbox("수량 열 (kg/10a)", num_cols, index=guess_idx(num_cols, ["수량", "생산량", "수확량"]), key="e_y")
        pr = c3.selectbox("단가 열 (원/kg)", num_cols, index=guess_idx(num_cols, ["단가", "가격", "판매가"], 1), key="e_p")
        if yq == pr:
            st.warning("⚠️ 수량 열과 단가 열이 같습니다. 서로 다른 열을 선택하세요.")
        byp_opts = ["(없음)"] + [c for c in num_cols if c not in (yq, pr)]
        byp = st.selectbox("부산물가액 열 (선택)", byp_opts,
                           index=guess_idx(byp_opts, ["부산물"]), key="e_by")

        _ac1, _ac2 = st.columns([1, 2])
        area_val = _ac1.number_input("자료의 기준 면적", 0.1, 10000.0, 10.0, 0.1,
                                     key="e_area",
                                     help="입력 자료가 몇 a 기준인지 적으세요. 10a면 그대로 둡니다.")
        area_factor = 10.0 / float(area_val) if area_val else 1.0
        if abs(area_factor - 1.0) > 1e-9:
            _ac2.info(f"📐 자료가 {area_val}a 기준이므로 모든 수입·비용을 "
                      f"**10a 기준으로 {area_factor:.3f}배 환산**합니다.")
        else:
            _ac2.caption("모든 지표는 10a(1,000㎡) 기준으로 계산합니다.")

        st.markdown("##### 1️⃣ 경영비 비목")
        _cand_all = [c for c in num_cols if c not in (yq, pr, byp)]
        # ⑦ 합계열·별도 계산 항목은 자동 선택에서 제외 (중복계상 방지)
        _blocked = [c for c in _cand_all if is_excluded_cost(c)]
        _safe = [c for c in _cand_all if c not in _blocked]
        _auto = [c for c in _safe if str(c).endswith("비") or "비용" in str(c)]
        cost_cols = st.multiselect("경영비에 포함할 열", _cand_all, default=_auto, key="e_c")
        if _blocked:
            st.warning("⚠️ **중복계상 위험으로 자동 선택에서 뺀 열**: "
                       + ", ".join(map(str, _blocked))
                       + " — 이미 합계이거나(경영비합계 등) 아래에서 따로 계산되는 항목"
                         "(자가노력비·자본용역비·토지용역비)입니다.")
        _dup_sel = [c for c in cost_cols if is_excluded_cost(c)]
        if _dup_sel:
            st.error("❗ 선택한 열에 **중복 가능성**이 있습니다: " + ", ".join(map(str, _dup_sel))
                     + " — 경영비가 실제보다 크게 계산될 수 있습니다.")
        with st.expander("📋 비용 구성 확인 (계산 전 점검)"):
            st.markdown("**경영비에 합산될 열** (" + str(len(cost_cols)) + "개)\n\n"
                        + (", ".join(map(str, cost_cols)) if cost_cols else "(선택 없음)")
                        + "\n\n**아래에서 따로 계산되는 항목** (위에 중복 선택 금지)\n"
                        + "- 자가노력비 = 자가노동시간 × 농촌임료금\n"
                        + "- 유동·고정자본용역비 = 자본액 × 이자율\n"
                        + "- 토지용역비 = 직접 입력값\n\n"
                        + "**제외된 열**: "
                        + (", ".join(map(str, _blocked)) if _blocked else "없음"))
        # ⑨ 부가가치 계산을 위한 비용 분류 (자동 결정하지 않고 확인 요청)
        with st.expander("🏷️ 비용 분류 확인 (부가가치·손익분기 계산용)", expanded=False):
            st.caption("부가가치 = 총수입 − 중간재비. 아래에서 **중간재비가 아닌 항목**을 "
                       "골라 주세요. 자동으로 판단하지 않습니다.")
            def _guess(keys):
                return [c for c in cost_cols if any(k in str(c) for k in keys)]
            # '임차'가 들어갔다고 모두 토지 임차료는 아니다. '스마트장비임차비'처럼
            # 기계·장비 임차료가 토지 임차료로 잡히면, 아래 토지용역비와 중복이라는
            # 경고가 잘못 뜨면서 분석이 막힌다.
            _MACH_HINT = ["농기계", "기계", "장비", "시설", "하우스", "트랙터",
                          "관리기", "드론", "로봇", "스마트", "설비", "차량"]
            _LAND_HINT = ["토지", "농지", "지대", "밭", "논", "경지", "부지", "전답"]
            _rentish = _guess(["임차", "임대", "리스"])
            _g_mach = [c for c in _rentish if any(k in str(c) for k in _MACH_HINT)]
            _g_rent = [c for c in _guess(["임차", "임대", "지대"])
                       if c not in _g_mach
                       and (any(k in str(c) for k in _LAND_HINT)
                            or str(c).strip() in ("임차료", "임대료", "지대", "임차비", "임대비"))]
            _g_hire = _guess(["고용", "노력", "노임", "인건"])
            _g_trust = _guess(["위탁", "대행"])
            _g_dep = _guess(["상각", "감가"])
            cls1, cls2 = st.columns(2)
            rent_col = cls1.multiselect("토지 임차료", cost_cols, default=_g_rent, key="e_rent")
            hire_col = cls2.multiselect("고용노력비", cost_cols, default=_g_hire, key="e_hire")
            cls3, cls4 = st.columns(2)
            mach_col = cls3.multiselect("농기계·시설 임차료", cost_cols,
                                        default=_g_mach, key="e_mach")
            trust_col = cls4.multiselect("위탁영농비", cost_cols, default=_g_trust, key="e_trust")
            dep_col = st.multiselect("감가상각비", cost_cols, default=_g_dep, key="e_dep")
            st.markdown("---")
            _g_yv = _guess(["수확", "선별", "포장", "상자", "박스", "운송", "운반",
                            "출하", "유통", "선과", "저장"])
            yield_var_col = st.multiselect(
                "📦 수량에 비례하는 비용 (손익분기 계산용)", cost_cols,
                default=_g_yv, key="e_yieldvar",
                help="수확·선별·포장·운송비처럼 '수량이 줄면 같이 줄어드는' 비용만 고르세요. "
                     "종묘비·비료비·토지용역비처럼 면적에 대해 정해지는 비용은 고르지 "
                     "않습니다. 아무것도 고르지 않으면 손익분기수량 = "
                     "(생산비 − 부산물가액) ÷ 단가 가 됩니다.")
            _excl_va = list(dict.fromkeys(rent_col + hire_col + mach_col + trust_col + dep_col))
            _inter_cols = [c for c in cost_cols if c not in _excl_va]
            _amb = [c for c in cost_cols
                    if c not in _excl_va and any(k in str(c) for k in
                                                 ("료", "임", "용역", "수수료"))]
            st.markdown("**중간재비로 계산될 항목** (" + str(len(_inter_cols)) + "개)\n\n"
                        + (", ".join(map(str, _inter_cols)) if _inter_cols else "(없음)")
                        + "\n\n**중간재비에서 제외될 항목**\n\n"
                        + (", ".join(map(str, _excl_va)) if _excl_va else "(없음)"))
            if _amb:
                st.warning("⚠️ 분류가 애매한 항목이 있습니다: " + ", ".join(map(str, _amb))
                           + " — 위 분류에 넣을지 직접 확인해 주세요.")
        c4, c5 = st.columns(2)


        with st.expander("2️⃣ 자가 요소 (생산비·순수익 계산용)"):
            c6, c7 = st.columns(2)
            _lab = ["(없음)"] + [c for c in num_cols if c not in (yq, pr)]
            labor_col = c6.selectbox("자가노동시간 열", _lab,
                                     index=guess_idx(_lab, ["노동시간", "자가노동", "노력시간"]), key="e_l")
            wage = c7.number_input("농촌임료금 (원/시간)", 0, 200000,
                                   int(get_price("농업노임(남, 시간)", 19190)), 500,
                                   key="e_wage",
                                   help="자가노력비 = 자가노동시간 × 농촌임료금. 연도·지역별 실제 임료금을 입력하세요.")
            c8, c9 = st.columns(2)
            rate = c9.number_input("자본 이자율 (%)", 0.0, 20.0,
                                   float(get_price("자본이자율", 5.0)), 0.1, key="e_rate")
            _fa_mode = c8.radio("고정자산 입력", ["부분현재가 직접 입력", "신조가에서 자동 계산"],
                                horizontal=False, key="e_fixed_mode")
            _fixed_input_error = None
            if _fa_mode == "부분현재가 직접 입력":
                _fa1, _fa2 = st.columns(2)
                fixed_asset = _fa1.number_input(
                    "고정자산 부분현재가/평가액 (원/10a)", 0, 500000000, 0, 100000,
                    key="e_fixedasset",
                    help="대농구·영농시설의 현재 가치(부분현재가)를 해당 자산의 10a 기준으로 입력하세요.")
                fixed_asset_use_rate = _fa2.number_input(
                    "해당 작목 부담률 (%)", 0.0, 100.0, 100.0, 5.0, key="e_fixed_use",
                    help="이 자산을 여러 작목에 함께 쓰면 해당 작목이 부담할 비율만 입력합니다.")
            else:
                _fa1, _fa2, _fa3, _fa4 = st.columns(4)
                _new_value = _fa1.number_input("신조가 (원/10a)", 0, 1000000000, 0, 100000,
                                              key="e_fixed_new")
                _residual = _fa2.number_input("잔존가치 (원/10a)", 0, 1000000000, 0, 100000,
                                             key="e_fixed_residual")
                _life = _fa3.number_input("내용연수(년)", 1, 50, 10, 1, key="e_fixed_life")
                _used = _fa4.number_input("사용연수(년)", 0, 50, 0, 1, key="e_fixed_used")
                _fixed_input_error = None
                if _residual > _new_value:
                    _fixed_input_error = "잔존가치는 신조가보다 클 수 없습니다."
                    st.error("❌ " + _fixed_input_error)
                    fixed_asset = 0.0
                elif _used > _life:
                    _fixed_input_error = "사용연수는 내용연수보다 클 수 없습니다."
                    st.error("❌ " + _fixed_input_error)
                    fixed_asset = 0.0
                else:
                    _annual_dep = ((_new_value - _residual) / float(_life)) if _life else 0.0
                    fixed_asset = max(float(_new_value) - _annual_dep * float(_used), float(_residual))
                    st.caption(f"부분현재가 = 신조가 − (연간 감가상각비 × 사용연수) = **{fixed_asset:,.0f}원/10a**")
                fixed_asset_use_rate = st.number_input(
                    "해당 작목 부담률 (%)", 0.0, 100.0, 100.0, 5.0, key="e_fixed_use_auto",
                    help="이 자산을 여러 작목에 함께 쓰면 해당 작목이 부담할 비율만 입력합니다.")
            months = st.slider("재포기간 (개월)", 1, 12, 6, key="e_months",
                               help="유동자본용역비 = 유동자본 기준액 × 연이자율 × 1/2 × 재포기간(월/12). "
                                    "감가상각비와 조성비상각은 유동자본 기준액에서 제외합니다.")
            st.caption("유동자본은 경영비 전액에 이자를 붙이지 않습니다. 감가상각성 비용을 제외하고 "
                       "농촌진흥청 방식의 산출계수 1/2와 재포기간을 적용합니다.")
            # ⑧ 토지 이용 형태에 따라 토지비 산정 방식을 나눔
            st.markdown("**토지 이용 형태**")
            land_type = st.radio("토지 이용", ["자가 소유", "임차", "자가·임차 혼합", "토지비 제외"],
                                 horizontal=True, key="e_landtype", label_visibility="collapsed")
            _base_land = int(get_price("토지용역비(밭)", 260) * 1000)
            if land_type == "토지비 제외":
                land_opp = 0
                land_cash = 0
                st.caption("토지 관련 비용을 별도로 추가하지 않습니다.")
            elif land_type == "자가 소유":
                land_opp = st.number_input("자가토지 기회비용 (원/10a)", 0, 20000000,
                                           _base_land, 10000, key="e_land_own",
                                           help="자가토지를 빌려줬다면 받을 수 있는 임차료 수준. 생산비에만 반영합니다.")
                land_cash = 0
            elif land_type == "임차":
                land_cash = st.number_input("실제 토지 임차료 (원/10a)", 0, 20000000,
                                            _base_land, 10000, key="e_land_rent",
                                            help="실제로 지급한 임차료이므로 경영비에 반영합니다.")
                land_opp = 0
                st.warning("⚠️ 경영비 항목에 **토지 임차료 열이 이미 포함**되어 있다면 여기서는 0으로 두세요. "
                           "그렇지 않으면 임차료가 두 번 계산됩니다.")
            else:   # 혼합
                _lc1, _lc2 = st.columns(2)
                _own_ratio = _lc1.slider("자가 비율(%)", 0, 100, 50, 5, key="e_ownratio")
                _own_cost = _lc2.number_input("자가 기회비용 (원/10a, 자가 100% 가정)", 0, 20000000,
                                              _base_land, 10000, key="e_owncost")
                _rent_cost = st.number_input("임차료 (원/10a, 임차 100% 가정)", 0, 20000000,
                                             _base_land, 10000, key="e_rentcost")
                land_opp = round_half_up(_own_cost * _own_ratio / 100)
                land_cash = round_half_up(_rent_cost * (100 - _own_ratio) / 100)
                st.caption(f"자가토지 용역비 **{land_opp:,}원/10a**(생산비에만 반영) + "
                           f"실제 임차료 **{land_cash:,}원/10a**(경영비에 반영)")
            land = land_opp

        # 작목별 특수 항목
        estab_amort = 0
        if is_perennial:
            with st.expander("🌳 과수·다년생 작물 — 과수원 조성비 상각 (중요)", expanded=True):
                st.caption("과수·약용작물은 심은 뒤 수확까지 수년이 걸립니다. 그동안 든 조성비를 "
                           "성목 이후 매년 나눠서 비용으로 반영(상각)합니다.")
                cc1, cc2 = st.columns(2)
                estab_total = cc1.number_input("총 조성비 (원/10a)", 0, 200000000, 0, 100000,
                                               key="e_estabtotal",
                                               help="묘목비 + 유목기(미결실기) 관리비 합계")
                useful_years = cc2.number_input("성목 이후 내용연수 (년)", 1, 40, 15, 1,
                                                key="e_usefulyears",
                                                help="과수원을 경제적으로 이용하는 총 기간")
                estab_amort = round_half_up(estab_total / useful_years) if useful_years else 0
                if estab_amort:
                    st.info(f"연간 조성비 상각액 = {estab_total:,.0f} ÷ {useful_years}년 "
                            f"= **{estab_amort:,.0f} 원/10a** (경영비에 매년 포함)")
        if is_facility:
            st.caption("🏠 시설재배는 하우스·난방기 감가상각비, 난방비, 전기료를 경영비 항목에 꼭 포함하세요.")

        ctrl = st.selectbox("대조구(비교 기준)", ["(없음)"] + df[trt].astype(str).unique().tolist(), key="e_ctrl")
        _unit_confirm = st.checkbox(
            "수량=kg/입력 기준면적, 단가=원/kg, 비용·부산물가액=원/입력 기준면적임을 확인했습니다.",
            key="e_unit_confirm",
            help="단위가 다르면 금액이 10배·1000배 틀릴 수 있습니다. 원/10kg, ton, 천원 단위는 먼저 환산하세요.")

        _econ_errors = validate_economic_inputs(
            df, trt, yq, pr, cost_cols,
            labor_col=(None if labor_col == "(없음)" else labor_col),
            byproduct_col=(None if byp == "(없음)" else byp),
            wage_per_hour=wage, interest_rate=rate, capital_months=months,
            fixed_asset_per_10a=fixed_asset, fixed_asset_use_rate_percent=fixed_asset_use_rate,
            establishment_amort_per_10a=estab_amort, depreciation_cost_cols=dep_col,
            land_cost_per_10a=land_opp, land_cash_rent_per_10a=land_cash,
            land_type=land_type, source_area_a=area_val)
        if _fixed_input_error:
            _econ_errors.append(_fixed_input_error)
        if not _unit_confirm:
            _econ_errors.append("입력 열의 단위를 확인해야 계산할 수 있습니다.")
        _econ_errors = list(dict.fromkeys(_econ_errors))
        # 계산 엔진(economic_core)은 열 이름에 '임차'가 있으면 토지 임차료로 보고
        # 별도 토지용역비와 중복이라고 알린다. 그러나 '스마트장비임차비'처럼 기계·장비
        # 임차료는 토지와 무관하다. 위 '비용 분류 확인'에서 사용자가 **토지 임차료로
        # 지정한 열**이 있을 때만 진짜 중복이므로, 그 기준으로 다시 판정한다.
        _kept = []
        for _err in _econ_errors:
            if (("토지용역비" in _err and ("임차" in _err or "임대" in _err))
                    or ("토지 임차료" in _err and "경영비" in _err)):
                _listed = [t.strip() for t in _err.split(":")[-1].split(",") if t.strip()]
                _hit = [t for t in _listed if t in rent_col]
                if not _hit:
                    continue                     # 기계·장비 임차료였다 → 중복 아님
                _err = ("경영비에 토지 임차료(" + ", ".join(_hit) + ")가 들어 있어 "
                        "별도 입력한 토지 임차료와 중복 계산됩니다. 별도 토지 임차료를 0으로 두거나, "
                        "해당 열을 경영비에서 빼 주세요.")
            _kept.append(_err)
        _econ_errors = _kept
        for _err in _econ_errors:
            st.error("❌ " + _err)

        if keep_running("econ", "소득분석 실행", disabled=bool(_econ_errors)):
            st.session_state.pop("cap_econ", None)
            for _key in ["_econ_rowlevel", "_econ_test_소득", "_econ_test_순수익",
                         "_econ_test_수량", "_econ_signature"]:
                st.session_state.pop(_key, None)
            try:
                _row, _summary = calculate_row_economics(
                    df, trt, yq, pr, cost_cols,
                    byproduct_col=(None if byp == "(없음)" else byp),
                    labor_col=(None if labor_col == "(없음)" else labor_col),
                    wage_per_hour=wage, interest_rate=rate, capital_months=months,
                    fixed_asset_per_10a=fixed_asset, fixed_asset_use_rate_percent=fixed_asset_use_rate,
                    land_cost_per_10a=land_opp, land_cash_rent_per_10a=land_cash,
                    establishment_amort_per_10a=estab_amort,
                    source_area_a=area_val, depreciation_cost_cols=dep_col)
            except ValueError as _ex:
                st.error(f"경제성 계산을 중단했습니다: {_ex}")
                st.stop()

            _validation_errors = validate_economic_results(_row, _summary)
            if _validation_errors:
                for _err in _validation_errors:
                    st.error("❌ " + _err)
                st.error("계산 결과 역산 검증에 실패하여 결과를 표시하지 않습니다.")
                st.stop()
            st.success("✅ 입력·면적 환산·수입·비용·소득·순수익 역산 검증을 통과했습니다.")

            _rep = _row.groupby(trt).size()
            _aggregated = bool((_rep > 1).any())
            if _aggregated:
                _n = int(_rep.max())
                st.session_state["_econ_rowlevel"] = _row[[trt, "_소득", "_순수익",
                                                           "_총수입", "_경영비"]].copy()
                st.session_state["_econ_signature"] = dataframe_signature(df)
                st.info(f"🔁 반복 자료가 확인되어(최대 {_n}반복) 각 반복의 수입·비용을 먼저 계산한 뒤 "
                        "처리구 평균으로 집계했습니다.")

            _rename = {
                "_수량10a": yq, "_주산물가액": "주산물가액",
                "_부산물가액": "부산물가액", "_경영비": "경영비",
                "_조성비상각": "조성비상각", "_토지임차료": "토지임차료(별도입력)",
                "_자가노력비": "자가노력비", "_유동자본기준액": "유동자본기준액",
                "_유동자본용역비": "유동자본용역비", "_고정자본기준액": "고정자본기준액",
                "_고정자본용역비": "고정자본용역비", "_토지용역비": "토지용역비",
                "_총수입": "총수입", "_생산비": "생산비",
                "_소득": "소득", "_순수익": "순수익",
            }
            e = _summary.rename(columns=_rename).reset_index(drop=True)
            if estab_amort and "조성비상각" not in e.columns:
                e["조성비상각"] = float(estab_amort)

            # ---------- 통계 유의성 연동 (수량 ANOVA, 블록 지정 가능) ----------
            yield_p, yield_src = None, ""
            yield_test = {"source": "미검정", "status": "unknown", "p_value": None}
            _blk_for_test = None
            if (_rep > 1).any() and _rep.size >= 2:
                _blk_opts = ["(없음)"] + [c for c in df.columns if c not in (trt, yq, pr)]
                _bsel = st.selectbox("수량 검정에 쓸 반복(블록) 열", _blk_opts,
                                     index=guess_idx(_blk_opts, ["반복", "블록", "block", "rep"]),
                                     key="econ_blk",
                                     help="난괴법이면 블록을 지정해야 정확한 검정이 됩니다.")
                _blk_for_test = None if _bsel == "(없음)" else _bsel
                # 반복별 소득·순수익 검정도 수량 ANOVA와 같은 블록을 사용하도록 보관한다.
                if _blk_for_test and st.session_state.get("_econ_rowlevel") is not None:
                    _erl = st.session_state["_econ_rowlevel"].copy()
                    try:
                        _erl[_blk_for_test] = df.loc[_erl.index, _blk_for_test]
                        st.session_state["_econ_rowlevel"] = _erl
                    except Exception:
                        st.caption("소득 검정용 블록 열을 연결하지 못해 소득 검정은 블록 없이 수행됩니다.")
                try:
                    _cols_t = [trt, yq] + ([_blk_for_test] if _blk_for_test else [])
                    _raw = df[_cols_t].copy()
                    _raw[yq] = (_raw[yq] if pd.api.types.is_numeric_dtype(_raw[yq])
                                else to_numeric_clean(_raw[yq]))
                    _raw = _raw.dropna()
                    _okv, _ = validate_anova_data(_raw, trt, yq)
                    if _okv:
                        _f = safe_formula(yq, [trt] + ([_blk_for_test] if _blk_for_test else []))
                        _m = ols(_f, data=_raw).fit()
                        _a = sm.stats.anova_lm(_m, typ=2)
                        _k = f"C({q_ref(trt)})"
                        yield_p = float(_a.loc[_k, "PR(>F)"]) if _k in _a.index else float(_a["PR(>F)"].iloc[0])
                        yield_src = "반복 자료로 자동 검정"
                        yield_test = {"source": "ANOVA 자동 검정",
                                      "status": "significant" if yield_p < 0.05 else "not_significant",
                                      "p_value": round(yield_p, 6),
                                      "model": ("수량 ~ 처리구 + 블록" if _blk_for_test
                                                else "수량 ~ 처리구")}
                except Exception as _ex:
                    yield_p = None
                    st.caption(f"수량 자동 검정을 하지 못했습니다: {str(_ex)[:60]}")
            if yield_p is None:
                st.markdown("###### 🔬 수량의 통계적 유의성")
                st.caption("반복 자료가 없어 자동 검정을 못 했습니다. "
                           "분산분석 결과를 알고 있다면 알려 주세요. (경제성 해석의 신뢰도에 영향)")
                _sig_choice = st.radio("처리 간 수량 차이가 통계적으로 유의했나요?",
                                       ["모름 / 검정 안 함", "유의함 (p < 0.05)", "유의하지 않음 (p ≥ 0.05)"],
                                       horizontal=True, key="econ_sig")
                # ⑩ 사용자가 고른 '상태'만 기록하고 가짜 p값은 만들지 않는다
                if _sig_choice.startswith("유의함"):
                    yield_test = {"source": "사용자 입력", "status": "significant", "p_value": None}
                elif _sig_choice.startswith("유의하지"):
                    yield_test = {"source": "사용자 입력", "status": "not_significant", "p_value": None}
                else:
                    yield_test = {"source": "미검정", "status": "unknown", "p_value": None}
            sig_warn = ""
            if yield_p is not None:
                if yield_p >= 0.05:
                    sig_warn = ("⚠️ 처리 간 수량 차이는 통계적으로 유의하지 않았습니다. "
                                "경제성 결과는 관측된 수량·가격·비용의 점추정치이며, "
                                "소득 차이 자체의 통계적 유의성은 별도로 검정하지 않았습니다.")
                    st.warning(f"{sig_warn} ({yield_src}"
                               + (f", p = {yield_p:.4f}" if yield_src.startswith("반복") else "") + ")")
                else:
                    st.success("✅ 수량에서 처리 간 유의한 차이가 확인되었습니다"
                               + (f" (p = {yield_p:.4f}, {yield_src})." if yield_src.startswith("반복")
                                  else f" ({yield_src})."))

            # 계산 엔진의 원값을 유지하고, 화면 표시 단계에서만 반올림한다.
            _rev = e["총수입"].where(e["총수입"] != 0)
            # 나눗셈 결과는 소수점이 끝없이 이어지므로(52.266666...), 만들 때 바로 반올림한다.
            # (표시 직전에만 반올림하면 다운로드·AI 해석 경로가 이 반올림을 매번 건너뛰기 쉽다)
            e["소득률(%)"] = (e["소득"] / _rev * 100).round(2)
            e["순수익률(%)"] = (e["순수익"] / _rev * 100).round(2)
            _va_excl = list(dict.fromkeys(rent_col + hire_col + mach_col + trust_col + dep_col))
            _va_excl = [c for c in _va_excl if c in e.columns]
            # 선택한 비용 열은 계산 엔진에서 이미 10a 기준으로 환산되어 있다.
            _non_intermediate = (e[_va_excl].sum(axis=1) if _va_excl else pd.Series(0.0, index=e.index))
            if "토지임차료(별도입력)" in e.columns:
                _non_intermediate = _non_intermediate + e["토지임차료(별도입력)"]
            if "조성비상각" in e.columns:
                _non_intermediate = _non_intermediate + e["조성비상각"]
            inter = (e["경영비"] - _non_intermediate).clip(lower=0)
            e["부가가치"] = e["총수입"] - inter
            e["가족노동보수"] = (e["총수입"] - e["경영비"] - e["유동자본용역비"]
                            - e["고정자본용역비"] - e["토지용역비"])
            e["단년도 총수입/생산비"] = (e["총수입"] / e["생산비"].where(e["생산비"] != 0)).round(2)
            e["kg당 생산비"] = (e["생산비"] / e[yq].where(e[yq] != 0)).round(0)
            # 손익분기수량 — CVP 공식 Q* = 고정비 ÷ (단가 − 단위당 변동비).
            # 핵심은 '무엇이 변동비인가'다. 변동비는 **산출량(수량)에 비례하는** 비용만
            # 해당한다. 10a 기준 작물 예산에서 경영비 대부분(종묘·비료·농약·고용노력비)과
            # 자가노력비·토지용역비는 면적에 대해 정해지는 비용이라, 그해 수량이 줄어도
            # 같이 줄지 않는다. 이걸 전부 변동비로 넣으면 단위당 변동비가 부풀고 고정비가
            # 거의 남지 않아 손익분기수량이 비현실적으로 작게 나온다
            # (예: 실제 318kg인데 손익분기 23kg → 어떤 처리든 무조건 흑자로 보임).
            # 그래서 기본은 '수량비례비 없음' = 생산비 전액을 회수 대상으로 보고,
            # 수확·선별·포장·운송처럼 실제로 수량에 비례하는 비용만 사용자가 지정한다.
            _yield_var_cost_cols = [c for c in yield_var_col if c in e.columns]
            _vc_total = (e[_yield_var_cost_cols].sum(axis=1) if _yield_var_cost_cols
                         else pd.Series(0.0, index=e.index))
            # 부산물가액은 주산물 수량과 무관하게 들어오는 수입이므로 회수 대상에서 뺀다.
            _fc = (e["생산비"] - _vc_total - e["부산물가액"]).clip(lower=0)
            e["손익분기수량"] = break_even_qty(_fc, _vc_total, e[yq], e[pr]).round(1)
            e["손익분기가격"] = ((e["생산비"] - e["부산물가액"])
                                / e[yq].where(e[yq] != 0)).round(0)
            e["손익분기수량 계산불가"] = (e[pr] - _vc_total / e[yq].where(e[yq] != 0)) <= 0
            if labor_col != "(없음)":
                e["시간당 소득"] = (e["소득"] / e[labor_col].where(e[labor_col] != 0)).round(0)
            _zero_den = []
            if e["총수입"].eq(0).any(): _zero_den.append("총수입")
            if e["생산비"].eq(0).any(): _zero_den.append("생산비")
            if e[yq].eq(0).any(): _zero_den.append("수량")
            if e[pr].eq(0).any(): _zero_den.append("단가")
            if labor_col != "(없음)" and e[labor_col].eq(0).any(): _zero_den.append("자가노동시간")
            if _zero_den:
                st.warning("⚠️ " + ", ".join(_zero_den)
                           + "이 0인 처리구의 비율 지표는 계산 불가로 표시합니다.")
            # ---- 대조구 대비 증수 분석 (시험연구보고서 핵심 지표) ----
            # ③ 열을 먼저 NaN으로 만들어 두어 KeyError·ZeroDivisionError 방지
            qty_effect_ok = False
            for _c in ["증수량", "증수율(%)", "증수액", "추가투입비", "경영비증가액",
                       "소득증가액", "순수익증가액", "총수입증가액",
                       "수량효과 기준 순증가수익", "소득지수", "투자효율", "대조구대비 소득(%)"]:
                e[_c] = np.nan
            has_ctrl = False
            if ctrl != "(없음)":
                cmask = e[trt].astype(str) == ctrl
                if not cmask.any():
                    st.warning(f"⚠️ 대조구 '{ctrl}'을(를) 자료에서 찾지 못해 증수 분석을 생략합니다.")
                else:
                    def _safe_mean(series):
                        v = pd.to_numeric(series, errors="coerce").mean()
                        return None if (v is None or pd.isna(v)) else float(v)
                    b_inc = _safe_mean(e.loc[cmask, "소득"])
                    b_yld = _safe_mean(e.loc[cmask, yq])
                    b_cost = _safe_mean(e.loc[cmask, "경영비"])
                    b_prof = _safe_mean(e.loc[cmask, "순수익"])
                    b_rev = _safe_mean(e.loc[cmask, "총수입"])
                    if b_inc not in (None, 0):
                        e["대조구대비 소득(%)"] = ((e["소득"] - b_inc) / abs(b_inc) * 100).round(1)
                        e["소득지수"] = (e["소득"] / b_inc * 100).round(1)
                    elif b_inc == 0:
                        st.warning("⚠️ 대조구의 소득이 0이어서 소득지수·대조구대비 소득은 계산하지 않았습니다.")
                    if b_yld in (None, 0):
                        st.warning("⚠️ 대조구의 수량이 0이거나 비어 있어 증수율을 계산할 수 없습니다. "
                                   "증수량·증수액만 표시합니다."
                                   if b_yld == 0 else
                                   "⚠️ 대조구 수량이 비어 있어 증수 분석을 생략합니다.")
                        if b_yld == 0:
                            has_ctrl = True
                            e["증수량"] = e[yq] - 0
                            e["증수액"] = e["증수량"] * e[pr]
                            if b_cost is not None:
                                e["경영비증가액"] = e["경영비"] - b_cost
                                e["추가투입비"] = e["경영비증가액"]
                            if b_inc is not None:
                                e["소득증가액"] = e["소득"] - b_inc
                    else:
                        has_ctrl = True
                        e["증수량"] = e[yq] - b_yld
                        e["증수율(%)"] = ((e[yq] - b_yld) / b_yld * 100).round(1)
                        e["증수액"] = e["증수량"] * e[pr]
                        # ④ 기본 판단지표 = 실제 차이 (가격·부산물 차이까지 모두 반영)
                        if b_inc is not None:
                            e["소득증가액"] = e["소득"] - b_inc
                        if b_prof is not None:
                            e["순수익증가액"] = e["순수익"] - b_prof
                        if b_rev is not None:
                            e["총수입증가액"] = e["총수입"] - b_rev
                        if b_cost is not None:
                            e["경영비증가액"] = e["경영비"] - b_cost
                            e["추가투입비"] = e["경영비증가액"]
                        # 보조지표: 단가·부산물이 모두 같을 때만 계산 (조건 불충족 시 미표시)
                        _same_price = e[pr].nunique(dropna=True) <= 1
                        _same_byp = (e["부산물가액"].nunique(dropna=True) <= 1
                                     if "부산물가액" in e.columns else True)
                        qty_effect_ok = bool(_same_price and _same_byp)
                        if qty_effect_ok and b_cost is not None:
                            e["수량효과 기준 순증가수익"] = e["증수액"] - e["경영비증가액"]
                            with np.errstate(divide="ignore", invalid="ignore"):
                                e["투자효율"] = np.where(
                                    pd.to_numeric(e["추가투입비"], errors="coerce") > 0,
                                    e["증수액"] / e["추가투입비"].replace(0, np.nan),
                                    np.nan)
                                e["투자효율"] = pd.Series(e["투자효율"], index=e.index).round(2)

            st.markdown("#### 1) 소득 · 순수익")
            m1 = [trt, "총수입", "경영비", "생산비", "소득", "순수익", "소득률(%)", "순수익률(%)"]
            smart_table(money_table(e[m1]), width="stretch", hide_index=True)

            st.markdown("#### 2) 생산비 구성")
            m2 = [trt, "경영비", "자가노력비", "유동자본용역비", "고정자본용역비", "토지용역비", "생산비"]
            smart_table(money_table(e[m2]), width="stretch", hide_index=True)
            _dep_txt = ", ".join(map(str, dep_col)) if dep_col else "지정 없음(열 이름으로 자동 판별)"
            st.caption(f"유동자본용역비는 감가상각비({ _dep_txt })와 조성비상각을 제외한 유동자본 기준액에 "
                       f"이자율 {rate:g}% × 1/2 × 재포기간 {months}/12를 적용했습니다. "
                       f"고정자본용역비는 부분현재가 {fixed_asset:,.0f}원/10a × 작목부담률 "
                       f"{fixed_asset_use_rate:g}% × 이자율 {rate:g}%로 계산했습니다.")

            st.markdown("#### 3) 경영 지표")
            m3 = [trt, "부가가치", "가족노동보수", "단년도 총수입/생산비",
                  "kg당 생산비", "손익분기수량", "손익분기가격"]
            if "시간당 소득" in e.columns: m3.append("시간당 소득")
            if "대조구대비 소득(%)" in e.columns: m3.append("대조구대비 소득(%)")
            smart_table(money_table(e[m3]), width="stretch", hide_index=True)
            if e["손익분기수량 계산불가"].any():
                _bad = ", ".join(str(g) for g in e.loc[e["손익분기수량 계산불가"], trt])
                st.warning(f"⚠️ {_bad}: 단위당 수량비례비가 판매단가 이상이라 "
                           "아무리 많이 팔아도 손익분기점에 도달할 수 없습니다(손익분기수량 계산 불가). "
                           "'수량에 비례하는 비용' 선택이 맞는지 확인해 주세요.")
            _mean_mgmt = float(e["경영비"].mean())
            _mgmt_bc = (float(e["총수입"].mean()) / _mean_mgmt) if _mean_mgmt > 0 else np.nan
            _mgmt_bc_txt = f"{_mgmt_bc:.2f}" if pd.notna(_mgmt_bc) else "계산 불가"
            _yv_txt = (", ".join(map(str, _yield_var_cost_cols)) if _yield_var_cost_cols else "지정 안 함")
            st.caption("**손익분기수량 = (생산비 − 부산물가액 − 수량비례비) ÷ "
                       "(단가 − 단위당 수량비례비)**입니다. "
                       f"수량에 비례하는 비용: **{_yv_txt}** — 지정하지 않으면 "
                       "**(생산비 − 부산물가액) ÷ 단가**가 되어 '그해 들어간 비용을 "
                       "회수하려면 몇 kg을 수확해야 하는가'를 뜻합니다. "
                       "실제 수량이 이보다 많으면 이익입니다. "
                       "**손익분기가격 = (생산비 − 부산물가액) ÷ 실제 수량**으로, 현재 수량에서 최소 얼마를 받아야 하는지 보여줍니다. "
                       "**단년도 총수입/생산비 = 총수입 ÷ 생산비**(자가노력비·용역비 포함)입니다. "
                       "1을 넘으면 해당 연도의 입력 가격·수량·비용 조건에서 총수입이 생산비보다 큰 상태입니다. "
                       "이 값은 여러 해의 현금흐름을 할인하는 시설투자용 B/C와는 다릅니다. "
                       "통계적 유의성이나 가격 변동 위험도 별도로 확인해야 합니다. "
                       f"참고로 처리 평균값 기준 총수입/경영비는 {_mgmt_bc_txt}입니다.")

            # ---------- ⑩ 반복별 소득·순수익 통계 검정 (선택) ----------
            _rowlv = st.session_state.get("_econ_rowlevel")
            if st.session_state.get("_econ_signature") != dataframe_signature(df):
                _rowlv = None
            if _rowlv is not None and len(_rowlv) > len(e):
                with st.expander("🔬 소득·순수익의 통계 검정 (반복 자료가 있을 때)", expanded=False):
                    st.caption("경제성 결과는 기본적으로 관측값의 점추정치입니다. "
                               "반복 자료가 있으면 소득·순수익 차이 자체를 검정할 수 있습니다.")
                    _mt = st.multiselect("검정할 지표", ["_소득", "_순수익"],
                                         default=["_소득"],
                                         format_func=lambda x: x.lstrip("_"),
                                         key="econ_test_metrics")
                    if st.button("검정 실행", key="econ_test_run"):
                        for _mcol in _mt:
                            _res = econ_metric_test(
                                _rowlv, trt, _mcol,
                                control=(None if ctrl == "(없음)" else ctrl),
                                blk_col=(_blk_for_test if _blk_for_test in _rowlv.columns else None))
                            _label = _mcol.lstrip("_")
                            st.markdown(f"**{_label}**")
                            if _res is None:
                                st.warning(f"{_label}: 반복이 부족해 검정할 수 없습니다.")
                                continue
                            _p = _res.get("anova_p")
                            if _p is not None:
                                st.metric(f"{_label} ANOVA p", f"{_p:.4f}",
                                          "유의함" if _p < .05 else "유의하지 않음")
                            if _res.get("dunnett") is not None:
                                smart_table(_res["dunnett"], width="stretch", hide_index=True)
                                st.caption("95% 구간이 0을 포함하지 않으면 대조구와 유의한 차이가 있습니다.")
                            if _res.get("bootstrap"):
                                _br = pd.DataFrame([
                                    {"처리구": k, "차이": round(v["diff"], 1),
                                     "95% 하한": round(v["low"], 1),
                                     "95% 상한": round(v["high"], 1)}
                                    for k, v in _res["bootstrap"].items()])
                                st.markdown("부트스트랩 신뢰구간 (소표본·비정규 대비)")
                                smart_table(_br, width="stretch", hide_index=True)
                            st.session_state[f"_econ_test_{_label}"] = {
                                "anova_p": _p,
                                "dunnett": (_res["dunnett"].to_dict("records")
                                            if _res.get("dunnett") is not None else None),
                                "bootstrap": _res.get("bootstrap")}
                            log_action(f"경제성 {_label} 통계 검정")
            elif _rowlv is None:
                st.caption("💡 반복 자료(처리구별 여러 행)를 넣으면 소득·순수익 차이도 검정할 수 있습니다.")

            # ---- 대조구 대비 증수 분석 ----
            inc_txt = ""
            if has_ctrl:
                st.markdown("#### 4) 대조구 대비 증수 분석 ⭐")
                st.caption(f"'{ctrl}'을(를) 기준으로 각 처리의 증수 효과와 경제성을 계산했습니다. "
                           "시험연구보고서 경제성 항목에 바로 쓸 수 있는 지표입니다.")
                mi = [trt, yq, "증수량", "증수율(%)", "총수입증가액", "경영비증가액",
                      "소득증가액", "순수익증가액", "소득지수"]
                if qty_effect_ok and "수량효과 기준 순증가수익" in e.columns:
                    mi.append("수량효과 기준 순증가수익")
                mi = [c for c in mi if c in e.columns]
                if "투자효율" in e.columns: mi.append("투자효율")
                inc_df = e[mi].copy()
                smart_table(money_table(inc_df), width="stretch", hide_index=True)
                if not qty_effect_ok:
                    st.info("ℹ️ 처리별 가격 또는 부산물 수입이 달라 "
                            "**수량 효과만을 이용한 순증가수익은 계산하지 않았습니다.** "
                            "경제성 판단은 실제 차이인 **소득증가액**을 사용하세요.")
                st.caption("**소득증가액 = 처리구 소득 − 대조구 소득** ← 경제성 판단의 기본 지표\n\n"
                           "**순수익증가액** = 처리구 순수익 − 대조구 순수익\n\n"
                           "**총수입·경영비증가액** = 각각 대조구와의 차이\n\n"
                           "**수량효과 기준 순증가수익**(보조) = 증수액 − 경영비증가액 — "
                           "처리별 가격·부산물이 모두 같을 때만 표시됩니다.\n\n"
                           "**소득지수** = 대조구를 100으로 본 상대값")
                _ni = pd.to_numeric(e["소득증가액"], errors="coerce")
                _cand = _ni[(e[trt].astype(str) != ctrl) & _ni.notna()]
                if len(_cand):
                    best_i = e.loc[_cand.idxmax()]
                    _rate = best_i.get("증수율(%)")
                    _rate_s = f"{_rate}%" if pd.notna(_rate) else "-"
                    _add = best_i.get("추가투입비")
                    _add_s = f"{_add:,.0f}원" if pd.notna(_add) else "산출 불가"
                    _inc_v = float(best_i["소득증가액"])
                    _prof_v = best_i.get("순수익증가액")
                    inc_txt = ("○ " + f"'{ctrl}' 대비 '{best_i[trt]}'의 경제성\n"
                               + f"  - 수량은 {_rate_s} 증수되었다.\n"
                               + f"  - 경영비는 {_add_s} 증가하였다.\n"
                               + f"  - 실제 소득증가액은 {_inc_v:,.0f}원/10a이었다"
                               + (f"(순수익증가액 {float(_prof_v):,.0f}원/10a)."
                                  if pd.notna(_prof_v) else ".") + "\n"
                               + "  - " + ("입력 조건에서 대조구 대비 소득증가액이 양수로 산출되었다."
                                           if _inc_v > 0
                                           else "대조구 대비 소득 증가가 확인되지 않았다."))
                    if sig_warn:
                        inc_txt += "\n  - " + sig_warn.replace("⚠️ ", "")
                    st.markdown("###### 📋 보고서용 문장 (복사해서 쓰세요)")
                    st.code(inc_txt, language=None)
                else:
                    st.caption("대조구 외 처리의 소득증가액을 계산할 수 없어 요약 문장을 생략합니다.")
                dl_table(money_table(inc_df), f"{ctrl} 대비 증수 및 경제성 분석", "econinc", "증수분석")

            st.markdown("#### 5) 경영비 비목 구성")
            _cost_component_cols = list(cost_cols)
            if "토지임차료(별도입력)" in e.columns and float(land_cash) > 0:
                _cost_component_cols.append("토지임차료(별도입력)")
            if "조성비상각" in e.columns and float(estab_amort) > 0:
                _cost_component_cols.append("조성비상각")
            _cost_component_cols = list(dict.fromkeys(_cost_component_cols))
            comp = e.groupby(trt)[_cost_component_cols].mean()
            _comp_total = comp.sum(axis=1)
            _comp_share = comp.div(_comp_total.where(_comp_total != 0), axis=0) * 100
            smart_table(_comp_share.round(1).reset_index(), width="stretch")
            _comp_means = comp.mean(axis=0)
            if not _comp_means.empty and _comp_means.notna().any() and float(_comp_means.fillna(0).sum()) > 0:
                st.caption(f"평균적으로 '{_comp_means.idxmax()}'가 경영비에서 가장 큰 비중을 차지합니다.")
            else:
                st.caption("모든 경영비가 0이어서 비용 비중을 계산할 수 없습니다.")

            st.markdown("#### 6) 가격 민감도 (단가 변동 시 소득)")
            rates = [-20, -10, 0, 10, 20]
            sens = pd.DataFrame({trt: e[trt].astype(str)})
            for r in rates:
                sens[f"{r:+d}%"] = ((e["주산물가액"]*(1+r/100) + e["부산물가액"]) - e["경영비"]).round(0).values
            smart_table(sens, width="stretch")

            # ⑪ 같은 행에서 처리구·소득·소득률을 가져와야 함(다른 행 값이 섞이면 안 됨)
            _bi = e["소득"].idxmax()
            best = e.loc[_bi, trt]
            best_income = float(e.loc[_bi, "소득"])
            best_rate = float(e.loc[_bi, "소득률(%)"])
            _best_rate_txt = f"{best_rate:.1f}%" if np.isfinite(best_rate) else "계산 불가"
            _ni = e["순수익"].idxmax()
            bestn = e.loc[_ni, trt]
            bestn_profit = float(e.loc[_ni, "순수익"])
            _valid_bc = pd.to_numeric(e["단년도 총수입/생산비"], errors="coerce").dropna()
            _bc_count = int((_valid_bc > 1).sum())
            _sig_note = ("\n  - " + sig_warn.replace("⚠️ ", "")) if sig_warn else ""
            txt = ("○ 처리구별 경제성 분석 결과\n"
                   f"  - 소득이 가장 높은 처리구는 '{best}'로 {best_income:,.0f}원/10a"
                   f"(소득률 {_best_rate_txt})였으며, "
                   f"순수익이 가장 높은 처리구는 '{bestn}'({bestn_profit:,.0f}원/10a)이었다.\n"
                   f"  - 입력한 가격·수량·비용 조건에서 전체 {len(e)}개 처리구 중 "
                   f"{_bc_count}개가 단년도 총수입/생산비 1을 초과하였다. "
                   "최종 보급 판단에는 통계적 유의성·가격 변동·현장 적용성을 함께 검토해야 한다."
                   + _sig_note)
            st.markdown("###### 📋 보고서용 문장")
            st.code(txt, language=None)

            fw, fh = figsize()
            x = e[trt].astype(str).reset_index(drop=True)
            _xp = np.arange(len(x), dtype=float)
            _bw = 0.36

            # ① 소득·순수익: 겹치지 않는 그룹 막대
            fig_income, ax0 = plt.subplots(figsize=(max(9.5, fw*1.6), max(4.4, fh*1.05)))
            _income_plot = e["소득"].astype(float).values / 10000.0
            _profit_plot = e["순수익"].astype(float).values / 10000.0
            b1 = ax0.bar(_xp - _bw/2, _income_plot, width=_bw,
                         label="소득", color="#3D6F9F", edgecolor="#000000", linewidth=.55)
            b2 = ax0.bar(_xp + _bw/2, _profit_plot, width=_bw,
                         label="순수익", color="#A3C4E2", edgecolor="#000000", linewidth=.55)
            ax0.set_xticks(_xp, x.tolist())
            ax0.set_ylabel("만원/10a")
            deco(ax0, "처리구별 소득 · 순수익")
            ax0.tick_params(axis="x", rotation=0 if len(x) <= 6 else 18)
            ax0.legend(loc="upper center", bbox_to_anchor=(0.5, -0.14), ncol=2,
                       frameon=False, borderaxespad=0)
            if len(x) <= 10:
                for bars in (b1, b2):
                    for _b in bars:
                        _v = float(_b.get_height())
                        _va, _off = ("bottom", 5) if _v >= 0 else ("top", -5)
                        ax0.annotate(f"{_v:,.0f}", (_b.get_x()+_b.get_width()/2, _v),
                                     xytext=(0, _off), textcoords="offset points", ha="center",
                                     va=_va, fontsize=8, color="#23394D", fontweight="bold")
                ax0.margins(y=.17)
            fig_income.subplots_adjust(bottom=.23, top=.88, left=.10, right=.98)
            png_income = fig_to_png(fig_income)

            st.markdown("##### 비용 구조와 가격 위험")
            _ec1, _ec2 = st.columns(2)
            with _ec1:
                fig_cost, ax1 = plt.subplots(figsize=(max(5.8, fw), max(4.5, fh)))
                # 경영비 구성: 오래된 Office식 다색 조합 대신, 서로 구분되면서도
                # 한 화면에서 튀지 않는 차분한 cool-tone 팔레트로 통일한다.
                _cost_palette = ["#25577A", "#3F7FA6", "#62A1C1", "#7DBBC9",
                                 "#8FC8BF", "#B8D3DE", "#6F94AA", "#4D728E"]
                _comp_plot = comp.astype(float) / 10000.0
                _comp_plot.plot(kind="bar", stacked=True, ax=ax1,
                          color=[_cost_palette[i % len(_cost_palette)] for i in range(len(comp.columns))],
                          width=.62, edgecolor="white", linewidth=.7)
                # 충분히 큰 구간에는 금액(만원/10a)과 구성비를 함께 표시한다.
                # 너무 작은 구간은 억지로 글자를 넣지 않아 가독성을 지킨다.
                _bottom = np.zeros(len(_comp_plot), dtype=float)
                for _ci, _cname in enumerate(_comp_plot.columns):
                    _vals = _comp_plot[_cname].to_numpy(dtype=float)
                    _shares = _comp_share[_cname].to_numpy(dtype=float)
                    _rgb = _cost_palette[_ci % len(_cost_palette)].lstrip('#')
                    _r, _g, _b = [int(_rgb[k:k+2], 16) for k in (0,2,4)]
                    _lum = 0.2126*_r + 0.7152*_g + 0.0722*_b
                    _tc = "white" if _lum < 145 else "#20384D"
                    for _xi, (_vv, _ss, _bb) in enumerate(zip(_vals, _shares, _bottom)):
                        if np.isfinite(_vv) and np.isfinite(_ss) and _vv > 0 and _ss >= 5.5:
                            ax1.text(_xi, _bb + _vv/2, f"{_vv:,.0f}\n({_ss:.0f}%)",
                                     ha="center", va="center", fontsize=6.9,
                                     color=_tc, fontweight="bold")
                    _bottom += np.nan_to_num(_vals, nan=0.0)
                ax1.set_ylabel("만원/10a")
                ax1.set_xlabel("")
                deco(ax1, "경영비 구성")
                ax1.tick_params(axis="x", rotation=0 if len(comp.index) <= 5 else 18)
                # 범례는 그래프 아래 한 줄/두 줄로 정돈해 막대와 겹치지 않게 한다.
                ax1.legend(fontsize=7.3, ncol=min(3, max(1, len(comp.columns))),
                           loc="upper center", bbox_to_anchor=(0.5, -0.15),
                           frameon=False, borderaxespad=0, columnspacing=1.1,
                           handlelength=1.5, handletextpad=.45)
                # 상단/우측 축선은 숨기고, 왼쪽·아래쪽만 옅게 유지한다.
                ax1.spines["top"].set_visible(False)
                ax1.spines["right"].set_visible(False)
                ax1.spines["left"].set_color("#B9C7D3")
                ax1.spines["bottom"].set_color("#B9C7D3")
                fig_cost.subplots_adjust(bottom=.27, top=.86, left=.14, right=.98)
                png_cost = fig_to_png(fig_cost)

            with _ec2:
                fig_sens, ax2 = plt.subplots(figsize=(max(5.8, fw), max(4.5, fh)))
                _line_palette = ["#274C77", "#4F83B6", "#79A9D1", "#4F8A8B",
                                 "#8D6E63", "#7D6AA5", "#5D8A66", "#B07D4F"]
                _sens_raw = np.array([
                    [float(sens.iloc[i][f"{r:+d}%"]) for r in rates]
                    for i in range(len(e))], dtype=float)
                _sens_plot = _sens_raw / 10000.0
                for i in range(len(e)):
                    _col = _line_palette[i % len(_line_palette)]
                    ax2.plot(rates, _sens_plot[i], marker="o", markersize=5, linewidth=2.0,
                             color=_col, label=str(x.iloc[i]))
                # 각 x지점(-20,-10,0,+10,+20)마다 숫자끼리 실제로 겹치는지 계산해
                # y좌표를 최소 간격만큼 벌린다. 처리별 고정 오프셋 방식보다 훨씬 자연스럽다.
                _all_y = _sens_plot[np.isfinite(_sens_plot)]
                _yr = float(np.nanmax(_all_y) - np.nanmin(_all_y)) if _all_y.size else 1.0
                _min_gap = max(_yr * 0.030, 5.0)
                _base_gap = max(_yr * 0.010, 2.5)
                for _j, _rx in enumerate(rates):
                    _ys = [(float(_sens_plot[_i, _j]), _i) for _i in range(len(e))
                           if np.isfinite(_sens_plot[_i, _j])]
                    _ys.sort(key=lambda z: z[0])
                    _placed = []
                    for _yval, _i in _ys:
                        _ly = _yval + _base_gap
                        if _placed and _ly - _placed[-1][0] < _min_gap:
                            _ly = _placed[-1][0] + _min_gap
                        _placed.append((_ly, _i, _yval))
                    for _ly, _i, _yval in _placed:
                        _col = _line_palette[_i % len(_line_palette)]
                        # 점과 숫자가 멀어진 경우에만 아주 얇은 연결선을 보여준다.
                        if abs(_ly - _yval) > _base_gap * 1.7:
                            ax2.plot([_rx, _rx], [_yval, _ly - _base_gap*0.25],
                                     color=_col, lw=.45, alpha=.55, zorder=2)
                        ax2.text(_rx, _ly, f"{_yval:,.0f}",
                                 ha="center", va="bottom", fontsize=7.2,
                                 color=_col, fontweight="bold",
                                 bbox=dict(facecolor="white", edgecolor="none",
                                           alpha=.78, pad=.35), zorder=5)
                ax2.margins(y=.16)
                ax2.axhline(0, color="#000000", lw=.8, linestyle="--")
                ax2.set_xlabel("단가 변동(%)")
                ax2.set_ylabel("소득(만원/10a)")
                deco(ax2, "가격 민감도")
                ax2.set_xticks(rates)
                ax2.spines["top"].set_visible(False)
                ax2.spines["right"].set_visible(False)
                ax2.spines["left"].set_color("#B9C7D3")
                ax2.spines["bottom"].set_color("#B9C7D3")
                ax2.legend(fontsize=7.5, ncol=min(4, max(1, len(e))),
                           loc="lower center", bbox_to_anchor=(0.5, 1.19),
                           frameon=False, borderaxespad=0, columnspacing=.9)
                fig_sens.subplots_adjust(top=.74, bottom=.16, left=.15, right=.98)
                png_sens = fig_to_png(fig_sens)

            _gd1, _gd2, _gd3 = st.columns(3)
            _gd1.download_button("🖼️ 소득·순수익 그래프", png_income, "econ_income.png", "image/png",
                                 key="dl_econ_income", width="stretch")
            _gd2.download_button("🖼️ 경영비 구성 그래프", png_cost, "econ_cost.png", "image/png",
                                 key="dl_econ_cost", width="stretch")
            _gd3.download_button("🖼️ 가격 민감도 그래프", png_sens, "econ_sensitivity.png", "image/png",
                                 key="dl_econ_sens", width="stretch")

            # ---------- 7) 수량·단가 동시 변동 ----------
            st.markdown("#### 7) 수량과 단가가 함께 변하면 소득은?")
            st.caption("풍흉으로 **수량**이 ±20% 변하고 시세로 **단가**가 ±20% 변하는 경우를 "
                       "조합해 소득을 계산한 표입니다. 두 가지 위험을 동시에 보기 때문에 "
                       "'최악의 경우에도 적자가 아닌지'를 확인할 수 있습니다. "
                       "초록색이면 흑자, 빨간색이면 적자입니다.")
            st.caption("수량과 단가가 **동시에** 변할 때 소득이 어떻게 달라지는지 봅니다. "
                       "앞에서 지정한 수확·선별·포장·운송 등 **수량비례비용은 수량 변동률에 맞춰 함께 변동**시키고, "
                       "나머지 경영비는 고정한 상태로 계산합니다. 붉을수록 위험, 푸를수록 안전입니다.")
            tw1, tw2 = st.columns(2)
            _tsel = tw1.selectbox("기준 처리구", e[trt].astype(str).tolist(), key="tw_trt")
            _tstep = tw2.selectbox("변동 간격", [5, 10], index=1, key="tw_step")
            _row = e[e[trt].astype(str) == _tsel].iloc[0]
            _row_yvc = (float(_row[_yield_var_cost_cols].sum())
                        if _yield_var_cost_cols else 0.0)
            hm = two_way_sensitivity(float(_row[yq]), float(_row[pr]),
                                     float(_row["경영비"]),
                                     float(_row["부산물가액"]) if "부산물가액" in e.columns else 0,
                                     yield_variable_cost=_row_yvc, step=int(_tstep))
            fig_h = plot_sensitivity_heatmap(hm, f"'{_tsel}' — 수량·단가 변동에 따른 소득")
            png_h = fig_to_png(fig_h)
            _neg = int((hm < 0).sum().sum()); _tot = hm.size
            _base = hm.loc["+0%", "+0%"]
            if _neg:
                st.warning(f"⚠️ 총 {_tot}개 경우 중 **{_neg}개에서 소득이 적자**가 됩니다. "
                           "가격·수량 하락 위험에 대비가 필요합니다.")
            else:
                st.success("✅ 수량·단가가 ±20% 범위에서 변해도 소득이 (+)로 유지됩니다. "
                           "가격 변동 위험에 비교적 안정적입니다.")
            _tw_txt = ("○ 수량·단가 동시 변동에 따른 소득 변화 분석\n"
                       f"  - 기준({_tsel}) 소득은 {_base:,.0f}원/10a이었다.\n"
                       "  - 수량과 단가가 각각 20% 하락할 경우 소득은 "
                       f"{hm.loc['-20%', '-20%']:,.0f}원/10a으로 감소하였다.\n"
                       f"  - 분석한 {_tot}개 시나리오 중 적자가 발생하는 경우는 {_neg}개였다.")
            st.markdown("###### 📋 보고서용 문장")
            st.code(_tw_txt, language=None)
            st.download_button("🖼️ 히트맵 다운로드", png_h, "sensitivity_heatmap.png",
                               "image/png", key="dl_hm")
            _hm_out = hm.round(0).astype(int).reset_index().rename(columns={"index": "수량변동"})
            _hm_show = _hm_out.copy()
            for _c in _hm_show.columns:
                if pd.api.types.is_numeric_dtype(_hm_show[_c]):
                    _hm_show[_c] = _hm_show[_c].map(lambda v: f"{v:,.0f}")
            log_action("소득분석(소득조사 방식) 실행")
            _econ_blocks = [
                {"text": txt},
                {"caption": "처리구별 소득·순수익",
                 "table": money_table(e[m1]), "image": png_income},
                {"caption": "처리구별 생산비 구성",
                 "table": money_table(e[m2])},
                {"caption": "경영비 구성 금액(만원/10a)",
                 "table": _comp_plot.round(2).reset_index(), "image": png_cost},
                {"caption": "경영비 구성 비율(%)",
                 "table": _comp_share.round(1).reset_index()},
                {"caption": "가격 민감도",
                 "table": sens, "image": png_sens},
                {"text": _tw_txt},
                {"caption": f"{_tsel} 수량·단가 변동별 소득",
                 "table": _hm_show, "image": png_h},
            ]
            report_capture("cap_econ", "경제성 분석", None, blocks=_econ_blocks)
            _ctx_cols = list(dict.fromkeys(
                m1 + m3 + [yq, pr, "주산물가액", "부산물가액", "소득증가액",
                           "순수익증가액", "총수입증가액", "경영비증가액"]))
            _ctx_cols = [c for c in _ctx_cols if c in e.columns]
            _used_price_db = st.session_state.get("price_db")
            _econ_ctx = build_econ_context(
                base_area="10a", control=(None if ctrl == "(없음)" else ctrl),
                treatments=e[_ctx_cols],
                prices={"농촌임료금(원/시간)": wage, "자본이자율(%)": rate,
                        "고정자산 부분현재가(원/10a)": fixed_asset,
                        "고정자산 작목부담률(%)": fixed_asset_use_rate,
                        "토지이용형태": land_type,
                        "자가토지 용역비(원/10a)": land_opp,
                        "토지 임차료 별도입력(원/10a)": land_cash,
                        "재배기간(개월)": months,
                        "사용 기준단가 DB": _used_price_db},
                cost_cols=list(cost_cols), excluded_cols=list(_blocked),
                sensitivity={"가격민감도": sens, "수량×단가민감도": hm,
                             "기준처리": _tsel},
                yield_test=yield_test,
                income_test=st.session_state.get("_econ_test_소득"),
                profit_test=st.session_state.get("_econ_test_순수익"),
                cautions=([sig_warn] if sig_warn else [])
                         + ([] if qty_effect_ok else
                            ["처리별 가격·부산물이 달라 수량효과 기준 순증가수익은 계산하지 않음"]))
            ai_interpret_advanced("econ", "경제성(소득) 분석", e[m1],
                                  "소득증가액이 경제성 판단의 기본 지표입니다.",
                                  context=_econ_ctx, capture_slot="cap_econ")

            # 경제성분석의 한글/Excel은 중간중간 흩어 놓지 않고 분석의 맨 아래에서
            # 현재 화면의 모든 표·그래프를 한 번에 내려받게 한다.
            st.markdown("---")
            st.markdown("### 📥 경제성 분석 전체 결과 다운로드")
            st.caption("처리구별 소득·순수익, 생산비 구성, 경영비 금액·비율, "
                       "가격 민감도, 수량×단가 민감도까지 한 파일에 모두 담습니다.")
            _ed1, _ed2 = st.columns(2)
            try:
                _econ_hwp = build_report_hwpx(
                    [{"heading": "경제성 분석", "blocks": _econ_blocks}],
                    doc_title="경제성 분석 결과")
                _ed1.download_button(
                    "📘 한글 전체 보고서(hwpx)", _econ_hwp,
                    "경제성분석_전체결과.hwpx",
                    key="dl_econ_all_hwp", width="stretch")
            except Exception as _ex:
                _ed1.caption(f"한글 파일 생성 실패 ({type(_ex).__name__})")
            try:
                _econ_xls = make_xlsx_multi(_econ_blocks, doc_title="경제성 분석 결과")
                _ed2.download_button(
                    "📈 Excel 전체 결과(xlsx)", _econ_xls,
                    "경제성분석_전체결과.xlsx",
                    key="dl_econ_all_xlsx", width="stretch",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            except Exception as _ex:
                _ed2.caption(f"Excel 파일 생성 실패 ({type(_ex).__name__})")

        report_button("cap_econ")

    # ---------------- 부분예산 · MRR ----------------
    elif emode.startswith("📘"):
        with st.expander("ℹ️ 부분예산분석이란? (CIMMYT 방식)"):
            st.markdown("""
새 기술(신품종·신자재 등)을 **농가에 권장할 만한가**를 판단하는 국제 표준 방법입니다.
1. **조정수량** — 시험포장 수량은 농가 조건보다 높게 나오므로 보통 **10% 낮춰** 보정합니다.
2. **순편익** = 조정 조수입 − 가변비용(처리마다 달라지는 비용만)
3. **지배분석** — 비용이 더 드는데 순편익이 낮은 처리는 탈락(D 표시)
4. **한계수익률(MRR)** = 순편익 증가분 ÷ 비용 증가분 × 100
5. **판정** — MRR을 사용자가 정한 최소수용기준과 비교합니다. 100%는 보수적으로 쓸 수 있는 기본값일 뿐 절대 기준은 아닙니다.
""")
        allc = df.columns.tolist()
        c1, c2, c3 = st.columns(3)
        trt = c1.selectbox("처리구 열", allc, index=guess_idx(allc, ["처리", "품종", "시험구"]), key="pb_t")
        yq = c2.selectbox("수량 열", num_cols, index=guess_idx(num_cols, ["수량", "생산량"]), key="pb_y")
        pr = c3.selectbox("단가 열", num_cols, index=guess_idx(num_cols, ["단가", "가격"], 1), key="pb_p")
        _pb_candidates = [c for c in num_cols if c not in (yq, pr) and not is_excluded_cost(c)]
        var_cols = st.multiselect("가변비용 항목 (처리마다 달라지는 비용만)",
                                  _pb_candidates, key="pb_v")
        _pb_by_opts = ["(없음)"] + [c for c in num_cols if c not in (yq, pr) + tuple(var_cols)]
        pb_by = st.selectbox("부산물 편익 열 (선택)", _pb_by_opts, key="pb_by")
        _pb_levels = df[trt].dropna().astype(str).unique().tolist()
        pb_control = st.selectbox("기준 처리(관행·대조구)", _pb_levels, key="pb_control")
        pb_area = st.number_input("부분예산 자료 기준 면적(a)", 0.1, 10000.0, 10.0, 0.1, key="pb_area")
        c4, c5 = st.columns(2)
        adj = c4.slider("수량 조정률(%)", 0, 30, 10, 5, key="pb_adj",
                       help="시험포장→농가 조건 보정. 보통 10%")
        minmrr = c5.number_input("최소수용 MRR 기준(%)", 0, 500, 100, 10, key="pb_minmrr",
                                 help="기본값 100%는 보수적인 출발점입니다. 기술의 위험도·추가자본·농가 여건에 따라 조정하세요.")
        _pb_unit_confirm = st.checkbox(
            "수량=kg/입력 기준면적, 단가=원/kg, 가변비용·부산물=원/입력 기준면적임을 확인했습니다.",
            key="pb_unit_confirm")
        _pb_errors = []
        if trt in (yq, pr):
            _pb_errors.append("처리구 열을 수량 또는 단가 열로 사용할 수 없습니다.")
        if yq == pr:
            _pb_errors.append("수량 열과 단가 열은 서로 달라야 합니다.")
        if not var_cols:
            _pb_errors.append("가변비용 항목을 하나 이상 선택해 주세요.")
        if any(is_excluded_cost(c) for c in var_cols):
            _pb_errors.append("합계·소득·생산비 등 계산 결과 열은 가변비용으로 사용할 수 없습니다.")
        if not _pb_unit_confirm:
            _pb_errors.append("입력 열의 단위를 확인해야 계산할 수 있습니다.")
        for _msg in _pb_errors:
            st.error("❌ " + _msg)

        if keep_running("partbudget", "부분예산 분석 실행", disabled=bool(_pb_errors)):
            st.session_state.pop("cap_pb", None)
            # ⑤ 행(반복)별 계산 → 처리 평균. 순수 함수로 분리해 자동 검산과 UI가 같은 로직을 사용한다.
            try:
                _row, b, _pb_excluded = calculate_partial_budget(
                    df, trt, yq, pr, var_cols,
                    byproduct_col=(None if pb_by == "(없음)" else pb_by),
                    source_area_a=pb_area, adjustment_percent=adj)
            except ValueError as _ex:
                st.error(f"부분예산 계산을 중단했습니다: {_ex}")
                st.stop()
            _pb_validation = validate_partial_budget_results(_row)
            if _pb_validation:
                for _err in _pb_validation:
                    st.error("❌ " + _err)
                st.error("부분예산 역산 검증에 실패하여 결과를 표시하지 않습니다.")
                st.stop()
            st.success("✅ 부분예산의 조정수량·총편익·순편익 역산 검증을 통과했습니다.")
            _nrep = _row.groupby(trt).size()
            if (_nrep > 1).any():
                st.info(f"🔁 반복 자료가 확인되어(최대 {int(_nrep.max())}반복) "
                        "**각 반복별 편익·비용을 계산한 뒤 처리 평균**으로 집계했습니다.")

            # 반올림 전에 지배분석과 MRR을 계산해야 작은 차이가 사라지지 않는다.
            b_raw = perform_dominance_analysis(
                b, trt, cost_col="가변비용", benefit_col="순편익", control=pb_control)
            und_raw = calculate_mrr_table(
                b_raw, trt, cost_col="가변비용", benefit_col="순편익",
                minimum_mrr=minmrr, control=pb_control)
            mrr_num = [None if pd.isna(v) else float(v) for v in und_raw["MRR(%)"]]

            b = b_raw.copy()
            und = und_raw.copy()
            for _c in ["조정 전 수량", "조정수량"]:
                if _c in b.columns:
                    b[_c] = b[_c].round(1)
            for _c in ["주산물 편익", "부산물 편익", "총편익", "가변비용", "순편익"]:
                if _c in b.columns:
                    b[_c] = b[_c].map(lambda v: round_half_up(v) if pd.notna(v) else np.nan)
            for _c in ["가변비용", "순편익", "비용 증가액", "순편익 증가액"]:
                if _c in und.columns:
                    und[_c] = und[_c].map(lambda v: round_half_up(v) if pd.notna(v) else np.nan)
            if "MRR(%)" in und.columns:
                und["MRR(%)"] = pd.to_numeric(und["MRR(%)"], errors="coerce").round(1)

            st.markdown("#### 1) 부분예산표 (지배분석 포함)")
            smart_table(money_table(b[[trt, "기준 처리", "조정 전 수량", "조정수량",
                                       "주산물 편익", "부산물 편익", "총편익",
                                       "가변비용", "순편익", "지배"]],
                                     dec_overrides={"조정 전 수량": 1, "조정수량": 1}),
                        width="stretch", hide_index=True)
            st.caption("'D'는 비용이 더 들면서 순편익은 낮아 탈락한 처리입니다.")

            st.markdown("#### 2) 한계수익률(MRR) 분석")
            _mcols = [trt, "가변비용", "순편익", "비용 증가액", "순편익 증가액",
                      "MRR(%)", "권장 여부 및 근거"]
            _mcols = [c for c in _mcols if c in und.columns]
            smart_table(money_table(und[_mcols]), width="stretch", hide_index=True)
            st.caption("**MRR(%) = 순편익 증가액 ÷ 비용 증가액 × 100** — "
                       "비용을 1원 더 썼을 때 순편익이 몇 % 늘어나는지를 뜻합니다.\n\n"
                       f"현재 사용자가 설정한 최소수용 기준: **{minmrr}%**. 이 값은 절대 기준이 아니며, 기준을 넘더라도 "
                       "① 반복수·자료 신뢰도 ② 가격·수량 변동 시에도 유지되는지 "
                       f"③ 조정수량 가정({adj}% 감액)의 타당성을 함께 확인하세요.")

            ok_idx = [i for i in range(1, len(und_raw))
                      if mrr_num[i] is not None and mrr_num[i] >= minmrr
                      and float(und_raw.loc[i, "순편익 증가액"]) > 0
                      and float(und_raw.loc[i, "순편익"]) > 0]
            _ctrl_rows = b_raw[b_raw[trt].astype(str) == str(pb_control)]
            _ctrl_dom = (str(_ctrl_rows.iloc[0]["지배"]) if len(_ctrl_rows) else "기준 처리 없음")
            if _ctrl_dom:
                _start = und.iloc[0][trt] if len(und) else "없음"
                txt = (f"선택한 기준 처리 '{pb_control}'은 지배분석에서 {_ctrl_dom}로 분류되었습니다. "
                       f"따라서 관행 유지를 자동 권장하지 않으며, 효율경계 시작 처리 '{_start}'부터 "
                       "비용·순편익과 현장 적용성을 다시 검토해야 합니다.")
            elif ok_idx:
                last = ok_idx[-1]
                pick = und_raw.loc[last, trt]
                txt = (f"지배분석 후 남은 처리 중 MRR이 기준({minmrr}%)을 넘는 최상위 처리는 '{pick}'입니다. "
                       f"추가 투입 1원당 약 {mrr_num[last]/100:.1f}원의 순편익 증가가 기대됩니다. "
                       "다만 반복수와 가격·수량 민감도도 함께 확인해야 합니다.")
            else:
                txt = (f"MRR이 기준({minmrr}%)을 넘는 추가 처리 단계가 없습니다. "
                       f"선택한 기준 처리 '{pb_control}'이 비지배 처리이므로 현재 자료에서는 기준 처리 유지가 "
                       "상대적으로 합리적이지만, 통계적 불확실성과 민감도 결과를 함께 확인해야 합니다.")
            st.info("💡 " + txt)

            fig, ax = plt.subplots(figsize=figsize())
            ax.plot(b_raw["가변비용"], b_raw["순편익"], "o--", color="#9AAABB", label="전체", linewidth=1.4, markersize=5)
            ax.plot(und_raw["가변비용"], und_raw["순편익"], "o-", color="#3D6F9F", lw=2.2, markersize=6, label="비지배(효율경계)")
            for _, r in b_raw.iterrows():
                ax.annotate(str(r[trt]), (r["가변비용"], r["순편익"]), fontsize=8,
                            xytext=(3, 4), textcoords="offset points")
            ax.set_xlabel("가변비용(원/10a)"); ax.set_ylabel("순편익(원/10a)")
            ax.legend(fontsize=8); deco(ax, "부분예산 효율경계")
            plt.tight_layout(); png = fig_to_png(fig)
            st.download_button("🖼️ 그래프 다운로드", png, "mrr.png", "image/png")
            out = und[[c for c in [trt, "가변비용", "순편익", "비용 증가액",
                                   "순편익 증가액", "MRR(%)", "권장 여부 및 근거"]
                       if c in und.columns]]
            out_show = money_table(out)
            dl_table(out_show, "부분예산 한계수익률 분석", "mrr16", "mrr")
            log_action("부분예산·MRR 분석 실행")
            ai_interpret_button("pb", "부분예산·한계수익률(MRR)", out, f"검증된 계산값입니다. 수치를 다시 계산하지 마세요. 사용자가 설정한 최소 MRR 기준은 {minmrr}%이고 기준 처리구는 {pb_control}입니다.", capture_slot="cap_pb")
            report_capture("cap_pb", "신기술 경제성(부분예산·MRR)", txt, out_show, png)
        report_button("cap_pb")

    # ---------------- 시설·장기투자 ----------------
    else:
        st.markdown("### 📙 시설·장기투자 경제성")
        st.caption("비가림시설·하우스·건조기·선별기·스마트팜 장비처럼 여러 해 사용하는 투자는 "
                   "단년도 소득분석이 아니라 미래 편익과 비용을 현재가치로 환산해 판단합니다.")
        with st.expander("ℹ️ 지표 읽는 법", expanded=True):
            st.markdown("""
- **NPV(순현재가치)** = 할인된 편익 − 할인된 비용. **0보다 크면** 입력한 할인율 기준 경제성이 있습니다.
- **할인 B/C** = 편익 현재가 ÷ 비용 현재가. **1보다 크면** 할인된 편익이 할인된 비용보다 큽니다.
- **IRR(내부수익률)** = NPV를 0으로 만드는 수익률. 기준 할인율보다 높을수록 유리합니다.
- **회수기간** = 누적 순현금흐름으로 최초투자비를 회수하는 데 걸리는 기간입니다.

※ 이 화면의 **할인 B/C**가 시설투자에서 사용하는 정식 B/C입니다. 📗 소득분석의 `단년도 총수입/생산비`와 구분합니다.
""")
        _iv1, _iv2, _iv3 = st.columns(3)
        _initial = _iv1.number_input("최초 투자비 (원)", 0, 5000000000, 50000000, 1000000,
                                    key="inv_initial")
        _life = _iv2.number_input("내용연수 (년)", 1, 50, 10, 1, key="inv_life")
        _disc = _iv3.number_input("할인율 (%)", 0.0, 30.0, 5.0, 0.1, key="inv_disc")
        _iv4, _iv5, _iv6 = st.columns(3)
        _annual_b = _iv4.number_input("연간 추가 편익/수입 (원/년)", 0, 5000000000, 10000000, 100000,
                                     key="inv_benefit",
                                     help="시설 도입으로 매년 추가되는 판매수입·비용절감 등 금전 편익")
        _annual_c = _iv5.number_input("연간 추가 운영비 (원/년)", 0, 5000000000, 2000000, 100000,
                                     key="inv_cost",
                                     help="유지보수·전기·소모품·추가노동비 등 매년 추가로 발생하는 비용")
        _salvage = _iv6.number_input("내용연수 말 잔존가치 (원)", 0, 5000000000, 0, 100000,
                                    key="inv_salvage")
        with st.expander("📈 연도별 증가율을 반영하려면"):
            _ig1, _ig2 = st.columns(2)
            _bg = _ig1.number_input("연간 편익 증가율 (%)", -50.0, 50.0, 0.0, 0.5, key="inv_bg")
            _cg = _ig2.number_input("연간 운영비 증가율 (%)", -50.0, 50.0, 0.0, 0.5, key="inv_cg")
            st.caption("가격상승·생산성 변화·운영비 상승을 가정할 때만 입력하세요. 모르면 0%로 두는 것이 안전합니다.")

        if keep_running("investment", "장기투자 경제성 분석 실행"):
            try:
                _inv_table, _inv = calculate_investment_analysis(
                    _initial, _life, _disc, _annual_b, _annual_c,
                    salvage_value=_salvage,
                    annual_benefit_growth_percent=_bg,
                    annual_cost_growth_percent=_cg)
            except ValueError as _ex:
                st.error(f"장기투자 계산을 중단했습니다: {_ex}")
                st.stop()

            _npv = float(_inv["NPV"])
            _bcr = _inv["할인 B/C"]
            _irr = _inv["IRR(%)"]
            _dpb = _inv["할인 회수기간(년)"]
            _mc1, _mc2, _mc3, _mc4 = st.columns(4)
            _mc1.metric("NPV", f"{_npv:,.0f}원")
            _mc2.metric("할인 B/C", f"{_bcr:.2f}" if pd.notna(_bcr) else "계산 불가")
            _mc3.metric("IRR", f"{_irr:.1f}%" if pd.notna(_irr) else "계산 불가")
            _mc4.metric("할인 회수기간", f"{_dpb:.1f}년" if pd.notna(_dpb) else "내용연수 내 미회수")
            if _npv > 0 and (pd.isna(_bcr) or _bcr > 1):
                st.success(f"✅ 할인율 {_disc:g}% 기준 NPV가 양수여서 입력 조건에서는 경제성이 있습니다.")
            elif _npv < 0:
                st.warning(f"⚠️ 할인율 {_disc:g}% 기준 NPV가 음수여서 입력 조건에서는 경제성이 없습니다.")
            else:
                st.info("NPV가 0에 가까워 경제성 판단이 경계에 있습니다.")

            _show_inv = _inv_table.copy()
            for _c in ["편익", "비용", "순현금흐름", "편익현재가", "비용현재가", "순현재가"]:
                _show_inv[_c] = _show_inv[_c].map(lambda v: round_half_up(v))
            _show_inv["할인계수"] = _show_inv["할인계수"].round(4)
            st.markdown("#### 연도별 할인 현금흐름")
            smart_table(money_table(_show_inv), width="stretch", hide_index=True)

            st.markdown("#### 민감도 — 편익·운영비가 달라지면 NPV는?")
            _inv_sens = investment_sensitivity_table(
                _initial, _life, _disc, _annual_b, _annual_c, _salvage,
                change_rates=(-20, -10, 0, 10, 20))
            _inv_sens_show = _inv_sens.copy()
            for _c in _inv_sens_show.columns[1:]:
                _inv_sens_show[_c] = _inv_sens_show[_c].map(lambda v: round_half_up(v))
            smart_table(money_table(_inv_sens_show), width="stretch", hide_index=True)

            _inv_txt = ("○ 시설·장기투자 경제성 분석 결과\n"
                        f"  - 최초투자비 {_initial:,.0f}원, 내용연수 {_life}년, 할인율 {_disc:g}%를 적용하였다.\n"
                        f"  - NPV는 {_npv:,.0f}원, 할인 B/C는 "
                        f"{(_bcr if pd.notna(_bcr) else float('nan')):.2f}로 산출되었다.\n"
                        + (f"  - IRR은 {_irr:.1f}%로 산출되었다.\n" if pd.notna(_irr)
                           else "  - 현금흐름 구조상 IRR은 계산되지 않았다.\n")
                        + (f"  - 할인 회수기간은 {_dpb:.1f}년이었다.\n" if pd.notna(_dpb)
                           else "  - 내용연수 안에는 할인 기준 투자비를 회수하지 못하였다.\n")
                        + ("  - 입력한 가정에서는 경제성이 있는 것으로 판단된다."
                           if _npv > 0 else "  - 입력한 가정에서는 경제성이 없는 것으로 판단된다."))
            st.markdown("###### 📋 보고서용 문장")
            st.code(_inv_txt, language=None)
            _inv_blocks = [
                {"text": _inv_txt},
                {"caption": "장기투자 연도별 할인 현금흐름", "table": _show_inv},
                {"caption": "장기투자 NPV 민감도", "table": _inv_sens_show},
            ]
            report_capture("cap_inv", "시설·장기투자 경제성", None, blocks=_inv_blocks)
            log_action("시설·장기투자 경제성 분석 실행")
            st.markdown("### 📥 장기투자 결과 다운로드")
            _id1, _id2 = st.columns(2)
            try:
                _h = build_report_hwpx([{"heading":"시설·장기투자 경제성", "blocks":_inv_blocks}],
                                       doc_title="시설·장기투자 경제성 분석 결과")
                _id1.download_button("📘 한글 보고서(hwpx)", _h, "장기투자_경제성분석.hwpx",
                                     key="dl_inv_hwp", width="stretch")
            except Exception as _ex:
                _id1.caption(f"한글 파일 생성 실패 ({type(_ex).__name__})")
            try:
                _x = make_xlsx_multi(_inv_blocks, doc_title="시설·장기투자 경제성 분석 결과")
                _id2.download_button("📈 Excel 결과(xlsx)", _x, "장기투자_경제성분석.xlsx",
                                     key="dl_inv_xlsx", width="stretch",
                                     mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            except Exception as _ex:
                _id2.caption(f"Excel 파일 생성 실패 ({type(_ex).__name__})")
        report_button("cap_inv")

# ================================================================ 설문 분석
elif menu == "📋 설문조사 분석":
    st.title("📋 설문조사 분석")
    with st.expander("ℹ️ 설문 분석이 뭔가요?"): st.markdown(EXPLAIN["survey"])
    stype = st.radio("분석 유형", ["🤖 자동 인식 (문항 유형 자동 판별)", "📊 리커트 척도(점수형)",
                                 "🔘 객관식(단일선택)", "☑️ 다중응답",
                                 "✍️ 주관식(서술형)", "🔀 교차분석(집단 비교)"],
                     key="svy_type")

    # ---------- 자동 인식 ----------
    if stype.startswith("🤖"):
        st.caption("각 열의 값을 보고 문항 유형을 자동으로 판별한 뒤, 유형에 맞는 그래프와 분석을 한 번에 만듭니다.")
        det = detect_question_types(df)
        st.markdown("#### 📋 문항 유형 자동 판별 결과")
        smart_table(det, width="stretch", hide_index=True)
        likert = det[det["추정 유형"] == "리커트 척도"]["열 이름"].tolist()
        single = det[det["추정 유형"] == "객관식(단일선택)"]["열 이름"].tolist()
        multi = det[det["추정 유형"] == "다중응답"]["열 이름"].tolist()
        openq = det[det["추정 유형"] == "주관식(서술형)"]["열 이름"].tolist()

        # 판별이 틀렸을 때 사용자가 직접 유형을 옮길 수 있게 한다.
        # 데이터가 바뀌면(새 파일 업로드 등) 자동판별 값으로 초기화한다.
        _dsig = dataframe_signature(df)
        if st.session_state.get("_svy_auto_sig") != _dsig:
            st.session_state["_svy_auto_sig"] = _dsig
            st.session_state["_svy_likert_ov"] = likert
            st.session_state["_svy_single_ov"] = single
            st.session_state["_svy_multi_ov"] = multi
            st.session_state["_svy_openq_ov"] = openq
        with st.expander("🛠️ 판별이 잘못됐으면 여기서 유형을 고쳐 주세요", expanded=False):
            st.caption("예: 이름에 '의향'이 들어간 문항이 실제로는 5점 척도인데 객관식으로 분류된 경우 "
                       "여기서 '📊 리커트 척도'로 옮기면 됩니다. 응답자 ID·빈 열은 목록에서 빠집니다.")
            _id_like = det[det["추정 유형"].isin(["응답자 ID", "빈 열"])]["열 이름"].tolist()
            _classifiable = [c for c in df.columns if c not in _id_like]
            ov1, ov2 = st.columns(2)
            likert = ov1.multiselect("📊 리커트 척도", _classifiable, key="_svy_likert_ov")
            single = ov2.multiselect("🔘 객관식(단일선택)", _classifiable, key="_svy_single_ov")
            ov3, ov4 = st.columns(2)
            multi = ov3.multiselect("☑️ 다중응답", _classifiable, key="_svy_multi_ov")
            openq = ov4.multiselect("✍️ 주관식(서술형)", _classifiable, key="_svy_openq_ov")
            _dup = [c for c, n in Counter(likert + single + multi + openq).items() if n > 1]
            if _dup:
                st.warning("⚠️ 같은 열이 두 유형 이상에 겹쳐 선택되었습니다: " + ", ".join(_dup)
                           + " — 그래프가 중복으로 나올 수 있으니 한 곳에서만 고르세요.")

        # 객관식이라고 전부 '응답자 특성'인 것은 아니다. 성별·경력·지역처럼 응답자를
        # 설명하는 항목만 응답자 특성으로 묶고, '가장 효과적인 기술'처럼 내용을 묻는
        # 객관식 문항은 '문항별 응답 결과'로 따로 정리한다.
        _DEMO_HINT = ["성별", "연령", "나이", "경력", "학력", "소득", "지역", "시군",
                      "시도", "소속", "직업", "직급", "구분", "유형", "규모", "면적",
                      "거주", "응답자", "농가", "재배형태", "작목"]
        if st.session_state.get("_svy_demo_sig") != _dsig:
            st.session_state["_svy_demo_sig"] = _dsig
            st.session_state["_svy_demo_ov"] = [c for c in single
                                                if any(k in str(c) for k in _DEMO_HINT)]
        # 위 '유형 고치기'에서 객관식 목록이 바뀌면 없는 열이 남지 않도록 정리한다.
        st.session_state["_svy_demo_ov"] = [
            c for c in st.session_state.get("_svy_demo_ov", []) if c in single]
        demo_q = st.multiselect(
            "👥 객관식 중 '응답자 특성'인 항목", single, key="_svy_demo_ov",
            help="성별·경력·지역처럼 응답자를 설명하는 항목만 고르세요. "
                 "여기서 뺀 객관식은 '문항별 응답 결과'로 따로 정리됩니다.")
        demo_q = [c for c in single if c in demo_q]
        single_q = [c for c in single if c not in demo_q]

        c1, c2, c3, c4, c5 = st.columns(5)
        c1.metric("📊 리커트", len(likert)); c2.metric("👥 응답자 특성", len(demo_q))
        c3.metric("🔘 객관식 문항", len(single_q))
        c4.metric("☑️ 다중응답", len(multi)); c5.metric("✍️ 주관식", len(openq))
        chart_style = st.radio("응답자 특성 그래프", ["도넛", "원형", "막대"],
                               horizontal=True, key="svy_chart")
        st.caption("판별이 잘못되었으면 위 탭에서 유형을 직접 골라 분석하세요.")

        if keep_running("autoan", "🚀 자동 분석 실행"):
            rep_blocks = []

            # 1) 응답자 특성 · 객관식 문항 (원형/도넛) — 같은 그리기 방식을 나눠서 쓴다
            def _draw_choice(cols_list, heading, cap_prefix):
                if not cols_list:
                    return
                st.markdown(heading)
                cols_ = st.columns(min(3, len(cols_list)))
                for i, c in enumerate(cols_list):
                    vc = df[c].value_counts()
                    with cols_[i % len(cols_)]:
                        if chart_style == "막대":
                            fig, ax = plt.subplots(figsize=(4.2, 3.4))
                            # 설문 화면은 문항별 색 구분(_survey_palette)이 이미 있어
                            # 그대로 둔다. 값 표시·격자 정리는 아래 deco 가 맡는다.
                            _bs = ax.bar(vc.index.astype(str), vc.values,
                                         color=_survey_palette(len(vc)), width=.62,
                                         edgecolor="white", linewidth=1.0)
                            for _b, _v in zip(_bs, vc.values):
                                ax.text(_b.get_x()+_b.get_width()/2, _v,
                                        f"{_v}명\n({_v/vc.sum()*100:.1f}%)",
                                        ha="center", va="bottom", fontsize=8, color="#333")
                            ax.set_ylim(0, max(vc.values)*1.3)
                            ax.set_ylabel("응답자 수(명)", fontsize=9)
                            for _s in ("top", "right"):
                                ax.spines[_s].set_visible(False)
                            ax.spines["left"].set_color("#bbb")
                            ax.spines["bottom"].set_color("#bbb")
                            ax.tick_params(colors="#555", labelsize=8)
                            ax.set_title(f"{c}  (n={int(vc.sum())})", fontsize=11,
                                         fontweight="bold", color="#333")
                            plt.xticks(rotation=20, fontsize=8)
                            plt.tight_layout()
                        else:
                            fig = pie_chart(vc, c, donut=(chart_style == "도넛"))
                        show_plot(fig); plt.close(fig)
                        _t = pd.DataFrame({c: vc.index.astype(str), "빈도(명)": vc.values,
                                           "비율(%)": (vc.values/vc.sum()*100).round(1)})
                        smart_table(_t, width="stretch", hide_index=True)
                        rep_blocks.append({"caption": f"{cap_prefix} - {c}",
                                           "table": _t,
                                           "image": fig_to_png(fig, show=False)})

            _draw_choice(demo_q, "## 👥 응답자 특성", "응답자 특성")
            _draw_choice(single_q, "## 🔘 문항별 응답 결과", "문항 응답")

            # 2) 리커트 (요약 + 다이버징 + 평균 막대)
            if likert:
                st.markdown("## 📊 리커트 문항 분석")
                sdat = df[likert]  # 문항별 기술통계는 문항 각자의 결측만 제외 (교집합 아님)
                cc = sdat.dropna()  # 크론바흐 α는 문항 간 상관을 봐야 하므로 전 문항 응답한 사람만
                smax = int(sdat.max().max())
                _pos_cut, _neg_cut = likert_cutoffs(smax)
                _pos_col = f"긍정({_pos_cut}점↑)%"
                _neg_col = f"부정({_neg_cut}점↓)%"
                summ = pd.DataFrame({
                    "문항": likert,
                    "응답자(명)": [int(sdat[q].notna().sum()) for q in likert],
                    "평균": [round(sdat[q].dropna().mean(), 2) for q in likert],
                    "표준편차": [round(sdat[q].dropna().std(), 2) for q in likert],
                    "긍정 응답(명)": [int((sdat[q].dropna() >= _pos_cut).sum()) for q in likert],
                    _pos_col: [round((sdat[q].dropna() >= _pos_cut).mean()*100, 1) for q in likert],
                    "부정 응답(명)": [int((sdat[q].dropna() <= _neg_cut).sum()) for q in likert],
                    _neg_col: [round((sdat[q].dropna() <= _neg_cut).mean()*100, 1) for q in likert]})
                a_ = cronbach_alpha(cc)
                lvl = (("매우 높음" if a_ >= .9 else "높음" if a_ >= .8 else "양호" if a_ >= .7 else "낮음")
                       if np.isfinite(a_) else "산출 불가")
                m1, m2, m3 = st.columns(3)
                m1.metric("문항 수", len(likert))
                m2.metric("평균 만족도", f"{pd.concat([sdat[q].dropna() for q in likert]).mean():.2f}")
                m3.metric("크론바흐 α", f"{a_:.3f}" if np.isfinite(a_) else "-", lvl)
                if len(cc) < len(df):
                    st.caption(f"※ 크론바흐 α는 {len(likert)}개 문항에 모두 응답한 {len(cc)}명 기준입니다. "
                               f"(문항별 응답자 수는 위 표의 '응답자(명)' 참고)")
                counts = {q: [int((sdat[q].dropna() == v).sum()) for v in range(1, smax+1)] for q in likert}
                st.markdown("##### 응답 분포 (다이버징 차트)")
                figd = likert_diverging(counts, [f"{v}점" for v in range(1, smax+1)], "문항별 응답 분포")
                dv_png = fig_to_png(figd)
                col_a, col_b = st.columns(2)
                _sv_h = max(3.8, len(likert) * 0.48 + 1.4)
                with col_a:
                    fig, ax = plt.subplots(figsize=(max(5.2, figsize()[0]), _sv_h))
                    order = summ.sort_values("평균")
                    _vals = order["평균"].astype(float).tolist()
                    bars = ax.barh(order["문항"], order["평균"],
                                   color=bar_colors(values=_vals), height=.58, edgecolor="none")
                    ax.set_xlim(0, smax); ax.set_xlabel("평균 점수"); deco(ax, "문항별 평균", ylabel_top=False)
                    ax.grid(axis="y", visible=False)
                    for bar, v in zip(bars, order["평균"]):
                        ax.text(min(float(v)+0.05, smax-0.02), bar.get_y()+bar.get_height()/2,
                                f"{v:.2f}", va="center", fontsize=8.5, color="#33495C")
                    show_plot(fig); plt.close(fig)
                with col_b:
                    fig, ax = plt.subplots(figsize=(max(5.2, figsize()[0]), _sv_h))
                    _bp = ax.barh(summ["문항"], summ[_pos_col], color="#4576AB",
                                      label="긍정", height=.58)
                    _bn = ax.barh(summ["문항"], -summ[_neg_col], color="#C96767",
                                      label="부정", height=.58)
                    _lim = max(25.0, float(max(summ[_pos_col].max(), summ[_neg_col].max()))) * 1.30
                    ax.set_xlim(-_lim, _lim)
                    ax.axvline(0, color="#000000", lw=.9); ax.set_xlabel("← 부정 응답률(%)     긍정 응답률(%) →")
                    ax.invert_yaxis(); deco(ax, "긍정·부정 응답률", ylabel_top=False)
                    ax.grid(axis="y", visible=False)
                    for _b, _v in zip(_bp, summ[_pos_col]):
                        ax.text(float(_v)+_lim*.025, _b.get_y()+_b.get_height()/2, f"{float(_v):.1f}%",
                                va="center", ha="left", fontsize=8, color="#31485E")
                    for _b, _v in zip(_bn, summ[_neg_col]):
                        ax.text(-float(_v)-_lim*.025, _b.get_y()+_b.get_height()/2, f"{float(_v):.1f}%",
                                va="center", ha="right", fontsize=8, color="#7A3F3F")
                    ax.legend(fontsize=8, loc="upper center", bbox_to_anchor=(0.5, -0.16),
                              ncol=2, frameon=False, borderaxespad=0)
                    fig.subplots_adjust(bottom=.23, top=.88, left=.27, right=.97)
                    show_plot(fig); plt.close(fig)
                smart_table(summ, width="stretch")
                rep_blocks.append({"caption": "리커트 문항 요약", "table": summ, "image": dv_png})

            # 3) 다중응답 (가로 막대)
            if multi:
                st.markdown("## ☑️ 다중응답 문항")
                for c in multi:
                    ser = df[c].dropna().astype(str)
                    sep = next((x for x in [";", ",", "/", "|"]
                                if ser.str.contains(x, regex=False).mean() > 0.3), ";")
                    items = []
                    for v in ser: items += [x.strip() for x in v.split(sep) if x.strip()]
                    vc = pd.Series(items).value_counts()
                    t = pd.DataFrame({"응답 항목": vc.index, "응답 수": vc.values,
                                      "응답률(%)": (vc.values/len(ser)*100).round(1)})
                    cc1, cc2 = st.columns([1, 1])
                    with cc1:
                        st.write(f"**{c}** (응답자 {len(ser)}명)")
                        smart_table(t, width="stretch")
                    with cc2:
                        fig, ax = plt.subplots(figsize=figsize())
                        _vals = t["응답률(%)"][::-1].tolist()
                        _bars = ax.barh(t["응답 항목"][::-1], t["응답률(%)"][::-1],
                                        color=bar_colors(values=_vals))
                        for _b, _v, _n in zip(_bars, t["응답률(%)"][::-1], t["응답 수"][::-1]):
                            ax.text(_v + 1, _b.get_y() + _b.get_height()/2,
                                    f"{_v}% ({int(_n)}명)", va="center", fontsize=8)
                        ax.set_xlim(0, max(t["응답률(%)"]) * 1.25)
                        ax.set_xlabel("응답률(%)"); deco(ax, c)
                        show_plot(fig); plt.close(fig)
                    rep_blocks.append({"caption": f"다중응답 - {c}", "table": t,
                                       "image": fig_to_png(fig, show=False)})

            # 4) 주관식 (의견 목록)
            if openq:
                st.markdown("## ✍️ 주관식 의견")
                for c in openq:
                    ser = df[c].dropna().astype(str).str.strip(); ser = ser[ser != ""]
                    st.write(f"**{c}** — 응답 {len(ser)}건 "
                             f"(전체 {len(df)}명 중 {len(ser)/max(len(df),1)*100:.1f}%)")
                    op_tbl = pd.DataFrame({"번호": range(1, len(ser)+1), "의견": ser.values})
                    smart_table(op_tbl, width="stretch", hide_index=True, height=260)
                    _lim = st.number_input(
                        f"보고서에 담을 '{c}' 의견 수", 1, max(len(ser), 1), max(len(ser), 1),
                        key=f"svy_openlim_{c}",
                        help="의견이 많으면 한글 문서에서 표가 여러 쪽에 걸칩니다. "
                             "표는 쪽을 넘길 때 자동으로 나뉘고 머리행이 반복됩니다. "
                             "줄여서 담고 싶으면 여기서 건수를 조절하세요.")
                    rep_blocks.append({"caption": f"주관식 의견 - {c}",
                                       "table": op_tbl.head(int(_lim))})
                    st.session_state["subj_text"] = "\n".join(f"- {v}" for v in ser.tolist()[:100])

            txt = (f"문항 유형을 자동 판별해 리커트 {len(likert)}개, "
                   f"응답자 특성 {len(demo_q)}개, 객관식 문항 {len(single_q)}개, "
                   f"다중응답 {len(multi)}개, 주관식 {len(openq)}개를 그래프와 함께 분석했습니다.")
            st.success("✅ " + txt)
            if rep_blocks:
                log_action("설문 자동 인식 분석(그래프)")
                report_capture("cap_auto", "설문조사 자동 분석",
                               text=txt, blocks=[{"text": txt}] + rep_blocks)
                _ai_tbl = next((b["table"] for b in rep_blocks
                                if b.get("table") is not None), None)
                if _ai_tbl is not None:
                    ai_interpret_button("svyauto", "설문조사 자동 분석", _ai_tbl,
                                        "설문 응답 분포표입니다. 통계 검정 결과가 아니므로 "
                                        "'유의하다'는 표현은 쓰지 마세요.",
                                        capture_slot="cap_auto")
        report_button("cap_auto")
        survey_download_panel("cap_auto", "auto", "설문_자동분석")

    # ---------- 리커트 ----------
    # ---------- 리커트 ----------
    elif stype.startswith("📊"):
        c1, c2 = st.columns(2)
        demo = c1.multiselect("응답자 특성 열", df.columns.tolist(),
                              default=[c for c in cat_cols if c != "응답자ID"][:3], key="s_d")
        qs = c2.multiselect("문항 열 (숫자형)", num_cols,
                            default=[c for c in num_cols if not _looks_like_nominal_code(c)], key="s_q")
        demo = reorder_by_rank(demo, "s_d", "↕️ 응답자 특성 표시 순서 바꾸기")
        qs = reorder_by_rank(qs, "s_q", "↕️ 문항 표시 순서 바꾸기")
        scale_max = st.number_input("척도 최대값 (5점 척도면 5)", 2, 10, 5)
        _pos_cut, _neg_cut = likert_cutoffs(scale_max)
        _pos_col = f"긍정({_pos_cut}점↑)%"
        _neg_col = f"부정({_neg_cut}점↓)%"
        if qs and keep_running("svlikert", "리커트 분석 실행"):
            s = df[qs]  # 문항별 기술통계는 문항 각자의 결측만 제외 (교집합으로 자르지 않음)
            cc = s.dropna()  # 크론바흐 α·문항 제외 α는 전 문항 응답자만 필요
            st.markdown("#### 1) 응답자 특성")
            if demo:
                for d_ in demo:
                    vc = df[d_].value_counts()
                    st.write(f"**{d_}** (n={int(vc.sum())})")
                    smart_table(pd.DataFrame({d_: vc.index.astype(str), "빈도": vc.values,
                                               "비율(%)": (vc.values/vc.sum()*100).round(1)}),
                                 width="stretch")
            else:
                st.caption("응답자 특성 열을 선택하면 빈도표가 표시됩니다.")

            st.markdown("#### 2) 문항별 기술통계")
            summ = pd.DataFrame({
                "문항": qs,
                "응답자(명)": [int(s[q].notna().sum()) for q in qs],
                "평균": [round(s[q].dropna().mean(), 2) for q in qs],
                "표준편차": [round(s[q].dropna().std(), 2) for q in qs],
                "중앙값": [round(s[q].dropna().median(), 1) for q in qs],
                "긍정 응답(명)": [int((s[q].dropna() >= _pos_cut).sum()) for q in qs],
                _pos_col: [round((s[q].dropna() >= _pos_cut).mean()*100, 1) for q in qs],
                "부정 응답(명)": [int((s[q].dropna() <= _neg_cut).sum()) for q in qs],
                _neg_col: [round((s[q].dropna() <= _neg_cut).mean()*100, 1) for q in qs]})
            smart_table(summ, width="stretch")

            st.markdown("#### 3) 신뢰도 분석")
            a_ = cronbach_alpha(cc)
            lvl = (("매우 높음" if a_ >= .9 else "높음" if a_ >= .8 else "양호" if a_ >= .7 else "낮음")
                   if np.isfinite(a_) else "산출 불가")
            c1, c2 = st.columns(2)
            c1.metric("크론바흐 α", f"{a_:.3f}" if np.isfinite(a_) else "-"); c2.metric("신뢰도 수준", lvl)
            if len(cc) < len(df):
                st.caption(f"※ 크론바흐 α는 {len(qs)}개 문항에 모두 응답한 {len(cc)}명 기준입니다.")
            drop = pd.DataFrame({"제외 문항": qs,
                                 "제외 시 α": [round(cronbach_alpha(cc.drop(columns=[q])), 3) for q in qs]})
            st.write("**문항 제외 시 신뢰도** (α가 크게 올라가면 그 문항은 재검토 대상)")
            smart_table(drop, width="stretch")

            st.markdown("#### 4) 응답 분포")
            dist = pd.DataFrame({q: [int((s[q].dropna() == v).sum()) for v in range(1, scale_max+1)] for q in qs},
                                index=[f"{v}점" for v in range(1, scale_max+1)]).T
            # 인원과 비율을 함께 보여준다 (예: 12명 (20.0%))
            _dist_show = dist.copy().astype(object)
            for q in qs:
                _tot = max(int(dist.loc[q].sum()), 1)
                for v in range(1, scale_max+1):
                    _cnt = int(dist.loc[q, f"{v}점"])
                    _dist_show.loc[q, f"{v}점"] = f"{_cnt}명 ({_cnt/_tot*100:.1f}%)"
            _dist_show["합계"] = [f"{int(dist.loc[q].sum())}명 (100.0%)" for q in qs]
            smart_table(_dist_show.reset_index().rename(columns={"index": "문항"}),
                         width="stretch")
            st.caption("각 칸은 **응답 인원(명)과 비율(%)** 입니다.")
            st.markdown("##### 다이버징 차트 (중립 기준 좌우 분리)")
            counts_d = {q: [int((s[q].dropna() == v).sum()) for v in range(1, scale_max+1)] for q in qs}
            figd = likert_diverging(counts_d, [f"{v}점" for v in range(1, scale_max+1)], "문항별 응답 분포")
            png = fig_to_png(figd)
            _lk_h = max(4.2, len(qs) * 0.55 + 1.5)
            fig, axes = plt.subplots(1, 2, figsize=(max(11.5, figsize()[0]*2), _lk_h))
            _avg_vals = summ["평균"].astype(float).tolist()
            axes[0].barh(summ["문항"], summ["평균"],
                         color=bar_colors(values=_avg_vals), height=.58)
            axes[0].invert_yaxis(); axes[0].set_xlim(0, scale_max)
            axes[0].set_xlabel("평균 점수"); deco(axes[0], "문항별 평균", ylabel_top=False)
            axes[0].grid(axis="y", visible=False)
            bottom = np.zeros(len(qs))
            _mid = scale_max // 2
            _neg_full = ["#F3D8D8", "#E8AAAA", "#C96767", "#A94D4D", "#8F3E3E"]
            _pos_full = ["#D6E7F4", "#9EC5E5", "#6291C2", "#4576AB", "#2D5A8E"]
            if scale_max % 2 == 1:
                _lk_cols = (_neg_full[-_mid:] + ["#E3E9EF"] + _pos_full[:_mid])
            else:
                _lk_cols = (_neg_full[-_mid:] + _pos_full[:_mid])
            for v in range(1, scale_max+1):
                vals = dist[f"{v}점"].values
                _bars = axes[1].barh(qs, vals, left=bottom, label=f"{v}점",
                                     color=_lk_cols[v-1], height=.58, edgecolor="white", linewidth=.7)
                for _b, _vv, _left, _q in zip(_bars, vals, bottom, qs):
                    _tot = max(float(dist.loc[_q].sum()), 1.0)
                    if float(_vv) / _tot >= .10:
                        axes[1].text(float(_left)+float(_vv)/2, _b.get_y()+_b.get_height()/2,
                                     f"{float(_vv)/_tot*100:.0f}%", ha="center", va="center",
                                     fontsize=7, color="#23394D", fontweight="bold")
                bottom += vals
            axes[1].invert_yaxis(); deco(axes[1], "응답 분포(누적)", ylabel_top=False)
            axes[1].grid(axis="y", visible=False)
            axes[1].legend(fontsize=7.5, ncol=min(scale_max, 5), loc="upper center",
                           bbox_to_anchor=(0.5, -0.12), frameon=False, borderaxespad=0)
            fig.subplots_adjust(bottom=.20, top=.90, left=.10, right=.98, wspace=.34)
            show_plot(fig); plt.close(fig)

            st.markdown("#### 5) 집단별 비교")
            cmp_df = pd.DataFrame()
            if demo:
                rows = []
                for d_ in demo:
                    for q in qs:
                        sub = df[[d_, q]].dropna()
                        grp = [sub[sub[d_] == lv][q] for lv in sub[d_].unique()]
                        if len(grp) == 2: stat, p = stats.ttest_ind(*grp); tname = "t-검정"
                        elif len(grp) > 2: stat, p = stats.f_oneway(*grp); tname = "ANOVA"
                        else: continue
                        rows.append({"특성": d_, "문항": q, "검정": tname,
                                     "통계량": round(stat, 3), "p": round(p, 4),
                                     "유의성": "*" if p < .05 else "n.s."})
                cmp_df = pd.DataFrame(rows)
                smart_table(cmp_df, width="stretch")
                sig = cmp_df[cmp_df["유의성"] == "*"]
                st.info(f"💡 응답자 특성에 따라 차이가 유의한 항목 {len(sig)}개"
                        + (": " + ", ".join(f"{r['특성']}×{r['문항']}" for _, r in sig.head(3).iterrows()) if len(sig) else ""))
            else:
                st.caption("응답자 특성을 선택하면 집단별 차이 검정을 수행합니다.")

            valid_summ = summ.dropna(subset=["평균"])
            if valid_summ.empty:
                st.warning("⚠️ 선택한 문항 중 유효한 응답이 있는 열이 없습니다. "
                           "빈 열이 섞여 있지 않은지 '문항 열' 선택을 확인해 주세요.")
                top = low = "-"
            else:
                top = valid_summ.loc[valid_summ["평균"].idxmax(), "문항"]
                low = valid_summ.loc[valid_summ["평균"].idxmin(), "문항"]
            txt = (f"평균이 가장 높은 문항은 '{top}', 가장 낮은 문항은 '{low}'입니다. "
                   f"전체 신뢰도(크론바흐 α)는 {a_:.3f}로 {lvl} 수준입니다.")
            st.download_button("🖼️ 그래프 다운로드", png, "survey.png", "image/png")
            log_action("설문 리커트 분석")
            blocks = [{"text": txt},
                      {"caption": "문항별 기술통계", "table": summ, "image": png},
                      {"caption": "문항 제외 시 신뢰도", "table": drop}]
            if demo and len(cmp_df):
                blocks.append({"caption": "응답자 특성별 차이 검정", "table": cmp_df})
            report_capture("cap_survey", "설문조사 분석(리커트)", text=txt, blocks=blocks)
            ai_interpret_button("svylk", "설문 리커트 척도 분석", summ,
                                "평균·표준편차·긍정률과 크론바흐 알파가 있는 표입니다. "
                                "알파 0.7 이상이면 신뢰할 만하다는 기준을 함께 언급하세요.",
                                capture_slot="cap_survey")
        report_button("cap_survey")
        survey_download_panel("cap_survey", "likert", "설문_리커트분석")

    # ---------- 객관식 ----------
    elif stype.startswith("🔘"):
        st.caption("성별·소속·사용목적처럼 하나만 고르는 문항의 빈도와 비율을 분석합니다.")
        cols = st.multiselect("객관식 문항 열", df.columns.tolist(),
                              default=[c for c in cat_cols if c != "응답자ID"][:3], key="sc_q")
        if cols and keep_running("svmc", "객관식 분석 실행"):
            all_tbl = []
            for c in cols:
                vc = df[c].value_counts()
                t = pd.DataFrame({"문항": c, "응답": vc.index.astype(str), "빈도": vc.values,
                                  "비율(%)": (vc.values/vc.sum()*100).round(1)})
                all_tbl.append(t)
                cc1, cc2 = st.columns([1, 1])
                with cc1:
                    st.write(f"**{c}** (n={int(vc.sum())})")
                    smart_table(t.drop(columns=["문항"]), width="stretch")
                with cc2:
                    fig = pie_chart(vc, c, donut=True)
                    show_plot(fig); plt.close(fig)
            res = pd.concat(all_tbl, ignore_index=True)
            top = res.loc[res["빈도"].idxmax()]
            txt = f"가장 많은 응답은 '{top['문항']}'의 '{top['응답']}'({top['비율(%)']}%)입니다."
            st.info("💡 " + txt)
            png = fig_to_png(fig, show=False)
            log_action("설문 객관식 분석")
            report_capture("cap_mc", "설문 객관식 분석", txt, res, png)
            ai_interpret_button("svymc", "설문 객관식 응답 분포", res,
                                "빈도(명)와 비율(%)만 있는 표입니다. 통계 검정 결과가 아니므로 "
                                "'유의하다'는 표현은 쓰지 말고, 응답이 몰린 항목과 그 뜻을 서술하세요.",
                                capture_slot="cap_mc")
        report_button("cap_mc")
        survey_download_panel("cap_mc", "mc", "설문_객관식분석")

    # ---------- 다중응답 ----------
    elif stype.startswith("☑️"):
        st.caption("'해당되는 것을 모두 고르세요' 문항처럼 한 칸에 여러 답이 들어간 경우를 분석합니다.")
        c1, c2 = st.columns(2)
        mcol = c1.selectbox("다중응답 열", df.columns.tolist(), key="mr_c")
        sep = c2.selectbox("구분 기호", [";", ",", "/", "|", " "], key="mr_s")
        if keep_running("svmr", "다중응답 분석 실행"):
            ser = df[mcol].dropna().astype(str)
            n_resp = len(ser)
            items = []
            for v in ser: items += [x.strip() for x in v.split(sep) if x.strip()]
            vc = pd.Series(items).value_counts()
            t = pd.DataFrame({"응답 항목": vc.index, "응답 수": vc.values,
                              "응답률(%)": (vc.values/n_resp*100).round(1),
                              "구성비(%)": (vc.values/vc.sum()*100).round(1)})
            st.write(f"응답자 {n_resp}명 / 총 응답 {int(vc.sum())}건 (1인 평균 {vc.sum()/n_resp:.1f}개)")
            smart_table(t, width="stretch")
            fig, ax = plt.subplots(figsize=figsize())
            ax.barh(t["응답 항목"], t["응답률(%)"], color=bar_colors(values=t["응답률(%)"].tolist())); ax.invert_yaxis()
            ax.set_xlabel("응답률(%)"); deco(ax, f"{mcol} 다중응답")
            png = fig_to_png(fig)
            txt = (f"가장 많이 선택된 항목은 '{t.iloc[0]['응답 항목']}'로 응답자의 "
                   f"{t.iloc[0]['응답률(%)']}%가 선택했습니다. (응답률 합계가 100%를 넘는 것은 정상입니다)")
            st.info("💡 " + txt)
            log_action("설문 다중응답 분석")
            report_capture("cap_mr", "설문 다중응답 분석", txt, t, png)
            ai_interpret_button("svymr", "설문 다중응답 분석", t,
                                "다중응답이므로 응답률 합계가 100%를 넘는 것이 정상입니다. "
                                "이를 오류로 지적하지 마세요.",
                                capture_slot="cap_mr")
        report_button("cap_mr")
        survey_download_panel("cap_mr", "mr", "설문_다중응답분석")

    # ---------- 주관식 ----------
    elif stype.startswith("✍️"):
        st.caption("자유롭게 적은 의견을 모아 응답 목록·주요 단어를 확인하고, AI로 요약할 수 있습니다.")
        tcol = st.selectbox("주관식 열", df.columns.tolist(), key="tx_c")
        if keep_running("svtx", "주관식 분석 실행"):
            ser = df[tcol].dropna().astype(str).str.strip()
            ser = ser[ser != ""]
            c1, c2 = st.columns(2)
            c1.metric("응답 건수", f"{len(ser)}건")
            c2.metric("응답률", f"{len(ser)/max(len(df),1)*100:.1f} %")
            st.markdown("#### 의견 목록")
            op_tbl = pd.DataFrame({"번호": range(1, len(ser)+1), "의견": ser.values})
            smart_table(op_tbl, width="stretch", hide_index=True, height=320)
            txt = (f"주관식 문항 '{tcol}'에 대해 전체 {len(df)}명 중 {len(ser)}명"
                   f"({len(ser)/max(len(df),1)*100:.1f}%)이 의견을 제시하였다.")
            st.info("💡 " + txt)
            st.session_state["subj_text"] = "\n".join(f"- {v}" for v in ser.tolist()[:100])
            log_action("설문 주관식 분석")
            report_capture("cap_tx", "설문 주관식 의견", txt, op_tbl, None)
        if st.session_state.get("subj_text"):
            st.markdown("#### 🧠 AI로 의견 요약하기 (선택)")
            if st.button("AI 요약 실행"):
                with st.spinner("AI가 의견을 정리하는 중..."):
                    st.markdown(ai_call(
                        "다음은 설문조사의 주관식 응답입니다. 주요 의견을 3~5개 주제로 묶어 "
                        "각 주제별 핵심 내용과 대표 의견을 한국어로 정리해 주세요. "
                        "마지막에 개선 우선순위를 제안해 주세요.\n\n" + st.session_state["subj_text"],
                        st.session_state.get("api_key"), st.session_state.get("ai_model_g"), max_tokens=1200))
                    log_action("AI 주관식 의견 요약")
        report_button("cap_tx")
        survey_download_panel("cap_tx", "text", "설문_주관식분석")

    # ---------- 교차분석 ----------
    else:
        st.caption("두 범주형 문항의 관계를 교차표와 카이제곱 검정으로 확인합니다. (예: 소속 × 재사용 의향)")
        c1, c2 = st.columns(2)
        rowv = c1.selectbox("행 변수", df.columns.tolist(), key="ct_r")
        colv = c2.selectbox("열 변수", [c for c in df.columns if c != rowv], key="ct_c")
        pct = st.radio("비율 기준", ["빈도만", "행 기준 %", "열 기준 %"], horizontal=True)
        if keep_running("svct", "교차분석 실행"):
            sub = df[[rowv, colv]].dropna()
            ct = pd.crosstab(sub[rowv], sub[colv])
            if pct == "행 기준 %": pct_tbl = (ct.div(ct.sum(axis=1), axis=0)*100).round(1)
            elif pct == "열 기준 %": pct_tbl = (ct.div(ct.sum(axis=0), axis=1)*100).round(1)
            else: pct_tbl = (ct / ct.values.sum() * 100).round(1)
            # 인원(빈도)과 비율(%)을 한 칸에 함께 표시 (예: 12명 (33.3%))
            show = ct.astype(object).copy()
            for r in ct.index:
                for c in ct.columns:
                    show.loc[r, c] = f"{int(ct.loc[r, c])}명 ({pct_tbl.loc[r, c]:.1f}%)"
            show["합계"] = [f"{int(ct.loc[r].sum())}명 (100.0%)" for r in ct.index]
            smart_table(show, width="stretch")
            try:
                chi2, p, dof, exp = stats.chi2_contingency(ct)
                low = (exp < 5).sum() / exp.size * 100
                if low > 20:
                    st.warning(f"⚠️ 기대빈도가 5 미만인 칸이 {low:.0f}%입니다(기준 20%). "
                               "카이제곱 결과가 부정확할 수 있으니 범주를 합치거나 Fisher 정확검정을 고려하세요.")
                c1, c2, c3 = st.columns(3)
                c1.metric("카이제곱", f"{chi2:.3f}"); c2.metric("자유도", dof); c3.metric("p-value", f"{p:.4f}")
                txt = (f"'{rowv}'와 '{colv}' 사이에 "
                       + ("통계적으로 유의한 관련성이 있습니다" if p < .05 else "유의한 관련성이 없습니다")
                       + f" (χ²={chi2:.2f}, p={p:.4f}).")
                st.info("💡 " + txt)
            except Exception as e:
                txt = "교차표를 생성했습니다."; st.warning(f"카이제곱 검정 불가: {e}")
            row_pct = (ct.div(ct.sum(axis=1), axis=0)*100).round(1)  # 막대 라벨은 항상 행 기준 %로 표기
            _w, _h = figsize()
            fig, ax = plt.subplots(figsize=(_w + 1.2, _h))
            bottoms = np.zeros(len(ct))
            # 칸이 좁은데 '9명\n(30.0%)'를 두 줄로 넣으면 글씨가 서로 겹친다.
            # 그렇다고 '명'이나 '%'를 통째로 빼면 뭘 나타내는 숫자인지 알기 어려우므로,
            # 세로 공간이 부족하면 '9명(30.0%)'처럼 한 줄로 합쳐서라도 단위·비율을 남긴다.
            _ymax = float(ct.sum(axis=1).max()) or 1.0
            _hidden = 0
            _ct_cols = _survey_palette(len(ct.columns))
            for _ci, col in enumerate(ct.columns):
                vals = ct[col].values
                ax.bar(ct.index.astype(str), vals, bottom=bottoms, label=str(col),
                       color=_ct_cols[_ci], edgecolor="white", linewidth=.6)
                for x, (v, b, pv) in enumerate(zip(vals, bottoms, row_pct[col].values)):
                    if v <= 0:
                        continue
                    lab_fs = crosstab_bar_label(v, pv, _ymax)
                    if lab_fs is None:
                        _hidden += 1
                        continue
                    lab, fs = lab_fs
                    ax.text(x, b + v/2, lab, ha="center", va="center",
                            fontsize=fs, color="white")
                bottoms += vals
            ax.set_ylabel("응답 수"); deco(ax, f"{rowv} × {colv}")
            ax.set_ylim(0, _ymax * 1.05)
            plt.xticks(rotation=20)
            # 범례를 그림 밖으로 빼서 막대·숫자를 가리지 않게 한다
            ax.legend(title=str(colv), fontsize=7, title_fontsize=8,
                      loc="upper left", bbox_to_anchor=(1.01, 1.0), borderaxespad=0)
            plt.tight_layout(); png = fig_to_png(fig)
            if _hidden:
                st.caption(f"※ 칸이 너무 좁은 {_hidden}곳은 글씨가 겹쳐 숫자를 생략했습니다. "
                           "정확한 값은 위 교차표를 보세요.")
            out = show.reset_index()
            log_action(f"교차분석: {rowv} × {colv}")
            report_capture("cap_ct", f"{rowv} × {colv} 교차분석", txt, out, png)
            ai_interpret_button("svyct", f"{rowv} × {colv} 교차분석", out,
                                f"카이제곱 검정 결과는 다음과 같습니다: {txt} "
                                "표의 값은 '인원(비율%)' 형식이며 비율은 행 기준입니다.",
                                capture_slot="cap_ct")
        report_button("cap_ct")
        survey_download_panel("cap_ct", "cross", "설문_교차분석")

# ================================================================ 사용설명서
elif menu == "👑 관리자":
    render_admin_dashboard()

elif menu == "📖 사용설명서":
    st.title("📖 사용설명서")
    _MANUAL = "# 스마트 통계 에이전트 — 사용설명서\n\n농업연구사·지도사를 위한 실험데이터 통계분석 자동화 도구\n\n---\n\n## 1. 웹에서 시작하기\n\n별도 설치 없이 **스마트 통계 에이전트 웹주소에 접속**해서 사용합니다.\n로그인 기능이 켜져 있으면 회원가입/로그인 후 아래 순서대로 진행하세요.\n\n```\n① 데이터 넣기  →  ② 분석하기  →  ③ 결과 내려받기·보고서 만들기\n```\n\n**① 데이터 올리기**\n- 왼쪽 사이드바 위쪽에서 엑셀(xlsx)·CSV를 올리거나, 이미지·카메라·음성 입력을 선택합니다.\n- 엑셀에 시트가 여러 개면 **시트별로 자동 분리**됩니다.\n- 파일이 없으면 **🧪 샘플 데이터** 버튼으로 먼저 체험해 보세요.\n\n**② 분석하기**\n- 왼쪽 메뉴에서 원하는 분석을 고릅니다.\n- 결과 아래 **➕ 이 결과를 보고서에 담기**를 누릅니다.\n\n**③ 결과 내려받기·보고서 만들기**\n- 각 분석 화면에서 한글(hwpx)·Excel(xlsx) 결과를 바로 내려받거나,\n- 📑 보고서 메뉴에 여러 분석 결과를 모아 한글(hwpx)·워드(docx) 보고서를 만듭니다.\n\n> 💡 **결과는 사라지지 않습니다.** 다른 메뉴에 갔다 와도 분석 결과가 그대로 남아 있습니다.\n\n---\n\n## 2. 메뉴 한눈에 보기\n\n| 메뉴 | 무엇을 하나요 |\n|---|---|\n| 📊 **통계분석** | 데이터 정리 · 분산분석 · 상관 · 회귀 등 |\n| 🧠 **AI 도우미** | 궁금한 걸 물어보고, 어떤 분석을 할지 추천받기 |\n| 💰 **경제성분석** | 소득·순수익 계산, 부분예산표, 증수 효과 |\n| 📋 **설문조사 분석** | 만족도·의견 조사 결과 정리 |\n| 📑 **보고서** | 담아둔 결과를 문서로 만들기 |\n\n---\n\n## 3. 통계분석 메뉴\n\n### 📋 데이터\n- 올린 데이터를 확인하고 **🩺 데이터 검진** 결과를 봅니다.\n- 숫자인데 문자로 읽힌 열(`1,200`, `120kg`)이 있으면 **원클릭 변환** 버튼이 나타납니다.\n- 실험설계(난괴법 등)를 자동으로 추정해 알려줍니다.\n\n### 🧹 전처리\n데이터를 정리합니다. **화면 위에 현재 데이터가 항상 보이므로** 작업 결과를 바로 확인할 수 있습니다.\n\n| 작업 | 언제 쓰나요 |\n|---|---|\n| 결측치 처리 | 조사 누락으로 빈칸이 있을 때 |\n| 이상값 처리 | 입력 실수(15.0 → 150)가 의심될 때 |\n| 중복 행 제거 | 같은 자료가 두 번 들어갔을 때 |\n| 자료형 변환 | 숫자가 문자로 읽혔을 때 |\n| 열 삭제/이름변경 | 필요 없는 열을 뺄 때 |\n| 표준화·정규화 | 단위가 다른 변수를 비교할 때 |\n\n> ↩️ **실행취소** 버튼으로 최대 10단계까지 되돌릴 수 있습니다.\n\n### 🧮 파생변수\n기존 열을 조합해 새 열을 만듭니다.\n- **두 열 사칙연산**: `수량 × 단가 = 조수입`\n- **조건 열**: `기온 ≥ 33` → 폭염일 표시\n- **그룹별 집계**: 연도별 합계·평균\n\n### 🔗 상관분석\n두 변수가 함께 변하는 정도를 봅니다.\n- **Pearson**(직선 관계) / **Spearman**(순위·비정규분포)\n- 논문용 **유의성 별표 표**(`0.826***`)가 자동 생성됩니다.\n\n### 📈 분산분석 — 가장 많이 쓰는 기능\n\n**분석 방식 6가지**\n\n| 방식 | 언제 |\n|---|---|\n| **일원배치** | 처리구 하나 비교 (가장 기본) |\n| **이원배치** | 두 요인 + 상호작용 (품종 × 시비량) |\n| **🌾 분할구법** | 관수·경운처럼 큰 구역 요인이 있을 때 |\n| **🔁 반복측정** | 같은 개체를 시기별로 반복 조사 |\n| **🎚️ ANCOVA** | 초기 생육 차이를 보정하고 싶을 때 |\n| **📊 여러 형질 요약표** | 여러 항목을 한 표로 (논문 표 형식) |\n\n**⚠️ 꼭 확인하세요 — 반복(블록) 열**\n포장시험에서 반복을 두었다면 **반복 열을 반드시 지정**하세요.\n지정하지 않으면 블록 간 토양·경사 차이가 오차에 섞여, 실제로는 있는 처리 효과를 놓칠 수 있습니다.\n\n**사후검정 4가지**\n\n| 방법 | 특징 |\n|---|---|\n| **Tukey HSD** | 국제 표준. 논문 투고에 안전 |\n| **던컨(DMRT)** | 농업 논문 관행. 차이를 잘 잡아냄 |\n| **Bonferroni** | 매우 엄격 |\n| **던넷(Dunnett)** | 대조구와만 비교 (신품종 vs 대비품종) |\n\n**결과 읽기**\n- **유의성 문자**: 같은 문자를 공유하면 차이 없음 (`a`, `ab`, `b`)\n- **CV(%)**: 시험 정밀도. 포장시험 10~20% 양호, 20% 초과 시 재검토\n- **LSD**: 두 평균의 차이가 이 값보다 크면 유의한 차이\n- **논문용 표 각주**가 자동 생성되니 복사해서 쓰세요.\n\n### 🧪 비모수검정\n정규성·등분산 가정이 깨졌을 때 사용합니다.\n(분산분석에서 가정이 위배되면 **자동으로 결과를 함께 보여줍니다.**)\n\n### 🧬 PCA (주성분분석)\n형질이 많을 때 2개 축으로 압축해 그림 하나로 봅니다.\n- 누적 설명분산 **70% 이상**이면 신뢰할 만합니다.\n\n### 📉 회귀분석\n- **단순/다중 회귀**: X로 Y를 설명 (VIF 다중공선성 진단 포함)\n- **로지스틱**: Y가 두 가지 값일 때 (발병/미발병)\n- **🧪 프로빗**: 농약 시험의 **LC50/LD50** 산출\n- **잔차 진단**으로 모형이 적절한지 확인할 수 있습니다.\n\n### 🤖 머신러닝\n수량 예측·등급 판정 등. **🔮 새 데이터 예측** 기능으로 값을 넣으면 바로 예측합니다.\n\n> ⚠️ 표본이 30개 미만이면 쓰지 마세요. 처리 효과 검정은 **분산분석**을 쓰세요.\n\n---\n\n## 4. 🧠 AI 도우미 (선택)\n\nAPI 키를 넣으면 쓸 수 있습니다. **키가 없어도 나머지 기능은 모두 정상 작동합니다.**\n\n**키 넣는 법**: 사이드바 → 🤖 AI 기능 켜기 → 제공사 선택(Claude·Gemini·ChatGPT) → 키 입력\n\n**할 수 있는 것**\n- 결과를 자연어로 질문하기\n- 데이터 자동 요약\n- 연구계획서를 올리면 어떤 분석을 할지 추천\n- 각 분석 결과 아래 **🤖 AI 해석**으로 보고서 문장(`○`·`-` 형식) 생성\n\n> ⚠️ AI가 만든 문장은 **반드시 연구자가 수치와 해석을 확인**한 뒤 사용하세요.\n\n---\n\n## 5. 💰 경제성분석\n\n### 🧭 경제성 분석 길잡이 v3.6 — 처음이면 여기부터\n경제성 분석을 배운 적이 없어도 **STEP 1~5를 한 단계씩** 답하면 적합한 분석방법을 규칙 기반으로 추천합니다. 처음 화면에서는 길잡이와 바로 분석하기 중 하나만 선택하므로 화면이 복잡하지 않습니다. API 키 없이 작동합니다.\n\n1. **STEP 1 연구목적** — 신기술 비교, 현재 수익성, 여러 대안 선택, 시설투자, 손익분기, 위험평가 등\n2. **STEP 2 변화요인** — 품종·방제·재배법, 투입수준, 시설·농기계, 가격·수량 등\n3. **STEP 3 비교구조** — 대조구 vs 신기술, 비용이 다른 여러 대안, 여러 처리의 한 해 성과, 비교대상 없음\n4. **STEP 4 분석기간** — 한 작기·1년, 2년 이상, 시설·기계 내용연수 전체\n5. **STEP 5 보유자료** — 처리구, 수량, 가격, 경영비, 변화비용, 반복, 투자비, 연도별 편익·비용, 할인율 등\n\n모르는 항목은 **잘 모르겠어요**를 골라도 나머지 답으로 판단하며, 답이 너무 불확실하면 임의로 분석을 연결하지 않고 **추천 보류 + 확인할 두 가지**를 안내합니다.\n- 현재 작목의 수익성 → **소득분석**\n- 대조구와 신품종·신기술 비교 → **부분예산법**\n- 비용이 다른 여러 기술 중 추천대안 선택 → **지배분석·MRR**\n- 시설·농기계 장기투자 → **NPV·할인 B/C·IRR**\n- 손익분기 가격·수량 / 가격·수량 변동 위험 → **소득분석 안의 손익분기점·민감도**\n- 정책사업의 사회적 효과(CBA)는 현재 직접 계산하지 않으며 별도 분석이 필요하다고 안내합니다.\n\n추천 결과에는 **추천 확신도, 추천 이유, 필요한 자료, 함께 볼 분석, 자료 준비도**가 표시됩니다. STEP 5 앞의 **📦 자료 준비 가이드**에서는 공통자료, 농촌진흥청 소득조사 체계에 맞춘 비용 항목, 분석별 추가자료와 빈 CSV 서식을 제공합니다. 올린 데이터가 있으면 처리구·수량·단가·비용·반복 열 후보를 자동으로 찾아 STEP 5와 데이터 점검에 활용합니다. **🚀 추천 분석 바로 시작하기**를 누르면 해당 분석 방식이 자동 선택됩니다.\n\n### 💾 기준단가 관리 (먼저 설정)\n노임·자재비·임차료를 한 번 넣어두면 분석에 자동 반영됩니다.\nCSV로 내려받아 보관하고, 다음 해에 갱신해 다시 올릴 수 있습니다.\n\n**기본으로 들어있는 공식 수치** (기준연도 확인 후 사용하세요)\n\n| 항목 | 단가 | 기준연도 | 출처 |\n|---|---|---|---|\n| 농업노임(남) | 153,520원/일 | 2025년 | 통계청 KOSIS 농가판매·구입가격조사 |\n| 농업노임(여) | 121,392원/일 | 2025년 | 통계청 KOSIS |\n| 농업노임(남·시간) | 19,190원/시간 | 2025년 | 일당 ÷ 8시간 |\n| 농업노임(여·시간) | 15,174원/시간 | 2025년 | 일당 ÷ 8시간 |\n| 요소비료(20kg) | 17,900원 | 2026년 | 농협 (보조금 적용 시 16,250원) |\n| 토지용역비(밭) | 260원/㎡ | 2024년 | 농지임차료실태조사 |\n| 토지용역비(논) | 275원/㎡ | 2024년 | 농지임차료실태조사 |\n\n**자동으로 받아올 수 있는 자료**\n- **🌐 KAMIS** : 농산물 가격 (일별) — kamis.or.kr에서 인증키 발급(무료)\n- **📊 KOSIS** : 농촌 일용노임·농가구입가격지수 (분기) — kosis.kr/openapi에서 인증키 발급(무료)\n  - 가장 쉬운 방법: KOSIS 통계표 화면에서 [OpenAPI] 버튼 → 주소 복사 → 앱에 붙여넣기\n\n**직접 확인해 입력해야 하는 자료**\n- 자본용역비 이자율, 감가상각 내용연수 → 농촌진흥청 「농축산물 소득자료집」 부록\n  또는 농산업경영과(063-238-1197) 문의\n- 농협 자재 실판매가(연 1회), 위탁영농비, 지역별 노임\n\n### 📕 부분예산표 (손실적·이익적 요소) — 가장 많이 쓰는 기능\n신기술 도입 시 **바뀐 것만** 모아 늘어난 비용(A)과 늘어난 이익(B)을 비교하고 **추정수익액(B−A)** 을 구합니다. 관행·신기술이 똑같이 쓰는 비용은 넣지 않습니다.\n\n| 자료에서 바뀐 것 | 어디로 |\n|---|---|\n| 수량이 늘었다 (× 단가) | 이익적 요소(B) |\n| 수량이 줄었다 (× 단가) | 손실적 요소(A) |\n| 비용이 늘었다 | 손실적 요소(A) |\n| 비용이 줄었다 | **이익적 요소(B)** — 절감은 번 것입니다 |\n| 노동시간이 늘었다 (× 시간당 노임) | 손실적 요소(A) |\n\n> ⚠️ **자가노동시간처럼 값이 '시간'인 열**은 비용 열이 아니라 **노동시간 열**에 넣으세요. 비용 열에 넣으면 10시간이 10원으로 계산됩니다. 이름에 '시간'이 들어가면 자동으로 골라 줍니다.\n\n**두 가지 방법으로 만들 수 있습니다.**\n1. 화면에서 항목·산출근거·금액을 직접 입력\n2. **📊 올린 데이터에서 자동으로 채우기** — 처리구 열, 수량 열, 대조구, 신기술구, 단가, 비용 열을 고르면 대조구 대비 달라진 값만 뽑아 표를 채워 줍니다(10a 기준 자동 환산). 채운 뒤 손으로 고치거나 항목을 더할 수 있습니다.\n\n> 💡 화면의 **🧮 이 숫자가 어떻게 나온 건가요?** 에 예시 숫자로 따라가는 계산 과정과 자주 하는 실수가 정리되어 있습니다.\n\n### 📗 소득분석\n```\n총수입 = 주산물가액 + 부산물가액\n소득   = 총수입 − 경영비\n순수익 = 총수입 − 생산비\n소득률(%) = 소득 ÷ 총수입 × 100\n```\n**어떤 엑셀을 올리나요** — 화면의 `📋 어떤 엑셀을 올려야 하나요?` 에서 예시 서식을 CSV로 내려받아 숫자만 바꿔 올리면 됩니다.\n- **한 줄 = 한 조사구**(처리구 × 반복)로 적습니다.\n- 모든 값은 **10a(1,000㎡) 기준**으로 환산해 적습니다.\n- 합계·소득같은 **계산 결과 열은 넣지 마세요.** 이중으로 잡힙니다.\n- **작목 유형**(식량작물·노지채소·시설채소·과수 등)을 고르면 그에 맞는 항목을 안내합니다.\n- **과수·다년생**은 과수원 조성비를 내용연수로 나눠 매년 상각합니다.\n- **대조구를 지정하면** 증수율·증수액·순증가소득이 자동 계산됩니다.\n\n### 📘 신기술 경제성 (부분예산·MRR)\nCIMMYT의 부분예산 원칙을 참고해 지배분석과 한계수익률을 계산합니다.\n권장 여부는 사용자가 설정한 최소 MRR과 반복수·가격·수량 민감도를 함께 확인합니다.\n\n---\n\n## 6. 📋 설문조사 분석\n\n**🤖 자동 인식**을 쓰면 문항 유형을 스스로 판별해 한 번에 분석합니다.\n\n| 유형 | 결과 |\n|---|---|\n| 리커트 척도 | 평균·표준편차, 긍정률, **크론바흐 α**, 다이버징 차트 |\n| 객관식 | 빈도·비율 표 + 원형/도넛 그래프 (%·인원 표시) |\n| 다중응답 | 응답률 (합계가 100%를 넘는 것이 정상) |\n| 주관식 | 의견 목록 표, AI 요약 |\n| 교차분석 | 교차표 + 카이제곱 검정 |\n\n**크론바흐 α**: 0.7 이상이면 신뢰할 만합니다.\n\n**🤖 AI 해석**: 객관식·다중응답·리커트·교차분석·자동인식 결과 아래에서 보고서 문장을 만들 수 있습니다.\n설문은 실험과 달라서, 검정하지 않은 결과에 '유의하다'고 쓰지 않도록 AI에게 미리 일러 둡니다.\n\n---\n\n## 7. 📑 보고서\n\n**만드는 순서**\n1. 각 분석에서 **➕ 보고서에 담기**\n2. 📑 보고서 메뉴로 이동\n3. 필요하면 **표지·재료및방법**, **적요**, **표·그림 목차** 추가\n4. 한글(hwpx) 또는 워드(docx)로 내려받기\n\n**자동으로 만들어지는 것**\n- `□` 항목 제목, `◦` 불릿 (시험연구보고서 양식)\n- `<표 1>`, `<그림 1>` 캡션 (왼쪽 정렬)\n- **통계처리 문구** — 어떤 분석·사후검정을 썼는지 자동 서술\n- **적요 초안** — 담긴 결과를 요약\n\n**계획서 첨부**: 연구계획서 파일(hwpx·docx·pdf)을 올리면 내용을 보고서 앞에 넣을 수 있습니다.\n\n---\n\n## 8. 자주 겪는 문제\n\n| 증상 | 해결 |\n|---|---|\n| 한글 파일이 안 열림 | 구버전 한글은 hwpx 미지원 → **워드(docx)로 받으세요** |\n| 워드(docx) 버튼이 비활성화됨 | 현재 배포 서버에서 Word 생성 기능이 꺼진 상태입니다. 관리자에게 문의하세요. |\n| 그래프 글자가 □로 깨짐 | 유의성 문자는 서버 글꼴에 의존하지 않는 수식 위첨자로 표시됩니다. 일반 한글이 깨지면 배포 서버 글꼴을 확인하세요. |\n| 결과가 안 보임 | 분석 실행 버튼을 눌렀는지 확인 |\n| 숫자 열인데 분석이 안 됨 | 데이터 탭에서 **숫자로 변환** 실행 |\n| 반복 열을 안 넣었는데 결과가 이상함 | 분산분석에서 **반복(블록) 열**을 지정하세요 |\n| 빨간 오류 상자가 떴음 | 그 아래 **🆘 이 오류, 도움받기**에서 `🤖 앱 안에서 바로 물어보기` |\n\n---\n\n## 9. 꼭 기억할 것\n\n1. **반복(블록) 열을 지정하세요** — 포장시험 결과가 달라집니다.\n2. **CV(%)를 확인하세요** — 20% 넘으면 시험 정밀도를 점검하세요.\n3. **AI 문장은 반드시 검토하세요** — 그대로 제출하지 마세요.\n4. **분석 결과는 바로 보고서에 담으세요** — 나중에 한 번에 문서가 됩니다.\n5. **기준단가는 기관 공식 자료로 입력하세요** — 자동으로 받아오지 않습니다.\n\n---\n\n*스마트 통계 에이전트*\n"
    def _manual_html(md_text):
        """설명서를 어느 컴퓨터에서나 열리는 HTML로 변환"""
        try:
            import markdown as _md
            body = _md.markdown(md_text, extensions=["tables", "fenced_code"])
        except Exception:
            body = "<pre>" + md_text.replace("<", "&lt;") + "</pre>"
        return """<!DOCTYPE html><html lang="ko"><head><meta charset="utf-8">
<title>스마트 통계 에이전트 사용설명서</title><style>
body{{font-family:'맑은 고딕','Malgun Gothic',sans-serif;max-width:900px;margin:40px auto;
padding:0 20px;line-height:1.7;color:#222}}
h1{{border-bottom:3px solid #6c8ebf;padding-bottom:8px}}
h2{{margin-top:36px;border-left:5px solid #6c8ebf;padding-left:10px}}
h3{{margin-top:24px;color:#365f91}}
table{{border-collapse:collapse;width:100%;margin:14px 0}}
th,td{{border:1px solid #ccc;padding:8px;text-align:left}}
th{{background:#eef3fa}}
code{{background:#f4f4f4;padding:2px 5px;border-radius:3px}}
pre{{background:#f7f7f7;padding:12px;border-radius:5px;overflow-x:auto}}
blockquote{{border-left:4px solid #d79b00;background:#fff8ec;margin:12px 0;padding:8px 14px}}
@media print{{body{{margin:0}}}}
</style></head><body>{body}</body></html>""".format(body=body)
    d1, d2 = st.columns(2)
    d1.download_button("🌐 설명서 내려받기 (HTML — 바로 열림)",
                       _manual_html(_MANUAL).encode("utf-8"),
                       "사용설명서.html", mime="text/html", width="stretch")
    d2.download_button("📄 설명서 내려받기 (텍스트)", _MANUAL.encode("utf-8"),
                       "사용설명서.txt", mime="text/plain", width="stretch")
    st.caption("HTML 파일은 더블클릭하면 웹브라우저에서 바로 열립니다. 인쇄도 가능해요.")
    st.markdown(_MANUAL)

# ================================================================ 보고서
else:
    st.title("📑 자동 보고서 생성")
    with st.expander("ℹ️ 이 기능이 뭔가요?"): st.markdown(EXPLAIN["report"])

    logs = st.session_state.get("log", [])
    with st.expander(f"🕘 분석 이력 ({len(logs)}건)"):
        if logs:
            log_df = pd.DataFrame(logs)
            smart_table(log_df, width="stretch")
            c1, c2 = st.columns(2)
            if c1.button("➕ 이력을 보고서에 담기", width="stretch"):
                st.session_state.report_items.append(
                    {"heading": "분석 수행 이력", "table": log_df,
                     "text": f"본 보고서 작성 과정에서 수행한 분석 {len(logs)}건의 기록입니다.",
                     "image": None})
                st.success("보고서에 담았습니다.")
            if c2.button("🗑️ 이력 지우기", width="stretch"):
                st.session_state.log = []; st.rerun()
        else:
            st.caption("아직 수행한 분석이 없습니다. 분석을 실행하면 자동으로 기록됩니다.")

    with st.expander("📝 표지 · 재료 및 방법 만들기 (시험연구보고서 표준 양식)", expanded=False):
        st.caption("아래를 채우면 보고서 맨 앞에 **표지**와 **재료 및 방법** 항목이 자동으로 만들어집니다. "
                   "빈칸은 생략됩니다.")
        st.markdown("###### 1) 표지 정보")
        cc1, cc2 = st.columns(2)
        pj_title = cc1.text_input("과제명", key="pj_title",
                                  placeholder="예) 고추 신품종의 생육 및 수량 특성 비교")
        pj_org = cc2.text_input("소속 기관", key="pj_org", placeholder="예) OO도농업기술원")
        cc3, cc4, cc5 = st.columns(3)
        pj_author = cc3.text_input("연구자", key="pj_author")
        pj_period = cc4.text_input("연구 기간", key="pj_period", placeholder="예) 2026. 1. ~ 2026. 12.")
        pj_keyword = cc5.text_input("색인용어", key="pj_kw", placeholder="예) 고추, 신품종, 수량")

        st.markdown("###### 2) 재료 및 방법")
        mm1, mm2 = st.columns(2)
        m_site = mm1.text_input("시험 장소", key="m_site", placeholder="예) OO연구소 시험포장")
        m_variety = mm2.text_input("공시 품종·재료", key="m_var", placeholder="예) 청양, 수비초 등 4품종")
        mm3, mm4 = st.columns(2)
        m_design = mm3.selectbox("실험 설계",
                                 ["(선택 안 함)", "난괴법(RCBD) 3반복", "난괴법(RCBD) 4반복",
                                  "완전임의배치(CRD) 3반복", "완전임의배치(CRD) 4반복",
                                  "분할구법(Split-plot) 3반복", "요인배치법"], key="m_design")
        m_area = mm4.text_input("구당 면적", key="m_area", placeholder="예) 10㎡ (3.3m × 3m)")
        m_treat = st.text_area("처리 내용", key="m_treat", height=80,
                               placeholder="예) 1. 대조구(관행)  2. 처리1(질소 20% 증시)  3. 처리2(질소 40% 증시)")
        m_manage = st.text_area("재배 관리", key="m_manage", height=80,
                                placeholder="예) 정식 5월 10일, 시비량 N-P-K = 19-11-14kg/10a, 관행 방제")
        m_survey = st.text_area("조사 항목 및 방법", key="m_survey", height=80,
                                placeholder="예) 초장·엽수는 정식 후 30일 간격 5주 조사, 수량은 수확기별 전량 계량")

        st.markdown("###### 3) 통계처리 문구 (자동 생성)")
        auto_cv = st.text_input("CV(%) — 분산분석에서 확인한 값 (선택)", key="m_cv", placeholder="예) 12.3")
        stat_text = build_stat_method_text(
            logs, {"design": None if m_design.startswith("(") else m_design,
                   "cv": auto_cv.strip() or None})
        st.code(stat_text, language=None)
        st.caption("분석 이력을 바탕으로 자동 작성되었습니다. 필요하면 복사해 수정하세요.")

        if st.button("➕ 표지·재료및방법을 보고서 맨 앞에 넣기", width="stretch"):
            new_items = []
            if pj_title or pj_org or pj_author:
                cover = []
                if pj_title: cover.append(f"◦ 과제명 : {pj_title}")
                if pj_org: cover.append(f"◦ 소속 : {pj_org}")
                if pj_author: cover.append(f"◦ 연구자 : {pj_author}")
                if pj_period: cover.append(f"◦ 연구기간 : {pj_period}")
                if pj_keyword: cover.append(f"◦ 색인용어 : {pj_keyword}")
                new_items.append({"heading": "시험연구", "text": "\n".join(cover),
                                  "table": None, "image": None})
            mm_parts = []
            if m_site: mm_parts.append(("시험 장소", m_site, False))
            if m_variety: mm_parts.append(("공시 품종·재료", m_variety, False))
            if not m_design.startswith("("): mm_parts.append(("시험 설계", m_design, False))
            if m_area: mm_parts.append(("구당 면적", m_area, False))
            if m_treat: mm_parts.append(("처리 내용", m_treat, True))
            if m_manage: mm_parts.append(("재배 관리", m_manage, True))
            if m_survey: mm_parts.append(("조사 항목 및 방법", m_survey, True))
            mm_parts.append(("통계 처리", stat_text, True))
            _KOR = "가나다라마바사아자차"
            mm_lines = []
            for _i, (t, v, nl) in enumerate(mm_parts):
                head = _KOR[_i] if _i < len(_KOR) else str(_i+1)
                if nl:
                    mm_lines.append(f"{head}. {t}")
                    for _ln in strip_md(v).split("\n"):
                        if _ln.strip():
                            _c = _ln.strip()
                            mm_lines.append(_c if _c.startswith(("○", "-", "•")) else f"  ○ {_c}")
                else:
                    mm_lines.append(f"{head}. {t} : {v}")
            new_items.append({"heading": "Ⅰ. 재료 및 방법", "text": "\n".join(mm_lines),
                              "table": None, "image": None})
            for it in reversed(new_items):
                st.session_state.report_items.insert(0, it)
            log_action("표지·재료및방법 작성")
            st.success("보고서 맨 앞에 넣었습니다!")
            st.rerun()

    with st.expander("📎 연구계획서·결과보고서 첨부 (보고서 앞부분에 넣기)"):
        st.caption("계획서나 기존 보고서 파일을 올리면 그 내용을 이 보고서 맨 앞에 넣을 수 있어요. "
                   "(hwpx·docx·pdf·txt·csv·xlsx 지원)")
        planf = st.file_uploader("파일 첨부", type=["hwpx", "docx", "pdf", "txt", "md", "csv", "xlsx"],
                                 key="report_plan")
        if planf is not None:
            ptext = read_uploaded_text(planf, limit=20000)
            if ptext.startswith("⚠️"):
                st.warning(ptext)
            else:
                st.success(f"'{planf.name}'에서 {len(ptext):,}자를 읽었습니다.")
                mode_p = st.radio("어떻게 넣을까요?",
                                  ["원문 그대로", "AI로 핵심만 정리 (키 필요)"], horizontal=True)
                sec_title = st.text_input("섹션 제목", value="Ⅰ. 연구 개요")
                with st.expander("읽어온 내용 미리보기"):
                    st.text(ptext[:2000] + ("..." if len(ptext) > 2000 else ""))
                if st.button("➕ 계획서 내용을 보고서 앞에 넣기"):
                    content = ptext
                    if mode_p.startswith("AI"):
                        with st.spinner("AI가 계획서를 정리하는 중..."):
                            content = ai_call(
                                "다음은 농업 연구계획서(또는 결과보고서)입니다. 시험 목적, 처리 내용, "
                                "조사 항목, 방법을 보고서 서두에 넣을 수 있도록 간결한 개요로 한국어로 "
                                "정리해 주세요. 없는 내용은 지어내지 마세요.\n\n" + ptext,
                                st.session_state.get("api_key"),
                                st.session_state.get("ai_model_g"),
                                max_tokens=1500)
                    st.session_state.report_items.insert(
                        0, {"heading": sec_title, "text": content, "table": None, "image": None})
                    st.success("보고서 맨 앞에 넣었습니다!")
                    log_action("계획서를 보고서에 첨부")
                    st.rerun()

    items = st.session_state.report_items

    with st.expander("📄 적요(요약) 자동 초안 만들기"):
        st.caption("보고서에 담긴 분석 결과를 읽어 시험연구보고서의 **적요** 초안을 만듭니다. "
                   "결과를 먼저 담은 뒤 눌러 주세요.")
        ab_purpose = st.text_input("시험 목적 (선택)", key="ab_purpose",
                                   placeholder="예) 고추 신품종의 생육 및 수량 특성을 구명하고자")
        ab1, ab2 = st.columns(2)
        ab_design = ab1.text_input("시험 설계 (선택)", key="ab_design",
                                   value=st.session_state.get("m_design", "") if not str(
                                       st.session_state.get("m_design", "")).startswith("(") else "",
                                   placeholder="예) 난괴법 3반복")
        ab_cv = ab2.text_input("CV(%) (선택)", key="ab_cv", placeholder="예) 8.5")
        draft = build_abstract(items, {"purpose": ab_purpose.strip() or None,
                                       "design": ab_design.strip() or None,
                                       "cv": ab_cv.strip() or None})
        st.markdown("###### 자동 생성된 적요 초안")
        st.code(draft, language=None)
        cbtn1, cbtn2 = st.columns(2)
        if cbtn1.button("➕ 적요를 보고서 맨 앞에 넣기", width="stretch"):
            st.session_state.report_items.insert(
                0, {"heading": "적요(要約)", "text": draft, "table": None, "image": None})
            log_action("적요 자동 생성")
            st.success("보고서 맨 앞에 넣었습니다!"); st.rerun()
        if st.session_state.get("api_key"):
            if cbtn2.button("✨ AI로 다듬기", width="stretch"):
                with st.spinner("AI가 문장을 다듬는 중..."):
                    polished = ai_call(
                        "다음은 농업 시험연구보고서의 적요 초안입니다. 학술 보고서 문체("
                        "'~하였다', '~로 나타났다')로 자연스럽게 다듬어 주세요. "
                        "새로운 사실을 추가하지 말고, 있는 내용만 정리하세요.\n\n" + draft,
                        st.session_state.get("api_key"), st.session_state.get("ai_model_g"),
                        max_tokens=900)
                    st.markdown(polished)
                    ai_disclaimer()
        else:
            cbtn2.caption("AI 키를 넣으면 문장을 더 다듬을 수 있어요.")

    if not items:
        st.info("아직 담긴 분석이 없어요. 각 분석에서 '➕ 이 결과를 보고서에 담기'를 눌러 추가하세요.")
    else:
        st.write(f"**현재 담긴 분석: {len(items)}개**")
        for i, it in enumerate(items):
            c1, c2 = st.columns([6, 1])
            blks = it.get("blocks")
            if blks:
                nt = sum(1 for b in blks if b.get("table") is not None)
                ni = sum(1 for b in blks if b.get("image"))
                info = f"　📊표 {nt}개" + (f"　🖼️그림 {ni}개" if ni else "")
            else:
                info = ("　📊표" if it.get("table") is not None else "") + \
                       ("　🖼️그림" if it.get("image") else "")
            c1.write(f"{i+1}. {it['heading']}{info}")
            if c2.button("삭제", key=f"rm_{i}"):
                st.session_state.report_items.pop(i); st.rerun()
        _tabs, _figs = collect_captions(items)
        if _tabs or _figs:
            with st.expander(f"📑 표·그림 목차 미리보기 (표 {len(_tabs)}개, 그림 {len(_figs)}개)"):
                if _tabs:
                    st.markdown("**표 목차**")
                    st.text("\n".join(_tabs))
                if _figs:
                    st.markdown("**그림 목차**")
                    st.text("\n".join(_figs))
                if st.button("➕ 목차를 보고서 앞에 넣기", width="stretch"):
                    toc = ""
                    if _tabs: toc += "표 목차\n" + "\n".join(_tabs) + "\n\n"
                    if _figs: toc += "그림 목차\n" + "\n".join(_figs)
                    st.session_state.report_items.insert(
                        0, {"heading": "표·그림 목차", "text": toc.strip(),
                            "table": None, "image": None})
                    log_action("표·그림 목차 생성")
                    st.success("넣었습니다!"); st.rerun()
        rtitle = st.text_input("보고서 제목", value="2026 실험 통계 분석 보고서")
        st.caption("한글이 안 열리는 컴퓨터에서는 워드(docx)로 받으세요.")
        if not _HAS_DOCX:
            st.warning("⚠️ 현재 배포 서버에서 워드(docx) 저장 기능이 비활성화되어 있습니다. "
                       "한글(hwpx)을 이용하거나 관리자에게 문의하세요.")
        gen = st.checkbox("📄 보고서 파일 만들기", key="gen_report",
                          help="체크하면 문서를 생성합니다. (체크 전에는 만들지 않아 화면이 빠릅니다)")
        c1, c2, c3 = st.columns(3)
        if gen:
            with st.spinner("보고서를 만드는 중..."):
                with c1:
                    st.download_button("📘 한글(hwpx)", build_report_hwpx(items, rtitle),
                                       "통계분석.hwpx", width="stretch")
                with c2:
                    if _HAS_DOCX:
                        st.download_button("📝 워드(docx)", build_report_docx(items, rtitle),
                                           "통계분석.docx", width="stretch")
                    else:
                        st.caption("워드(docx): 현재 배포 환경에서 비활성화")
        with c3:
            if st.button("🗑️ 전체 비우기", width="stretch"):
                st.session_state.report_items = []; st.rerun()
